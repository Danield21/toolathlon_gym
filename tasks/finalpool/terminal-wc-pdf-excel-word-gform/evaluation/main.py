"""Evaluation for terminal-wc-pdf-excel-word-gform.
Checks:
1. Product_Quality_Audit.xlsx with 4 sheets and correct data
2. Quality_Audit_Report.docx with required sections
3. Google Form "Supplier Quality Feedback" with 5+ questions
4. quality_audit.py script exists
"""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
BLOCKING_FAIL_COUNT = 0


def check(name, condition, detail="", runtime_only=False):
    global PASS_COUNT, FAIL_COUNT, BLOCKING_FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        if not runtime_only:
            BLOCKING_FAIL_COUNT += 1
        suffix = " (runtime-only)" if runtime_only else ""
        print(f"  [FAIL] {name}{suffix}: {str(detail)[:200]}")


EXPECTED_CATEGORIES = {"Audio", "Cameras", "Electronics", "Home Appliances", "TV & Home Theater", "Watches"}

# GT-derived expected scores (tol 1.0 on score; exact risk level)
EXPECTED_SCORES = {
    "Audio":            {"score": 81.7, "risk": "Medium"},
    "Cameras":          {"score": 93.8, "risk": "Low"},
    "Electronics":      {"score": 83.34, "risk": "Medium"},
    "Home Appliances":  {"score": 94.4, "risk": "Low"},
    "TV & Home Theater":{"score": 95.6, "risk": "Low"},
    "Watches":          {"score": 81.46, "risk": "Medium"},
}


def check_excel(workspace):
    print("\n=== Check 1: Product_Quality_Audit.xlsx ===")
    path = os.path.join(workspace, "Product_Quality_Audit.xlsx")
    if not os.path.exists(path):
        check("Excel file exists", False, f"Not found at {path}")
        return
    check("Excel file exists", True)

    wb = openpyxl.load_workbook(path)
    sheets = wb.sheetnames
    check("Has at least 4 sheets", len(sheets) >= 4, f"Found {len(sheets)}: {sheets}")

    sheets_lower = [s.lower() for s in sheets]

    # Review_Summary sheet
    rs_idx = next((i for i, s in enumerate(sheets_lower) if "review" in s or "summary" in s), 0)
    ws = wb[sheets[rs_idx]]
    rows = list(ws.iter_rows(values_only=True))
    data_rows = [r for r in rows[1:] if any(c for c in r)]
    check("Review_Summary has exactly 6 category rows", len(data_rows) == 6, f"Found {len(data_rows)}")

    # All 6 categories present
    found_categories = set()
    for row in data_rows:
        if row and row[0]:
            name = str(row[0]).strip()
            for ec in EXPECTED_CATEGORIES:
                if name.lower() == ec.lower():
                    found_categories.add(ec)
                    break
    check("Review_Summary has all 6 expected categories",
          found_categories == EXPECTED_CATEGORIES,
          f"Missing: {EXPECTED_CATEGORIES - found_categories}")

    # Refund_Analysis sheet
    ra_idx = next((i for i, s in enumerate(sheets_lower) if "refund" in s or "order" in s), 1)
    if ra_idx < len(sheets):
        ws2 = wb[sheets[ra_idx]]
        rows2 = list(ws2.iter_rows(values_only=True))
        data_rows2 = [r for r in rows2[1:] if any(c for c in r)]
        check("Refund_Analysis has at least 5 status rows", len(data_rows2) >= 5, f"Found {len(data_rows2)}")
        all_text2 = " ".join(str(c) for r in rows2 for c in r if c).lower()
        check("Contains refunded status", "refunded" in all_text2, f"Text: {all_text2[:120]}")
        check("Contains completed status", "completed" in all_text2, f"Text: {all_text2[:120]}")

        # Per-status order_count + total_amount validation (GT-derived)
        EXPECTED_REFUND = {
            "cancelled": {"order_count": 8,  "total_amount": 4547.42},
            "completed": {"order_count": 72, "total_amount": 30296.82},
            "failed":    {"order_count": 5,  "total_amount": 2130.43},
            "on-hold":   {"order_count": 15, "total_amount": 6845.7},
            "pending":   {"order_count": 13, "total_amount": 4514.01},
        }
        refund_map = {}
        for row in data_rows2:
            if row and row[0]:
                refund_map[str(row[0]).strip().lower()] = row
        for status, exp in EXPECTED_REFUND.items():
            row = refund_map.get(status)
            if row is None:
                check(f"Refund_Analysis has '{status}' status", False,
                      f"Statuses: {list(refund_map.keys())}")
                continue
            # order_count (col 1) exact
            try:
                ac = int(float(row[1]))
                check(f"Refund '{status}' order_count=={exp['order_count']}",
                      ac == exp['order_count'],
                      f"Got {ac}")
            except (TypeError, ValueError, IndexError):
                check(f"Refund '{status}' order_count parseable", False, f"Row: {row}")
            # total_amount (col 2) within tol
            try:
                at = float(row[2])
                check(f"Refund '{status}' total_amount ~{exp['total_amount']}",
                      abs(at - exp['total_amount']) <= 1.0,
                      f"Got {at}")
            except (TypeError, ValueError, IndexError):
                check(f"Refund '{status}' total_amount parseable", False, f"Row: {row}")

    # Quality_Scorecard sheet
    qs_idx = next((i for i, s in enumerate(sheets_lower) if ("quality" in s and "score" in s) or "scorecard" in s), 2)
    if qs_idx < len(sheets):
        ws3 = wb[sheets[qs_idx]]
        rows3 = list(ws3.iter_rows(values_only=True))
        data_rows3 = [r for r in rows3[1:] if any(c for c in r)]
        check("Quality_Scorecard has exactly 6 rows", len(data_rows3) == 6, f"Found {len(data_rows3)}")

        # Per-category scores + risk levels
        score_map = {}
        for row in data_rows3:
            if row and row[0]:
                name = str(row[0]).strip()
                score_map[name.lower()] = row
        for cat, exp in EXPECTED_SCORES.items():
            row = score_map.get(cat.lower())
            if row is None:
                check(f"Quality_Scorecard has category '{cat}'", False,
                      f"Categories: {list(score_map.keys())}")
                continue
            # Score (col 1) within ±1.0
            try:
                actual_score = float(row[1])
                check(f"Quality_Scorecard {cat} score ~{exp['score']}",
                      abs(actual_score - exp['score']) <= 1.0,
                      f"Got {actual_score}")
            except (TypeError, ValueError, IndexError):
                check(f"Quality_Scorecard {cat} score parseable", False, f"Row: {row}")
            # Risk level (col 2) exact case-insensitive
            try:
                actual_risk = str(row[2]).strip().lower()
                check(f"Quality_Scorecard {cat} risk_level == {exp['risk']}",
                      actual_risk == exp['risk'].lower(),
                      f"Got {row[2]}")
            except IndexError:
                check(f"Quality_Scorecard {cat} risk_level present", False, f"Row: {row}")

    # Survey_Questions sheet
    sq_idx = next((i for i, s in enumerate(sheets_lower) if "survey" in s or "question" in s), 3)
    if sq_idx < len(sheets):
        ws4 = wb[sheets[sq_idx]]
        rows4 = list(ws4.iter_rows(values_only=True))
        data_rows4 = [r for r in rows4[1:] if any(c for c in r)]
        check("Survey_Questions has at least 5 rows", len(data_rows4) >= 5, f"Found {len(data_rows4)}")
        # Topics coverage (use root-form tolerance)
        all_q_text = " ".join(str(c).lower() for r in data_rows4 for c in r if c)
        required_topics = [
            ("defect", ["defect"]),
            ("response", ["respon"]),  # response/respond
            ("improvement", ["improv"]),
            ("packaging", ["packag"]),
            ("satisfaction", ["satisf", "overall quality", "overall rating"]),  # overall quality counts
        ]
        for topic, roots in required_topics:
            ok = any(root in all_q_text for root in roots)
            check(f"Survey covers '{topic}' topic",
                  ok,
                  f"Q text sample: {all_q_text[:200]}")


def check_word(workspace):
    print("\n=== Check 2: Quality_Audit_Report.docx ===")
    path = os.path.join(workspace, "Quality_Audit_Report.docx")
    if not os.path.exists(path):
        check("Word file exists", False, f"Not found at {path}")
        return
    check("Word file exists", True)

    try:
        from docx import Document
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        # Executive summary explicit
        check("Contains 'executive summary' phrase", "executive summary" in all_text,
              f"Text: {all_text[:200]}")
        # Review analysis by category
        check("Contains review analysis by category",
              ("review" in all_text and "category" in all_text) or "review analysis" in all_text,
              f"Text: {all_text[:200]}")
        # Refund section + concrete terms
        check("Contains refund patterns/analysis", "refund" in all_text,
              f"Text: {all_text[:200]}")
        # Recommendations with at least 3 items
        # Check recommendations by looking at all paragraphs (not just concatenated text)
        rec_idx = all_text.find("recommend")
        if rec_idx >= 0:
            # Count paragraphs after 'recommendations' heading
            rec_para_idx = None
            for i, p in enumerate(doc.paragraphs):
                if "recommend" in p.text.lower():
                    rec_para_idx = i
                    break
            if rec_para_idx is not None:
                # Count non-empty paragraphs after the heading (each is an action item)
                following = [p for p in doc.paragraphs[rec_para_idx+1:] if p.text.strip()]
                # Also consider paragraphs with numeric/bullet markers or meaningful length
                items = [p for p in following if len(p.text.strip()) > 15]
                check("Recommendations section has >=3 action items",
                      len(items) >= 3,
                      f"Found {len(items)} items after 'recommendations'")
            else:
                check("Recommendations section has >=3 action items", False,
                      "No paragraph with 'recommendations'")
        else:
            check("Contains 'recommendations' section", False,
                  f"Text: {all_text[:200]}")
    except ImportError:
        check("python-docx available", False, "python-docx not installed")


def check_gform():
    print("\n=== Check 3: Supplier Quality Feedback Form ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gform.forms WHERE lower(title) LIKE '%%supplier%%quality%%'")
    forms = cur.fetchall()
    check("Supplier Quality Feedback form exists", len(forms) >= 1,
          f"Found forms: {[f[1] for f in forms]}", runtime_only=True)

    if forms:
        form_id = forms[0][0]
        cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (form_id,))
        qcount = cur.fetchone()[0]
        check("Form has at least 5 questions", qcount >= 5,
              f"Found {qcount} questions", runtime_only=True)

        cur.execute("SELECT title, question_type FROM gform.questions WHERE form_id = %s", (form_id,))
        questions = cur.fetchall()
        types = set(q[1] for q in questions)
        check("Form has multiple question types", len(types) >= 2,
              f"Types: {types}", runtime_only=True)

        # Must cover topics: defect, response(respond), improvement, packaging, satisfaction
        q_titles = " ".join((q[0] or "").lower() for q in questions)
        topic_roots = {
            "defect": ["defect"],
            "response": ["respon"],  # response or respond
            "improvement": ["improv"],
            "packaging": ["packag"],
            "satisfaction": ["satisf", "overall quality", "overall rating"],
        }
        for topic, roots in topic_roots.items():
            ok = any(root in q_titles for root in roots)
            check(f"Form covers '{topic}' topic",
                  ok,
                  f"Questions: {[q[0] for q in questions]}",
                  runtime_only=True)

    cur.close()
    conn.close()


def check_script(workspace):
    print("\n=== Check 4: quality_audit.py ===")
    path = os.path.join(workspace, "quality_audit.py")
    if not os.path.exists(path):
        check("quality_audit.py exists", False, f"Not at {path}")
        return
    check("quality_audit.py exists", True)
    # Must have meaningful non-trivial content
    try:
        content = open(path, "r", encoding="utf-8", errors="replace").read()
    except Exception as e:
        check("quality_audit.py readable", False, str(e))
        return
    check("quality_audit.py is not trivial (>=500 chars)",
          len(content) >= 500,
          f"Length={len(content)}")
    # Must reference key operations
    c_lower = content.lower()
    check("Script mentions category/categories",
          "categor" in c_lower,
          "Missing category reference")
    check("Script computes quality score formula (rating & penalty)",
          ("avg_rating" in c_lower or "average_rating" in c_lower or "rating" in c_lower)
          and ("pct" in c_lower or "percent" in c_lower or "below" in c_lower),
          "Missing rating/percentage references")
    check("Script classifies risk levels (low/medium/high)",
          all(lvl in c_lower for lvl in ["low", "medium", "high"]),
          "Risk levels not all mentioned")


def check_reverse_validation(workspace):
    """Verify things that should NOT exist in the output."""
    print("\n=== Reverse Validation ===")
    path = os.path.join(workspace, "Product_Quality_Audit.xlsx")
    if os.path.exists(path):
        wb = openpyxl.load_workbook(path, data_only=True)
        # No unexpected sheets beyond the 4 required - collect ALL unexpected, no break
        expected_keywords = {"review", "summary", "refund", "order", "quality", "scorecard", "survey", "question"}
        unexpected = []
        for s in wb.sheetnames:
            s_lower = s.lower()
            if not any(kw in s_lower for kw in expected_keywords):
                unexpected.append(s)
        check("No unexpected sheets in Excel", len(unexpected) == 0,
              f"Unexpected: {unexpected}")

        # Review ratings should not be negative
        sheets_lower = [s.lower() for s in wb.sheetnames]
        rs_idx = next((i for i, s in enumerate(sheets_lower) if "review" in s or "summary" in s), 0)
        ws = wb[wb.sheetnames[rs_idx]]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        has_negative = False
        for r in rows:
            for c in r:
                try:
                    v = float(str(c).replace(',', '').replace('%', '').replace('$', '').strip()) if c else None
                except (ValueError, TypeError):
                    v = None
                if v is not None and v < 0:
                    has_negative = True
                    break
            if has_negative:
                break
        check("No negative values in Review_Summary", not has_negative,
              "Found negative value in summary data")

    # Google Form: no duplicate forms with same name
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT title, COUNT(*) FROM gform.forms
            WHERE lower(title) LIKE '%%supplier%%quality%%'
            GROUP BY title HAVING COUNT(*) > 1
        """)
        dupes = cur.fetchall()
        check("No duplicate supplier quality forms", len(dupes) == 0,
              f"Duplicates: {dupes}")
        cur.close()
        conn.close()
    except Exception:
        pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace)
    check_word(args.agent_workspace)
    check_gform()
    check_script(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if BLOCKING_FAIL_COUNT == 0:
        print(f"PASS (runtime-only fails: {FAIL_COUNT - BLOCKING_FAIL_COUNT})")
        sys.exit(0)
    else:
        print(f"FAIL ({BLOCKING_FAIL_COUNT} blocking checks failed; "
              f"runtime-only: {FAIL_COUNT - BLOCKING_FAIL_COUNT})")
        sys.exit(1)


if __name__ == "__main__":
    main()
