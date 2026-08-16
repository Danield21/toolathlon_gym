"""Evaluation for student-success-tracking-multi-touch.

The task asks the agent to integrate multiple support channels and produce
deliverables across:
  - learning management system (canvas) -> engagement extraction
  - email logs (advisor touchpoints)
  - data warehouse (snowflake) -> outcomes
  - knowledge management system (notion) -> intervention documentation
  - calendar (google_calendar) -> advisor meeting events
  - report document (word) -> final synthesis
  - email -> stakeholder communication

This evaluator verifies concrete artifacts that a meaningful execution of
the workflow MUST produce. It does NOT just check file existence.
"""
import argparse
import json
import os
import sys
from pathlib import Path

import psycopg2

DB = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


# Required topical keywords for the synthesis report
REPORT_KEYWORDS = [
    "engagement",
    "retention",
    "intervention",
]


def check_report(agent_workspace):
    print("\n=== Checking final report (Word) ===")
    ws = Path(agent_workspace)
    docx_files = list(ws.glob("*.docx")) + list(ws.glob("**/*.docx"))
    check("At least one Word document produced", len(docx_files) >= 1,
          f"found {[str(p) for p in docx_files][:5]}")
    if not docx_files:
        return
    # Check the largest doc has substantive content + topical keywords
    largest = max(docx_files, key=lambda p: p.stat().st_size)
    try:
        from docx import Document
        doc = Document(str(largest))
    except Exception as e:
        check("Word file readable", False, str(e))
        return
    full_text = ""
    for p in doc.paragraphs:
        full_text += p.text + "\n"
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"
    text_lower = full_text.lower()
    check(f"Report '{largest.name}' has >= 500 characters of content",
          len(full_text) >= 500, f"got {len(full_text)} chars")
    for kw in REPORT_KEYWORDS:
        check(f"Report mentions topical keyword '{kw}'", kw in text_lower,
              f"missing in {largest.name}")
    check("Report mentions student/students", "student" in text_lower)
    # Synthesis: must mention at least one specific intervention/recommendation
    check("Report contains recommendation/insight section",
          ("recommend" in text_lower or "insight" in text_lower or "strategy" in text_lower),
          "missing recommend/insight/strategy")

    # Value-level: synthesis report should have multiple section-level headings
    # (mirrors GT structure: Engagement / Advisor / Outcome / Recommendations).
    heading_paras = sum(1 for p in doc.paragraphs
                        if p.style and p.style.name and 'heading' in p.style.name.lower()
                        and p.text.strip())
    # Topical heading-like markers in flat text as fallback (some agents use bold instead of styles)
    heading_topics = sum(1 for kw in ('engagement', 'advisor', 'outcome', 'recommend',
                                       'intervention', 'finding', 'analysis')
                         if kw in text_lower)
    check("Report has structured sections (>=2 headings or >=4 topical headers)",
          heading_paras >= 2 or heading_topics >= 4,
          f"heading_paras={heading_paras}, heading_topics={heading_topics}")

    # Phase-5 statistical-analysis content: task asks for correlation /
    # regression / quantitative analysis. Report should reference at least one.
    stat_terms = ('correlat', 'regression', 'statistic', 'coefficient',
                  'quantitative', 'percentage', 'percent', 'rate', 'ratio')
    stat_hits = sum(1 for term in stat_terms if term in text_lower)
    check("Report references statistical/quantitative analysis (>=1 term)",
          stat_hits >= 1,
          f"matched terms (need >=1 of {stat_terms}): {stat_hits}")

    # Numeric content: meaningful synthesis report should contain numeric figures
    # (counts/rates/percentages tied to retention, intervention, or engagement).
    import re
    numeric_tokens = re.findall(r'\b\d+\.?\d*%?\b', full_text)
    check("Report contains numeric figures (counts/percentages)",
          len(numeric_tokens) >= 3,
          f"found {len(numeric_tokens)} numeric tokens")


def check_notion(is_gt_self_test=False):
    print("\n=== Checking Notion knowledge base ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
    except psycopg2.Error as e:
        check("DB connect", False, str(e))
        return
    cur.execute("SELECT id, properties FROM notion.pages")
    pages = cur.fetchall()
    cur.close()
    conn.close()
    # In GT self-test mode: notion is not populated by GT files alone, so skip notion checks entirely.
    if is_gt_self_test:
        print("  [WARN] Notion checks skipped (GT self-test, non-blocking)")
        return
    check("At least 1 Notion page created", len(pages) >= 1, f"found {len(pages)} pages")
    if not pages:
        return
    # Look for student-success related content in titles/properties
    matched = 0
    for pid, props in pages:
        ptext = json.dumps(props or {}).lower()
        if any(kw in ptext for kw in ("student", "success", "intervention", "support", "advis", "retention")):
            matched += 1
    check("At least 1 Notion page has student-success/intervention/advisor content",
          matched >= 1, f"matched {matched}/{len(pages)} pages")


def check_gcal(is_gt_self_test=False):
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
    except psycopg2.Error as e:
        check("DB connect", False, str(e))
        return
    cur.execute("SELECT id, summary, description, start_datetime FROM gcal.events")
    events = cur.fetchall()
    cur.close()
    conn.close()
    # In GT self-test mode: gcal is not populated by GT files alone, so skip gcal checks entirely.
    if is_gt_self_test:
        print("  [WARN] Calendar checks skipped (GT self-test, non-blocking)")
        return
    check("At least 1 calendar event created", len(events) >= 1, f"found {len(events)}")
    if not events:
        return
    matched = 0
    for eid, summary, desc, start in events:
        text = f"{summary or ''} {desc or ''}".lower()
        if any(kw in text for kw in ("advis", "meeting", "student", "support", "interven")):
            matched += 1
    check("At least 1 calendar event is advisor/meeting/student related",
          matched >= 1, f"matched {matched}/{len(events)}")


def check_email(is_gt_self_test=False):
    print("\n=== Checking Email (sent communications) ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
    except psycopg2.Error as e:
        check("DB connect", False, str(e))
        return
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages WHERE folder_id IN (SELECT id FROM email.folders WHERE LOWER(name) IN ('sent', 'drafts'))")
    sent = cur.fetchall()
    cur.close()
    conn.close()
    # In GT self-test mode: email is not populated by GT files alone, so skip email checks entirely.
    if is_gt_self_test:
        print("  [WARN] Email checks skipped (GT self-test, non-blocking)")
        return
    check("At least 1 email composed/sent by agent", len(sent) >= 1, f"found {len(sent)} sent")
    if not sent:
        return
    matched = 0
    for subj, frm, to, body in sent:
        text = f"{subj or ''} {body or ''}".lower()
        if any(kw in text for kw in ("student", "success", "interven", "support", "retention", "advis")):
            matched += 1
    check("At least 1 sent email is student-success topical",
          matched >= 1, f"matched {matched}/{len(sent)}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    # Detect GT self-test (V1 parity test).
    try:
        gt_canon = os.path.realpath(args.groundtruth_workspace) if args.groundtruth_workspace else ""
        ag_canon = os.path.realpath(args.agent_workspace) if args.agent_workspace else ""
        is_gt_self_test = bool(gt_canon) and (gt_canon == ag_canon)
    except Exception:
        is_gt_self_test = False

    check_report(args.agent_workspace)
    check_notion(is_gt_self_test)
    check_gcal(is_gt_self_test)
    check_email(is_gt_self_test)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    all_ok = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": all_ok}, f)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
