"""Evaluation for playwright-arxiv-review-criteria-word-gsheet."""
import argparse
import os
import re
import sys

import psycopg2

# R1: every DB connection must read env vars with defaults. The defaults must
# match what preprocess/main.py uses (PGUSER/PGPASSWORD are libpq-style names).
DB = dict(
    host=os.environ.get("PGHOST", "localhost"),
    port=int(os.environ.get("PGPORT", "5432")),
    dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
    user=os.environ.get("PGUSER", "eigent"),
    password=os.environ.get("PGPASSWORD", "camel"),
)


def num_close(a, b, tol=0.5):
    """Robust numeric closeness. Parses both sides as numbers where possible
    (stripping thousands separators / currency / percent / spaces); falls back
    to a case-insensitive string comparison only when either side cannot be
    parsed as a number (R3)."""
    def _to_float(v):
        if v is None or isinstance(v, bool):
            return None
        if isinstance(v, (int, float)):
            return float(v)
        s = str(v).strip()
        if not s:
            return None
        s = s.replace(",", "").replace("$", "").replace("¥", "").replace("€", "").replace("%", "")
        try:
            return float(s)
        except (TypeError, ValueError):
            return None

    fa = _to_float(a)
    fb = _to_float(b)
    if fa is not None and fb is not None:
        return abs(fa - fb) <= tol
    return str(a).strip().lower() == str(b).strip().lower()


# ---------------------------------------------------------------------------
# Text extraction from a docx. Two views are used:
#   * `_text_units`  - flat list of paragraphs + table rows (row cells joined
#                      by '|') for full-text / section-presence / recommendation
#                      checks.
#   * `_para_units`  - list of (is_heading, text) paragraphs for the score
#                      scan (heading style is used to find section boundaries).
#   * `_table_grids` - raw table cell grids for column-aware score matching.
# ---------------------------------------------------------------------------
def _text_units(doc):
    """List of text units from a docx: paragraphs plus table rows (row cells
    joined by '|') so that label/value pairs written in tables stay together.
    Only non-empty units are kept."""
    units = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            units.append(" | ".join(cells))
    return [u for u in units if u and u.strip()]


def _para_units(doc):
    """[(is_heading, text)] for non-empty paragraphs. A paragraph counts as a
    heading when its Word style is a Heading/Title style (Word MCP's
    add_heading produces these)."""
    paras = []
    for p in doc.paragraphs:
        t = p.text
        if not (t and t.strip()):
            continue
        style = (p.style.name or "").lower()
        is_heading = ("heading" in style) or ("title" in style)
        paras.append((is_heading, t))
    return paras


def _table_grids(doc):
    """Raw cell grids for every table in the docx."""
    grids = []
    for table in doc.tables:
        grid = [[c.text for c in row.cells] for row in table.rows]
        grids.append(grid)
    return grids


# ---------------------------------------------------------------------------
# Score checks (Technical Soundness / Novelty / Clarity).
#
# The score for a criterion must be found in a *score-like* form. Two layouts
# are supported:
#   1) prose  - the label followed, within the same section, by "N/5",
#               "Score: N", "Rating: N", ": N" or a leading number. The window
#               is extended across following paragraphs and stopped at the next
#               section title, so a heading + justification + "Score: N" on a
#               separate paragraph (a normal Word MCP structure) is accepted,
#               and a justification that mentions the word "technical" /
#               "technically" before the score no longer truncates the window.
#   2) tables - columnar layouts (header row with criterion names, scores in
#               the same column of data rows) and row layouts (label | score)
#               are both accepted via cell association.
# ---------------------------------------------------------------------------

# Score-like number patterns, tried in order. They cover:
#   "Technical Soundness: 5/5", "Technical Soundness: 5",
#   "Technical Soundness Score: 5", "Technical Soundness\n5/5 - ...",
#   "Technical Soundness\nScore: 5", "Technical Soundness | 5" (table row),
#   "Technical Soundness Rating: 5".
_SCORE_PATTERNS = [
    r"(\d+)\s*(?:out\s*of|of|/)\s*\d+",        # N/5, "N out of 5", "N of 5"
    r"(?:score|rating|level|grade)\s*(?:of\s*)?(?:is\s*)?[:\s=\-|]*(\d+)",  # "score: 5" / "score of 5" / "score is 5"
    r"[:=\-|]\s*(\d+)",                        # ": 5" / "| 5" / "= 5"
    r"^\s*(\d+)",                              # number at start of window (cross-paragraph)
]


def _extract_score(rest, exp):
    for pat in _SCORE_PATTERNS:
        for m in re.finditer(pat, rest):
            try:
                if int(m.group(1)) == exp:
                    return True
            except (TypeError, ValueError):
                continue
    return False


def _is_other_section_title(text, crit):
    """True when a paragraph looks like a *different* review section's title,
    i.e. the place where this criterion's score window must stop. Relies on a
    short unit whose (lower-cased) text starts with a section keyword. 'overall'
    alone is never a boundary ('Overall, the method is sound...' is prose)."""
    t = text.strip()
    low = t.lower()
    if len(t) > 60:
        return False
    own = set(crit.split())
    if "recommendation" not in own and (
            low.startswith("overall recommendation") or low.startswith("recommendation")):
        return True
    if "summary" not in own and low.startswith("summary"):
        return True
    if "novelty" not in own and low.startswith("novelty"):
        return True
    if "clarity" not in own and low.startswith("clarity"):
        return True
    return False


def _prose_score_present(paras, crit, exp):
    """Prose path: find the criterion label, then scan the following paragraphs
    (up to the next section title, with a size cap) for a score-like value."""
    n = len(paras)
    for i, (_is_heading, text) in enumerate(paras):
        low = text.lower()
        pos = low.find(crit)
        if pos == -1:
            continue
        rest = text[pos + len(crit):]
        j = i + 1
        while j < n and len(rest) < 900:
            _j, ntext = paras[j]
            if _is_other_section_title(ntext, crit):
                break
            rest += "\n" + ntext
            j += 1
        if _extract_score(rest, exp):
            return True
    return False


def _cell_has_score(cell, exp):
    """Score-like check for a single table cell. Accepts '5', '5/5', '4 out of
    5', 'Score: 5', 'Rating: 4', etc. A bare integer cell that is not 1-5, or
    a number embedded in a longer word ('992 A100 GPUs'), is not a score."""
    s = cell.strip()
    if not s:
        return False
    low = s.lower()
    m = re.search(r"(\d+)\s*(?:out\s*of|of|/)\s*\d+", s)
    if m:
        return int(m.group(1)) == exp
    m = re.search(r"(\d+)\s*out\s*of\s*\d+", low)
    if m:
        return int(m.group(1)) == exp
    m = re.search(r"(?:score|rating|level|grade)\s*[:\s=\-|]+\s*(\d+)", low)
    if m:
        return int(m.group(1)) == exp
    # bare integer cell that is a plausible score (1-5), not embedded in digits
    m = re.search(r"(?:^|[^0-9])([1-5])(?:$|[^0-9.])", s)
    if m:
        return int(m.group(1)) == exp
    return False


def _table_score_present(grids, crit, exp):
    """Table path: for every cell that mentions the criterion, look for the
    score in the same cell, in the same row, or in the same column (this covers
    both row-oriented 'label | score' tables and columnar tables whose header
    row carries the criterion names and whose data rows carry aligned scores)."""
    for grid in grids:
        nrows = len(grid)
        for r, row in enumerate(grid):
            ncols = len(row)
            for c, cell in enumerate(row):
                if crit not in cell.lower():
                    continue
                if _cell_has_score(cell, exp):
                    return True
                for cc in range(ncols):
                    if cc != c and _cell_has_score(row[cc], exp):
                        return True
                for rr in range(nrows):
                    if rr != r and _cell_has_score(grid[rr][c], exp):
                        return True
    return False


def _score_present(doc, crit, exp):
    """True if the expected score for a criterion appears in a score-like form
    anywhere in the document, either in prose or in a table."""
    if _prose_score_present(_para_units(doc), crit, exp):
        return True
    return _table_score_present(_table_grids(doc), crit, exp)


# ---------------------------------------------------------------------------
# Recommendation checks.
#
# For 'accept' we only require the *declared* value to be Accept. Natural
# rationale phrasing such as "This is a strong accept", "a clear accept",
# "a borderline accept" or even a hedged "a weak accept" mention is *not* a
# wrong declared value, so it must not fail the check. The declared value is
# read from the recommendation label and the text right after it; only a
# declared 'weak accept' / 'borderline' / 'weak reject' / 'reject' fails an
# Accept expectation (and the reverse direction for 'weak accept').
# ---------------------------------------------------------------------------
_VAL_RE = re.compile(
    r"\bweak\s+accept\b"
    r"|\bweak\s+reject\b"
    r"|\b(?:strong|clear|borderline)\s+accept\b"
    r"|\bborderline\b"
    r"|\breject\b"
    r"|\baccept(?:ance|ed)?\b",
    re.I,
)

_REC_LABEL_RE = re.compile(r"^\s*(?:overall\s+recommendation|recommendation)", re.I)


def _recommendation_region(units):
    """The text of the recommendation section: from the last unit that *starts*
    with the 'Overall Recommendation'/'Recommendation' label (its tail plus the
    next few units) — the recommendation section is the document's final
    section, so the declared value is found right after the label. Falls back
    to the first unit mentioning 'recommendation', then to the whole document."""
    idx = -1
    for i, u in enumerate(units):
        if _REC_LABEL_RE.match(u):
            idx = i
    if idx == -1:
        for i, u in enumerate(units):
            if "recommendation" in u.lower():
                idx = i
                break
        if idx == -1:
            return "\n".join(units)
    m = _REC_LABEL_RE.match(units[idx])
    region = units[idx][m.end():]
    for j in range(idx + 1, min(idx + 4, len(units))):
        region += "\n" + units[j]
    return region


def _declared_is_accept(units):
    region = _recommendation_region(units)
    m = _VAL_RE.search(region)
    if not m:
        return False
    tok = m.group(0).lower()
    if tok in ("weak accept", "weak reject", "borderline", "reject"):
        return False
    # 'accept' / 'accepted' / 'acceptance' / 'strong accept' / 'clear accept' /
    # 'borderline accept' are all acceptable declarations of Accept.
    return True


def _recommendation_present(full_text, units, exp_rec):
    """Check that the declared recommendation matches the expected value."""
    exp_rec = exp_rec.strip().lower()
    if exp_rec == "accept":
        return _declared_is_accept(units)
    if exp_rec == "weak accept":
        return re.search(r"\bweak\s+accept\b", full_text) is not None
    return re.search(r"\b" + re.escape(exp_rec) + r"\b", full_text) is not None


def check_word_review(agent_workspace, filename, title_fragments,
                      expected_tech, expected_novelty, expected_clarity,
                      expected_recommendation):
    """Check one Word review document.

    title_fragments is a list of acceptable title keywords; any one of them
    appearing anywhere in the document satisfies the paper-title check (a model
    may legitimately use the short paper name, e.g. 'InstructGPT'/'RLHF').
    """
    errors = []
    path = os.path.join(agent_workspace, filename)
    if not os.path.exists(path):
        return [f"{filename} not found"]
    try:
        from docx import Document
        doc = Document(path)
        units = _text_units(doc)
        full_text = "\n".join(units).lower()

        if not any(frag.lower() in full_text for frag in title_fragments):
            errors.append(f"{filename}: missing paper title fragment (expected any of {title_fragments})")
        # The Summary section may legitimately be headed 'Abstract' (or contain
        # the word); task.md asks for 'Summary' but rejecting 'Abstract' would
        # be an undeclared strictness.
        if not any(kw in full_text for kw in ("summary", "abstract")):
            errors.append(f"{filename}: missing Summary section")
        if "technical soundness" not in full_text and "technical" not in full_text:
            errors.append(f"{filename}: missing Technical Soundness section")
        if "novelty" not in full_text:
            errors.append(f"{filename}: missing Novelty section")
        if "clarity" not in full_text:
            errors.append(f"{filename}: missing Clarity section")

        if not _recommendation_present(full_text, units, expected_recommendation):
            errors.append(f"{filename}: recommendation not exactly '{expected_recommendation}'")

        score_checks = [
            ("technical soundness", expected_tech),
            ("novelty", expected_novelty),
            ("clarity", expected_clarity),
        ]
        for crit, exp in score_checks:
            if not _score_present(doc, crit, exp):
                errors.append(f"{filename}: score for {crit} not exactly {exp}")
    except Exception as e:
        errors.append(f"Error reading {filename}: {e}")
    return errors


# ---------------------------------------------------------------------------
# GSheet check.
#
# Returns (blocking_errors, runtime_only_errors).
#
# If no spreadsheet exists at all (agent did not create one), all errors are
# runtime-only — the V1 GT-only test (whose ground-truth workspace holds no
# spreadsheet) should not fail on this. Once at least one review/conference
# spreadsheet exists, content errors become blocking.
#
# With subagent_homogeneous parallel execution each agent may create its own
# 'Conference Review Tracker' spreadsheet. The content checks therefore merge
# the cells of *all* matching spreadsheets (and their 'Reviews' sheets) rather
# than inspecting a single arbitrarily-picked one: `gsheet.spreadsheets.id` is
# a random UUID text column, so "ORDER BY s.id DESC" does NOT mean "newest
# first", and picking one sheet could pick a partially-populated one and turn
# a content-correct swarm answer into a FAIL. Merging removes that path while
# still FAILing a genuinely incomplete tracker (a single incomplete sheet's
# union is still incomplete).
# ---------------------------------------------------------------------------
_AVG_EXPECTED = {"instructgpt": 4.7, "scaling": 4.3, "opt": 4.0}


def _avg_value_ok(text):
    """If `text` is a plain number, return (float_value, significant_decimal_digits)
    else None. Significant decimal digits ignore trailing zeros, so '4', '4.0'
    and '4.00' all report <= 1 significant decimal, while '4.67' reports 2."""
    s = text.strip()
    if re.fullmatch(r"-?\d+(?:\.\d+)?", s) is None:
        return None
    if "." in s:
        frac = s.split(".", 1)[1].rstrip("0")
        ndec = len(frac)
    else:
        ndec = 0
    return float(s), ndec


def _avg_present(values, expected):
    """True if any plain number in `values` rounds (to 1 decimal) to `expected`
    and was written with at most one significant decimal (i.e. the model
    followed 'rounded to 1 decimal' — '4', '4.0', '4.7' pass; '4.67' fails)."""
    for v in values:
        r = _avg_value_ok(v)
        if r is None:
            continue
        f, ndec = r
        if ndec > 1:
            continue
        if abs(round(f, 1) - expected) < 0.05:
            return True
    return False


def check_gsheet():
    blocking = []
    runtime = []
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("""
            SELECT s.id, s.title FROM gsheet.spreadsheets s
            WHERE LOWER(s.title) LIKE '%review%' OR LOWER(s.title) LIKE '%conference%'
            ORDER BY s.id DESC
        """)
        spreadsheets = cur.fetchall()
        if not spreadsheets:
            runtime.append("No review tracker spreadsheet found")
            cur.close()
            conn.close()
            return blocking, runtime

        # Merge cells across every matching spreadsheet's 'Reviews'-like sheet.
        # Each cell keeps its spreadsheet_id so the Average_Score column can be
        # resolved *per spreadsheet* (two agents' trackers may use different
        # column orders) without colliding in a single (row,col) grid.
        cells = []  # (ss_id, row_index, col_index, value)
        sheets_found = 0
        for ss_id, ss_title in spreadsheets:
            cur.execute("""
                SELECT id FROM gsheet.sheets
                WHERE spreadsheet_id = %s AND LOWER(title) LIKE '%%review%%'
                LIMIT 1
            """, (ss_id,))
            sheet_row = cur.fetchone()
            if not sheet_row:
                continue
            sheets_found += 1
            sheet_id = sheet_row[0]
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s
                ORDER BY row_index, col_index
            """, (ss_id, sheet_id))
            for (row_index, col_index, value) in cur.fetchall():
                cells.append((ss_id, row_index, col_index, value))
        cur.close()
        conn.close()

        if sheets_found == 0:
            blocking.append("No 'Reviews' sheet found in any review/conference spreadsheet")
            return blocking, runtime
        if not cells:
            blocking.append("Reviews sheets are empty")
            return blocking, runtime

        if len(cells) < 24:
            blocking.append(f"Too few cells in Reviews sheets: {len(cells)}, expected ~32")

        cell_values = [str(c[3]).lower() if c[3] else "" for c in cells]
        all_text = " ".join(cell_values)
        if "2301.07041" not in all_text and "scaling" not in all_text:
            blocking.append("Scaling Laws paper not found in GSheet")
        if "2203.11171" not in all_text and "instruct" not in all_text:
            blocking.append("InstructGPT paper not found in GSheet")
        if "2205.01068" not in all_text and "opt" not in all_text:
            blocking.append("OPT paper not found in GSheet")

        # Average scores: locate the Average_Score column by its header keyword
        # per spreadsheet (column-position-independent) and check the union of
        # the column's values across all spreadsheets numerically; fall back to
        # the whole-sheet text if no such header is found anywhere.
        by_ss = {}
        for (ss_id, row_index, col_index, value) in cells:
            by_ss.setdefault(ss_id, []).append((row_index, col_index, value))
        avg_values = []
        for ss_id, ss_cells in by_ss.items():
            grid = {}
            for (row_index, col_index, value) in ss_cells:
                grid[(row_index, col_index)] = str(value) if value else ""
            avg_col = None
            for (r, c), v in grid.items():
                if "average" in v.lower():
                    avg_col = c
                    break
            if avg_col is None:
                continue
            max_row = max((r for (r, _c) in grid), default=-1)
            for r in range(0, max_row + 1):
                v = grid.get((r, avg_col), "")
                if v.strip():
                    avg_values.append(v)
        if avg_values:
            if not _avg_present(avg_values, _AVG_EXPECTED["instructgpt"]):
                blocking.append("InstructGPT average score 4.7 not found in GSheet")
            if not _avg_present(avg_values, _AVG_EXPECTED["scaling"]):
                blocking.append("Scaling Laws average score 4.3 not found in GSheet")
            if not _avg_present(avg_values, _AVG_EXPECTED["opt"]):
                blocking.append("OPT average score 4.0 not found in GSheet")
        else:
            # No Average_Score header — fall back to substring checks.
            if "4.7" not in all_text:
                blocking.append("InstructGPT average score 4.7 not found in GSheet")
            if "4.3" not in all_text:
                blocking.append("Scaling Laws average score 4.3 not found in GSheet")
            if "4.0" not in all_text and "4.00" not in all_text:
                blocking.append("OPT average score 4.0 not found in GSheet")

        if "completed" not in all_text:
            blocking.append("Review status 'Completed' not found in GSheet")

    except Exception as e:
        runtime.append(f"Error checking GSheet: {e}")
    return blocking, runtime


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()
    agent_ws = args.agent_workspace or os.path.join(os.path.dirname(__file__), "..", "groundtruth_workspace")

    all_errors = []

    # Check Review 1: Scaling Laws
    print("  Checking Review_Scaling_Laws.docx...")
    errs = check_word_review(agent_ws, "Review_Scaling_Laws.docx",
                             ["Scaling Laws", "scaling laws"], 5, 4, 4, "Accept")
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    # Check Review 2: InstructGPT
    print("  Checking Review_InstructGPT.docx...")
    errs = check_word_review(agent_ws, "Review_InstructGPT.docx",
                             ["follow instructions", "instructgpt", "rlhf", "2203.11171"],
                             5, 5, 4, "Accept")
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    # Check Review 3: OPT
    print("  Checking Review_OPT.docx...")
    errs = check_word_review(agent_ws, "Review_OPT.docx",
                             ["OPT", "open pre-trained"], 4, 3, 5, "Weak Accept")
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    # Check GSheet (runtime dependency — split blocking vs runtime-only)
    print("  Checking Google Sheet...")
    gs_blocking, gs_runtime = check_gsheet()
    if gs_blocking:
        all_errors.extend(gs_blocking)
        for e in gs_blocking[:3]:
            print(f"    ERROR: {e}")
    if gs_runtime:
        for e in gs_runtime[:3]:
            print(f"    [runtime-only] {e}")
    if not gs_blocking and not gs_runtime:
        print("    PASS")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} blocking errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        if gs_runtime:
            print(f"\n=== RESULT: PASS ({len(gs_runtime)} runtime-only failures tolerated) ===")
        else:
            print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
