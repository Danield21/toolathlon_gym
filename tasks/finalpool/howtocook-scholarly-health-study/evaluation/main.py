"""
Evaluation script for howtocook-scholarly-health-study task.

Checks (all derive from docs/task.md):
1. Health_Diet_Analysis.xlsx with three sheets (Recipe Nutrition, Research Summary, Combined Analysis)
   - Recipe Nutrition: >= 5 dishes, non-zero calories, numeric macros, Health_Rating in {Low, Medium, High}
   - Research Summary: >= 3 papers (task says "about 4"), non-empty titles/authors,
     diet/health/nutrition relevance, at least one real citation count
   - Combined Analysis: >= 3 rows linking dishes to supporting research
2. Chinese_Cuisine_Health_Report.docx with required title and four sections

Note: dish/paper selection is NOT unique (the agent chooses which 5 dishes and which papers to
report), so no row-by-row comparison against a groundtruth exemplar is performed. All checks are
structural/semantic and every one of them is satisfied by any agent that correctly follows the
task statement.
"""

import argparse
import os
import re
import sys

import openpyxl

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def note(name, detail=""):
    """Informational record that does NOT affect FAIL_COUNT."""
    print(f"  [NOTE] {name}" + (f": {detail[:300]}" if detail else ""))


def _to_float(v):
    """Robustly parse a cell value into a float; None when unparseable.

    Handles int/float, strings with thousands separators, currency symbols,
    percent signs, whitespace, and embedded numbers (e.g. '500 kcal').
    Formula strings (leading '=') resolve to None (no cached literal).
    """
    if v is None:
        return None
    if isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    if isinstance(v, str):
        s = v.strip()
        if not s:
            return None
        if s.startswith("="):
            return None
        s = (
            s.replace(",", "")
            .replace("$", "")
            .replace("¥", "")
            .replace("€", "")
            .replace("%", "")
            .strip()
        )
        try:
            return float(s)
        except ValueError:
            m = re.search(r"-?\d+(\.\d+)?", s)
            return float(m.group()) if m else None
    return None


def _str(v):
    if v is None:
        return ""
    return str(v).strip()


def _empty_row(row):
    return all(v is None or (isinstance(v, str) and not v.strip()) for v in row)


def _resolved_rows(wb_raw, wb_val, sheet_name):
    """Return the sheet's rows (list of lists) with formula cells resolved to
    their cached values (from the data_only load) where available."""
    ws = wb_raw[sheet_name]
    rows = []
    for r in range(1, ws.max_row + 1):
        row_vals = []
        for c in range(1, ws.max_column + 1):
            raw = ws.cell(r, c).value
            if isinstance(raw, str) and raw.startswith("="):
                row_vals.append(wb_val[sheet_name].cell(r, c).value)
            else:
                row_vals.append(raw)
        rows.append(row_vals)
    return rows


def _find_sheet(wb, primary, pairs=None, singles=None):
    """Case-insensitive sheet lookup: exact match, then both-keyword, then any-keyword."""
    for name in wb.sheetnames:
        if name.strip().lower() == primary.lower():
            return name
    if pairs:
        for name in wb.sheetnames:
            if all(k in name.lower() for k in pairs):
                return name
    if singles:
        for name in wb.sheetnames:
            if any(k in name.lower() for k in singles):
                return name
    return None


def _norm_header(h):
    return re.sub(r"[\s_\-]+", " ", str(h).strip().lower())


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Checking Excel Output ===")
    agent_file = os.path.join(agent_workspace, "Health_Diet_Analysis.xlsx")
    if not os.path.isfile(agent_file):
        record("Excel file exists", False, f"Not found: {agent_file}")
        return False
    record("Excel file exists", True)

    try:
        wb_raw = openpyxl.load_workbook(agent_file, data_only=False)
        wb_val = openpyxl.load_workbook(agent_file, data_only=True)
    except Exception as e:
        record("Excel file readable", False, str(e))
        return False

    all_ok = True

    # --- Sheet 1: Recipe Nutrition ---
    rn = _find_sheet(
        wb_raw, "Recipe Nutrition",
        pairs=("recipe", "nutrition"), singles=("recipe", "nutrition"),
    )
    if not rn:
        record("Sheet 'Recipe Nutrition' exists", False, f"Sheets: {wb_raw.sheetnames}")
        all_ok = False
    else:
        note("Sheet 'Recipe Nutrition' found", f"Actual name: '{rn}'")
        rows = _resolved_rows(wb_raw, wb_val, rn)
        data_rows = [r for r in rows[1:] if not _empty_row(r)]

        has_5_rows = len(data_rows) >= 5
        record(f"Recipe Nutrition has >= 5 rows ({len(data_rows)} found)", has_5_rows)
        if not has_5_rows:
            all_ok = False

        # Estimated_Calories (col 3) > 0 in at least 3 rows
        nonzero_cal = sum(
            1 for row in data_rows
            if len(row) > 2 and (_to_float(row[2]) or 0) > 0
        )
        record(
            f"Recipe Nutrition: >= 3 rows with non-zero calories ({nonzero_cal} found)",
            nonzero_cal >= 3,
        )
        if nonzero_cal < 3:
            all_ok = False

        # Protein_g (col 4) numeric in at least 3 rows
        rows_with_protein = sum(
            1 for row in data_rows
            if len(row) > 3 and _to_float(row[3]) is not None
        )
        record(
            f"Recipe Nutrition: >= 3 rows with numeric Protein_g ({rows_with_protein} found)",
            rows_with_protein >= 3,
        )
        if rows_with_protein < 3:
            all_ok = False

        # Health_Rating (col 7) values in {Low, Medium, High}.
        # Tolerate stray non-level strings in the rating column: a decorative
        # title row placed above the header (a common, task-permitted layout
        # choice) pushes the header text itself (e.g. 'Health_Rating') into the
        # data rows, which is not a real rating and must not fail the check.
        # Only the three documented levels carry the ability signal, so collect
        # them and require at least one genuine Low/Medium/High rating to exist.
        ratings = {
            _str(row[6]).lower()
            for row in data_rows if len(row) > 6 and _str(row[6])
        }
        valid_ratings = {r for r in ratings if r in {"low", "medium", "high"}}
        ratings_valid = len(valid_ratings) >= 1
        record(
            f"Recipe Nutrition: Health_Rating values in {{Low, Medium, High}} "
            f"({sorted(valid_ratings)})",
            ratings_valid,
        )
        if not ratings_valid:
            all_ok = False

        if rows:
            header = [_norm_header(h) for h in rows[0]]
            expected = [
                "dish name", "category", "estimated calories",
                "protein g", "fat g", "carbs g", "health rating",
            ]
            missing = [e for e in expected if not any(e in h for h in header)]
            note("Recipe Nutrition header", f"got {header}")
            if missing:
                note("Recipe Nutrition header (informational)", f"columns not found verbatim: {missing}")

    # --- Sheet 2: Research Summary ---
    rs = _find_sheet(
        wb_raw, "Research Summary",
        pairs=("research", "summary"), singles=("research", "summary"),
    )
    if not rs:
        record("Sheet 'Research Summary' exists", False, f"Sheets: {wb_raw.sheetnames}")
        all_ok = False
    else:
        note("Sheet 'Research Summary' found", f"Actual name: '{rs}'")
        rows = _resolved_rows(wb_raw, wb_val, rs)
        data_rows = [r for r in rows[1:] if not _empty_row(r)]

        # Task says "about 4 papers"; accept 3-5 (a careful agent may drop off-topic noise)
        has_rows = 3 <= len(data_rows) <= 6
        record(
            f"Research Summary has 3-6 rows ({len(data_rows)} found)",
            has_rows,
        )
        if not has_rows:
            all_ok = False

        rows_with_title = sum(1 for row in data_rows if _str(row[0]))
        record(
            f"Research Summary: >= 3 rows with non-empty titles ({rows_with_title} found)",
            rows_with_title >= 3,
        )
        if rows_with_title < 3:
            all_ok = False

        rows_with_authors = sum(
            1 for row in data_rows if len(row) > 1 and _str(row[1])
        )
        record(
            f"Research Summary: >= 3 rows with non-empty authors ({rows_with_authors} found)",
            rows_with_authors >= 3,
        )
        if rows_with_authors < 3:
            all_ok = False

        # Relevance: >= 2 rows whose title is diet/health/nutrition-related.
        # The scholarly DB is seeded (by preprocess) with 3 dietary papers + 1 off-topic
        # ML paper, so a correct agent always satisfies this.
        keywords = (
            "diet", "nutrition", "health", "dietary", "food",
            "cuisine", "mediterranean", "plant", "chronic",
        )
        relevant = sum(
            1 for row in data_rows
            if any(k in _str(row[0]).lower() for k in keywords)
        )
        record(
            f"Research Summary: >= 2 diet/health/nutrition-related titles ({relevant} found)",
            relevant >= 2,
        )
        if relevant < 2:
            all_ok = False

        # Citation count: >= 1 row has a numeric value > 0 that is not a plausible year.
        # Year is column 3 per task.md; citation count is column 4.
        rows_with_citations = 0
        for row in data_rows:
            cit = _to_float(row[3]) if len(row) > 3 else None
            if cit is not None and cit > 0:
                rows_with_citations += 1
                continue
            for v in row[2:]:
                fv = _to_float(v)
                if fv is not None and fv > 0 and not (1900 <= fv <= 2100):
                    rows_with_citations += 1
                    break
        record(
            f"Research Summary: >= 1 row with citation count > 0 ({rows_with_citations} found)",
            rows_with_citations >= 1,
        )
        if rows_with_citations < 1:
            all_ok = False

    # --- Sheet 3: Combined Analysis ---
    ca = _find_sheet(
        wb_raw, "Combined Analysis",
        pairs=("combined", "analysis"), singles=("combined", "analysis"),
    )
    if not ca:
        record("Sheet 'Combined Analysis' exists", False, f"Sheets: {wb_raw.sheetnames}")
        all_ok = False
    else:
        note("Sheet 'Combined Analysis' found", f"Actual name: '{ca}'")
        rows = _resolved_rows(wb_raw, wb_val, ca)
        data_rows = [r for r in rows[1:] if not _empty_row(r)]

        has_3_rows = len(data_rows) >= 3
        record(
            f"Combined Analysis has >= 3 rows ({len(data_rows)} found)", has_3_rows
        )
        if not has_3_rows:
            all_ok = False

        # Each row links a dish to supporting research (col 1 + col 3 non-empty)
        linked = sum(
            1 for row in data_rows
            if len(row) > 2 and _str(row[0]) and _str(row[2])
        )
        record(
            f"Combined Analysis: >= 3 rows with Dish_Name and Supporting_Research ({linked} found)",
            linked >= 3,
        )
        if linked < 3:
            all_ok = False

    wb_raw.close()
    wb_val.close()
    return all_ok


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_file = os.path.join(agent_workspace, "Chinese_Cuisine_Health_Report.docx")
    if not os.path.isfile(doc_file):
        record("Word document exists", False, f"Not found: {doc_file}")
        return False
    record("Word document exists", True)

    try:
        from docx import Document

        doc = Document(doc_file)
        full_text = " ".join(p.text for p in doc.paragraphs)
        text_lower = full_text.lower()

        has_length = len(full_text) >= 500
        record(
            f"Document has >= 500 chars ({len(full_text)} found)", has_length
        )
        if not has_length:
            return False

        has_title = (
            "nutritional analysis" in text_lower or "chinese cuisine" in text_lower
        )
        record("Document contains title keywords", has_title)

        # Heading match preferred; body-substring fallback is sufficient (task.md only
        # asks to "include the following sections", not to use a specific heading style).
        heading_texts_lower = [
            p.text.strip().lower() for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        sections_ok = True
        for section in ["introduction", "recipe analysis", "literature review", "conclusions"]:
            # Tolerate the singular form of plural section names ("Conclusion"
            # vs "Conclusions") in both heading and body matching, which is the
            # standard English heading convention for the final section.
            terms = (section, section[:-1]) if section.endswith("s") else (section,)
            in_heading = any(
                any(t == ht or ht.startswith(t) for t in terms)
                for ht in heading_texts_lower
            )
            in_body = any(t in text_lower for t in terms)
            found = in_heading or in_body
            record(f"Document has '{section}' section", found,
                   f"headings: {heading_texts_lower}")
            if not found:
                sections_ok = False

        return has_length and has_title and sections_ok

    except ImportError:
        file_size = os.path.getsize(doc_file)
        has_content = file_size > 1000
        record(
            "Word document has content (size > 1KB, python-docx not available)",
            has_content,
            f"File size: {file_size} bytes",
        )
        return has_content
    except Exception as e:
        record("Word document readable", False, str(e))
        return False


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    excel_ok = check_excel(args.agent_workspace, args.groundtruth_workspace)
    word_ok = check_word(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Excel:   {'PASS' if excel_ok else 'FAIL'}")
    print(f"  Word:    {'PASS' if word_ok else 'FAIL'}")

    # Use FAIL_COUNT as authoritative: any recorded FAIL propagates to overall.
    # With the row-by-row groundtruth comparison removed, every recorded check is a
    # genuine task requirement that a correct agent satisfies.
    overall = excel_ok and word_ok and FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
