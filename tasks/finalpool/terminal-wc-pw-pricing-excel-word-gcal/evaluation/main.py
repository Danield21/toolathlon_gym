"""Evaluation for terminal-wc-pw-pricing-excel-word-gcal."""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=5.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


# Mock-page competitor prices (sourced from tmp/mock_pages/pricing.html which is a static fixture).
# Used as a fallback only; primary path computes competitor avg from Excel file the agent produced
# (so updates to the mock page automatically flow through).
_COMPETITOR_PRICES_FALLBACK = {
    "audio": [62.99, 49.99, 95.00],
    "cameras": [45.00, 28.99, 65.50],
    "electronics": [45.99, 38.50, 22.99, 35.00, 15.99],
    "headphones": [89.99, 65.00, 42.50, 120.00],
    "speakers": [55.99, 199.00, 89.99, 75.00],
}


def _our_prices_from_db():
    """Compute our category-level avg effective price (price column) from WC catalog,
    matching GT semantics (price = sale_price or regular_price).
    Returns {category_lower: avg_price} or {} on failure."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            """
            SELECT cat_name, AVG(price)::float
            FROM (
              SELECT price,
                     jsonb_array_elements(categories::jsonb)->>'name' AS cat_name
              FROM wc.products
              WHERE regular_price > 0 AND categories IS NOT NULL
            ) sub
            WHERE LOWER(cat_name) IN ('audio','cameras','electronics','headphones','speakers')
            GROUP BY cat_name
            """
        )
        out = {r[0].strip().lower(): float(r[1]) for r in cur.fetchall() if r[1] is not None}
        cur.close()
        conn.close()
        return out
    except Exception as e:
        print(f"[WARN] our prices DB query failed: {e}")
        return {}


def get_expected_gaps(comp_avgs=None):
    """Compute expected gap_pct per category using DB-backed our_prices and
    either competitor averages from the Excel/output (preferred) or the static
    mock-page fallback. Returns {category_lower: gap_pct}."""
    our = _our_prices_from_db()
    comp = comp_avgs or {
        k: sum(v) / len(v) for k, v in _COMPETITOR_PRICES_FALLBACK.items()
    }
    out = {}
    for cat, our_avg in our.items():
        cv = comp.get(cat)
        if cv is None or cv == 0:
            continue
        out[cat] = (our_avg - cv) / cv * 100
    return out


def check_excel(agent_workspace):
    print("\n=== Checking Excel Output ===")
    fpath = os.path.join(agent_workspace, "Competitive_Pricing_Report.xlsx")
    if not os.path.isfile(fpath):
        record("Excel file exists", False, f"Not found: {fpath}")
        return False
    record("Excel file exists", True)

    wb = openpyxl.load_workbook(fpath, data_only=True)

    # Sheet 1: Our_Products
    our_sheet = None
    for name in wb.sheetnames:
        if "our" in name.lower() and "product" in name.lower():
            our_sheet = name
            break
    if not our_sheet:
        record("Our_Products sheet exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Our_Products sheet exists", True)
        ws = wb[our_sheet]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        # Dynamically get expected product count from DB
        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cur = conn.cursor()
            cur.execute("SELECT COUNT(*) FROM wc.products WHERE regular_price > 0")
            expected_products = cur.fetchone()[0]
            cur.close()
            conn.close()
            record(f"Our_Products has >= {expected_products} rows", len(rows) >= expected_products,
                   f"Found {len(rows)}")
        except Exception:
            record("Our_Products has >= 30 rows", len(rows) >= 30, f"Found {len(rows)}")

    # Sheet 2: Competitor_Prices
    comp_sheet = None
    for name in wb.sheetnames:
        if "competitor" in name.lower():
            comp_sheet = name
            break
    comp_avgs_from_excel = {}
    if not comp_sheet:
        record("Competitor_Prices sheet exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Competitor_Prices sheet exists", True)
        ws = wb[comp_sheet]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        # 19 is the size of the static mock-page; verify exact match.
        record("Competitor_Prices has 19 rows", len(rows) == 19, f"Found {len(rows)}")

        # Try to extract competitor avg by category from this sheet for downstream gap math.
        hdr_row = list(next(ws.iter_rows(values_only=True), []))
        hdr_n = [str(c or "").strip().lower().replace(" ", "_") for c in hdr_row]
        cat_idx = next((i for i, h in enumerate(hdr_n) if h in ("category", "cat")), None)
        price_idx = next((i for i, h in enumerate(hdr_n)
                          if "price" in h), None)
        if cat_idx is not None and price_idx is not None:
            buckets = {}
            for r in rows:
                if not r or len(r) <= max(cat_idx, price_idx):
                    continue
                cat = str(r[cat_idx] or "").strip().lower()
                try:
                    p = float(str(r[price_idx]).replace("$", "").replace(",", ""))
                except (TypeError, ValueError):
                    continue
                if not cat:
                    continue
                buckets.setdefault(cat, []).append(p)
            comp_avgs_from_excel = {k: sum(v) / len(v) for k, v in buckets.items() if v}

    # Sheet 3: Price_Gap_Analysis
    gap_sheet = None
    for name in wb.sheetnames:
        if "gap" in name.lower() or "analysis" in name.lower():
            gap_sheet = name
            break
    if not gap_sheet:
        record("Price_Gap_Analysis sheet exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Price_Gap_Analysis sheet exists", True)
        ws = wb[gap_sheet]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        record("Price_Gap_Analysis has 5 rows", len(rows) == 5, f"Found {len(rows)}")

        # Header validation
        hdr = list(next(ws.iter_rows(values_only=True), []))
        hdr_norm = [str(c or "").strip().lower().replace(" ", "_") for c in hdr]
        for h in ["category", "our_avg_price", "competitor_avg_price", "gap_pct", "recommendation"]:
            record(f"Price_Gap_Analysis has '{h}' header", h in hdr_norm, f"got {hdr_norm}")

        # Sort: alphabetical by category
        cat_names = [str(r[0] or "").strip() for r in rows if r and r[0]]
        sorted_cats = sorted(cat_names, key=lambda s: s.lower())
        record(
            "Price_Gap_Analysis sorted by category alphabetically",
            cat_names == sorted_cats,
            f"got {cat_names}",
        )

        # Per-row Gap_Pct + Recommendation validation
        try:
            gap_idx = hdr_norm.index("gap_pct")
            rec_idx = hdr_norm.index("recommendation")
        except ValueError:
            gap_idx = 3
            rec_idx = 4
        # Compute expected gaps dynamically using DB our-prices and competitor
        # averages derived from the Competitor_Prices sheet (mock-page fallback).
        expected_gaps = get_expected_gaps(comp_avgs_from_excel or None)
        # Tolerance scales with magnitude (prices/gaps can be large).
        for row in rows:
            if not row or not row[0]:
                continue
            cat = str(row[0]).strip().lower()
            if cat in expected_gaps:
                gap_val = row[gap_idx] if len(row) > gap_idx else None
                exp = expected_gaps[cat]
                tol = max(2.0, abs(exp) * 0.05)  # 5% relative tol, min 2 percentage points
                record(
                    f"Gap_Pct for {cat} is correct (~{exp:.1f}, tol {tol:.1f})",
                    num_close(gap_val, exp, tol=tol),
                    f"Got {gap_val}, expected ~{exp:.2f}",
                )
                # Recommendation rule
                rec_val = str(row[rec_idx] or "").strip().lower() if len(row) > rec_idx else ""
                if exp > 10:
                    expected_rec = "consider price reduction"
                elif exp < -10:
                    expected_rec = "maintain competitive advantage"
                else:
                    expected_rec = "monitor pricing"
                record(
                    f"Recommendation for {cat} is '{expected_rec}'",
                    expected_rec in rec_val,
                    f"got '{rec_val}'",
                )

    wb.close()
    return True


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    fpath = os.path.join(agent_workspace, "Pricing_Strategy_Report.docx")
    if not os.path.isfile(fpath):
        record("Word document exists", False, f"Not found: {fpath}")
        return False
    record("Word document exists", True)

    from docx import Document
    doc = Document(fpath)
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    headings_or_titles = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style.name.startswith("Heading") or p.style.name.startswith("Title")
    ]
    htl = [h.lower() for h in headings_or_titles]
    record(
        "Document title 'Competitive Pricing Analysis Report'",
        any("competitive pricing analysis report" in h for h in htl)
        or "competitive pricing analysis report" in full_text,
        f"headings: {headings_or_titles[:5]}",
    )
    record("Document has executive summary", "executive summary" in full_text or "summary" in full_text)
    record("Document has conclusion / strategic next steps", "conclusion" in full_text or "next step" in full_text)
    # Each of the 5 expected categories should appear
    for cat in ["audio", "cameras", "electronics", "headphones", "speakers"]:
        record(f"Document mentions '{cat}' section", cat in full_text)
    return True


def check_terminal_output(agent_workspace):
    print("\n=== Checking Terminal Output ===")
    fpath = os.path.join(agent_workspace, "price_analysis_output.txt")
    if not os.path.isfile(fpath):
        record("price_analysis_output.txt exists", False)
        return False
    record("price_analysis_output.txt exists", True)
    with open(fpath) as f:
        content = f.read().lower()
    # Must mention gap and recommendation logic
    record("Output mentions 'gap'", "gap" in content)
    record("Output mentions percent (%)", "%" in content)
    record("Output mentions 'recommendation' or 'recommend'", "recommend" in content)
    return True


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            "SELECT summary, description, start_datetime, end_datetime "
            "FROM gcal.events WHERE summary ILIKE '%price review%' "
            "ORDER BY start_datetime"
        )
        events = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Calendar DB accessible", False, str(e))
        return False

    EXPECTED_CATEGORIES = {"audio", "cameras", "electronics", "headphones", "speakers"}

    categories_found = set()
    for summary, description, start_dt, end_dt in events:
        summary_lower = (summary or "").lower()
        if "price review" in summary_lower:
            for cat in EXPECTED_CATEGORIES:
                if cat in summary_lower:
                    categories_found.add(cat)
    record(
        "Price review events for all 5 categories",
        categories_found == EXPECTED_CATEGORIES,
        f"Found events for: {categories_found}",
    )

    # Validate each event: 14:00 start, 1h duration, on/after March 10 2026, dates differ
    from datetime import timedelta
    seen_dates = set()
    for summary, description, start_dt, end_dt in events:
        if start_dt is None or end_dt is None:
            record(f"Event '{summary}' has start/end", False)
            continue
        record(
            f"Event '{summary}' starts at 14:00",
            start_dt.hour == 14 and start_dt.minute == 0,
            f"start={start_dt}",
        )
        record(
            f"Event '{summary}' duration is 1 hour",
            (end_dt - start_dt) == timedelta(hours=1),
            f"duration={end_dt - start_dt}",
        )
        record(
            f"Event '{summary}' on/after 2026-03-10",
            start_dt.date() >= __import__("datetime").date(2026, 3, 10),
            f"date={start_dt.date()}",
        )
        seen_dates.add(start_dt.date())
        # Description must mention gap percentage
        desc_l = (description or "").lower()
        record(
            f"Event '{summary}' description mentions gap %",
            "%" in desc_l or "gap" in desc_l,
            f"description: {desc_l[:80]}",
        )
    record(
        "Events scheduled on different days",
        len(seen_dates) == len(events) and len(events) > 0,
        f"unique dates={len(seen_dates)} events={len(events)}",
    )
    return categories_found == EXPECTED_CATEGORIES


def check_reverse_validation(workspace):
    """Verify things that should NOT exist in the output."""
    print("\n=== Reverse Validation ===")
    fpath = os.path.join(workspace, "Competitive_Pricing_Report.xlsx")
    if os.path.isfile(fpath):
        wb = openpyxl.load_workbook(fpath, data_only=True)
        # No unexpected sheets
        expected_keywords = {"our", "product", "competitor", "price", "gap", "analysis"}
        unexpected = [s for s in wb.sheetnames
                      if not any(kw in s.lower() for kw in expected_keywords)]
        record("No unexpected sheets in Excel", len(unexpected) == 0,
               f"Unexpected: {unexpected}")

        # Prices should not be negative
        for sname in wb.sheetnames:
            ws = wb[sname]
            for row in ws.iter_rows(min_row=2, values_only=True):
                for c in row:
                    try:
                        v = float(str(c).replace(',', '').replace('$', '').strip()) if c else None
                    except (ValueError, TypeError):
                        v = None
                    if v is not None and v < 0 and "gap" not in sname.lower():
                        record(f"No negative prices in {sname}", False, f"Found {v}")
                        break
                else:
                    continue
                break
            else:
                continue
            break
        else:
            record("No negative prices in product sheets", True)

    # Calendar: no price review events before March 2026
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(*) FROM gcal.events
            WHERE summary ILIKE '%%price review%%'
              AND start_datetime < '2026-03-10'
        """)
        early_events = cur.fetchone()[0]
        record("No price review events before March 10", early_events == 0,
               f"Found {early_events} early events")
        cur.close()
        conn.close()
    except Exception:
        pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace)
    check_word(args.agent_workspace)
    check_terminal_output(args.agent_workspace)
    check_calendar()
    check_reverse_validation(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
