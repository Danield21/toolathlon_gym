"""Evaluation script for pw-wc-price-optimization-excel-word."""
import os
import argparse, json, os, sys
import openpyxl

# --- verify_v2 smart primitives ---
_EVAL_DIR = os.path.dirname(os.path.abspath(__file__))
_PACK_ROOT = os.path.abspath(os.path.join(_EVAL_DIR, "..", "..", "..", ".."))
if _PACK_ROOT not in sys.path:
    sys.path.insert(0, _PACK_ROOT)
try:
    from utils.verify_v2 import smart_column_exists
    from utils.verify_v2.eval_helpers import get_sheet_rows_as_dicts, get_gt_column_values
    _HAS_VERIFY_V2 = True
except Exception:
    _HAS_VERIFY_V2 = False

TASK_NAME = "pw-wc-price-optimization-excel-word"


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    
    excel_path = os.path.join(agent_workspace, "Price_Optimization_Report.xlsx")
    check("Price_Optimization_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Price_Optimization_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        def check_columns(sheet_name, expected_cols, min_rows):
            """Verify sheet exists, >= min_rows, contains required columns.
            Uses LLM semantic mapping via verify_v2, falls back to strict
            header match."""
            check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
            if sheet_name not in wb.sheetnames:
                return
            _ws = wb[sheet_name]
            _data_rows = list(_ws.iter_rows(min_row=2, values_only=True))
            check(f"{sheet_name} has >= {min_rows} rows",
                  len(_data_rows) >= min_rows, f"got {len(_data_rows)}")
            if _HAS_VERIFY_V2 and gt_wb is not None:
                _raw_headers, _agent_rows = get_sheet_rows_as_dicts(wb, sheet_name)
                for _exp in expected_cols:
                    _gt_vals = get_gt_column_values(gt_wb, sheet_name, _exp)
                    _ok, _matched, _reason = smart_column_exists(
                        expected_col=_exp, agent_headers=_raw_headers,
                        gt_samples=_gt_vals[:3], agent_rows=_agent_rows,
                        task_name=TASK_NAME,
                    )
                    _detail = _reason
                    if _ok and _matched and _matched.lower() != _exp.lower():
                        _detail = f"LLM-mapped to {_matched!r}"
                    check(f"{sheet_name} has {_exp} column", _ok, _detail)
            else:
                _headers = [str(c.value).strip().lower() if c.value else "" for c in _ws[1]]
                for _exp in expected_cols:
                    check(f"{sheet_name} has {_exp} column",
                          _exp.lower() in _headers, f"headers: {_headers[:8]}")

        check("Price_Comparison sheet exists", "Price_Comparison" in wb.sheetnames)
        if "Price_Comparison" in wb.sheetnames:
            ws = wb["Price_Comparison"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            data_rows = [r for r in data_rows if r and r[0] is not None and str(r[0]).strip()]
            check("Price_Comparison has >= 10 rows", len(data_rows) >= 10, f"got {len(data_rows)}")

            check_columns('Price_Comparison', ['Product_Name', 'Our_Price', 'Competitor_Price', 'Price_Difference', 'Difference_Pct', 'Recommendation'], 10)

            # Per-row value comparison vs GT (key by first 30 chars of Product_Name)
            if gt_wb is not None and "Price_Comparison" in gt_wb.sheetnames:
                gt_ws = gt_wb["Price_Comparison"]
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if r and r[0]]
                gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                check(f"Price_Comparison row count == {len(gt_rows)}",
                      len(data_rows) == len(gt_rows), f"got {len(data_rows)}")
                hmap = {h: i for i, h in enumerate(headers)}
                gt_by = {str(r[0]).strip().lower()[:30]: r for r in gt_rows}
                ag_by = {str(r[0]).strip().lower()[:30]: r for r in data_rows}
                for k, gt_row in gt_by.items():
                    found = k in ag_by
                    check(f"Price_Comparison product '{k[:25]}' present", found)
                    if found:
                        agent_row = ag_by[k]
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row) or gt_h == "product_name":
                                continue
                            gv = gt_row[ci]
                            agent_ci = hmap.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                tol = max(0.5, abs(gf) * 0.05)
                                check(f"Price_Comparison '{k[:20]}' {gt_h} ~{gf}",
                                      abs(gf - af) <= tol, f"got {av}")
                            elif gv is not None and av is not None:
                                check(f"Price_Comparison '{k[:20]}' {gt_h} == {gv}",
                                      str(gv).strip().lower() == str(av).strip().lower(),
                                      f"got {av}")

                # Sort order: ASCENDING by Product_Name
                names = [str(r[0]).strip() for r in data_rows]
                sorted_names = sorted(names, key=lambda s: s.lower())
                check("Price_Comparison sorted alphabetically by Product_Name",
                      [n.lower() for n in names] == [n.lower() for n in sorted_names],
                      f"first 3: {names[:3]}")
        check("Category_Summary sheet exists", "Category_Summary" in wb.sheetnames)
        if "Category_Summary" in wb.sheetnames:
            ws = wb["Category_Summary"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Category_Summary has >= 5 rows", len(data_rows) >= 5, f"got {len(data_rows)}")

            check_columns('Category_Summary', ['Category', 'Product_Count', 'Avg_Our_Price', 'Avg_Stock'], 5)
        check("Executive_Summary sheet exists", "Executive_Summary" in wb.sheetnames)
        if "Executive_Summary" in wb.sheetnames:
            ws = wb["Executive_Summary"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Executive_Summary has >= 5 rows", len(data_rows) >= 5, f"got {len(data_rows)}")

            check_columns('Executive_Summary', ['Metric', 'Value'], 5)
        word_path = os.path.join(agent_workspace, "Pricing_Strategy.docx")
        check("Pricing strategy Word exists", os.path.exists(word_path))
        if os.path.exists(word_path):
            from docx import Document
            doc = Document(word_path)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            check("Word mentions pricing", "pric" in text)
            check("Word mentions recommendation", "recommend" in text)
            # Validate required headings explicitly
            heading_texts = [p.text.strip().lower() for p in doc.paragraphs
                             if p.style and p.style.name and "heading" in p.style.name.lower()]
            heading_blob = " | ".join(heading_texts)
            check("Word has 'Market Position Analysis' heading",
                  "market position" in heading_blob,
                  f"headings: {heading_texts[:6]}")
            check("Word has 'Product-Level Recommendations' heading",
                  "product-level recommendation" in heading_blob or
                  "product level recommendation" in heading_blob,
                  f"headings: {heading_texts[:6]}")
            check("Word has 'Implementation Timeline' heading",
                  "implementation timeline" in heading_blob,
                  f"headings: {heading_texts[:6]}")
        # Validate Recommendation values are from expected set
        if "Price_Comparison" in wb.sheetnames:
            ws = wb["Price_Comparison"]
            headers_raw = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
            if "recommendation" in headers_raw:
                rec_idx = headers_raw.index("recommendation")
                allowed_vals = {"maintain", "consider increase", "consider decrease",
                                "increase", "decrease", "raise", "reduce",
                                "consider raising", "consider lowering",
                                "no change"}
                rec_vals = [str(r[rec_idx]).strip().lower() for r in ws.iter_rows(min_row=2, values_only=True)
                            if r and r[rec_idx] is not None]
                bad = [v for v in rec_vals if not any(av in v or v in av for av in allowed_vals)]
                check("Price_Comparison Recommendation values valid",
                      len(rec_vals) >= 5 and len(bad) == 0,
                      f"bad values: {bad[:3]} of {len(rec_vals)} total")
        # Validate Executive_Summary has required metrics
        if "Executive_Summary" in wb.sheetnames:
            ws = wb["Executive_Summary"]
            metric_pairs = {}
            for r in ws.iter_rows(min_row=2, values_only=True):
                if r and r[0]:
                    key = str(r[0]).strip().lower()
                    metric_pairs[key] = r[1] if len(r) > 1 else None
            required_metrics = ["total_products_compared", "products_overpriced",
                                "products_competitive", "products_underpriced"]
            missing = [m for m in required_metrics if m not in metric_pairs]
            check("Executive_Summary has required metric keys",
                  len(missing) == 0, f"missing: {missing}")
        check("price_optimizer.py exists", os.path.exists(os.path.join(agent_workspace, "price_optimizer.py")))
        # All three JSON files explicitly named in task.md
        for jname in ("competitor_prices.json", "our_products.json", "price_recommendations.json"):
            jpath = os.path.join(agent_workspace, jname)
            check(f"{jname} exists", os.path.exists(jpath))
            if os.path.exists(jpath):
                try:
                    with open(jpath, "r") as _f:
                        _d = json.load(_f)
                    check(f"{jname} is valid JSON and non-empty", bool(_d), "empty/null")
                except Exception as _e:
                    check(f"{jname} is valid JSON and non-empty", False, str(_e))


    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
