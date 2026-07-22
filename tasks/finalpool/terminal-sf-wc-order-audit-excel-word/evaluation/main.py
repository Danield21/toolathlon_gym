"""Evaluation for terminal-sf-wc-order-audit-excel-word.
Checks:
1. Order_Audit_Report.xlsx with 4 sheets and correct data (compared to GT)
2. Audit_Findings.docx with required sections
3. audit_analysis.py script exists
"""
import argparse
import json
import os
import sys
import openpyxl
import psycopg2


DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432,
          dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
          user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def num_close(a, b, tol):
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except Exception:
        return False


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('$', '').replace('%', '').strip())
    except Exception:
        return default


def _norm_header(s):
    return str(s or "").strip().lower().replace(" ", "_")


def _load_sheet(wb, name):
    for s in wb.sheetnames:
        if s.lower().replace(" ", "_") == name.lower().replace(" ", "_"):
            return wb[s]
    # fallback: substring match
    for s in wb.sheetnames:
        if name.lower().split("_")[0] in s.lower():
            return wb[s]
    return None


def check_excel(workspace, gt_dir):
    print("\n=== Check 1: Order_Audit_Report.xlsx ===")
    path = os.path.join(workspace, "Order_Audit_Report.xlsx")
    if not os.path.exists(path):
        check("Excel file exists", False, f"Not found at {path}")
        return
    check("Excel file exists", True)

    wb = openpyxl.load_workbook(path, data_only=True)
    gt_path = os.path.join(gt_dir, "Order_Audit_Report.xlsx")
    gt_wb = openpyxl.load_workbook(gt_path, data_only=True) if os.path.exists(gt_path) else None
    check("Has 4 sheets", len(wb.sheetnames) >= 4, f"Found {len(wb.sheetnames)}: {wb.sheetnames}")

    # ---------- DW_Summary ----------
    ws_dw = _load_sheet(wb, "DW_Summary")
    if ws_dw is None:
        check("DW_Summary sheet present", False, f"sheets: {wb.sheetnames}")
    else:
        check("DW_Summary sheet present", True)
        rows_dw = list(ws_dw.iter_rows(values_only=True))
        headers = [_norm_header(c) for c in (rows_dw[0] if rows_dw else [])]
        for h in ["status", "order_count", "total_revenue"]:
            check(f"DW_Summary header '{h}'", h in headers, f"got {headers}")
        if gt_wb:
            gt_ws = _load_sheet(gt_wb, "DW_Summary")
            gt_rows = list(gt_ws.iter_rows(values_only=True))
            gt_headers = [_norm_header(c) for c in gt_rows[0]]
            gt_data = [r for r in gt_rows[1:] if r and r[0]]
            a_lookup = {str(r[0]).strip().lower(): r for r in rows_dw[1:] if r and r[0]}
            for g_row in gt_data:
                key = str(g_row[0]).strip().lower()
                a_row = a_lookup.get(key)
                if a_row is None:
                    check(f"DW_Summary row '{g_row[0]}'", False)
                    continue
                # Order_Count
                if "order_count" in headers:
                    oc_idx = headers.index("order_count")
                    g_oc = g_row[gt_headers.index("order_count")]
                    check(
                        f"DW_Summary '{g_row[0]}' order_count",
                        num_close(safe_float(a_row[oc_idx]), safe_float(g_oc), max(20, safe_float(g_oc, 1) * 0.005)),
                        f"a={a_row[oc_idx]} g={g_oc}",
                    )
                # Total_Revenue (tol 0.5%)
                if "total_revenue" in headers:
                    tr_idx = headers.index("total_revenue")
                    g_tr = g_row[gt_headers.index("total_revenue")]
                    check(
                        f"DW_Summary '{g_row[0]}' total_revenue",
                        num_close(safe_float(a_row[tr_idx]), safe_float(g_tr), max(100, safe_float(g_tr, 1) * 0.005)),
                        f"a={a_row[tr_idx]} g={g_tr}",
                    )

    # ---------- Store_Summary ----------
    ws_store = _load_sheet(wb, "Store_Summary")
    if ws_store is None:
        check("Store_Summary sheet present", False)
    else:
        check("Store_Summary sheet present", True)
        rows_store = list(ws_store.iter_rows(values_only=True))
        headers = [_norm_header(c) for c in (rows_store[0] if rows_store else [])]
        for h in ["metric", "value"]:
            check(f"Store_Summary header '{h}'", h in headers, f"got {headers}")
        if gt_wb:
            gt_ws = _load_sheet(gt_wb, "Store_Summary")
            gt_rows = list(gt_ws.iter_rows(values_only=True))
            gt_data = [r for r in gt_rows[1:] if r and r[0]]
            a_lookup = {str(r[0]).strip().lower(): r for r in rows_store[1:] if r and r[0]}
            for g_row in gt_data:
                key = str(g_row[0]).strip().lower()
                a_row = a_lookup.get(key)
                if a_row is None:
                    check(f"Store_Summary metric '{g_row[0]}'", False)
                    continue
                # Value (tol 0; integer counts and currency)
                v_idx = headers.index("value") if "value" in headers else 1
                g_v = g_row[1]
                check(
                    f"Store_Summary '{g_row[0]}' value",
                    num_close(safe_float(a_row[v_idx]), safe_float(g_v), 0.5),
                    f"a={a_row[v_idx]} g={g_v}",
                )

    # ---------- ShipMode_Breakdown ----------
    ws_ship = _load_sheet(wb, "ShipMode_Breakdown")
    if ws_ship is None:
        check("ShipMode_Breakdown sheet present", False)
    else:
        check("ShipMode_Breakdown sheet present", True)
        rows_ship = list(ws_ship.iter_rows(values_only=True))
        headers = [_norm_header(c) for c in (rows_ship[0] if rows_ship else [])]
        for h in ["ship_mode", "order_count", "total_revenue"]:
            check(f"ShipMode_Breakdown header '{h}'", h in headers, f"got {headers}")
        if gt_wb:
            gt_ws = _load_sheet(gt_wb, "ShipMode_Breakdown")
            gt_rows = list(gt_ws.iter_rows(values_only=True))
            gt_headers = [_norm_header(c) for c in gt_rows[0]]
            gt_data = [r for r in gt_rows[1:] if r and r[0]]
            a_lookup = {str(r[0]).strip().lower(): r for r in rows_ship[1:] if r and r[0]}
            # Validate all 4 ship modes (Economy, Express, Next Day, Standard)
            for g_row in gt_data:
                key = str(g_row[0]).strip().lower()
                a_row = a_lookup.get(key)
                if a_row is None:
                    check(f"ShipMode '{g_row[0]}' present", False)
                    continue
                if "order_count" in headers:
                    oc_idx = headers.index("order_count")
                    g_oc = g_row[gt_headers.index("order_count")]
                    check(
                        f"ShipMode '{g_row[0]}' order_count",
                        num_close(safe_float(a_row[oc_idx]), safe_float(g_oc), max(20, safe_float(g_oc, 1) * 0.01)),
                        f"a={a_row[oc_idx]} g={g_oc}",
                    )
                if "total_revenue" in headers:
                    tr_idx = headers.index("total_revenue")
                    g_tr = g_row[gt_headers.index("total_revenue")]
                    check(
                        f"ShipMode '{g_row[0]}' total_revenue",
                        num_close(safe_float(a_row[tr_idx]), safe_float(g_tr), max(100, safe_float(g_tr, 1) * 0.01)),
                        f"a={a_row[tr_idx]} g={g_tr}",
                    )

    # ---------- Reconciliation ----------
    ws_recon = _load_sheet(wb, "Reconciliation")
    if ws_recon is None:
        check("Reconciliation sheet present", False)
    else:
        check("Reconciliation sheet present", True)
        rows_recon = list(ws_recon.iter_rows(values_only=True))
        headers = [_norm_header(c) for c in (rows_recon[0] if rows_recon else [])]
        for h in ["metric", "dw_value", "store_value", "difference"]:
            check(f"Reconciliation header '{h}'", h in headers, f"got {headers}")
        if gt_wb:
            gt_ws = _load_sheet(gt_wb, "Reconciliation")
            gt_rows = list(gt_ws.iter_rows(values_only=True))
            gt_headers = [_norm_header(c) for c in gt_rows[0]]
            gt_data = [r for r in gt_rows[1:] if r and r[0]]
            a_lookup = {str(r[0]).strip().lower(): r for r in rows_recon[1:] if r and r[0]}
            for g_row in gt_data:
                key = str(g_row[0]).strip().lower()
                a_row = a_lookup.get(key)
                if a_row is None:
                    check(f"Reconciliation metric '{g_row[0]}' present", False)
                    continue
                # dw_value, store_value, difference (tol 0.5% or 5)
                for col in ["dw_value", "store_value", "difference"]:
                    if col in headers:
                        idx = headers.index(col)
                        g_v = g_row[gt_headers.index(col)]
                        check(
                            f"Reconciliation '{g_row[0]}' {col}",
                            num_close(safe_float(a_row[idx]), safe_float(g_v), max(5, abs(safe_float(g_v, 1)) * 0.01)),
                            f"a={a_row[idx]} g={g_v}",
                        )


def check_word(workspace):
    print("\n=== Check 2: Audit_Findings.docx ===")
    path = os.path.join(workspace, "Audit_Findings.docx")
    if not os.path.exists(path):
        check("Word document exists", False, f"Not found at {path}")
        return
    check("Word document exists", True)

    try:
        from docx import Document
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name.startswith("Title")]

        check(
            "Doc title 'Order Reconciliation Audit Report' present",
            any("order reconciliation audit report" in h for h in headings)
            or "order reconciliation audit report" in full_text,
            f"headings: {headings[:5]}",
        )
        # Required sections (require phrase, not single keyword)
        check("Doc has data warehouse summary section",
              "data warehouse" in full_text and "summary" in full_text)
        check("Doc has online store summary section",
              ("online store" in full_text or "store summary" in full_text) and "summary" in full_text)
        check("Doc has ship mode breakdown section", "ship" in full_text and "mode" in full_text)
        check("Doc has reconciliation findings section", "reconcil" in full_text and "finding" in full_text)
        check("Doc has recommendations", "recommend" in full_text)
        check("Doc mentions all 4 status values (delivered, cancelled, processing, shipped)",
              all(s in full_text for s in ["delivered", "cancelled", "processing", "shipped"]))
        check("Doc mentions all 4 ship modes", all(s in full_text for s in ["economy", "express", "next day", "standard"]))
    except Exception as e:
        check("Word document readable", False, str(e))


def check_script(workspace):
    print("\n=== Check 3: audit_analysis.py ===")
    path = os.path.join(workspace, "audit_analysis.py")
    if not os.path.exists(path):
        check("audit_analysis.py exists", False, f"Not found at {path}")
        return
    check("audit_analysis.py exists", True)
    # Verify script content references reconciliation logic
    with open(path) as f:
        content = f.read().lower()
    check("audit_analysis.py mentions order count comparison",
          "order" in content and ("count" in content or "total" in content))
    check("audit_analysis.py mentions ship mode breakdown",
          ("ship_mode" in content or "ship mode" in content) and any(t in content for t in ["economy", "express", "standard", "next day"]))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace, args.groundtruth_workspace)
    check_word(args.agent_workspace)
    check_script(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed")

    result = {"total_passed": PASS_COUNT, "total_checks": total}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} failures)")
        sys.exit(1)


if __name__ == "__main__":
    main()
