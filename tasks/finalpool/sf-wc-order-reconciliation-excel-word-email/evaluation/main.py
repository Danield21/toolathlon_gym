"""Evaluation script for sf-wc-order-reconciliation-excel-word-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Order_Reconciliation_Report.xlsx
    excel_path = os.path.join(agent_workspace, "Order_Reconciliation_Report.xlsx")
    check("Order_Reconciliation_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Order_Reconciliation_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            # Per-column tolerance (financial values rounded to 2 decimals -> 0.02)
            COL_TOL = {
                "total_orders": 0,           # integer count
                "order_count": 0,            # integer count
                "total_revenue": 0.5,        # rounded to 2 decimals; allow ~0.5 for rounding
                "avg_order_value": 0.05,     # rounded to 2 decimals
                "total_amount": 0.5,
                "dw_value": 0.5,
                "store_value": 0.5,
                "difference": 0.5,
            }

            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Strict row count parity
                    gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if r and r[0] is not None]
                    data_rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if r and r[0] is not None]
                    check(
                        f"{sheet_name} has exactly {len(gt_rows)} data rows",
                        len(data_rows) == len(gt_rows),
                        f"got {len(data_rows)}",
                    )

                    header_map = {h: i for i, h in enumerate(headers)}

                    # Build agent lookup keyed by first column (case-insensitive)
                    if data_rows:
                        agent_lookup = {str(r[0]).strip().lower(): r for r in data_rows}
                    else:
                        agent_lookup = {}

                    for ri, gt_row in enumerate(gt_rows):
                        key = str(gt_row[0]).strip().lower()
                        agent_row = agent_lookup.get(key)
                        if agent_row is None:
                            check(f"{sheet_name} row '{gt_row[0]}' present", False)
                            continue
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                tol = COL_TOL.get(gt_h, max(0.05, abs(gf) * 0.005))  # 0.5% otherwise
                                check(
                                    f"{sheet_name} '{gt_row[0]}' {gt_h} ~{gf:.2f}",
                                    abs(gf - af) <= tol,
                                    f"got {af} (tol={tol})",
                                )
                            elif gv is not None and av is not None:
                                # Strict equality (case-insensitive, trim) for strings
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    check(
                                        f"{sheet_name} '{gt_row[0]}' {gt_h} text equals",
                                        gs == avs,
                                        f"expected '{gs[:80]}', got '{avs[:80]}'",
                                    )

    # Check Reconciliation_Audit.docx
    docx_path = os.path.join(agent_workspace, "Reconciliation_Audit.docx")
    check("Reconciliation_Audit.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        full_text = " ".join([p.text for p in doc.paragraphs])
        check("Reconciliation_Audit.docx has content", len(full_text) > 50, f"text length: {len(full_text)}")
        # Title or top-level heading
        title_or_headings = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading") or p.style.name in ("Title",)
        ]
        thl = [h.lower() for h in title_or_headings]
        # Section headings only
        section_headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        # Required: title (Heading or Title) + 3 section Headings
        EXPECTED_TITLE = "q1 order reconciliation audit report"
        EXPECTED_SECTIONS = [
            "data warehouse analysis",
            "online store analysis",
            "reconciliation findings",
        ]
        check(
            f"Reconciliation_Audit.docx has title '{EXPECTED_TITLE}'",
            any(EXPECTED_TITLE == h or EXPECTED_TITLE in h for h in thl),
            f"titles+headings: {title_or_headings[:8]}",
        )
        for eh in EXPECTED_SECTIONS:
            check(
                f"Reconciliation_Audit.docx has section heading '{eh}'",
                any(eh == h or eh in h for h in section_headings),
                f"section headings: {section_headings}",
            )

    # Check Python script exists (specifically reconciler.py)
    if os.path.isdir(agent_workspace):
        py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
        check(
            "reconciler.py exists in workspace",
            "reconciler.py" in py_files,
            f"found .py files: {py_files}",
        )
        # Intermediate JSON files mentioned in task — must exist AND parse as JSON with non-empty content
        for jname in ["dw_orders.json", "store_orders.json", "reconciliation_results.json"]:
            jpath = os.path.join(agent_workspace, jname)
            check(f"{jname} exists in workspace", os.path.exists(jpath), f"missing: {jpath}")
            if os.path.exists(jpath):
                try:
                    with open(jpath, "r", encoding="utf-8") as fjson:
                        jdata = json.load(fjson)
                    if jname == "reconciliation_results.json":
                        # Should be a non-empty dict mentioning revenue/aov/order keys
                        is_dict = isinstance(jdata, dict)
                        keys_lower = {str(k).lower() for k in jdata.keys()} if is_dict else set()
                        has_required = any(
                            any(tok in k for k in keys_lower)
                            for tok in ("revenue", "aov", "avg_order", "order")
                        )
                        check(
                            f"{jname} contains expected keys (revenue/AOV/order)",
                            is_dict and has_required,
                            f"keys: {sorted(keys_lower)[:10]}",
                        )
                    else:
                        # Other JSONs must be a non-empty list or dict
                        non_empty = bool(jdata) and (isinstance(jdata, (list, dict)))
                        check(f"{jname} is non-empty list/dict", non_empty,
                              f"type={type(jdata).__name__}, len={len(jdata) if hasattr(jdata,'__len__') else 'NA'}")
                except Exception as e:
                    check(f"{jname} parses as valid JSON", False, str(e)[:200])

    # Database checks (email)
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute(
            "SELECT subject, to_addr, body_text FROM email.messages "
            "WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1)"
        )
        sent = cur.fetchall()
        # Find email with subject 'Q1 Order Reconciliation Complete' (case-insensitive exact)
        target_subj = "q1 order reconciliation complete"
        match = None
        for subj, to_addr, body in sent:
            if (subj or "").strip().lower() == target_subj:
                match = (subj, to_addr, body)
                break
        check(
            "Email subject 'Q1 Order Reconciliation Complete' present",
            match is not None,
            f"sent subjects: {[s[0] for s in sent][:10]}",
        )
        if match is not None:
            subj, to_addr, body = match
            # Recipient must be finance@company.com
            ta = str(to_addr or "").lower()
            check(
                "Email recipient is finance@company.com",
                "finance@company.com" in ta,
                f"to_addr: {to_addr}",
            )
            # Body should mention reconciliation/findings
            body_l = (body or "").lower()
            keys = ["reconciliation", "revenue", "order"]
            missing = [k for k in keys if k not in body_l]
            check(
                "Email body summarises key findings",
                not missing,
                f"missing terms: {missing}",
            )
        # Reverse verification
        cur.execute(
            "SELECT COUNT(*) FROM email.messages "
            "WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) "
            "AND subject ILIKE '%newsletter%'"
        )
        noise_sent = cur.fetchone()[0]
        check("No noise emails in Sent folder", noise_sent == 0, f"found {noise_sent} noise emails in Sent")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()