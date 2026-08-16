"""Evaluation script for fetch-sf-kpi-dashboard-excel-word-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
LOCAL_FAIL_COUNT = 0
CURRENT_CATEGORY = "local"

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT, LOCAL_FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        if CURRENT_CATEGORY == "local":
            LOCAL_FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check KPI_Dashboard_Report.xlsx
    excel_path = os.path.join(agent_workspace, "KPI_Dashboard_Report.xlsx")
    check("KPI_Dashboard_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "KPI_Dashboard_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if r and any(c is not None for c in r)]
                    data_rows = [r for r in ws.iter_rows(min_row=2, values_only=True) if r and any(c is not None for c in r)]
                    # Require exactly len(gt_rows) data rows (was len(gt_rows)-2).
                    check(f"{sheet_name} has {len(gt_rows)} data rows", len(data_rows) == len(gt_rows), f"got {len(data_rows)}, expected {len(gt_rows)}")

                    # Sort-order validation per task.md requirements.
                    header_map_sort = {h: i for i, h in enumerate(headers)}
                    if sheet_name == "KPI_Scorecard":
                        ach_i = header_map_sort.get("achievement_pct")
                        if ach_i is not None:
                            vals = [safe_float(r[ach_i]) for r in data_rows if ach_i < len(r)]
                            vals_nn = [v for v in vals if v is not None]
                            check("KPI_Scorecard sorted by Achievement_Pct ascending",
                                  vals_nn == sorted(vals_nn), f"{vals[:8]}")
                    elif sheet_name == "Revenue_Detail":
                        reg_i = header_map_sort.get("region")
                        if reg_i is not None:
                            vals = [str(r[reg_i]).strip().lower() for r in data_rows if reg_i < len(r) and r[reg_i] is not None]
                            check("Revenue_Detail sorted by Region",
                                  vals == sorted(vals), f"{vals[:8]}")

                    # Cell value comparison against groundtruth (ALL rows)
                    header_map = {h: i for i, h in enumerate(headers)}
                    gt_header_map = {h: i for i, h in enumerate(gt_headers)}
                    for ri in range(min(len(gt_rows), len(data_rows))):
                        gt_row = gt_rows[ri]
                        agent_row = data_rows[ri]
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                # Tighter tolerance: 5% or 0.5 abs
                                tol = max(0.5, abs(gf) * 0.05)
                                check(f"{sheet_name} R{ri+2} {gt_h} ~{gf:.2f}",
                                      abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    # Exact equality only (avoid Met/Missed/Near substring collisions)
                                    check(f"{sheet_name} R{ri+2} {gt_h} text",
                                          gs == avs,
                                          f"expected {gs[:50]}, got {avs[:50]}")

    # Check KPI_Review.docx
    docx_path = os.path.join(agent_workspace, "KPI_Review.docx")
    check("KPI_Review.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("KPI_Review.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        gt_doc_path = os.path.join(groundtruth_workspace, "KPI_Review.docx")
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if p.style.name.startswith("Heading")]
            for gh in gt_headings:
                if gh:
                    found = any(gh in h or h in gh for h in headings)
                    check(f"KPI_Review.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("KPI_Review.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Check Python script with specific filename - runtime
    global CURRENT_CATEGORY
    CURRENT_CATEGORY = "runtime"
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("kpi_reconciler.py exists", "kpi_reconciler.py" in py_files, f"found: {py_files}")
    check("kpi_reconciliation.json output exists",
          os.path.exists(os.path.join(agent_workspace, "kpi_reconciliation.json")),
          f"workspace files: {os.listdir(agent_workspace)[:20]}")

    # Database checks - runtime
    try:
        conn = get_conn()
        cur = conn.cursor()
        # Task.md: subject exactly "Q1 KPI Performance Summary" to executive-team@company.com
        # Drop Sent-folder filter: check by subject + recipient only.
        cur.execute("""
            SELECT subject, to_addr FROM email.messages
            WHERE subject ILIKE '%%q1 kpi performance summary%%'
              AND to_addr::text ILIKE '%%executive-team@company.com%%'
        """)
        email_row = cur.fetchone()
        check("Email with subject 'Q1 KPI Performance Summary' to executive-team@company.com",
              email_row is not None, "no matching email found")
        # Reverse verification: noise subject emails (should not be injected by agent as sent)
        cur.execute("""
            SELECT COUNT(*) FROM email.messages
            WHERE subject ILIKE '%%newsletter%%'
              AND to_addr::text ILIKE '%%executive-team@company.com%%'
        """)
        noise_sent = cur.fetchone()[0]
        check("No newsletter emails to executive-team", noise_sent == 0, f"found {noise_sent}")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    print(f"Local FAIL_COUNT: {LOCAL_FAIL_COUNT}, Total FAIL_COUNT: {FAIL_COUNT}")
    return LOCAL_FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks (local_fails={LOCAL_FAIL_COUNT})"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()