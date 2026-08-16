"""
Evaluation for yt-transcript-afrobeat-song-list-excel-email task.

Checks:
1. Afrobeat_Tracklist.xlsx exists in agent_workspace
2. "Tracklist" sheet has >= 10 data rows
3. "Tracklist" sheet has Track_Number, Song_Title, Artist columns
4. "Artist_Summary" sheet exists with >= 4 rows
5. Curator_Notes.docx exists with >= 3 headings; required section titles
6. Curator_Notes.docx contains music-related keywords
7. Notion page exists with "Afrobeat"/"Mix"/"Tracklist" in title
8. Email sent to music@label.com with required subject and body
9. Per-row validation: >=80% Tracklist rows match GT, >=80% Artist_Summary artists match
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Check 1-4: Afrobeat_Tracklist.xlsx ===")
    xlsx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".xlsx") and ("afrobeat" in fname.lower() or "tracklist" in fname.lower()):
            xlsx_path = os.path.join(agent_workspace, fname)
            break

    record("Afrobeat_Tracklist.xlsx exists", xlsx_path is not None,
           f"No matching xlsx found in {agent_workspace}")

    if not xlsx_path:
        record("Tracklist sheet has >= 10 data rows", False, "xlsx not found")
        record("Tracklist sheet has required columns", False, "xlsx not found")
        record("Artist_Summary sheet has >= 4 rows", False, "xlsx not found")
        return

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path)

        # Check Tracklist sheet
        tracklist_sheet = None
        for name in wb.sheetnames:
            if "tracklist" in name.lower() or "track" in name.lower():
                tracklist_sheet = wb[name]
                break
        if tracklist_sheet is None and wb.sheetnames:
            tracklist_sheet = wb[wb.sheetnames[0]]

        if tracklist_sheet:
            data_rows = [r for r in tracklist_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Tracklist sheet has >= 10 data rows", len(data_rows) >= 10,
                   f"Found {len(data_rows)} data rows, expected at least 10")

            headers = [str(c.value).strip() if c.value else "" for c in next(tracklist_sheet.iter_rows(max_row=1))]
            headers_lower = [h.lower() for h in headers]
            has_track_num = any("track" in h and "num" in h or h == "track_number" for h in headers_lower)
            has_song = any("song" in h or "title" in h for h in headers_lower)
            has_artist = any("artist" in h for h in headers_lower)
            record("Tracklist sheet has Track_Number, Song_Title, Artist columns",
                   has_track_num and has_song and has_artist,
                   f"Headers found: {headers}")
        else:
            record("Tracklist sheet has >= 10 data rows", False, "No Tracklist sheet found")
            record("Tracklist sheet has required columns", False, "No Tracklist sheet found")

        # Check Artist_Summary sheet
        summary_sheet = None
        for name in wb.sheetnames:
            if "artist" in name.lower() or "summary" in name.lower():
                summary_sheet = wb[name]
                break

        record("Artist_Summary sheet exists", summary_sheet is not None,
               f"Sheets: {wb.sheetnames}")

        if summary_sheet:
            data_rows = [r for r in summary_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Artist_Summary sheet has >= 4 data rows", len(data_rows) >= 4,
                   f"Found {len(data_rows)} rows")
        else:
            record("Artist_Summary sheet has >= 4 data rows", False, "Sheet not found")

        # --- Groundtruth XLSX value comparison ---
        gt_path = os.path.join(groundtruth_workspace, "Afrobeat_Tracklist.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                # Tracklist: extracted song count is fuzzy; allow >= GT row count for Tracklist; Artist_Summary should match GT
                if "track" in gt_sname.lower():
                    record(f"GT '{gt_sname}' row count >= {len(gt_rows)}",
                           len(a_rows) >= max(10, len(gt_rows) - 2),
                           f"Expected at least {max(10, len(gt_rows) - 2)}, got {len(a_rows)}")
                else:
                    record(f"GT '{gt_sname}' row count >= {len(gt_rows) - 2}",
                           len(a_rows) >= max(4, len(gt_rows) - 2),
                           f"Expected at least {max(4, len(gt_rows) - 2)}, got {len(a_rows)}")
                # For Tracklist (fuzzy from transcript extraction): require >= 80% of GT
                # tracks to appear in agent output (matched by Track_Number OR Song_Title).
                # For Artist_Summary (deterministic from tracklist): require ALL GT artists to
                # appear with correct Song_Count and Total_Duration_Sec.
                if "track" in gt_sname.lower():
                    a_by_num = {}
                    a_by_song = {}
                    for ar in a_rows:
                        if ar and len(ar) > 0 and ar[0] is not None:
                            try:
                                a_by_num[int(ar[0])] = ar
                            except (TypeError, ValueError):
                                pass
                        if ar and len(ar) > 1 and ar[1] is not None:
                            a_by_song[str(ar[1]).strip().lower()] = ar
                    matched_count = 0
                    for gr in gt_rows:
                        if not gr or len(gr) < 2:
                            continue
                        gnum = gr[0]
                        gsong = str(gr[1]).strip().lower() if gr[1] is not None else ""
                        ar = None
                        try:
                            ar = a_by_num.get(int(gnum)) if gnum is not None else None
                        except (TypeError, ValueError):
                            ar = None
                        if ar is None and gsong:
                            ar = a_by_song.get(gsong)
                        # Tighter song match: substring match only when the
                        # shorter string is at least 5 chars (avoid 'love' matching
                        # 'lovely', 'on' matching anything).
                        if ar is None and gsong and len(gsong) >= 5:
                            for asong, arow in a_by_song.items():
                                if not asong or len(asong) < 5:
                                    continue
                                shorter = gsong if len(gsong) <= len(asong) else asong
                                longer = asong if shorter == gsong else gsong
                                if len(shorter) >= 5 and shorter in longer:
                                    ar = arow
                                    break
                        if ar is not None:
                            matched_count += 1
                    threshold = max(int(len(gt_rows) * 0.8), 8)
                    record(f"GT '{gt_sname}' >= {threshold}/{len(gt_rows)} rows matched in agent",
                           matched_count >= threshold,
                           f"matched {matched_count}/{len(gt_rows)}")
                else:
                    # Artist_Summary. The transcript carries NO artist metadata
                    # (only one "it's Kel Vib" drop), so artist attribution is
                    # not derivable from the data source — grading agents
                    # against GT artist NAMES would be arbitrary (c4
                    # case-study 2026-08-15: agent song list matched the
                    # transcript word-for-word but 0/9 artist overlap with the
                    # stale GT). Grade the sheet's INTERNAL CONSISTENCY with the
                    # agent's own Tracklist instead: every summary row's
                    # Song_Count and Total_Duration_Sec must be reproducible
                    # from the agent's track rows.
                    a_track_rows = []
                    for asn in wb.sheetnames:
                        if "track" in asn.lower():
                            a_track_rows = [r for r in wb[asn].iter_rows(min_row=2, values_only=True)
                                            if any(c is not None for c in r)]
                            break
                    # agent tracklist: artist col idx 2, duration col idx 4
                    expected_summary = {}
                    for tr in a_track_rows:
                        if tr and len(tr) >= 5 and tr[2] is not None:
                            akey = str(tr[2]).strip().lower()
                            try:
                                dur = float(tr[4] or 0)
                            except (TypeError, ValueError):
                                dur = 0.0
                            e = expected_summary.setdefault(akey, [0, 0.0])
                            e[0] += 1
                            e[1] += dur
                    artist_match_count = 0
                    artist_value_match = 0
                    for ar in a_rows:
                        if not ar or ar[0] is None:
                            continue
                        akey = str(ar[0]).strip().lower()
                        exp = expected_summary.get(akey)
                        if exp is None:
                            continue
                        artist_match_count += 1
                        try:
                            cnt_ok = abs(float(ar[1]) - exp[0]) < 0.5
                            dur_ok = abs(float(ar[2]) - exp[1]) <= max(exp[1] * 0.10, 2.0)
                        except (TypeError, ValueError, IndexError):
                            cnt_ok = dur_ok = False
                        if cnt_ok and dur_ok:
                            artist_value_match += 1
                    # threshold relative to the AGENT's own sheet size
                    n_agent_artists = max(len([r for r in a_rows if r and r[0] is not None]), 1)
                    artist_threshold = max(int(n_agent_artists * 0.8), 4)
                    record(f"GT '{gt_sname}' >= {artist_threshold} artists present in agent",
                           artist_match_count >= min(artist_threshold, n_agent_artists),
                           f"matched {artist_match_count}/{n_agent_artists}")
                    record(f"GT '{gt_sname}' matched artist rows have correct values (>= 80%)",
                           artist_value_match >= min(artist_threshold, n_agent_artists),
                           f"value-match {artist_value_match}/{n_agent_artists}")
            gt_wb.close()

    except Exception as e:
        record("Tracklist sheet has >= 10 data rows", False, str(e))
        record("Tracklist sheet has required columns", False, str(e))
        record("Artist_Summary sheet exists", False, str(e))
        record("Artist_Summary sheet has >= 4 data rows", False, str(e))


def check_word(agent_workspace):
    print("\n=== Check 5-6: Curator_Notes.docx ===")
    docx_path = None
    # Strict filename match first.
    for fname in os.listdir(agent_workspace):
        if fname.lower() == "curator_notes.docx":
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        # Permissive fallback: filename must contain BOTH 'curator' AND 'notes'.
        for fname in os.listdir(agent_workspace):
            fl = fname.lower()
            if fl.endswith(".docx") and "curator" in fl and "notes" in fl:
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Curator_Notes.docx exists", docx_path is not None,
           f"No matching docx in {agent_workspace}")

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found")
        record("Word doc contains music-related keywords", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings")

        # Validate the three specific section titles required by task.md
        all_para_text = " ".join(p.text for p in doc.paragraphs).lower()
        heading_text = " ".join(p.text for p in headings).lower()
        # Allow the section to appear either as a heading or as text in body
        # Required: main title 'Afrobeat Mix 2024 - Curator Notes' (heading), 'Notable Tracks', 'Genre Analysis'
        has_main = "afrobeat mix 2024" in all_para_text and "curator notes" in all_para_text
        has_notable = "notable tracks" in all_para_text
        has_genre = "genre analysis" in all_para_text
        record("Word doc has 'Afrobeat Mix 2024 - Curator Notes' title",
               has_main, f"Heading text: {heading_text[:200]}")
        record("Word doc has 'Notable Tracks' section",
               has_notable, f"Found in text: {has_notable}")
        record("Word doc has 'Genre Analysis' section",
               has_genre, f"Found in text: {has_genre}")

        keywords = ["track", "song", "artist", "afrobeat", "mix", "music"]
        found = [k for k in keywords if k in all_para_text]
        record("Word doc contains music-related keywords", len(found) >= 3,
               f"Found keywords: {found}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e))
        record("Word doc contains music-related keywords", False, str(e))


def check_notion():
    print("\n=== Check 7: Notion page 'Afrobeat Mix 2024 Tracklist' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT id, properties FROM notion.pages
        ORDER BY created_time DESC
    """)
    rows = cur.fetchall()

    # Title must contain BOTH 'afrobeat' AND 'tracklist' (mix optional but expected).
    matched_id = None
    for row_id, props in rows:
        props_str = json.dumps(props).lower() if props else ""
        if "afrobeat" in props_str and "tracklist" in props_str:
            matched_id = row_id
            break
    record(
        "Notion page exists with title containing 'Afrobeat' AND 'Tracklist'",
        matched_id is not None,
        f"Total notion pages: {len(rows)}",
    )

    if matched_id is not None:
        # Validate content blocks: must reference at least 5 tracks (track number,
        # song title, or artist) so that the page contains the actual tracklist
        # rather than an empty stub.
        try:
            cur.execute(
                """
                SELECT COALESCE(
                    string_agg(block_data::text, ' '),
                    ''
                )
                FROM notion.blocks
                WHERE parent_id = %s
                """,
                (matched_id,),
            )
            block_text = (cur.fetchone() or [""])[0] or ""
        except Exception:
            block_text = ""
        block_text_lower = block_text.lower()
        # Count distinct track-like tokens. We expect at least track numbers 1..5
        # to appear AND at least 3 of these keywords: "track", "song", "artist".
        has_track_numbers = sum(
            1 for n in range(1, 11) if f"{n}." in block_text or f"track {n}" in block_text_lower
        )
        keyword_hits = sum(1 for k in ["track", "song", "artist"] if k in block_text_lower)
        record(
            "Notion page content references multiple tracks (>=5 numbered + >=2 keywords)",
            has_track_numbers >= 5 and keyword_hits >= 2,
            f"Numbered tracks found: {has_track_numbers}, keywords: {keyword_hits}, content len: {len(block_text)}",
        )

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 8: Email sent to music@label.com ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else to_addr.lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "music@label.com" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to music@label.com", matching is not None,
           f"Total messages found: {len(messages)}")
    if matching:
        subj = (matching[0] or "").lower()
        body = matching[3] or ""
        # Subject must contain 'Afrobeat Mix 2024 Tracklist Ready' or close approximation
        subj_ok = ("afrobeat" in subj) and ("tracklist" in subj) and ("2024" in subj or "ready" in subj)
        record("Email subject 'Afrobeat Mix 2024 Tracklist Ready'", subj_ok,
               f"Subject: {matching[0]}")
        # Body must mention number of tracks and reference spreadsheet/curator notes
        body_lower = body.lower()
        body_ok = ("track" in body_lower) and ("spreadsheet" in body_lower or "xlsx" in body_lower or
                                                "curator" in body_lower or "notes" in body_lower)
        record("Email body mentions tracks + spreadsheet/curator notes", body_ok,
               f"Body excerpt: {body[:200]}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    agent_ws = args.agent_workspace

    check_excel(agent_ws, args.groundtruth_workspace)
    check_word(agent_ws)
    check_notion()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} failed checks)")
        sys.exit(1)


if __name__ == "__main__":
    main()
