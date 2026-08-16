"""
Evaluation script for canvas-submission-timeline-gcal task.

Checks:
1. Word file (Late_Submission_Report.docx) - per-course late counts/percentages
2. Google Calendar events - exactly the assignments with due_at, with summary in '[CODE]: [NAME] Due' format and date matching due_at
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

import psycopg2
from zoneinfo import ZoneInfo

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# Source of truth: docs/task.md. The Calendar MCP writes events as UTC
# instants into the DB (`TIMESTAMPTZ`), so all comparisons anchor to the
# timezone that task.md declares for the event's wall-clock semantics.
# Never use bare `st.date()` on rows read from `gcal.events`: psycopg2
# returns them in the PG session `TimeZone` (case-study: compute node
# default was Asia/Shanghai), so `st.date()` can shift a day. Use
# `utils.evaluation.gcal_helpers` instead (session-tz-independent).
# ──────────────────────────────────────────────────────────────────────────
# task.md: "Schedule each event on the assignment due date." Canvas default
# timezone is ET (America/New_York), and the event must match the due_at
# calendar day in that timezone. We compare both sides in ET to stay
# session-tz-independent and self-consistent.
EXPECTED_TIMEZONE = ZoneInfo("America/New_York")

from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

FALL_2013_CODES = ["AAA-2013J", "BBB-2013J", "DDD-2013J", "EEE-2013J", "FFF-2013J", "GGG-2013J"]

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def fetch_canvas_stats():
    """Return per-course totals/late counts and assignment list."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT c.course_code, c.name,
                   COUNT(s.id) AS total,
                   SUM(CASE WHEN s.late THEN 1 ELSE 0 END) AS late_n
            FROM canvas.courses c
            JOIN canvas.assignments a ON a.course_id = c.id
            JOIN canvas.submissions s ON s.assignment_id = a.id
            WHERE c.course_code IN %s
            GROUP BY c.course_code, c.name
            ORDER BY c.course_code
        """, (tuple(FALL_2013_CODES),))
        stats = {}
        for code, name, total, late_n in cur.fetchall():
            total = int(total or 0)
            late_n = int(late_n or 0)
            pct = (late_n / total * 100) if total else 0.0
            stats[code] = {"name": name, "total": total, "late": late_n, "pct": pct}

        # Assignments with due_at
        cur.execute("""
            SELECT c.course_code, a.id, a.name, a.due_at
            FROM canvas.courses c
            JOIN canvas.assignments a ON a.course_id = c.id
            WHERE c.course_code IN %s AND a.due_at IS NOT NULL
            ORDER BY c.course_code, a.due_at
        """, (tuple(FALL_2013_CODES),))
        assignments = [(code, aid, an, due) for code, aid, an, due in cur.fetchall()]
        cur.close()
        conn.close()
        return stats, assignments
    except Exception as e:
        print(f"[fallback] Canvas DB query failed: {e}")
        return None, None


def check_word(agent_workspace, stats):
    print("\n=== Checking Word Output ===")
    docx_path = os.path.join(agent_workspace, "Late_Submission_Report.docx")
    check("Word file exists", os.path.isfile(docx_path), f"Expected {docx_path}")
    if not os.path.isfile(docx_path):
        return

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        check("Word file readable", False, str(e))
        return

    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"
    text_lower = all_text.lower()

    check("Word doc has title 'Fall 2013 Late Submission Analysis'",
          "fall 2013 late submission analysis" in text_lower,
          f"Sample: {all_text[:300]}")

    # Each course: code + name + total + late + pct
    if stats:
        for code, info in stats.items():
            check(f"Word contains course code {code}",
                  code.lower() in text_lower,
                  f"Not found")
            # Match total/late as standalone tokens to avoid '156' matching '15600'
            total_re = re.compile(r"(?<!\d)" + re.escape(str(info['total'])) + r"(?!\d)")
            late_re = re.compile(r"(?<!\d)" + re.escape(str(info['late'])) + r"(?!\d)")
            check(f"Word contains total submissions {info['total']} for {code}",
                  total_re.search(all_text) is not None,
                  f"Expected total {info['total']} as standalone number")
            check(f"Word contains late count {info['late']} for {code}",
                  late_re.search(all_text) is not None,
                  f"Expected late {info['late']} as standalone number")
            # Percentage tolerance (allow either rounded to 1 or 2 decimals)
            pct = info['pct']
            pct_strs = [f"{pct:.0f}", f"{pct:.1f}", f"{pct:.2f}"]
            found_pct = any(p in all_text for p in pct_strs)
            check(f"Word contains late pct ~{pct:.1f}% for {code}",
                  found_pct,
                  f"Expected one of {pct_strs}")
    else:
        # Fallback: at minimum mentions all 6 codes
        for code in FALL_2013_CODES:
            check(f"Word contains course code {code}",
                  code.lower() in text_lower,
                  "Not found")


def check_calendar(assignments):
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, start_datetime, description
            FROM gcal.events
            ORDER BY start_datetime
        """)
        events = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar DB query", False, str(e))
        return

    print(f"  [INFO] Found {len(events)} total calendar events.")

    # Filter to events whose summary mentions a Fall 2013 course code
    course_events = []
    for summary, st, desc in events:
        if not summary:
            continue
        for code in FALL_2013_CODES:
            if code.upper() in summary.upper():
                course_events.append((summary, st, code))
                break

    expected_n = len(assignments) if assignments else 51
    check(f"Calendar has exactly {expected_n} events with Fall 2013 course codes",
          len(course_events) == expected_n,
          f"Found {len(course_events)} of expected {expected_n}")

    # Per-course coverage (all 6)
    by_code = {}
    for summary, st, code in course_events:
        by_code.setdefault(code, []).append((summary, st))

    for code in FALL_2013_CODES:
        n = len(by_code.get(code, []))
        # If we have ground-truth assignments, require matching count per course
        if assignments:
            expected_for_code = sum(1 for c, _, _, _ in assignments if c == code)
            check(f"Course {code} has exactly {expected_for_code} calendar events",
                  n == expected_for_code,
                  f"Got {n}")
        else:
            check(f"Course {code} has at least 1 event",
                  n >= 1,
                  f"Got {n}")

    # Validate summary format: '[CODE]: [NAME] Due'
    pattern = re.compile(r"^\s*([A-Z]{3}-2013J)\s*:\s*.+?\s+Due\s*$", re.IGNORECASE)
    matched = sum(1 for summary, _, _ in course_events if summary and pattern.match(summary.strip()))
    if course_events:
        ratio = matched / len(course_events)
        check("All calendar event summaries follow '[CODE]: [NAME] Due' format",
              ratio == 1.0,
              f"{matched}/{len(course_events)} matched the format")

    # Check event start_datetime matches assignment due_at
    if assignments and course_events:
        # Build map: code -> set of due_at dates, expressed in
        # EXPECTED_TIMEZONE (ET) so the comparison is session-tz-independent
        # and self-consistent with the gcal event side below.
        due_dates = {}
        for code, aid, aname, due in assignments:
            if due:
                due_d, _, _ = get_zone_components(due, EXPECTED_TIMEZONE)
                due_dates.setdefault(code, set()).add(due_d)

        # Check each event start date (in ET) matches some due_at.
        match_count = 0
        for summary, st, code in course_events:
            if st and code in due_dates:
                ev_date, _, _ = get_zone_components(st, EXPECTED_TIMEZONE)
                if ev_date in due_dates[code]:
                    match_count += 1
        ratio = match_count / len(course_events) if course_events else 0
        check("All events match an assignment due_at date",
              ratio == 1.0,
              f"{match_count}/{len(course_events)} match due dates")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    stats, assignments = fetch_canvas_stats()

    check_word(args.agent_workspace, stats)
    check_calendar(assignments)

    all_ok = FAIL_COUNT == 0
    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": all_ok}, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
