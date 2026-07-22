"""Evaluation for sf-sales-customer-segment-gsheet-word.

Checks:
1. Google Sheet "Customer Segment Analysis" with 4 segment rows
   - Contains Consumer, Enterprise, Government, SMB
   - Revenue data matches actual DB values (within tolerance)
2. Word doc Customer_Segment_Report.docx with:
   - Heading "Customer Segment Analysis Report"
   - Table with segment data
   - Recommendations section
3. Email to sales-strategy@company.example.com
   Subject: "Customer Segment Analysis Report"
"""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

# Expected segment data (actual from sf_data query) - dynamic fallback
_SEGMENT_FALLBACK = [
    {"name": "Consumer", "customers": 532, "orders": 5423, "revenue": 839609.20},
    {"name": "Enterprise", "customers": 513, "orders": 5058, "revenue": 793741.69},
    {"name": "Government", "customers": 474, "orders": 4679, "revenue": 712686.66},
    {"name": "SMB", "customers": 481, "orders": 4840, "revenue": 702960.78},
]


def _fetch_segments():
    try:
        conn = psycopg2.connect(**DB, connect_timeout=5)
        cur = conn.cursor()
        cur.execute("""
            SELECT c."SEGMENT", COUNT(DISTINCT c."CUSTOMER_ID"),
                   COUNT(o."ORDER_ID"), ROUND(SUM(o."TOTAL_AMOUNT")::numeric, 2)
            FROM sf_data."SALES_DW__PUBLIC__CUSTOMERS" c
            JOIN sf_data."SALES_DW__PUBLIC__ORDERS" o ON c."CUSTOMER_ID" = o."CUSTOMER_ID"
            GROUP BY c."SEGMENT"
        """)
        result = []
        for name, cust, orders, rev in cur.fetchall():
            result.append({"name": name, "customers": int(cust), "orders": int(orders), "revenue": float(rev)})
        cur.close(); conn.close()
        return result if result else None
    except Exception as e:
        print(f"[fallback] {e}")
        return None


SEGMENTS = _fetch_segments() or _SEGMENT_FALLBACK
_TOTAL_REV = sum(s["revenue"] for s in SEGMENTS)


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=500.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE LOWER(title) LIKE '%customer%' AND LOWER(title) LIKE '%segment%'
    """)
    sheets = cur.fetchall()
    check("Customer Segment Analysis spreadsheet exists", len(sheets) >= 1,
          f"Found {len(sheets)} matching spreadsheets")

    if sheets:
        ss_id = sheets[0][0]
        # Build row-wise cell grid
        cur.execute("""
            SELECT c.sheet_id, c.row_index, c.col_index, c.value FROM gsheet.cells c
            WHERE c.spreadsheet_id = %s
            ORDER BY c.sheet_id, c.row_index, c.col_index
        """, (ss_id,))
        cells = cur.fetchall()
        grids = {}
        for sheet_id, ri, ci, v in cells:
            grids.setdefault(sheet_id, {}).setdefault(ri, {})[ci] = v

        # Take the first sheet with data
        sheet_id = next(iter(grids)) if grids else None
        grid = grids.get(sheet_id, {})

        header_row = grid.get(0, {})
        headers = {}  # name -> col_idx
        for ci, v in header_row.items():
            if v is not None:
                headers[str(v).strip().lower()] = ci

        # Find col indices
        seg_col = headers.get("segment")
        cust_col = headers.get("customer_count") or headers.get("customers")
        order_col = headers.get("order_count") or headers.get("orders")
        rev_col = headers.get("total_revenue") or headers.get("revenue")
        pct_col = headers.get("revenue_pct") or headers.get("revenue %") or headers.get("revenue_pct_of_total")

        check("GSheet has Segment column", seg_col is not None, f"Headers: {list(headers)}")
        check("GSheet has Customer_Count column", cust_col is not None, f"Headers: {list(headers)}")
        check("GSheet has Order_Count column", order_col is not None, f"Headers: {list(headers)}")
        check("GSheet has Total_Revenue column", rev_col is not None, f"Headers: {list(headers)}")
        check("GSheet has Revenue_Pct column", pct_col is not None, f"Headers: {list(headers)}")

        # Data rows (rows >= 1)
        data_rows = [grid[r] for r in sorted(grid) if r >= 1 and any(grid[r].values())]
        check("GSheet has exactly 4 segment rows", len(data_rows) == 4,
              f"Found {len(data_rows)} data rows")

        # Map agent data by segment name
        agent_map = {}
        if seg_col is not None:
            for row in data_rows:
                seg_val = row.get(seg_col)
                if seg_val:
                    agent_map[str(seg_val).strip().lower()] = row

        def parse_num(s):
            try:
                return float(str(s).replace(",", "").replace("%", "").replace("$", "").strip())
            except (ValueError, TypeError):
                return None

        for seg in SEGMENTS:
            a_row = agent_map.get(seg["name"].lower())
            if a_row is None:
                check(f"Segment '{seg['name']}' present", False, f"Rows: {list(agent_map)}")
                continue
            check(f"Segment '{seg['name']}' present", True)
            # Customer_Count exact
            if cust_col is not None:
                v = parse_num(a_row.get(cust_col))
                check(f"{seg['name']} Customer_Count == {seg['customers']}",
                      v == seg["customers"], f"Got {v}")
            # Order_Count exact
            if order_col is not None:
                v = parse_num(a_row.get(order_col))
                check(f"{seg['name']} Order_Count == {seg['orders']}",
                      v == seg["orders"], f"Got {v}")
            # Total_Revenue tol 5.0
            if rev_col is not None:
                v = parse_num(a_row.get(rev_col))
                ok = v is not None and abs(v - seg["revenue"]) <= 5.0
                check(f"{seg['name']} Total_Revenue ~{seg['revenue']}",
                      ok, f"Got {v}")
            # Revenue_Pct tol 0.05
            if pct_col is not None:
                v = parse_num(a_row.get(pct_col))
                expected_pct = round(seg["revenue"] / _TOTAL_REV * 100, 2)
                ok = v is not None and abs(v - expected_pct) <= 0.05
                check(f"{seg['name']} Revenue_Pct ~{expected_pct}",
                      ok, f"Got {v}, expected {expected_pct}")

        # Sort order: rows by Total_Revenue DESC
        if rev_col is not None:
            rev_values = [parse_num(r.get(rev_col)) for r in data_rows]
            rev_values = [v for v in rev_values if v is not None]
            check("GSheet sorted by Total_Revenue DESC",
                  rev_values == sorted(rev_values, reverse=True),
                  f"Values: {rev_values}")

    cur.close()
    conn.close()


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "Customer_Segment_Report.docx")
    check("Customer_Segment_Report.docx exists", os.path.isfile(doc_path),
          f"Expected at {doc_path}")
    if not os.path.isfile(doc_path):
        return

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        return

    full_text = " ".join(p.text.lower() for p in doc.paragraphs)
    headings = [p.text.lower() for p in doc.paragraphs
                if p.style.name.startswith("Heading")]

    check("Word doc has 'Customer Segment Analysis Report' heading",
          any("customer segment analysis" in h for h in headings) or
          "customer segment analysis report" in full_text,
          "Heading not found")

    check("Word doc has a table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    if doc.tables:
        # Check first table has required columns (Segment, Customer_Count, Order_Count, Total_Revenue)
        table = doc.tables[0]
        if len(table.rows) > 0:
            header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
            for col_name, keywords in [
                ("Segment", ["segment"]),
                ("Customer_Count", ["customer_count", "customer count", "customers"]),
                ("Order_Count", ["order_count", "order count", "orders"]),
                ("Total_Revenue", ["total_revenue", "total revenue", "revenue"]),
            ]:
                ok = any(any(kw in hc for hc in header_cells) for kw in keywords)
                check(f"Word table has {col_name} column", ok, f"Headers: {header_cells}")

        table_text = " ".join(
            cell.text.lower()
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        for seg in SEGMENTS:
            check(f"Table contains segment '{seg['name']}'",
                  seg["name"].lower() in table_text,
                  f"Not found in table")
            # Numeric value present: customer count
            check(f"Table contains Customer_Count for '{seg['name']}' ({seg['customers']})",
                  str(seg["customers"]) in table_text,
                  f"Count {seg['customers']} not found in table")

    check("Word doc has Recommendations section",
          "recommendation" in full_text or "recommend" in full_text,
          "Recommendations not found")

    check("Word doc mentions Consumer as top revenue segment",
          "consumer" in full_text, "Consumer segment not mentioned")


def check_emails():
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    conn.close()

    def parse_recipients(to_addr):
        if to_addr is None:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        to_str = str(to_addr).strip()
        try:
            parsed = json.loads(to_str)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [to_str.lower()]
        except (json.JSONDecodeError, TypeError):
            return [to_str.lower()]

    target = "sales-strategy@company.example.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if target in recipients:
            found = (subj, from_addr, to_addr, body)
            break

    check("Email sent to sales-strategy@company.example.com", found is not None)
    if found:
        subj, from_addr, to_addr, body = found
        check("Email from analytics@company.example.com",
              "analytics@company.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}")
        check("Subject is exactly 'Customer Segment Analysis Report'",
              (subj or "").strip().lower() == "customer segment analysis report",
              f"Subject: {subj}")
        # Body must mention top-revenue segment and top-customer segment
        body_lower = (body or "").lower()
        top_rev = max(SEGMENTS, key=lambda s: s["revenue"])
        top_cust = max(SEGMENTS, key=lambda s: s["customers"])
        check(f"Body mentions top-revenue segment '{top_rev['name']}'",
              top_rev["name"].lower() in body_lower,
              f"Body sample: {body_lower[:200]}")
        check(f"Body mentions top-customer segment '{top_cust['name']}'",
              top_cust["name"].lower() in body_lower,
              f"Body sample: {body_lower[:200]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SALES CUSTOMER SEGMENT GSHEET WORD - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_emails()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
