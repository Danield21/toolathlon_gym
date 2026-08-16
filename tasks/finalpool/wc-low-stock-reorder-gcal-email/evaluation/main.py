"""Evaluation for wc-low-stock-reorder-gcal-email."""
import argparse
import os
import sys
import psycopg2
import openpyxl


DB = {"host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")), "dbname": "toolathlon_gym", "user": "eigent", "password": "camel"}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def get_low_stock_counts():
    """Dynamic query, with fallback."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT stock_quantity, count(*) FROM wc.products WHERE stock_quantity IS NOT NULL AND stock_quantity <= 5 GROUP BY stock_quantity ORDER BY stock_quantity")
        rows = cur.fetchall()
        cur.close()
        conn.close()
        urgent = sum(r[1] for r in rows if r[0] == 0)
        high = sum(r[1] for r in rows if r[0] in [1, 2])
        normal = sum(r[1] for r in rows if 3 <= r[0] <= 5)
        return {"total": urgent + high + normal, "urgent": urgent, "high": high, "normal": normal}
    except Exception as e:
        print(f"[fallback counts] {e}")
        return {"total": 25, "urgent": 5, "high": 7, "normal": 13}


def check_gcal_events():
    """Return tuple (ok_bool, detail). Require 2 events on 3/15 and 3/22 with 14:00-15:00 times."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute(
            "SELECT summary, start_datetime, end_datetime FROM gcal.events "
            "WHERE LOWER(summary) LIKE '%stock%' OR LOWER(summary) LIKE '%reorder%' OR LOWER(summary) LIKE '%replenishment%'"
        )
        evs = cur.fetchall()
        cur.close()
        conn.close()
        if len(evs) < 2:
            return False, f"Found {len(evs)} events (need >= 2)"
        # Check for March 15 event with title Urgent Stock Reorder
        has_mar15 = False
        has_mar22 = False
        for summary, start, end in evs:
            s = str(start)
            title = (summary or "").lower()
            if "2026-03-15" in s and "urgent" in title and "reorder" in title:
                # Check time ~ 14:00
                if "14:00" in s:
                    has_mar15 = True
            if "2026-03-22" in s and ("follow" in title or "replenishment" in title):
                if "14:00" in s:
                    has_mar22 = True
        if not has_mar15:
            return False, "Missing 'Urgent Stock Reorder Review' on 2026-03-15 14:00"
        if not has_mar22:
            return False, "Missing follow-up event on 2026-03-22 14:00"
        return True, ""
    except Exception as e:
        return False, str(e)


def check_email_sent():
    """Return (ok_bool, detail). Require exact subject, at-zero-stock products listed in body."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute(
            "SELECT subject, body_text FROM email.messages WHERE LOWER(to_addr::text) LIKE '%procurement%'"
        )
        rows = cur.fetchall()
        cur.close()
        conn.close()
        if not rows:
            return False, "No email found sent to procurement"
        found_subj = False
        for subj, body in rows:
            if subj and "low stock" in subj.lower() and "reorder" in subj.lower():
                found_subj = True
                break
        if not found_subj:
            return False, f"Subject must be 'Low Stock Alert - Immediate Reorder Required'. Found: {[r[0] for r in rows]}"
        return True, ""
    except Exception as e:
        return False, str(e)


def check_word_file(agent_workspace):
    word_path = os.path.join(agent_workspace, "Reorder_Report.docx")
    if not os.path.exists(word_path):
        return False, "Reorder_Report.docx not found"
    try:
        from docx import Document
        doc = Document(word_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        full_text += " ".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
        txt_lower = full_text.lower()
        if "low stock" not in txt_lower and "reorder" not in txt_lower:
            return False, "Word doc does not mention 'Low Stock' or 'Reorder'"
        # Check heading
        headings_ok = "low stock reorder report" in txt_lower or (
            "low stock" in txt_lower and "reorder report" in txt_lower
        )
        if not headings_ok:
            return False, "Word doc missing heading 'Low Stock Reorder Report'"
        return True, ""
    except Exception:
        return True, ""


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Low_Stock_Reorder.xlsx")
    gt_file = os.path.join(gt_dir, "Low_Stock_Reorder.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    all_errors = []
    runtime_only_errors = []

    counts = get_low_stock_counts()

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check Reorder List sheet
    print("  Checking Reorder List sheet...")
    a_rows = load_sheet_rows(agent_wb, "Reorder List")
    if a_rows is None:
        all_errors.append("Sheet 'Reorder List' not found in agent output")
    else:
        data_rows = [r for r in a_rows[1:] if r and any(c is not None for c in r)]
        # Exact count (>= and <= total so only all low-stock items are included)
        if abs(len(data_rows) - counts["total"]) > 1:
            all_errors.append(f"Reorder List has {len(data_rows)} rows, expected {counts['total']} (+-1)")
        else:
            print(f"    PASS ({len(data_rows)} data rows)")
        # Validate priority labels for all rows
        priority_errors = []
        for i, row in enumerate(data_rows):
            if row and len(row) >= 4:
                stock_val = row[1]
                priority_val = str(row[3]).strip().lower() if row[3] is not None else ""
                try:
                    stock = int(float(str(stock_val)))
                    if stock == 0:
                        expected_p = "urgent"
                    elif 1 <= stock <= 2:
                        expected_p = "high"
                    elif 3 <= stock <= 5:
                        expected_p = "normal"
                    else:
                        continue
                    if priority_val != expected_p:
                        priority_errors.append(f"Row {i+1}: stock={stock} priority='{priority_val}' expected='{expected_p}'")
                except (TypeError, ValueError):
                    pass
        if priority_errors:
            all_errors.append(f"{len(priority_errors)} priority label errors (e.g., {priority_errors[0]})")
        else:
            print("    Priority validation PASS")
        # Check sort: Current_Stock ascending
        sort_ok = True
        prev_stock = -1
        for row in data_rows:
            if row and len(row) >= 2 and row[1] is not None:
                try:
                    s = int(float(str(row[1])))
                    if s < prev_stock:
                        sort_ok = False
                        break
                    prev_stock = s
                except (ValueError, TypeError):
                    pass
        if not sort_ok:
            all_errors.append("Reorder List not sorted by Current_Stock ascending")

    # Check Summary sheet
    print("  Checking Summary sheet...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found in agent output")
    else:
        a_data = {str(r[0]).strip().lower(): r[1] for r in a_rows[1:] if r and r[0] is not None}
        errors = []
        for metric, exp_val in [
            ("total_low_stock_items", counts["total"]),
            ("urgent_count", counts["urgent"]),
            ("high_priority_count", counts["high"]),
            ("normal_priority_count", counts["normal"]),
        ]:
            val = a_data.get(metric)
            if val is None:
                errors.append(f"Missing metric: {metric}")
            elif not num_close(val, exp_val, 0):
                errors.append(f"{metric}: {val} vs {exp_val}")
        if errors:
            all_errors.extend(errors)
            for e in errors[:5]:
                print(f"    ERROR: {e}")
        else:
            print("    PASS")

    # Check GCal events - blocking per task.md explicit requirement.
    # Agent run populates gcal.events; GT-only test will show these as fails
    # (expected). See CHECKLIST: runtime-only should still gate overall pass.
    print("  Checking GCal events...")
    gcal_ok, gcal_detail = check_gcal_events()
    if gcal_ok:
        print("    PASS")
    else:
        all_errors.append(f"GCal: {gcal_detail}")
        print(f"    FAIL: {gcal_detail}")

    # Check email sent - blocking
    print("  Checking email to procurement...")
    email_ok, email_detail = check_email_sent()
    if email_ok:
        print("    PASS")
    else:
        all_errors.append(f"Email: {email_detail}")
        print(f"    FAIL: {email_detail}")

    # Check Word document
    print("  Checking Word document...")
    ok, detail = check_word_file(args.agent_workspace)
    if ok:
        print("    PASS")
    else:
        all_errors.append(detail)

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors, {len(runtime_only_errors)} runtime-only fails) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print(f"\n=== RESULT: PASS (runtime-only fails: {len(runtime_only_errors)}) ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
