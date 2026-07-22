"""Evaluation script for pw-sf-hr-skills-gap-excel-word."""
import os
import argparse, json, os, sys
import openpyxl
from docx import Document as DocxDocument

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    
    # Check 1: Excel file exists and has correct structure
    excel_path = os.path.join(agent_workspace, "Skills_Gap_Analysis.xlsx")
    check("Excel file exists", os.path.exists(excel_path))

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)

        # Sheet 1: Department_Overview
        check("Department_Overview sheet exists", "Department_Overview" in wb.sheetnames)
        if "Department_Overview" in wb.sheetnames:
            ws = wb["Department_Overview"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Department_Overview has 7 departments", len(rows) == 7, f"got {len(rows)}")

            # Load groundtruth
            gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Analysis.xlsx")
            gt_wb = openpyxl.load_workbook(gt_path)
            gt_ws = gt_wb["Department_Overview"]
            gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))

            # GT columns: Department, Employee_Count, Avg_Salary, Avg_Experience,
            #             Avg_Performance, Industry_Benchmark_Salary, Salary_Gap
            COL_NAMES = ["Department", "Employee_Count", "Avg_Salary",
                         "Avg_Experience", "Avg_Performance",
                         "Industry_Benchmark_Salary", "Salary_Gap"]
            TOL = [None, 1, 10.0, 0.05, 0.02, 500, 500]  # per-column absolute tolerance
            for gt_row in gt_rows:
                dept = gt_row[0]
                agent_row = None
                for r in rows:
                    if r[0] and str(r[0]).strip().lower() == str(dept).strip().lower():
                        agent_row = r
                        break
                if agent_row:
                    # Validate ALL 7 columns, not just salary + gap
                    for ci in range(1, 7):
                        gv = safe_float(gt_row[ci])
                        av = safe_float(agent_row[ci])
                        if gv is None:
                            continue
                        tol = TOL[ci]
                        if av is None:
                            check(f"{dept} {COL_NAMES[ci]}", False,
                                  f"expected ~{gv}, got None")
                        else:
                            check(f"{dept} {COL_NAMES[ci]}",
                                  abs(gv - av) <= tol,
                                  f"expected ~{gv}, got {av}")

        # Sheet 2: Skills_Matrix
        check("Skills_Matrix sheet exists", "Skills_Matrix" in wb.sheetnames)
        if "Skills_Matrix" in wb.sheetnames:
            ws = wb["Skills_Matrix"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            # Match GT row count and cell values (Role+Skill key)
            gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Analysis.xlsx")
            try:
                gt_wb2 = openpyxl.load_workbook(gt_path)
                gt_sm = gt_wb2["Skills_Matrix"]
                gt_sm_rows = [r for r in gt_sm.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                expected_cnt = len(gt_sm_rows)
            except Exception:
                expected_cnt = 17
                gt_sm_rows = []
            check(f"Skills_Matrix has {expected_cnt} skills",
                  len(rows) == expected_cnt, f"got {len(rows)}")

            # GT Skills_Matrix cols: Role, Skill, Proficiency_Level, Demand_Score,
            #                        Industry_Salary, Our_Avg_Salary, Gap_Pct
            # Match by (Role, Skill) key; validate Demand_Score, Industry_Salary,
            # Our_Avg_Salary, Gap_Pct per row
            if gt_sm_rows:
                # Build agent lookup by (role, skill)
                agent_lookup = {}
                for ar in rows:
                    if ar and len(ar) >= 2 and ar[0] and ar[1]:
                        k = (str(ar[0]).strip().lower(), str(ar[1]).strip().lower())
                        agent_lookup[k] = ar
                for gr in gt_sm_rows[:8]:  # check first 8 rows to limit output
                    if not (gr and len(gr) >= 2 and gr[0] and gr[1]):
                        continue
                    k = (str(gr[0]).strip().lower(), str(gr[1]).strip().lower())
                    ar = agent_lookup.get(k)
                    if not ar:
                        check(f"Skills_Matrix row '{gr[0]}/{gr[1]}' present", False,
                              "Missing in agent")
                        continue
                    # Demand score (col 3), Industry Salary (col 4), Our Avg Salary (col 5), Gap_Pct (col 6)
                    tol_map = {3: 2.0, 4: 500, 5: 100, 6: 1.0}
                    for ci, tol in tol_map.items():
                        if ci >= len(gr) or ci >= len(ar):
                            continue
                        gv = safe_float(gr[ci])
                        av = safe_float(ar[ci])
                        if gv is None:
                            continue
                        if av is None:
                            check(f"Skills_Matrix '{gr[0]}/{gr[1]}' col{ci+1}",
                                  False, "Agent value None")
                        else:
                            check(f"Skills_Matrix '{gr[0]}/{gr[1]}' col{ci+1}",
                                  abs(gv - av) <= tol,
                                  f"expected ~{gv}, got {av}")

        # Sheet 3: Summary - validate ALL rows against GT
        check("Summary sheet exists", "Summary" in wb.sheetnames)
        if "Summary" in wb.sheetnames:
            ws = wb["Summary"]
            agent_rows = {str(r[0]).strip(): r[1] for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}
            try:
                gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Analysis.xlsx")
                gt_wb2 = openpyxl.load_workbook(gt_path)
                gt_sum = gt_wb2["Summary"]
                gt_rows = {str(r[0]).strip(): r[1] for r in gt_sum.iter_rows(min_row=2, values_only=True) if r[0]}
            except Exception:
                gt_rows = {"Total_Departments": 7}
            for metric, exp_val in gt_rows.items():
                actual = agent_rows.get(metric)
                if exp_val is None:
                    continue
                exp_f = safe_float(exp_val)
                ag_f = safe_float(actual)
                if exp_f is not None and ag_f is not None:
                    tol = max(1.0, abs(exp_f) * 0.05)
                    check(f"Summary '{metric}' ~ {exp_f}",
                          abs(exp_f - ag_f) <= tol, f"got {actual}")
                elif exp_val is not None and actual is not None:
                    check(f"Summary '{metric}' matches",
                          str(exp_val).strip().lower() in str(actual).strip().lower()
                          or str(actual).strip().lower() in str(exp_val).strip().lower(),
                          f"expected {exp_val}, got {actual}")

    # Check 2: Word document
    word_path = os.path.join(agent_workspace, "Skills_Gap_Report.docx")
    check("Word report exists", os.path.exists(word_path))
    if os.path.exists(word_path):
        from docx import Document
        doc = Document(word_path)
        full_text = " ".join([p.text for p in doc.paragraphs]).lower()
        check("Word contains 'skills gap'", "skills gap" in full_text)
        check("Word contains 'recommendation'", "recommend" in full_text)
        check("Word contains 'benchmark'", "benchmark" in full_text)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        check("Word has at least 3 headings", len(headings) >= 3, f"got {len(headings)}")

        # Load GT headings and require all of them exist in agent headings (text equality after normalization)
        try:
            gt_doc_path = os.path.join(groundtruth_workspace, "Skills_Gap_Report.docx")
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs
                           if p.style and p.style.name and p.style.name.startswith("Heading")
                           and p.text.strip()]
            agent_headings_norm = [h.strip().lower() for h in headings if h.strip()]
            missing = [gh for gh in gt_headings if gh not in agent_headings_norm]
            check("Word headings match GT exactly",
                  len(missing) == 0,
                  f"Missing: {missing[:3]}")
        except Exception as e:
            check("Word GT headings accessible", False, str(e))

    # Check 3: Terminal artifacts - only check script when running on real agent ws
    try:
        gt_canon = os.path.realpath(groundtruth_workspace)
        ag_canon = os.path.realpath(agent_workspace)
    except Exception:
        gt_canon, ag_canon = groundtruth_workspace, agent_workspace

    if gt_canon != ag_canon:
        script_path = os.path.join(agent_workspace, "skills_analysis.py")
        check("Python script exists", os.path.exists(script_path))
    json_path = os.path.join(agent_workspace, "raw_skills_data.json")
    check("Raw JSON data exists", os.path.exists(json_path))
    processed_path = os.path.join(agent_workspace, "processed_analysis.json")
    check("Processed JSON exists", os.path.exists(processed_path))
    # Validate processed_analysis.json structure (not just existence)
    if os.path.exists(processed_path):
        try:
            with open(processed_path) as f:
                data = json.load(f)
            check("processed_analysis.json is non-empty JSON",
                  bool(data) and (isinstance(data, dict) or isinstance(data, list)),
                  f"type={type(data).__name__}")
        except Exception as e:
            check("processed_analysis.json readable JSON", False, str(e))
    

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
