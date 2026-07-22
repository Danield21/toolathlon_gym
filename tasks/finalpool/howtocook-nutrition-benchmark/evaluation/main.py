"""
Evaluation for howtocook-nutrition-benchmark task.
Checks Excel and Word document structure. Recipe names are dynamic (from HowToCook MCP).
"""
import argparse
import os
import sys

import openpyxl

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try: return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError): return False


def str_match(a, b):
    if a is None or b is None: return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower().replace(" ", "_") == sheet_name.strip().lower().replace(" ", "_"):
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
        if name.strip().lower().replace("_", " ") == sheet_name.strip().lower().replace("_", " "):
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def find_col(header, names):
    if not header:
        return None
    for i, cell in enumerate(header):
        if cell is None:
            continue
        cl = str(cell).strip().lower().replace(" ", "_")
        for n in names:
            if n.lower().replace(" ", "_") == cl:
                return i
    return None


def check_excel(workspace, groundtruth_workspace="."):
    print("\n=== Checking Excel ===")
    path = os.path.join(workspace, "Nutrition_Benchmark.xlsx")
    if not os.path.isfile(path):
        record("Excel exists", False, f"Not found: {path}")
        return False
    record("Excel exists", True)

    wb = openpyxl.load_workbook(path, data_only=True)

    # Recipe Analysis
    ra_rows = load_sheet_rows(wb, "Recipe Analysis") or load_sheet_rows(wb, "Recipe_Analysis")
    if ra_rows is None:
        record("Sheet 'Recipe Analysis' exists", False, f"Sheets: {wb.sheetnames}")
        return False
    record("Sheet 'Recipe Analysis' exists", True)

    header = ra_rows[0]
    data = [r for r in ra_rows[1:] if r and r[0] is not None and str(r[0]).strip()]
    record("Recipe Analysis has >= 8 rows", len(data) >= 8, f"Found {len(data)}")

    cal_col = find_col(header, ["Est_Calories", "Est Calories", "Estimated_Calories", "Calories"])
    if cal_col is not None:
        # Validate ALL rows have a calories value in plausible per-meal range
        valid_count = 0
        invalid = []
        for r in data:
            if cal_col < len(r) and r[cal_col] is not None:
                try:
                    cal = float(r[cal_col])
                    if 50 <= cal <= 3000:
                        valid_count += 1
                    else:
                        invalid.append((r[0], cal))
                except (TypeError, ValueError):
                    invalid.append((r[0], r[cal_col]))
        record(f"All {len(data)} rows have valid calorie values (50-3000)",
               valid_count == len(data),
               f"{len(invalid)} invalid: {invalid[:3]}")
    else:
        record("Est_Calories column exists", False, f"Header: {header}")

    meets_col = find_col(header, ["Meets_Guidelines", "Meets Guidelines", "Guidelines"])
    if meets_col is not None:
        # Validate every row has a Yes or No value
        valid_yn = 0
        invalid = []
        for r in data:
            if meets_col < len(r) and r[meets_col] is not None:
                v = str(r[meets_col]).strip().lower()
                if v in ("yes", "no"):
                    valid_yn += 1
                else:
                    invalid.append((r[0], v))
            else:
                invalid.append((r[0], None))
        record(f"All {len(data)} rows have Yes/No in Meets_Guidelines",
               valid_yn == len(data),
               f"{len(invalid)} invalid: {invalid[:3]}")

        # Cross-check: must have BOTH Yes and No
        all_vals = {str(r[meets_col]).strip().lower() for r in data if meets_col < len(r) and r[meets_col]}
        record("Meets_Guidelines includes both Yes and No values",
               "yes" in all_vals and "no" in all_vals,
               f"vals={all_vals}")
    else:
        record("Meets_Guidelines column exists", False, f"Header: {header}")

    # Meal Plan Suggestions
    mp_rows = load_sheet_rows(wb, "Meal Plan Suggestions") or load_sheet_rows(wb, "Meal_Plan_Suggestions")
    if mp_rows is None:
        record("Sheet 'Meal Plan Suggestions' exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Sheet 'Meal Plan Suggestions' exists", True)
        data2 = [r for r in mp_rows[1:] if r and r[0] is not None]
        record("Meal Plan has == 9 rows (3 days x 3 meals)", len(data2) == 9,
               f"Found {len(data2)}")
        # All 3 days x 3 meals pairs covered
        mp_pairs = set()
        for r in data2:
            if len(r) >= 2 and r[0] and r[1]:
                mp_pairs.add((str(r[0]).strip().lower(), str(r[1]).strip().lower()))
        record(f"Meal Plan has 9 unique Day/Meal pairs",
               len(mp_pairs) == 9, f"got {len(mp_pairs)}")

    # Summary
    sum_rows = load_sheet_rows(wb, "Summary")
    if sum_rows is None:
        record("Sheet 'Summary' exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Sheet 'Summary' exists", True)
        metrics = {}
        for row in sum_rows[1:]:
            if row and row[0]:
                metrics[str(row[0]).strip().lower().replace(" ", "_")] = row[1] if len(row) > 1 else None

        # Required summary keys
        for required in ["total_recipes_analyzed", "recipes_meeting_guidelines",
                         "average_calories", "highest_calorie_recipe", "lowest_calorie_recipe"]:
            record(f"Summary has '{required}'",
                   required in metrics,
                   f"Keys: {list(metrics.keys())}")

        tp_key = next((k for k in metrics if "total" in k and "recip" in k and "analyz" in k), None)
        if tp_key:
            try:
                record("Total_Recipes_Analyzed >= 8",
                       float(metrics[tp_key]) >= 8, f"Got {metrics[tp_key]}")
            except (TypeError, ValueError):
                record("Total_Recipes_Analyzed is numeric", False)

        # Recipes_Meeting_Guidelines must equal count of "Yes" rows in Recipe Analysis
        rmg_key = next((k for k in metrics if "recip" in k and "meet" in k), None)
        if rmg_key and meets_col is not None:
            try:
                rmg_val = int(float(metrics[rmg_key]))
                yes_count = sum(1 for r in data if meets_col < len(r) and str(r[meets_col]).strip().lower() == "yes")
                record(f"Recipes_Meeting_Guidelines == count of Yes rows ({yes_count})",
                       rmg_val == yes_count, f"reported {rmg_val}, actual yes={yes_count}")
            except (TypeError, ValueError):
                pass

        # Highest_Calorie_Recipe must be the row with the largest calories
        hc_key = next((k for k in metrics if "highest" in k and "calor" in k), None)
        if hc_key and cal_col is not None:
            try:
                expected = max(
                    (r for r in data if cal_col < len(r) and r[cal_col] is not None),
                    key=lambda r: float(r[cal_col])
                )
                expected_name = str(expected[0]).strip().lower()
                reported = str(metrics[hc_key]).strip().lower()
                record(f"Highest_Calorie_Recipe == '{expected[0]}'",
                       reported == expected_name, f"reported '{metrics[hc_key]}'")
            except (TypeError, ValueError, StopIteration):
                pass

        lc_key = next((k for k in metrics if "lowest" in k and "calor" in k), None)
        if lc_key and cal_col is not None:
            try:
                expected = min(
                    (r for r in data if cal_col < len(r) and r[cal_col] is not None),
                    key=lambda r: float(r[cal_col])
                )
                expected_name = str(expected[0]).strip().lower()
                reported = str(metrics[lc_key]).strip().lower()
                record(f"Lowest_Calorie_Recipe == '{expected[0]}'",
                       reported == expected_name, f"reported '{metrics[lc_key]}'")
            except (TypeError, ValueError, StopIteration):
                pass

    # --- Groundtruth XLSX comparison ---
    gt_path = os.path.join(groundtruth_workspace, "Nutrition_Benchmark.xlsx")
    if os.path.isfile(gt_path):
        gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
        for gt_sname in gt_wb.sheetnames:
            gt_ws = gt_wb[gt_sname]
            a_ws = None
            for asn in wb.sheetnames:
                if asn.strip().lower() == gt_sname.strip().lower():
                    a_ws = wb[asn]; break
            if a_ws is None:
                record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                continue
            gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
            a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
            record(f"GT '{gt_sname}' row count", len(a_rows) == len(gt_rows),
                   f"Expected {len(gt_rows)}, got {len(a_rows)}")
            # Recipe names are dynamic (from HowToCook MCP), so we only validate types/ranges.
            # The Summary sheet metric *values* are validated above against derived totals,
            # which provides cross-sheet consistency.
        gt_wb.close()

    return True


def check_word(workspace):
    print("\n=== Checking Word Document ===")
    path = os.path.join(workspace, "Nutrition_Report.docx")
    if not os.path.isfile(path):
        record("Word document exists", False, f"Not found: {path}")
        return False
    record("Word document exists", True)

    try:
        from docx import Document
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs).lower()

        # At least 4 paragraphs of substantive content per task.md
        substantive_paras = [p for p in doc.paragraphs if p.text and len(p.text.strip()) > 50]
        record("Report has at least 4 substantive paragraphs",
               len(substantive_paras) >= 4, f"got {len(substantive_paras)} substantive paras")
        record("Report has substantial content (>500 chars)", len(text) > 500, f"Only {len(text)} chars")
        record("Mentions nutrition/calories", "calori" in text or "nutrition" in text)
        record("Mentions benchmark/guideline", "benchmark" in text or "guideline" in text)
        record("Mentions meal plan", "meal" in text and "plan" in text)
        # Specific content per task.md
        record("Mentions sodium (specifically called out in task)",
               "sodium" in text)
        return True
    except Exception as e:
        record("Word readable", False, str(e))
        return False


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace, args.groundtruth_workspace)
    check_word(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
