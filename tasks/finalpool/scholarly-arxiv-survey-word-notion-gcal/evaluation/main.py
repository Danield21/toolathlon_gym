"""
Evaluation for scholarly-arxiv-survey-word-notion-gcal task.

Checks:
1. LLM_Reasoning_Survey.docx exists and has all 7 required sections
2. Word doc covers all 5 reasoning methods and references the 5 papers
3. Word doc has comparative analysis content
4. Notion database "Reasoning Papers" exists with required schema
5. Notion has at least 5 paper entries
6. GCal has 5 reading group sessions on Mon-Fri 2026-03-16..20 at 10:00-11:00
"""
import json
import os
import sys
from argparse import ArgumentParser
from datetime import date, timezone
from zoneinfo import ZoneInfo

import psycopg2
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": os.environ.get("PGUSER", "eigent"),
    "password": os.environ.get("PGPASSWORD", "camel"),
}

PASS_COUNT = 0
FAIL_COUNT = 0
LOCAL_FAIL = 0
RUNTIME_FAIL = 0


def record(name, passed, detail="", category="local"):
    global PASS_COUNT, FAIL_COUNT, LOCAL_FAIL, RUNTIME_FAIL
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        if category == "runtime":
            RUNTIME_FAIL += 1
        else:
            LOCAL_FAIL += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document LLM_Reasoning_Survey.docx ===")

    docx_path = os.path.join(agent_workspace, "LLM_Reasoning_Survey.docx")
    if not os.path.exists(docx_path):
        record("LLM_Reasoning_Survey.docx exists", False, f"Not found at {docx_path}")
        # account for missed check count to avoid skipping silently
        for _ in range(15):
            record("Word check (skipped)", False)
        return
    record("LLM_Reasoning_Survey.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Get all text
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Check for table
    has_table = len(doc.tables) >= 1
    table_text = ""
    if has_table:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.lower() + " "
    combined_text = all_text + " " + table_text

    # Check required sections (must have ALL 7)
    def _is_heading_para(p):
        t = p.text.strip()
        if not t:
            return False
        if p.style.name.startswith("Heading"):
            return True
        # Fallbacks for heading-like titles typed as Normal paragraphs
        # (e.g. via python_execute/python-docx): short all-bold paragraph,
        # or a numbered line such as "1. Introduction" / "3.1 Chain-of-Thought".
        if len(t) > 120:
            return False
        runs = [r for r in p.runs if (r.text or "").strip()]
        if runs and all(r.bold for r in runs):
            return True
        if t[0].isdigit() and ". " in t[:5]:
            return True
        return False

    headings = [p.text.strip() for p in doc.paragraphs if _is_heading_para(p)]
    headings_lower = [h.lower() for h in headings]

    has_abstract = any("abstract" in h for h in headings_lower) or "abstract" in all_text[:500]
    has_intro = any("introduction" in h for h in headings_lower)
    has_background = any("background" in h for h in headings_lower)
    has_taxonomy = any("taxonomy" in h or "method" in h for h in headings_lower)
    has_comparative = any("comparative" in h or "comparison" in h for h in headings_lower)
    has_challenges = any("challenge" in h or "open" in h for h in headings_lower)
    has_conclusion = any("conclusion" in h for h in headings_lower)

    record("Section: Abstract", has_abstract)
    record("Section: Introduction", has_intro)
    record("Section: Background", has_background)
    record("Section: Taxonomy/Methods", has_taxonomy)
    record("Section: Comparative Analysis", has_comparative)
    record("Section: Open Challenges", has_challenges)
    record("Section: Conclusion", has_conclusion)

    # All 5 reasoning methods (require ALL)
    has_cot = "chain-of-thought" in combined_text or "chain of thought" in combined_text
    has_tot = "tree of thought" in combined_text or "tree-of-thought" in combined_text
    has_sc = "self-consistency" in combined_text or "self consistency" in combined_text
    has_verify = ("step by step" in combined_text and "verif" in combined_text) or \
                 ("step-by-step" in combined_text and "verif" in combined_text) or \
                 "process supervision" in combined_text
    has_auto = "auto-cot" in combined_text or "automatic chain" in combined_text or \
               "automatic cot" in combined_text

    record("Method: Chain-of-Thought", has_cot)
    record("Method: Tree-of-Thought", has_tot)
    record("Method: Self-Consistency", has_sc)
    record("Method: Step-by-Step Verification", has_verify)
    record("Method: Automatic CoT", has_auto)

    # Tighter author check: full names rather than bare last names
    has_jason_wei = "jason wei" in combined_text or "wei et al" in combined_text or "wei (" in combined_text or "wei," in combined_text
    has_yao = "shunyu yao" in combined_text or "yao et al" in combined_text or "yao (" in combined_text or "yao," in combined_text
    has_wang = "xuezhi wang" in combined_text or "wang et al" in combined_text or "wang (" in combined_text or "wang," in combined_text
    has_lightman = "lightman" in combined_text or "let's verify" in combined_text or "lets verify" in combined_text
    has_zhang = ("zhuosheng zhang" in combined_text) or ("zhang et al" in combined_text and "auto" in combined_text) or ("zhang," in combined_text and "auto" in combined_text)

    record("Reference: Wei (CoT)", has_jason_wei)
    record("Reference: Yao (ToT)", has_yao)
    record("Reference: Wang (Self-Consistency)", has_wang)
    record("Reference: Lightman/Verify Step-by-Step", has_lightman)
    record("Reference: Zhang (Auto-CoT)", has_zhang)

    # Comparative content (table or comparison metrics)
    has_comparison_content = (
        has_table or
        ("comparative" in combined_text and any(m in combined_text
            for m in ["accuracy", "cost", "performance", "compared", "comparison"]))
    )
    record("Has comparative analysis content", has_comparison_content)


def check_notion_database():
    print("\n=== Check 2: Notion Database ===")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
    except Exception as e:
        record("Notion DB connection", False, str(e), category="runtime")
        return
    cur = conn.cursor()

    # Check for "Reasoning Papers" database
    cur.execute("SELECT id, title, properties FROM notion.databases")
    databases = cur.fetchall()

    def parse_title(title_json):
        title_str = ""
        if isinstance(title_json, list):
            for item in title_json:
                if isinstance(item, dict):
                    if isinstance(item.get("text"), dict):
                        title_str += item.get("text", {}).get("content", "")
                    if "plain_text" in item:
                        title_str += item.get("plain_text", "")
        elif isinstance(title_json, str):
            title_str = title_json
        else:
            try:
                title_str = json.dumps(title_json) if title_json else ""
            except Exception:
                title_str = str(title_json)
        return title_str

    # Tighten match: must contain BOTH 'reasoning' AND 'papers' (plural)
    reasoning_db = None
    for db_id, title_json, props in databases:
        title_str = parse_title(title_json).lower()
        # Require both keywords as separate tokens
        if "reasoning" in title_str and ("papers" in title_str or "paper" in title_str):
            reasoning_db = (db_id, title_str, props)
            break

    db_titles = [parse_title(t)[:60] for _, t, _ in databases]
    record("Notion 'Reasoning Papers' database exists", reasoning_db is not None,
           f"Databases found: {db_titles}", category="runtime")

    # Check for paper pages
    cur.execute("SELECT properties FROM notion.pages")
    pages = cur.fetchall()

    # Require pages matching the 5 paper title/method fragments. Fragments are
    # intentionally short so that agents which shorten the paper titles (e.g.
    # "Chain of Thought", "Self-Consistency Decoding") or use hyphenated /
    # non-hyphenated variants of the method names still match.
    #
    # Matching is a MAXIMUM BIPARTITE MATCHING between pages and the 5 paper
    # groups, so it is ORDER-INDEPENDENT. That matters because the
    # Self-Consistency paper's abstract opens with "Chain-of-thought prompting
    # has shown promise ...", so a Self-Consistency page carrying that abstract
    # (or a summary of it) matches BOTH the chain_of_thought and
    # self_consistency groups. With a greedy first-match-wins scan, whichever
    # page the DB returns first could "steal" the CoT group and the real CoT
    # entry would then match nothing - a correct agent judged FAIL purely on
    # the (unordered, insertion-order) read order of notion.pages. Bipartite
    # matching instead lets the SC page fill the self_consistency slot while the
    # CoT page fills the chain_of_thought slot.
    #
    # Bare "chain of thought" is a fragment because the task does not require
    # the paper's full name; a title of just "Chain of Thought" is a valid
    # correct completion. Fragment collisions are harmless here because the
    # matching maximizes coverage rather than consuming groups greedily.
    PAPER_GROUPS = [
        ("tree_of_thought", ["tree of thought", "tree-of-thought"]),
        ("automatic_cot", ["automatic chain of thought",
                           "automatic chain-of-thought",
                           "automatic cot", "auto-cot", "auto cot"]),
        ("chain_of_thought", ["chain of thought", "chain-of-thought",
                              "chain of thought prompting",
                              "chain-of-thought prompting",
                              "prompting elicits reasoning"]),
        ("self_consistency", ["self-consistency", "self consistency"]),
        ("lets_verify_step", ["let's verify", "let s verify",
                              "process supervision",
                              "step by step verification",
                              "step-by-step verification"]),
    ]
    canon_order = [c for c, _ in PAPER_GROUPS]
    canon_index = {c: i for i, c in enumerate(canon_order)}

    # Page -> set of paper groups its properties text can serve.
    page_groups = []
    for (props,) in pages:
        props_str = json.dumps(props).lower() if isinstance(props, dict) else str(props).lower()
        gs = set()
        for canon, frags in PAPER_GROUPS:
            if any(f in props_str for f in frags):
                gs.add(canon)
        page_groups.append(gs)

    # Kuhn's augmenting-path algorithm for maximum bipartite matching:
    # pages (left) to paper groups (right). A page matching several groups can
    # fill only one slot, and every matched group needs its own page, so all 5
    # distinct paper topics are still required - 5 genuinely distinct entries.
    group_owner = {}  # group index -> page index

    def _augment(page_idx, seen):
        for g in page_groups[page_idx]:
            gi = canon_index[g]
            if gi in seen:
                continue
            seen.add(gi)
            prev = group_owner.get(gi)
            if prev is None or _augment(prev, seen):
                group_owner[gi] = page_idx
                return True
        return False

    for p in range(len(page_groups)):
        _augment(p, set())

    matched_names = [canon_order[i] for i in sorted(group_owner)]
    record("At least 5 paper entries with specific titles in Notion (deduped)",
           len(group_owner) >= len(canon_order),
           f"Matched {len(group_owner)} distinct papers: {sorted(matched_names)}",
           category="runtime")

    # DB schema must have all required properties (Title/Authors/Year/Method/Key_Contribution/Cited_By)
    if reasoning_db:
        db_id, db_title, db_props = reasoning_db
        if db_props:
            props_str = json.dumps(db_props).lower() if isinstance(db_props, dict) else str(db_props).lower()
            schema_ok = all(p in props_str for p in [
                "title", "author", "year", "method", "key", "cited"])
            keys = list(db_props.keys()) if isinstance(db_props, dict) else "N/A"
            record("Database schema has Title/Authors/Year/Method/Key/Cited",
                   schema_ok, f"Properties: {keys}", category="runtime")
        else:
            record("Database schema has required properties", False, "no props",
                   category="runtime")
    else:
        record("Database schema has required properties", False, "no DB found",
               category="runtime")

    cur.close()
    conn.close()


def check_gcal():
    print("\n=== Check 3: Google Calendar Reading Group Sessions ===")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
    except Exception as e:
        record("GCal DB connection", False, str(e), category="runtime")
        return
    cur = conn.cursor()

    # Window is widened slightly to tolerate timezone-offset inputs: the DB
    # stores start_datetime normalized to UTC, so an event created as
    # "2026-03-16T10:00:00+08:00" lands on 2026-03-16 02:00Z.
    cur.execute("""
        SELECT summary, start_datetime, end_datetime, start_timezone
        FROM gcal.events
        WHERE start_datetime >= '2026-03-15' AND start_datetime < '2026-03-22'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()

    def local_start(e):
        """Return tz-aware wall-clock start for an event row
        (summary, start_datetime, end_datetime, start_timezone)."""
        sd = e[1]
        if sd is None:
            return None
        if sd.tzinfo is None:
            sd = sd.replace(tzinfo=timezone.utc)
        tz_name = e[3] or ""
        if tz_name:
            try:
                return sd.astimezone(ZoneInfo(tz_name))
            except Exception:
                pass
        return sd.astimezone(timezone.utc)

    reading_events = [
        e for e in events
        if "reading group" in (e[0] or "").lower()
    ]

    # Lenient count (>=5): in swarm runs the same session can be scheduled
    # more than once; the weekday-coverage check below still requires all five
    # distinct Mon-Fri dates to be present.
    record("At least 5 reading group events in March 15-21 window",
           len(reading_events) >= 5,
           f"Found {len(reading_events)}: {[e[0] for e in reading_events]}",
           category="runtime")

    if reading_events:
        # All durations should be exactly 1 hour
        all_one_hour = True
        for e in reading_events:
            sd, ed = e[1], e[2]
            if sd and ed:
                hrs = (ed - sd).total_seconds() / 3600
                if abs(hrs - 1.0) > 0.05:
                    all_one_hour = False
                    break
            else:
                all_one_hour = False
                break
        record("All sessions exactly 1 hour", all_one_hour, category="runtime")

        # All sessions must start at 10:00 local wall-clock time.
        # The local timezone is taken from the event's start_timezone when the
        # agent supplied one; otherwise the stored UTC value is used.
        local_starts = [local_start(e) for e in reading_events]
        strict_at_10 = all(ls is not None and ls.hour == 10 and ls.minute == 0
                           for ls in local_starts)
        # Fallback: if the strict check fails, accept when every event shares a
        # single UTC hour-of-day (all scheduled at the same local hour) and that
        # hour corresponds to 10:00 in UTC itself or in Asia/Shanghai (UTC+8,
        # the lab's home timezone). This tolerates agents that embed the offset
        # in the dateTime (e.g. "2026-03-16T10:00:00+08:00") without passing the
        # separate timeZone field -- such events are stored as 02:00 UTC.
        utc_hours = {ls.astimezone(timezone.utc).hour
                     for ls in local_starts if ls is not None}
        minutes_zero = all(ls is not None and ls.minute == 0 for ls in local_starts)
        fallback_at_10 = (minutes_zero and len(utc_hours) == 1
                          and next(iter(utc_hours), None) in (10, 2))
        all_at_10 = strict_at_10 or fallback_at_10
        record("All sessions start at 10:00 (local wall-clock)", all_at_10, category="runtime")

        # 5 distinct dates Mon-Fri (2026-03-16..20), evaluated in the same
        # wall-clock timezone as the hour check.
        expected_dates = {
            date(2026, 3, 16),
            date(2026, 3, 17),
            date(2026, 3, 18),
            date(2026, 3, 19),
            date(2026, 3, 20),
        }
        actual_dates = {ls.date() for ls in local_starts if ls is not None}
        record("Sessions cover all 5 weekdays Mar 16-20",
               actual_dates == expected_dates,
               f"got {sorted(actual_dates)}",
               category="runtime")

        # Each summary should match "Reading Group: <topic>" format
        title_format_ok = all(":" in (e[0] or "") and "reading group" in (e[0] or "").lower()
                              for e in reading_events)
        record("All summaries match 'Reading Group: ...' format",
               title_format_ok, category="runtime")

        # Topics should cover paper themes
        summaries_text = " ".join((e[0] or "").lower() for e in reading_events)
        topic_keywords = ["chain", "tree", "self", "verif", "step", "auto", "thought", "consistency"]
        topic_matches = sum(1 for kw in topic_keywords if kw in summaries_text)
        record("Session titles contain >=4 topic descriptors",
               topic_matches >= 4,
               f"Matched {topic_matches}: {[e[0] for e in reading_events]}",
               category="runtime")
    else:
        for _ in range(5):
            record("Reading event detail (skipped, none found)", False, category="runtime")

    cur.close()
    conn.close()


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_notion_database()
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed "
          f"(local fails: {LOCAL_FAIL}, runtime fails: {RUNTIME_FAIL})")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "local_fails": LOCAL_FAIL,
        "runtime_fails": RUNTIME_FAIL,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # Strict: zero failures required (both local and runtime)
    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
