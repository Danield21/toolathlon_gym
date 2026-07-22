"""Evaluation script for canvas-risk-intervention-excel-word-gcal-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Student_Risk_Assessment.xlsx
    excel_path = os.path.join(agent_workspace, "Student_Risk_Assessment.xlsx")
    check("Student_Risk_Assessment.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Student_Risk_Assessment.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Row counts and values: validate by structural rules,
                    # not by exact GT match. The GT was authored with
                    # illustrative/fabricated course names that don't match
                    # the canvas read-only data, so any cell value comparison
                    # would create false negatives for honest agents.
                    gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True)
                               if r and any(c is not None for c in r)]
                    data_rows = [r for r in ws.iter_rows(min_row=2, values_only=True)
                                 if r and any(c is not None for c in r)]
                    check(f"{sheet_name} has >= {len(gt_rows)} data rows",
                          len(data_rows) >= len(gt_rows),
                          f"got {len(data_rows)}, expected at least {len(gt_rows)}")

                    sn_lower = sheet_name.strip().lower()
                    header_map = {h: i for i, h in enumerate(headers)}

                    if "risk_overview" in sn_lower or "risk overview" in sn_lower:
                        # Validate arithmetic and structure.
                        ts_idx = header_map.get("total_students", -1)
                        cc_idx = header_map.get("critical_count", -1)
                        ar_idx = header_map.get("at_risk_count", -1)
                        ot_idx = header_map.get("on_track_count", -1)
                        rr_idx = header_map.get("risk_rate_pct", -1)

                        arith_ok = True
                        sum_count_ok = True
                        rates = []
                        for r in data_rows:
                            try:
                                ts = float(r[ts_idx]) if ts_idx >= 0 and ts_idx < len(r) else None
                                cc = float(r[cc_idx]) if cc_idx >= 0 and cc_idx < len(r) else None
                                ar = float(r[ar_idx]) if ar_idx >= 0 and ar_idx < len(r) else None
                                ot = float(r[ot_idx]) if ot_idx >= 0 and ot_idx < len(r) else None
                                rr = float(r[rr_idx]) if rr_idx >= 0 and rr_idx < len(r) else None
                                if ts is not None and cc is not None and ar is not None and ot is not None:
                                    # Total = sum of three categories
                                    if abs(ts - (cc + ar + ot)) > 0.5:
                                        sum_count_ok = False
                                if ts is not None and cc is not None and ar is not None and rr is not None:
                                    expected_rate = (cc + ar) / ts * 100 if ts > 0 else 0
                                    if abs(rr - expected_rate) > 0.2:
                                        arith_ok = False
                                if rr is not None:
                                    rates.append(rr)
                            except (TypeError, ValueError):
                                pass
                        check(f"{sheet_name} Total_Students = Critical + At_Risk + On_Track per row",
                              sum_count_ok, "")
                        check(f"{sheet_name} Risk_Rate_Pct = (Critical + At_Risk) / Total * 100",
                              arith_ok, "")
                        # Sort check: rates should be in descending order
                        sort_ok = all(rates[i] >= rates[i+1] - 0.01 for i in range(len(rates)-1)) if len(rates) >= 2 else True
                        check(f"{sheet_name} sorted by Risk_Rate_Pct descending",
                              sort_ok, f"rates: {rates[:6]}")

                    elif "summary" in sn_lower:
                        # Summary has Metric/Value rows. Verify required metrics
                        # exist by name (case-insensitive).
                        metric_idx = header_map.get("metric", 0)
                        val_idx = header_map.get("value", 1)
                        metrics = {}
                        for r in data_rows:
                            if metric_idx < len(r) and r[metric_idx] is not None:
                                metrics[str(r[metric_idx]).strip().lower()] = (
                                    r[val_idx] if val_idx < len(r) else None
                                )
                        required = [
                            "total_students_assessed", "critical_students",
                            "at_risk_students", "on_track_students",
                            "overall_risk_rate_pct", "highest_risk_course",
                        ]
                        for m in required:
                            present = any(m in k or k in m for k in metrics.keys())
                            check(f"{sheet_name} contains metric '{m}'",
                                  present, f"metrics found: {list(metrics.keys())[:8]}")
                        # Arithmetic: total = critical + at_risk + on_track
                        try:
                            total_v = next((float(v) for k, v in metrics.items() if "total" in k), None)
                            crit_v = next((float(v) for k, v in metrics.items() if "critical" in k and "course" not in k), None)
                            risk_v = next((float(v) for k, v in metrics.items() if "at_risk" in k or "at risk" in k), None)
                            track_v = next((float(v) for k, v in metrics.items() if "on_track" in k or "on track" in k), None)
                            if all(x is not None for x in [total_v, crit_v, risk_v, track_v]):
                                check(f"{sheet_name} Total_Students_Assessed = Critical + At_Risk + On_Track",
                                      abs(total_v - (crit_v + risk_v + track_v)) <= 0.5, "")
                        except Exception:
                            pass

                    elif "intervention_plan" in sn_lower or "intervention plan" in sn_lower:
                        # Validate enum values for Risk_Level and that all required cols are non-empty.
                        rl_idx = header_map.get("risk_level", -1)
                        rp_idx = header_map.get("responsible_party", -1)
                        ra_idx = header_map.get("recommended_action", -1)
                        dl_idx = header_map.get("deadline", -1)
                        valid_levels = {"high", "medium", "low", "critical"}
                        levels_ok = True
                        non_empty_ok = True
                        for r in data_rows:
                            if rl_idx >= 0 and rl_idx < len(r) and r[rl_idx] is not None:
                                v = str(r[rl_idx]).strip().lower()
                                if v and v not in valid_levels:
                                    levels_ok = False
                            for idx in [rp_idx, ra_idx, dl_idx]:
                                if idx >= 0:
                                    if idx >= len(r) or r[idx] is None or str(r[idx]).strip() == "":
                                        non_empty_ok = False
                        check(f"{sheet_name} Risk_Level values in {{High,Medium,Low,Critical}}",
                              levels_ok, "")
                        check(f"{sheet_name} all rows non-empty in Recommended_Action/Responsible_Party/Deadline",
                              non_empty_ok, "")

                    elif "critical_students" in sn_lower or "critical students" in sn_lower:
                        # Validate Avg_Score < 50 for all rows (Critical = score < 50 per task)
                        avg_idx = header_map.get("avg_score", -1)
                        sid_idx = header_map.get("student_id", -1)
                        all_critical = True
                        non_empty = True
                        for r in data_rows:
                            if avg_idx >= 0 and avg_idx < len(r) and r[avg_idx] is not None:
                                try:
                                    sc = float(r[avg_idx])
                                    if sc >= 50:
                                        all_critical = False
                                except (TypeError, ValueError):
                                    pass
                            if sid_idx >= 0 and (sid_idx >= len(r) or r[sid_idx] is None):
                                non_empty = False
                        check(f"{sheet_name} all Avg_Score values < 50 (Critical threshold per task)",
                              all_critical, "")
                        check(f"{sheet_name} Student_ID populated for all rows",
                              non_empty, "")

    # Check Intervention_Report.docx
    docx_path = os.path.join(agent_workspace, "Intervention_Report.docx")
    check("Intervention_Report.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Intervention_Report.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        gt_doc_path = os.path.join(groundtruth_workspace, "Intervention_Report.docx")
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if p.style.name.startswith("Heading")]
            for gh in gt_headings:
                if gh:
                    # Tighten: require GT heading to appear as a substring of
                    # an agent heading (one-directional). Bidirectional
                    # `gh in h or h in gh` allowed micro-headings like 'risk'
                    # to pass when GT was 'risk assessment methodology'.
                    found = any(gh in h for h in headings)
                    check(f"Intervention_Report.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("Intervention_Report.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Check Python script exists (terminal usage). Task specifies the
    # script name 'risk_identifier.py' and that it reads
    # 'student_performance.json' and writes 'risk_assessment.json'.
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    has_risk_id = "risk_identifier.py" in py_files
    check("risk_identifier.py exists", has_risk_id, f"py files: {py_files}")
    workspace_files = set(os.listdir(agent_workspace))
    check("student_performance.json exists",
          "student_performance.json" in workspace_files,
          f"workspace: {sorted(workspace_files)[:25]}")
    check("risk_assessment.json exists",
          "risk_assessment.json" in workspace_files,
          f"workspace: {sorted(workspace_files)[:25]}")

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        # Email: AND on exact subject and to_addr
        cur.execute("""
            SELECT subject, to_addr, body_text FROM email.messages
            WHERE LOWER(subject) = LOWER(%s)
              AND to_addr::text ILIKE %s
        """, ("Urgent: Student Risk Assessment Results",
              "%academic-affairs@university.edu%"))
        email_row = cur.fetchone()
        check("Email with exact subject 'Urgent: Student Risk Assessment Results' to academic-affairs@university.edu sent",
              email_row is not None,
              "no matching email found")
        if email_row:
            check("Email recipient is academic-affairs@university.edu",
                  "academic-affairs@university.edu" in str(email_row[1] or "").lower(),
                  f"to_addr: {email_row[1]}")
            body = str(email_row[2] or "").lower()
            # Body should mention high-risk content. Task: "highlighting
            # courses with risk rates above 40%". Accept: explicit '40' OR
            # mention of risk/critical/intervention with substantive length.
            has_risk_mention = (
                "risk" in body or "critical" in body or "intervention" in body
            )
            has_threshold_or_courses = (
                "40" in body or "course" in body or "%" in body
            )
            check(
                "Email body mentions risk content (risk/critical/intervention) and threshold/courses",
                has_risk_mention and has_threshold_or_courses and len(body) >= 80,
                f"body[:200]: {body[:200]}",
            )

        # Calendar: AND on summary, date, time, duration
        cur.execute("""
            SELECT summary, start_datetime, end_datetime, description
            FROM gcal.events
            WHERE LOWER(summary) = LOWER(%s)
        """, ("Academic Intervention Planning Meeting",))
        events = cur.fetchall()
        check("Calendar event 'Academic Intervention Planning Meeting' exists",
              len(events) >= 1,
              "no matching event found")
        target_event = None
        for e in events:
            sd = e[1]
            try:
                if sd is not None and sd.year == 2026 and sd.month == 3 and sd.day == 13 \
                   and sd.hour == 15 and sd.minute == 0:
                    target_event = e
                    break
            except AttributeError:
                pass
        check("Event scheduled on 2026-03-13 at 15:00 UTC",
              target_event is not None,
              f"event start_datetimes: {[e[1] for e in events]}")
        if target_event and target_event[2] and target_event[1]:
            try:
                delta = (target_event[2] - target_event[1]).total_seconds() / 60.0
                check("Event duration is 90 minutes",
                      abs(delta - 90) < 5,
                      f"got {delta} min")
            except Exception:
                pass

        # Reverse verification: noise emails should not be in Sent folder
        cur.execute("SELECT COUNT(*) FROM email.messages WHERE folder_id = (SELECT id FROM email.folders WHERE name = 'Sent' LIMIT 1) AND subject ILIKE '%newsletter%'")
        noise_sent = cur.fetchone()[0]
        check("No noise emails in Sent folder", noise_sent == 0, f"found {noise_sent} noise emails in Sent")
        # Reverse verification: noise events should not match task keyword
        cur.execute("SELECT COUNT(*) FROM gcal.events WHERE summary ILIKE '%standup%' OR summary ILIKE '%lunch%'")
        noise_events = cur.fetchone()[0]
        check("Noise events exist (not deleted by agent)", noise_events >= 1, f"noise events: {noise_events}")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()