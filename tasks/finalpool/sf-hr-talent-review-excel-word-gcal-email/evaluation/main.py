"""Evaluation for sf-hr-talent-review-excel-word-gcal-email.

Checks:
1. Excel file: Talent_Review_Report.xlsx with 3 sheets and correct data
2. Word file: Talent_Review_Summary.docx with correct structure and content
3. GCal: 7 department review meetings scheduled ~14-20 days from launch
4. Email: summary email sent to hr_leadership@company.com
"""
import argparse
import json
import os
import sys
from datetime import datetime, timedelta, timezone

import openpyxl
import psycopg2

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)

from docx import Document

DB = dict(
    host=os.environ.get("PGHOST", "localhost"), port=5432,
    dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
    user="eigent", password="camel"
)

PASS_COUNT = 0
FAIL_COUNT = 0

DEPARTMENTS = ["Engineering", "Finance", "HR", "Operations", "R&D", "Sales", "Support"]

# Fallback values (only used if DB query fails)
DEPT_EXPECTED_FALLBACK = {
    "Engineering": {"headcount": 7096, "avg_perf": 3.21, "avg_salary": 58991.61,
                    "min_salary": 0, "max_salary": 0, "promo": 1578},
    "Finance":     {"headcount": 7148, "avg_perf": 3.21, "avg_salary": 57878.19,
                    "min_salary": 0, "max_salary": 0, "promo": 1588},
    "HR":          {"headcount": 7077, "avg_perf": 3.20, "avg_salary": 58920.45,
                    "min_salary": 0, "max_salary": 0, "promo": 1588},
    "Operations":  {"headcount": 7120, "avg_perf": 3.18, "avg_salary": 57808.74,
                    "min_salary": 0, "max_salary": 0, "promo": 1541},
    "R&D":         {"headcount": 7083, "avg_perf": 3.20, "avg_salary": 57905.93,
                    "min_salary": 0, "max_salary": 0, "promo": 1519},
    "Sales":       {"headcount": 7232, "avg_perf": 3.19, "avg_salary": 58864.79,
                    "min_salary": 0, "max_salary": 0, "promo": 1596},
    "Support":     {"headcount": 7244, "avg_perf": 3.20, "avg_salary": 58400.48,
                    "min_salary": 0, "max_salary": 0, "promo": 1540},
}

# From Compensation_Benchmarks.pdf
INDUSTRY_BENCHMARKS = {
    "Engineering": 62000, "Finance": 60000, "HR": 55000, "Operations": 54000,
    "R&D": 63000, "Sales": 61000, "Support": 53000,
}


def load_dept_expected():
    """Runtime DB query for department stats; fallback to constants if it fails."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute('''
            SELECT "DEPARTMENT",
                   COUNT(*) AS headcount,
                   ROUND(AVG("PERFORMANCE_RATING")::numeric, 2) AS avg_perf,
                   ROUND(AVG("SALARY")::numeric, 2) AS avg_salary,
                   MIN("SALARY") AS min_salary,
                   MAX("SALARY") AS max_salary
            FROM sf_data."HR_ANALYTICS__PUBLIC__EMPLOYEES"
            GROUP BY "DEPARTMENT"
        ''')
        dept_stats = {r[0]: {"headcount": int(r[1]), "avg_perf": float(r[2]),
                             "avg_salary": float(r[3]), "min_salary": float(r[4]),
                             "max_salary": float(r[5])} for r in cur.fetchall()}
        # Promotion candidates per PDF: rating>=4 AND years_experience>=3
        cur.execute('''
            SELECT "DEPARTMENT", COUNT(*) AS promo
            FROM sf_data."HR_ANALYTICS__PUBLIC__EMPLOYEES"
            WHERE "PERFORMANCE_RATING" >= 4 AND "YEARS_EXPERIENCE" >= 3
            GROUP BY "DEPARTMENT"
        ''')
        for r in cur.fetchall():
            if r[0] in dept_stats:
                dept_stats[r[0]]["promo"] = int(r[1])
        cur.close(); conn.close()
        # Ensure all 7 depts covered
        for d in DEPARTMENTS:
            if d not in dept_stats:
                dept_stats[d] = DEPT_EXPECTED_FALLBACK.get(d, {})
            dept_stats[d].setdefault("promo", 0)
        return dept_stats
    except Exception as e:
        print(f"[eval] load_dept_expected fallback: {e}")
        return dict(DEPT_EXPECTED_FALLBACK)


DEPT_EXPECTED = load_dept_expected()


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def safe_float(val, default=None):
    try:
        return float(val)
    except (TypeError, ValueError):
        return default


def parse_recipients(to_addr):
    if to_addr is None:
        return []
    if isinstance(to_addr, list):
        return [str(r).strip().lower() for r in to_addr]
    to_str = str(to_addr).strip()
    try:
        parsed = json.loads(to_str)
        if isinstance(parsed, list):
            return [str(r).strip().lower() for r in parsed]
        return [to_str.lower()]
    except (json.JSONDecodeError, TypeError):
        return [to_str.lower()]


def check_excel(agent_workspace):
    print("\n=== Checking Excel: Talent_Review_Report.xlsx ===")
    xlsx_path = os.path.join(agent_workspace, "Talent_Review_Report.xlsx")
    check("Excel file exists", os.path.isfile(xlsx_path), f"Not found: {xlsx_path}")
    if not os.path.isfile(xlsx_path):
        return

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    sheet_names = [s.lower() for s in wb.sheetnames]

    # Sheet 1: Department_Summary
    has_dept_summary = any("department" in s and "summary" in s for s in sheet_names)
    check("Sheet 'Department_Summary' exists", has_dept_summary, f"Sheets: {wb.sheetnames}")

    if has_dept_summary:
        ws = None
        for name in wb.sheetnames:
            if "department" in name.lower() and "summary" in name.lower():
                ws = wb[name]
                break

        headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        rows = [r for r in rows if r[0] is not None]

        check("Department_Summary has >= 7 data rows", len(rows) >= 7, f"Found {len(rows)} rows")

        # Find column indices
        dept_col = next((i for i, h in enumerate(headers) if "department" in h), 0)
        hc_col = next((i for i, h in enumerate(headers) if "headcount" in h), 1)
        perf_col = next((i for i, h in enumerate(headers) if "perf" in h), 2)
        sal_col = next((i for i, h in enumerate(headers) if "avg" in h and "sal" in h), 3)
        promo_col = next((i for i, h in enumerate(headers) if "promo" in h), 5)

        depts_found = set()
        for row in rows:
            dept_name = str(row[dept_col]).strip() if row[dept_col] else ""
            matched_dept = None
            for d in DEPARTMENTS:
                if d.lower() == dept_name.lower():
                    matched_dept = d
                    break
            if not matched_dept:
                continue
            depts_found.add(matched_dept)
            exp = DEPT_EXPECTED[matched_dept]

            hc = safe_float(row[hc_col])
            if hc is not None:
                check(f"{matched_dept} Headcount", abs(hc - exp['headcount']) <= 5,
                      f"Expected ~{exp['headcount']}, got {hc}")

            perf = safe_float(row[perf_col])
            if perf is not None:
                check(f"{matched_dept} Avg_Performance", abs(perf - exp['avg_perf']) <= 0.05,
                      f"Expected ~{exp['avg_perf']}, got {perf}")

            promo = safe_float(row[promo_col])
            if promo is not None:
                check(f"{matched_dept} Promo_Candidates", abs(promo - exp['promo']) <= 10,
                      f"Expected ~{exp['promo']}, got {promo}")

        check("All 7 departments in Department_Summary", len(depts_found) == 7,
              f"Found {len(depts_found)}: {depts_found}")

    # Sheet 2: Promotion_Candidates
    has_promo = any("promotion" in s or "promo" in s for s in sheet_names)
    check("Sheet 'Promotion_Candidates' exists", has_promo, f"Sheets: {wb.sheetnames}")
    if has_promo:
        ws2 = None
        for name in wb.sheetnames:
            if "promotion" in name.lower() or "promo" in name.lower():
                ws2 = wb[name]
                break
        rows2 = list(ws2.iter_rows(min_row=2, values_only=True))
        rows2 = [r for r in rows2 if r[0] is not None]
        check("Promotion_Candidates has 21 rows (3 per dept)", abs(len(rows2) - 21) <= 2,
              f"Found {len(rows2)} rows")

        # Validate candidates match actual top-3 per department using PDF criteria
        # (rating = 5 AND years_experience >= 3, ordered by salary DESC within dept).
        all_text = " ".join(str(c) for row in rows2 for c in row if c)
        try:
            conn_ck = psycopg2.connect(**DB)
            cur_ck = conn_ck.cursor()
            # PDF criteria: rating>=4 AND years>=3, ordered by rating DESC
            # then salary DESC as tiebreaker per PDF guidance
            cur_ck.execute('''
                SELECT "EMPLOYEE_NAME" FROM (
                  SELECT *, ROW_NUMBER() OVER (
                    PARTITION BY "DEPARTMENT"
                    ORDER BY "PERFORMANCE_RATING" DESC, "SALARY" DESC, "EMPLOYEE_ID" ASC
                  ) AS rn
                  FROM sf_data."HR_ANALYTICS__PUBLIC__EMPLOYEES"
                  WHERE "PERFORMANCE_RATING" >= 4 AND "YEARS_EXPERIENCE" >= 3
                ) t WHERE rn <= 3
            ''')
            top_names = [r[0] for r in cur_ck.fetchall()]
            cur_ck.close(); conn_ck.close()
            matched = sum(1 for n in top_names if n in all_text)
            check("Promotion_Candidates matches top-3 per dept (>=17 of 21)",
                  matched >= 17, f"Matched {matched}/{len(top_names)} top candidates")
        except Exception as e:
            # Fallback: at least verify rows exist
            check("Promotion_Candidates non-empty (fallback)", len(rows2) > 0,
                  f"DB query failed: {e}")

    # Sheet 3: Salary_Analysis
    has_salary = any("salary" in s for s in sheet_names)
    check("Sheet 'Salary_Analysis' exists", has_salary, f"Sheets: {wb.sheetnames}")
    if has_salary:
        ws3 = None
        for name in wb.sheetnames:
            if "salary" in name.lower():
                ws3 = wb[name]
                break
        headers3 = [str(c.value).strip().lower() if c.value else "" for c in ws3[1]]
        rows3 = list(ws3.iter_rows(min_row=2, values_only=True))
        rows3 = [r for r in rows3 if r[0] is not None]
        check("Salary_Analysis has >= 7 rows", len(rows3) >= 7, f"Found {len(rows3)} rows")

        # Check for benchmark columns
        has_benchmark = any("benchmark" in h for h in headers3)
        check("Salary_Analysis has benchmark column", has_benchmark, f"Headers: {headers3}")
        has_gap = any("gap" in h for h in headers3)
        check("Salary_Analysis has gap column", has_gap, f"Headers: {headers3}")

        # Per-row value checks for Avg_Salary and Benchmark_Gap
        dept_col3 = next((i for i, h in enumerate(headers3) if "department" in h), 0)
        avg_sal_col3 = next((i for i, h in enumerate(headers3) if "avg" in h and "sal" in h), 1)
        min_sal_col3 = next((i for i, h in enumerate(headers3) if "min" in h), None)
        max_sal_col3 = next((i for i, h in enumerate(headers3) if "max" in h), None)
        bench_col3 = next((i for i, h in enumerate(headers3) if "benchmark" in h and "gap" not in h), None)
        gap_col3 = next((i for i, h in enumerate(headers3) if "gap" in h), None)
        for row in rows3:
            dept_name = str(row[dept_col3]).strip() if row[dept_col3] else ""
            matched_dept = next((d for d in DEPARTMENTS if d.lower() == dept_name.lower()), None)
            if not matched_dept:
                continue
            exp = DEPT_EXPECTED[matched_dept]
            # Avg_Salary
            av = safe_float(row[avg_sal_col3])
            if av is not None:
                check(f"Salary_Analysis {matched_dept} Avg_Salary",
                      abs(av - exp["avg_salary"]) <= max(1.0, exp["avg_salary"] * 0.005),
                      f"Expected ~{exp['avg_salary']}, got {av}")
            # Min_Salary
            if min_sal_col3 is not None and exp.get("min_salary"):
                mn = safe_float(row[min_sal_col3])
                if mn is not None:
                    check(f"Salary_Analysis {matched_dept} Min_Salary",
                          abs(mn - exp["min_salary"]) <= 1.0,
                          f"Expected {exp['min_salary']}, got {mn}")
            # Max_Salary
            if max_sal_col3 is not None and exp.get("max_salary"):
                mx = safe_float(row[max_sal_col3])
                if mx is not None:
                    check(f"Salary_Analysis {matched_dept} Max_Salary",
                          abs(mx - exp["max_salary"]) <= 1.0,
                          f"Expected {exp['max_salary']}, got {mx}")
            # Benchmark_Gap = Avg_Salary - Industry_Benchmark
            if gap_col3 is not None and matched_dept in INDUSTRY_BENCHMARKS:
                expected_gap = round(
                    exp["avg_salary"] - INDUSTRY_BENCHMARKS[matched_dept], 2)
                gv = safe_float(row[gap_col3])
                if gv is not None:
                    check(f"Salary_Analysis {matched_dept} Benchmark_Gap",
                          abs(gv - expected_gap) <= 1.0,
                          f"Expected ~{expected_gap}, got {gv}")
            # Industry_Benchmark
            if bench_col3 is not None and matched_dept in INDUSTRY_BENCHMARKS:
                bv = safe_float(row[bench_col3])
                if bv is not None:
                    check(f"Salary_Analysis {matched_dept} Industry_Benchmark",
                          abs(bv - INDUSTRY_BENCHMARKS[matched_dept]) <= 1.0,
                          f"Expected {INDUSTRY_BENCHMARKS[matched_dept]}, got {bv}")

    wb.close()


def check_word(agent_workspace):
    print("\n=== Checking Word: Talent_Review_Summary.docx ===")
    docx_path = os.path.join(agent_workspace, "Talent_Review_Summary.docx")
    check("Word file exists", os.path.isfile(docx_path), f"Not found: {docx_path}")
    if not os.path.isfile(docx_path):
        return

    doc = Document(docx_path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()

    # Check headings (accept both "Heading X" and "Title" styles)
    headings = [p.text for p in doc.paragraphs
                if p.style.name.startswith("Heading") or p.style.name == "Title"]
    heading_text = " ".join(headings).lower()
    check("Word has 'Talent Review' heading", "talent review" in heading_text,
          f"Headings: {headings[:5]}")
    check("Word has 'Department' section", "department" in heading_text,
          f"Headings: {headings[:5]}")
    check("Word has 'Salary' section",
          "salary" in heading_text or "competitiveness" in heading_text,
          f"Headings: {headings[:5]}")

    # Check content: total headcount = 50000
    check("Word mentions total headcount ~50000",
          "50000" in full_text or "50,000" in full_text,
          "50000/50,000 not found in text")

    # Check content: mentions promotion candidates
    total_promo = sum(d["promo"] for d in DEPT_EXPECTED.values())
    check("Word mentions promotion candidates",
          "promo" in full_text_lower or str(total_promo) in full_text or "10950" in full_text or "10,950" in full_text,
          f"Expected mention of ~{total_promo} promo candidates")

    # Check content: mentions all departments
    dept_count = sum(1 for d in DEPARTMENTS if d.lower() in full_text_lower)
    check("Word mentions all 7 departments", dept_count == 7,
          f"Found {dept_count}/7 departments")

    # Check: has at least one table
    check("Word contains at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    # Check table has all 7 department rows
    if doc.tables:
        all_tables_text = ""
        for t in doc.tables:
            all_tables_text += " " + " ".join(cell.text for row in t.rows for cell in row.cells)
        all_tables_text_lower = all_tables_text.lower()
        found_depts = sum(1 for d in DEPARTMENTS if d.lower() in all_tables_text_lower)
        check("Word tables include all 7 departments", found_depts == 7,
              f"Found {found_depts}/7 in tables")


def check_gcal(launch_time_str):
    print("\n=== Checking Google Calendar ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT summary, start_datetime, end_datetime, description
        FROM gcal.events
        WHERE LOWER(summary) LIKE '%talent review%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 7 Talent Review events created", len(events) >= 7,
          f"Found {len(events)} events")

    if launch_time_str:
        try:
            launch_dt = datetime.strptime(launch_time_str, "%Y-%m-%d %H:%M:%S")
            if launch_dt.tzinfo is None:
                launch_dt = launch_dt.replace(tzinfo=timezone.utc)
        except Exception:
            launch_dt = datetime(2026, 3, 7, 10, 0, 0, tzinfo=timezone.utc)
    else:
        launch_dt = datetime(2026, 3, 7, 10, 0, 0, tzinfo=timezone.utc)

    target_start = launch_dt + timedelta(days=14)

    # Check department coverage
    depts_found = set()
    for summary, start_dt, end_dt, desc in events:
        for dept in DEPARTMENTS:
            if dept.lower() in summary.lower():
                depts_found.add(dept)
    check("Calendar events cover all 7 departments", len(depts_found) == 7,
          f"Found {len(depts_found)}: {depts_found}")

    # Check event timing: first event should be ~14 days from launch
    if events:
        first_start = events[0][1]
        if first_start.tzinfo is None:
            first_start = first_start.replace(tzinfo=timezone.utc)
        diff_days = abs((first_start.date() - target_start.date()).days)
        check("First meeting ~14 days from launch", diff_days <= 3,
              f"First event at {first_start.date()}, expected ~{target_start.date()}")

    # Check event duration: should be ~90 minutes
    if events:
        for summary, start_dt, end_dt, desc in events[:1]:
            if start_dt and end_dt:
                duration_min = (end_dt - start_dt).total_seconds() / 60
                check("Meeting duration ~90 minutes", abs(duration_min - 90) <= 15,
                      f"Duration: {duration_min} min")

    # Check descriptions mention headcount
    if events:
        descs_with_count = sum(1 for _, _, _, desc in events if desc and any(c.isdigit() for c in str(desc)))
        check("Event descriptions include numeric data", descs_with_count >= 5,
              f"{descs_with_count}/7 events have numeric descriptions")


def check_email():
    print("\n=== Checking Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    # Find the talent review email
    target_email = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if "hr_leadership@company.com" in recipients:
            subj_lower = (subj or "").lower()
            if "talent review" in subj_lower or "q1 2026" in subj_lower:
                target_email = (subj, from_addr, to_addr, body)
                break

    check("Summary email sent to hr_leadership@company.com", target_email is not None,
          "No matching email found")

    if target_email:
        subj, from_addr, to_addr, body = target_email
        body_lower = (body or "").lower()

        check("Email from talent_review@company.com",
              "talent_review@company.com" in (from_addr or "").lower(),
              f"From: {from_addr}")

        check("Email subject contains 'Talent Review'",
              "talent review" in (subj or "").lower(),
              f"Subject: {subj}")

        # Check body mentions key metrics
        check("Email body mentions headcount",
              "50000" in body or "50,000" in body or "headcount" in body_lower,
              "No headcount mention")

        check("Email body mentions performance or rating",
              "performance" in body_lower or "rating" in body_lower,
              "No performance/rating mention")

        check("Email body mentions promotion",
              "promo" in body_lower,
              "No promotion mention")

        check("Email body mentions salary",
              "salary" in body_lower or "compensation" in body_lower,
              "No salary/compensation mention")

    # Reverse check: noise email should still exist
    noise_found = any("renovation" in (subj or "").lower() for subj, _, _, _ in all_emails)
    check("Noise email not deleted", noise_found, "Office Renovation email missing")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF HR TALENT REVIEW EXCEL WORD GCAL EMAIL - EVALUATION")
    print("=" * 70)

    check_excel(args.agent_workspace)
    check_word(args.agent_workspace)
    check_gcal(args.launch_time)
    check_email()

    print(f"\n{'='*70}")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT, "result": "PASS" if overall else "FAIL"}, f)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
