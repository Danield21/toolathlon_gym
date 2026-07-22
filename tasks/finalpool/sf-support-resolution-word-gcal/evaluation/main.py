"""Evaluation for sf-support-resolution-word-gcal."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Resolution_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "resolution" in p.text.lower() and "analysis" in p.text.lower():
            has_heading = True
            break
    check("Document has resolution analysis heading", has_heading)

    # Check tables exist
    check("Document has at least 2 tables", len(doc.tables) >= 2,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 2:
        return False

    def find_table_by_header(doc, required_columns):
        """Find a table whose header row contains all of required_columns (case-insensitive)."""
        rc_lower = [c.lower() for c in required_columns]
        for t in doc.tables:
            if not t.rows:
                continue
            header_cells = [cell.text.strip().lower() for cell in t.rows[0].cells]
            header_blob = " | ".join(header_cells)
            if all(any(rc in hc for hc in header_cells) for rc in rc_lower):
                return t
        return None

    # Locate priority and category tables explicitly (not by index)
    priority_table = find_table_by_header(doc, ["priority", "total_tickets", "avg_hours", "min_hours", "max_hours"])
    category_table = find_table_by_header(doc, ["category", "total_tickets", "avg_hours"])
    # Avoid double-counting when priority_table matches "category" table check
    if category_table is priority_table:
        # Try to find a different table that has 'category' but not 'priority'
        for t in doc.tables:
            if t is priority_table:
                continue
            header_cells = [cell.text.strip().lower() for cell in t.rows[0].cells]
            if any("category" in hc for hc in header_cells):
                category_table = t
                break

    check("Priority table found by header", priority_table is not None)
    check("Category table found by header", category_table is not None)

    # Load expected data from groundtruth Excel (authoritative)
    gt_priority = []
    gt_category = []
    gt_file = os.path.join(groundtruth_workspace, "Resolution_Data.xlsx")
    if os.path.isfile(gt_file):
        gt_wb = openpyxl.load_workbook(gt_file, data_only=True)
        if "By Priority" in gt_wb.sheetnames:
            for row in gt_wb["By Priority"].iter_rows(min_row=2, values_only=True):
                if row and row[0]:
                    gt_priority.append(row)
        if "By Category" in gt_wb.sheetnames:
            for row in gt_wb["By Category"].iter_rows(min_row=2, values_only=True):
                if row and row[0]:
                    gt_category.append(row)

    # Validate priority table
    if priority_table is not None:
        rows = []
        for row in priority_table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        check(f"Priority table has {len(gt_priority)} rows", len(rows) == len(gt_priority),
              f"Got {len(rows)} rows, expected {len(gt_priority)}")

        for prio, total, avg_h, min_h, max_h in gt_priority:
            matched = None
            for r in rows:
                if r and r[0].strip().lower() == prio.strip().lower():
                    matched = r
                    break
            if not matched:
                check(f"Priority '{prio}' present", False)
                continue

            # Check count exact
            found_count = False
            for cell in matched[1:]:
                try:
                    val = int(cell.replace(",", ""))
                    if val == total:
                        found_count = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Priority {prio} ticket count exact", found_count, f"Expected {total} got {matched[1:]}")

            # Check avg hours, min, max
            for label, expected, tol in [("avg_hours", avg_h, 0.05),
                                          ("min_hours", min_h, 0.05),
                                          ("max_hours", max_h, 0.05)]:
                found = False
                for cell in matched[1:]:
                    try:
                        val = float(cell.replace(",", ""))
                        if num_close(val, float(expected), tol):
                            found = True
                            break
                    except (ValueError, AttributeError):
                        continue
                check(f"Priority {prio} {label}", found, f"Expected ~{expected}")

    # Validate category table
    if category_table is not None:
        rows = []
        for row in category_table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        check(f"Category table has {len(gt_category)} rows", len(rows) == len(gt_category),
              f"Got {len(rows)} rows, expected {len(gt_category)}")

        for cat, total, avg_h in gt_category:
            matched = None
            for r in rows:
                if r and r[0].strip().lower() == cat.strip().lower():
                    matched = r
                    break
            if not matched:
                check(f"Category '{cat}' present", False)
                continue
            found_count = False
            for cell in matched[1:]:
                try:
                    val = int(cell.replace(",", ""))
                    if val == total:
                        found_count = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Category {cat} ticket count exact", found_count, f"Expected {total}")
            found_avg = False
            for cell in matched[1:]:
                try:
                    val = float(cell.replace(",", ""))
                    if num_close(val, float(avg_h), 0.05):
                        found_avg = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Category {cat} avg hours", found_avg, f"Expected ~{avg_h}")

        # Check sort: descending by avg_hours
        avg_vals = []
        for r in rows:
            if r and len(r) > 2:
                try:
                    avg_vals.append(float(r[2].replace(",", "")))
                except (ValueError, AttributeError):
                    pass
        is_desc = all(avg_vals[i] >= avg_vals[i+1] for i in range(len(avg_vals)-1)) if len(avg_vals) > 1 else True
        check("Category table sorted by avg_hours descending", is_desc, f"Got {avg_vals}")

    # Check summary paragraph
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    # Require an overall total resolved tickets count somewhere in summary
    total_resolved = sum(int(r[1]) for r in gt_priority) if gt_priority else 0
    has_overall = ("overall" in full_text and "resolved" in full_text)
    check("Summary mentions 'overall' and 'resolved'", has_overall)
    has_total_count = str(total_resolved) in full_text if total_resolved else True
    check(f"Summary contains overall ticket count {total_resolved}", has_total_count)

    return True


def check_gcal():
    """Check Google Calendar event ."""
    print("\n=== Checking Google Calendar  ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Find specifically the "Support Resolution Review Meeting" event on 2026-03-12 14:00-15:00
    matching_event = None
    for e in events:
        summary, desc, start_dt, end_dt = e
        if not summary:
            continue
        s_lower = summary.lower()
        # Require both 'resolution' AND 'review' AND 'meeting' (per task title)
        if "resolution" in s_lower and "review" in s_lower:
            if "2026-03-12" in str(start_dt or ""):
                matching_event = e
                break

    check("Support Resolution Review Meeting event exists on 2026-03-12",
          matching_event is not None,
          f"Events: {[(e[0], str(e[2])) for e in events]}")

    if matching_event:
        summary, desc, start_dt, end_dt = matching_event
        # Check exact start time 14:00
        start_str = str(start_dt or "")
        check("Event start time is 14:00 (2 PM)",
              "14:00" in start_str,
              f"start_datetime={start_str}")
        # Check exact end time 15:00
        end_str = str(end_dt or "")
        check("Event end time is 15:00 (3 PM)",
              "15:00" in end_str,
              f"end_datetime={end_str}")
        # Check description mentions findings/analysis
        desc_lower = (desc or "").lower()
        has_desc = ("review" in desc_lower or "finding" in desc_lower or "analysis" in desc_lower or "resolution" in desc_lower)
        check("Event description mentions review/findings/analysis",
              has_desc, f"description={desc[:120] if desc else '(empty)'}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("SF SUPPORT RESOLUTION WORD GCAL - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_gcal()

    total_pass = PASS_COUNT
    total_fail = FAIL_COUNT
    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": total_pass, "failed": total_fail, "success": all_ok}, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
