"""Evaluation for terminal-arxiv-canvas-gsheet-word-gcal."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
RUNTIME_ONLY_FAIL = 0


def check(name, condition, detail="", runtime_only=False):
    global PASS_COUNT, FAIL_COUNT, RUNTIME_ONLY_FAIL
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        if runtime_only:
            RUNTIME_ONLY_FAIL += 1
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def num_close(a, b, tol=2.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except Exception:
        return False


# ---------- Word Document ----------
def check_word(ws):
    print("\n=== Checking Word Document ===")
    fpath = os.path.join(ws, "Curriculum_Modernization_Proposal.docx")
    if not os.path.isfile(fpath):
        check("Word document exists", False, f"Not found: {fpath}")
        return
    check("Word document exists", True)

    from docx import Document
    doc = Document(fpath)
    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    check("Title mentions curriculum modernization",
          ("curriculum modernization proposal" in full_text) or
          ("curriculum" in full_text and "modernization" in full_text))
    check("Has executive summary section", "executive summary" in full_text)
    check("Has current state analysis",
          ("current state" in full_text) or ("current" in full_text and "analysis" in full_text))
    check("Has research landscape section",
          ("research landscape" in full_text) or ("research" in full_text and "landscape" in full_text))
    check("Has identified gaps section",
          ("gap" in full_text and "identified" in full_text) or ("gaps" in full_text))
    check("Has proposed updates section",
          "proposed" in full_text and ("update" in full_text or "recommendation" in full_text))
    check("Has implementation timeline",
          "implementation" in full_text or "timeline" in full_text or "phase" in full_text)
    check("Mentions applied analytics", "applied analytics" in full_text or "analytics" in full_text and "algorithms" in full_text)
    check("Mentions at least 2 paper titles",
          sum(1 for t in ["deep learning", "graph neural", "reinforcement learning", "analytics methods", "benchmarking"]
              if t in full_text) >= 2,
          "Expected references to research papers")
    # Count structured recommendations: numbered list entries or explicit "Recommendation N" headings.
    import re
    rec_matches = (
        len(re.findall(r"\brecommendation\s*\d+\b", full_text)) +
        len(re.findall(r"^\s*[1-9]\.[^\d]", full_text, re.MULTILINE)) +
        full_text.count("- recommendation")
    )
    # Fallback to coarse count if structured markers absent.
    coarse = full_text.count("recommendation") + full_text.count("propose")
    check("At least 3 recommendations",
          rec_matches >= 3 or coarse >= 3,
          f"structured={rec_matches}, coarse={coarse}")


# ---------- Google Sheet ----------
def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Check spreadsheet exists
        cur.execute("SELECT id, title FROM gsheet.spreadsheets WHERE title ILIKE '%curriculum%gap%'")
        rows = cur.fetchall()
        check("Spreadsheet 'Curriculum Gap Analysis' exists", len(rows) >= 1,
              f"Found {len(rows)} matching spreadsheets",
              runtime_only=True)
        if not rows:
            cur.close()
            conn.close()
            return

        ss_id = rows[0][0]

        # Check sheets
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
        sheets = cur.fetchall()
        sheet_names = [s[1].lower() for s in sheets]
        sheet_map = {s[1].lower(): s[0] for s in sheets}

        has_current = any("current" in n and "topic" in n for n in sheet_names)
        has_research = any("research" in n and "topic" in n for n in sheet_names)
        has_gap = any("gap" in n and "analysis" in n for n in sheet_names)

        check("Has Current_Topics sheet", has_current, f"Sheets: {sheet_names}")
        check("Has Research_Topics sheet", has_research, f"Sheets: {sheet_names}")
        check("Has Gap_Analysis sheet", has_gap, f"Sheets: {sheet_names}")

        # Check Current_Topics content
        ct_id = None
        for n, sid in sheet_map.items():
            if "current" in n and "topic" in n:
                ct_id = sid
                break
        if ct_id:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, ct_id))
            row_count = cur.fetchone()[0]
            check("Current_Topics has data rows (>= 5)", row_count >= 5, f"Got {row_count}",
                  runtime_only=True)

            # Verify Last_Updated column: Course_ID 1 → 'Fall 2013', Course_ID 2 → 'Fall 2014'.
            # Find column indices by reading header row (row_index = 0).
            cur.execute("""
                SELECT col_index, LOWER(value) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index = 0
            """, (ss_id, ct_id))
            hdr_cells = dict(cur.fetchall())
            cid_col = next((i for i, v in hdr_cells.items() if "course" in (v or "") and "id" in (v or "")), None)
            lu_col = next((i for i, v in hdr_cells.items() if "last" in (v or "") and "updated" in (v or "")), None)
            if cid_col is not None and lu_col is not None:
                # Fetch (row_index, course_id, last_updated) for each data row.
                cur.execute("""
                    SELECT row_index, col_index, value FROM gsheet.cells
                    WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
                      AND col_index IN (%s, %s)
                """, (ss_id, ct_id, cid_col, lu_col))
                row_vals = {}
                for ri, ci, v in cur.fetchall():
                    row_vals.setdefault(ri, {})[ci] = v
                mismatches = []
                for ri, cells in row_vals.items():
                    cid_v = str(cells.get(cid_col, "")).strip()
                    lu_v = str(cells.get(lu_col, "")).strip().lower()
                    if cid_v in ("1", "1.0"):
                        if "fall 2013" not in lu_v:
                            mismatches.append((ri, cid_v, lu_v))
                    elif cid_v in ("2", "2.0"):
                        if "fall 2014" not in lu_v:
                            mismatches.append((ri, cid_v, lu_v))
                check("Last_Updated: Course 1→Fall 2013, Course 2→Fall 2014",
                      len(mismatches) == 0, f"Mismatches: {mismatches[:5]}",
                      runtime_only=True)

        # Check Research_Topics content
        rt_id = None
        for n, sid in sheet_map.items():
            if "research" in n and "topic" in n:
                rt_id = sid
                break
        if rt_id:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, rt_id))
            row_count = cur.fetchone()[0]
            check("Research_Topics has >= 4 paper rows", row_count >= 4, f"Got {row_count}")

        # Check Gap_Analysis content
        ga_id = None
        for n, sid in sheet_map.items():
            if "gap" in n and "analysis" in n:
                ga_id = sid
                break
        if ga_id:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, ga_id))
            row_count = cur.fetchone()[0]
            check("Gap_Analysis has data rows (>= 2)", row_count >= 2, f"Got {row_count}")

            # Check gap types
            cur.execute("""
                SELECT DISTINCT LOWER(value) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND col_index = 0 AND row_index > 0
            """, (ss_id, ga_id))
            gap_types = set(r[0] for r in cur.fetchall() if r[0])
            check("Gap_Analysis has 'Curriculum Gap' or 'Missing Research' types",
                  any("curriculum" in g or "gap" in g or "missing" in g for g in gap_types),
                  f"Found: {gap_types}")

        cur.close()
        conn.close()
    except Exception as e:
        check("Google Sheet DB check", False, str(e))


# ---------- Google Calendar ----------
def check_gcal(launch_time_str=None):
    print("\n=== Checking Google Calendar ===")
    # Parse launch_time to compute expected start_datetimes (+14d 10:00, +21d 13:00).
    exp_review_start = None
    exp_workshop_start = None
    if launch_time_str:
        try:
            from datetime import datetime, timedelta
            lt = datetime.fromisoformat(launch_time_str.replace(" ", "T"))
            exp_review_start = (lt + timedelta(days=14)).replace(hour=10, minute=0, second=0, microsecond=0)
            exp_workshop_start = (lt + timedelta(days=21)).replace(hour=13, minute=0, second=0, microsecond=0)
        except Exception as e:
            print(f"  [warn] launch_time parse: {e}")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Check for Curriculum Review Meeting
        cur.execute("SELECT summary, start_datetime, end_datetime, description FROM gcal.events WHERE summary ILIKE '%curriculum%review%meeting%'")
        rows = cur.fetchall()
        check("'Curriculum Review Meeting' event exists", len(rows) >= 1, f"Found {len(rows)}",
              runtime_only=True)
        if rows:
            check("Review meeting has description",
                  rows[0][3] is not None and len(str(rows[0][3])) > 10,
                  f"Description: {str(rows[0][3])[:100]}",
                  runtime_only=True)
            # Duration check (2h)
            if rows[0][1] and rows[0][2]:
                dur_h = (rows[0][2] - rows[0][1]).total_seconds() / 3600
                check("Review meeting duration ~2 hours",
                      abs(dur_h - 2.0) <= 0.25, f"Got {dur_h} hours",
                      runtime_only=True)
            # Start_datetime: launch_time+14d at 10:00
            if exp_review_start is not None and rows[0][1]:
                actual = rows[0][1]
                # Strip tz if needed, compare date + hour
                delta = abs((actual.replace(tzinfo=None) - exp_review_start).total_seconds())
                check(f"Review meeting starts near launch_time+14d 10:00 (within 2h)",
                      delta <= 2 * 3600, f"Expected {exp_review_start}, got {actual}",
                      runtime_only=True)

        # Check for Faculty Workshop
        cur.execute("SELECT summary, start_datetime, end_datetime, description FROM gcal.events WHERE summary ILIKE '%faculty%workshop%'")
        rows = cur.fetchall()
        check("'Faculty Workshop on New Topics' event exists", len(rows) >= 1, f"Found {len(rows)}",
              runtime_only=True)
        if rows:
            check("Workshop has description",
                  rows[0][3] is not None and len(str(rows[0][3])) > 10,
                  f"Description: {str(rows[0][3])[:100]}",
                  runtime_only=True)
            # Duration check (3h)
            if rows[0][1] and rows[0][2]:
                dur_h = (rows[0][2] - rows[0][1]).total_seconds() / 3600
                check("Workshop duration ~3 hours",
                      abs(dur_h - 3.0) <= 0.5, f"Got {dur_h} hours",
                      runtime_only=True)
            if exp_workshop_start is not None and rows[0][1]:
                actual = rows[0][1]
                delta = abs((actual.replace(tzinfo=None) - exp_workshop_start).total_seconds())
                check(f"Workshop starts near launch_time+21d 13:00 (within 2h)",
                      delta <= 2 * 3600, f"Expected {exp_workshop_start}, got {actual}",
                      runtime_only=True)

        cur.close()
        conn.close()
    except Exception as e:
        check("Google Calendar check", False, str(e))


# ---------- Terminal outputs ----------
def check_terminal_outputs(ws):
    print("\n=== Checking Terminal Outputs ===")

    # current_topics.json
    ct_path = os.path.join(ws, "current_topics.json")
    if os.path.isfile(ct_path):
        check("current_topics.json exists", True)
        with open(ct_path) as f:
            data = json.load(f)
        check("current_topics.json is a list with >= 5 items",
              isinstance(data, list) and len(data) >= 5, f"Got {type(data).__name__} with {len(data) if isinstance(data, list) else 'N/A'} items")
        if isinstance(data, list) and len(data) > 0:
            first = data[0]
            # Per triage: current_topics items have all 4 keys
            for k in ["topic", "course_id", "source_type", "source_name"]:
                check(f"current_topics items have '{k}' key", k in first, f"Keys: {list(first.keys())}")
    else:
        check("current_topics.json exists", False, runtime_only=True)

    # topic_gaps.json
    tg_path = os.path.join(ws, "topic_gaps.json")
    if os.path.isfile(tg_path):
        check("topic_gaps.json exists", True)
        with open(tg_path) as f:
            data = json.load(f)
        # Per triage: require all three keys
        for k in ["covered_topics", "gap_topics", "missing_research_areas"]:
            check(f"topic_gaps.json has '{k}' key",
                  isinstance(data, dict) and k in data,
                  f"Keys: {list(data.keys()) if isinstance(data, dict) else 'not a dict'}")
    else:
        check("topic_gaps.json exists", False, runtime_only=True)

    # Script files required per task. Scripts not bundled in GT, so runtime_only for V1.
    for script in ["extract_topics.py", "gap_analysis.py", "generate_summary.py"]:
        path = os.path.join(ws, script)
        check(f"{script} exists", os.path.isfile(path), f"Not found: {path}",
              runtime_only=True)


def check_reverse_validation(workspace):
    """Check that noise/irrelevant data is NOT present in outputs."""
    print("\n=== Reverse Validation ===")

    # Noise arxiv paper titles that should NOT appear in gsheet
    noise_titles = ["marine biodiversity", "paleolithic cave art", "arctic ecosystems", "uranium-thorium"]
    noise_ids = ["2025.20001", "2025.20002"]

    # Check gsheet does not contain noise papers
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        cur.execute("SELECT id FROM gsheet.spreadsheets WHERE title ILIKE '%curriculum%gap%'")
        rows = cur.fetchall()
        if rows:
            ss_id = rows[0][0]
            cur.execute("SELECT LOWER(value) FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,))
            all_values = " ".join(r[0] for r in cur.fetchall() if r[0])

            no_noise = not any(nt in all_values for nt in noise_titles)
            check("No noise arxiv papers in gsheet (marine biodiversity, cave art)",
                  no_noise,
                  f"Found noise content in gsheet values")
        else:
            check("No noise arxiv papers in gsheet", True, "No spreadsheet to check")

        cur.close()
        conn.close()
    except Exception as e:
        check("Reverse validation (gsheet noise)", False, str(e))

    # Check Word document does not reference noise papers as legitimate research.
    # Allow brief context-mentions (e.g. "identified as irrelevant"); flag only if
    # a noise title appears in a recommendation/proposal/adopt context.
    fpath = os.path.join(workspace, "Curriculum_Modernization_Proposal.docx")
    if os.path.isfile(fpath):
        from docx import Document
        doc = Document(fpath)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()

        # For each noise title, check whether it appears in a problematic context.
        # "identified as irrelevant", "noise papers", "excluded" context is OK.
        problematic = []
        for nt in noise_titles:
            idx = full_text.find(nt)
            if idx < 0:
                continue
            ctx = full_text[max(0, idx - 80):idx + 80]
            legit_markers = ["noise", "irrelevant", "exclude", "not related", "unrelated",
                             "out of scope"]
            if any(m in ctx for m in legit_markers):
                continue  # mentioned in rejection context = OK
            problematic.append(nt)
        check("No noise arxiv papers recommended in Word document",
              len(problematic) == 0,
              f"Found noise titles {problematic} outside rejection context")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_gcal(args.launch_time)
    check_terminal_outputs(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nOverall: {PASS_COUNT}/{total} ({accuracy:.1f}%)")
    # Gate on non-runtime-only FAIL_COUNT == 0 (was 70% accuracy)
    non_runtime_fail = FAIL_COUNT - RUNTIME_ONLY_FAIL
    print(f"Non-runtime fails: {non_runtime_fail}, runtime-only fails: {RUNTIME_ONLY_FAIL}")
    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)
    sys.exit(0 if non_runtime_fail == 0 else 1)


if __name__ == "__main__":
    main()
