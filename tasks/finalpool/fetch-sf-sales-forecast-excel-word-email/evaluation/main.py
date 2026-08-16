"""Evaluation script for fetch-sf-sales-forecast-excel-word-email."""
import os
import argparse
import json
import sys
import openpyxl


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default


def num_close(a, b, abs_tol=1.0, rel_tol=0.0):
    af, bf = safe_float(a), safe_float(b)
    if af is None or bf is None:
        return False
    return abs(af - bf) <= max(abs_tol, abs(bf) * rel_tol)


def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)


def load_gt_lookup(gt_wb, sheet_name, key_col=0, val_col_offset=1):
    """Build dict from GT sheet keyed by first column."""
    if sheet_name not in gt_wb.sheetnames:
        return {}
    ws = gt_wb[sheet_name]
    rows = list(ws.iter_rows(values_only=True))
    out = {}
    for r in rows[1:]:
        if r and r[0] is not None:
            out[str(r[0]).strip()] = r
    return out


def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Excel checks ===")
    excel_path = os.path.join(agent_workspace, "Sales_Forecast_Q2_2026.xlsx")
    if not os.path.exists(excel_path):
        check("Sales_Forecast_Q2_2026.xlsx exists", False)
        # Account for missing-file check count
        for _ in range(40):
            check("Excel sub-check (skipped, missing file)", False)
        return
    check("Sales_Forecast_Q2_2026.xlsx exists", True)
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
    except Exception as e:
        check("Sales_Forecast_Q2_2026.xlsx readable", False, str(e))
        return

    gt_path = os.path.join(groundtruth_workspace, "Sales_Forecast_Q2_2026.xlsx")
    gt_wb = openpyxl.load_workbook(gt_path, data_only=True) if os.path.exists(gt_path) else None

    # === Regional_Forecast ===
    check("Regional_Forecast sheet exists", "Regional_Forecast" in wb.sheetnames)
    if "Regional_Forecast" in wb.sheetnames:
        ws = wb["Regional_Forecast"]
        all_rows = list(ws.iter_rows(values_only=True))
        if not all_rows:
            check("Regional_Forecast has data", False)
        else:
            headers = [str(c).strip().lower() if c else "" for c in all_rows[0]]
            expected_hdrs = ['region', 'current_orders', 'current_revenue',
                             'growth_rate_pct', 'forecasted_revenue', 'revenue_increase']
            for col in expected_hdrs:
                check(f"Regional_Forecast has '{col}' column",
                      col in headers, f"got: {headers[:8]}")
            data_rows = [r for r in all_rows[1:] if any(c is not None for c in r)]
            check("Regional_Forecast has exactly 5 rows",
                  len(data_rows) == 5, f"got {len(data_rows)}")

            # Sort alphabetically by Region
            regions = [str(r[0]).strip() for r in data_rows if r and r[0]]
            check("Regional_Forecast sorted alphabetically by Region",
                  regions == sorted(regions), f"got {regions}")

            if gt_wb and "Regional_Forecast" in gt_wb.sheetnames:
                gt_lookup = load_gt_lookup(gt_wb, "Regional_Forecast")
                a_lookup = {str(r[0]).strip(): r for r in data_rows if r and r[0]}
                for region, gt_row in gt_lookup.items():
                    a_row = a_lookup.get(region)
                    if a_row is None:
                        check(f"Regional_Forecast region '{region}' present", False)
                        continue
                    check(f"Region '{region}' present", True)
                    # Current_Orders (exact)
                    check(f"'{region}' Current_Orders exact",
                          num_close(a_row[1], gt_row[1], abs_tol=0),
                          f"got {a_row[1]}, expected {gt_row[1]}")
                    # Current_Revenue (1 cent tol)
                    check(f"'{region}' Current_Revenue ~",
                          num_close(a_row[2], gt_row[2], abs_tol=0.5),
                          f"got {a_row[2]}, expected {gt_row[2]}")
                    # Growth_Rate (0.05 tol)
                    check(f"'{region}' Growth_Rate_Pct",
                          num_close(a_row[3], gt_row[3], abs_tol=0.05),
                          f"got {a_row[3]}, expected {gt_row[3]}")
                    # Forecasted_Revenue (1 dollar tol)
                    check(f"'{region}' Forecasted_Revenue",
                          num_close(a_row[4], gt_row[4], abs_tol=1.0),
                          f"got {a_row[4]}, expected {gt_row[4]}")
                    # Revenue_Increase
                    check(f"'{region}' Revenue_Increase",
                          num_close(a_row[5], gt_row[5], abs_tol=1.0),
                          f"got {a_row[5]}, expected {gt_row[5]}")

    # === Forecast_Summary ===
    check("Forecast_Summary sheet exists", "Forecast_Summary" in wb.sheetnames)
    if "Forecast_Summary" in wb.sheetnames:
        ws = wb["Forecast_Summary"]
        all_rows = list(ws.iter_rows(values_only=True))
        if not all_rows:
            check("Forecast_Summary has data", False)
        else:
            headers = [str(c).strip().lower() if c else "" for c in all_rows[0]]
            for col in ['metric', 'value']:
                check(f"Forecast_Summary has '{col}' column",
                      col in headers, f"got: {headers}")

        a_lookup = {}
        for r in all_rows[1:]:
            if r and r[0] is not None:
                a_lookup[str(r[0]).strip().lower()] = r[1] if len(r) > 1 else None

        if gt_wb and "Forecast_Summary" in gt_wb.sheetnames:
            gt_rows = list(gt_wb["Forecast_Summary"].iter_rows(values_only=True))
            gt_lookup = {}
            for r in gt_rows[1:]:
                if r and r[0] is not None:
                    gt_lookup[str(r[0]).strip().lower()] = r[1] if len(r) > 1 else None

            # Required metrics
            req_metrics = ['total_current_revenue', 'total_forecasted_revenue',
                           'total_revenue_increase', 'avg_growth_rate',
                           'highest_growth_region', 'lowest_growth_region',
                           'total_orders']
            for m in req_metrics:
                a_val = a_lookup.get(m)
                gt_val = gt_lookup.get(m)
                check(f"Forecast_Summary has metric '{m}'",
                      a_val is not None, f"keys: {list(a_lookup.keys())}")
                if a_val is None or gt_val is None:
                    continue
                if isinstance(gt_val, (int, float)):
                    if 'avg_growth' in m:
                        ok = num_close(a_val, gt_val, abs_tol=0.1)
                    elif 'orders' in m:
                        ok = num_close(a_val, gt_val, abs_tol=0)
                    else:
                        # Tightened: cumulative revenue should be deterministic from per-region rows
                        ok = num_close(a_val, gt_val, abs_tol=1.0)
                    check(f"Forecast_Summary '{m}' value",
                          ok, f"got {a_val}, expected {gt_val}")
                else:
                    check(f"Forecast_Summary '{m}' string",
                          str(a_val).strip().lower() == str(gt_val).strip().lower(),
                          f"got {a_val}, expected {gt_val}")

    # === Growth_Ranking ===
    check("Growth_Ranking sheet exists", "Growth_Ranking" in wb.sheetnames)
    if "Growth_Ranking" in wb.sheetnames:
        ws = wb["Growth_Ranking"]
        all_rows = list(ws.iter_rows(values_only=True))
        if all_rows:
            headers = [str(c).strip().lower() if c else "" for c in all_rows[0]]
            for col in ['rank', 'region', 'growth_rate_pct', 'current_revenue', 'market_share_pct']:
                check(f"Growth_Ranking has '{col}' column", col in headers,
                      f"got: {headers}")
        data_rows = [r for r in all_rows[1:] if any(c is not None for c in r)]
        check("Growth_Ranking has exactly 5 rows", len(data_rows) == 5,
              f"got {len(data_rows)}")

        # Verify sort: descending by Growth_Rate_Pct
        if data_rows and len(data_rows[0]) > 2:
            rates = [safe_float(r[2]) for r in data_rows if r and r[2] is not None]
            check("Growth_Ranking sorted descending by Growth_Rate_Pct",
                  rates == sorted(rates, reverse=True), f"got {rates}")

        # Compare to GT
        if gt_wb and "Growth_Ranking" in gt_wb.sheetnames:
            gt_rows = list(gt_wb["Growth_Ranking"].iter_rows(values_only=True))[1:]
            for i, gt_r in enumerate(gt_rows):
                if i >= len(data_rows):
                    check(f"Growth_Ranking rank {i+1} present", False)
                    continue
                a_r = data_rows[i]
                check(f"Growth_Ranking rank {i+1} = {gt_r[1]}",
                      str(a_r[1]).strip() == str(gt_r[1]).strip(),
                      f"got {a_r[1]}")
                check(f"Growth_Ranking rank {i+1} Market_Share_Pct",
                      num_close(a_r[4], gt_r[4], abs_tol=0.2),
                      f"got {a_r[4]}, expected {gt_r[4]}")


def check_word(agent_workspace, groundtruth_workspace):
    print("\n=== Word checks ===")
    word_path = os.path.join(agent_workspace, "Q2_Forecast_Report.docx")
    if not os.path.exists(word_path):
        check("Q2_Forecast_Report.docx exists", False)
        for _ in range(7):
            check("Word sub-check (skipped)", False)
        return
    check("Q2_Forecast_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(word_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        return

    text = "\n".join(p.text for p in doc.paragraphs)
    text_lower = text.lower()
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
    headings_text = " ".join(headings).lower()

    check("Word title 'Q2 2026 Sales Forecast Report'",
          "q2 2026 sales forecast report" in text_lower)
    check("Word section: Market Overview",
          "market overview" in headings_text or "market overview" in text_lower)
    check("Word section: Regional Analysis",
          "regional analysis" in headings_text or "regional analysis" in text_lower)
    check("Word section: Strategic Recommendations",
          "strategic recommendations" in headings_text or "strategic recommendations" in text_lower)
    # All 5 regions referenced
    regions = ["asia pacific", "europe", "latin america", "middle east", "north america"]
    region_hits = sum(1 for r in regions if r in text_lower)
    check("Word references all 5 regions", region_hits == 5,
          f"hit {region_hits}/5")

    # At least 3 recommendations (count enumerated lines)
    rec_lines = 0
    in_rec = False
    for p in doc.paragraphs:
        if "recommendation" in p.text.lower() and (p.style.name.startswith("Heading") or p.style.name == "Title"):
            in_rec = True
            continue
        if in_rec and p.style.name.startswith("Heading"):
            in_rec = False
        if in_rec and p.text.strip():
            rec_lines += 1
    check("At least 3 recommendation lines under Strategic Recommendations",
          rec_lines >= 3, f"got {rec_lines}")


def check_email(agent_workspace):
    print("\n=== Email checks ===")
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, to_addr, body_text
            FROM email.messages
            WHERE subject ILIKE '%Q2 2026 Sales Forecast Ready%'
        """)
        emails = cur.fetchall()
        conn.close()
    except Exception as e:
        check("Email DB connection", False, str(e))
        return

    check("Email with subject 'Q2 2026 Sales Forecast Ready' sent",
          len(emails) >= 1, f"found {len(emails)} matching email(s)")
    if emails:
        subj, to_addr, body = emails[0]
        check("Email subject exact match",
              str(subj).strip() == "Q2 2026 Sales Forecast Ready",
              f"got '{subj}'")
        check("Email to sales-team@company.com",
              "sales-team@company.com" in str(to_addr).lower(),
              f"got '{to_addr}'")
        body_lower = (body or "").lower()
        check("Email body mentions total forecasted revenue",
              ("total" in body_lower and "forecast" in body_lower) or
              "3,219" in (body or "") or "3219" in (body or ""),
              f"body sample: {(body or '')[:200]}")
        check("Email body mentions highest growth region (Asia + descriptor)",
              "asia" in body_lower and ("highest" in body_lower or "top" in body_lower or "leading" in body_lower or "fastest" in body_lower),
              f"body sample: {(body or '')[:200]}")


def check_scripts(agent_workspace):
    print("\n=== Script & data files ===")
    files = [
        "forecast_builder.py",
        "historical_sales.json",
        "market_outlook.json",
        "sales_forecast.json",
    ]
    for f in files:
        check(f"{f} exists",
              os.path.exists(os.path.join(agent_workspace, f)),
              f"workspace: {agent_workspace}")


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    check_excel(agent_workspace, groundtruth_workspace)
    check_word(agent_workspace, groundtruth_workspace)
    check_email(agent_workspace)
    check_scripts(agent_workspace)

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
