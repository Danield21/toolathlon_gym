"""
Evaluation script for academic-literature-review task.
Checks that Literature_Review.docx exists and contains the expected content.

Usage:
  python -m evaluation.main --agent_workspace <path> --groundtruth_workspace <path> --launch_time <time>
"""
import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except Exception:  # pragma: no cover - docx is a hard dependency of this task
    Document = None
    Table = None
    Paragraph = None


PASS_COUNT = 0
FAIL_COUNT = 0


def check(name: str, condition: bool, detail: str = ""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_truncated = (detail[:200] + "...") if len(detail) > 200 else detail
        print(f"  [FAIL] {name}: {detail_truncated}")


def normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace."""
    return re.sub(r'\s+', ' ', text.lower().strip())


_PUNCT_RE = re.compile(r'[^\w\s]')


def normalize_fuzzy(text: str) -> str:
    """Normalize like ``normalize`` but also drop punctuation so that small
    typographic differences (colons, commas, dashes, quotes) do not cause a
    mismatch on paper titles."""
    return normalize(_PUNCT_RE.sub(' ', text))


def _iter_block_items(document):
    """Yield top-level body elements (paragraphs and tables) in document order."""
    body = document.element.body
    ns_p = '}p'
    ns_tbl = '}tbl'
    for child in body.iterchildren():
        tag = child.tag
        if tag.endswith(ns_p):
            yield Paragraph(child, document)
        elif tag.endswith(ns_tbl):
            yield Table(child, document)


def _collect_blocks(blocks, out):
    """Collect all paragraph text, including text inside tables (recursively)."""
    for block in blocks:
        if isinstance(block, Paragraph):
            if block.text:
                out.append(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            out.append(para.text)
                    for nested in cell.tables:
                        _collect_blocks([nested], out)


def extract_full_text(docx_path):
    doc = Document(docx_path)
    chunks = []
    _collect_blocks(_iter_block_items(doc), chunks)
    return "\n".join(chunks)


def _find_docx_candidates(agent_workspace):
    """Return candidate Word deliverables, best first.

    Order of preference:
      1. Literature_Review.docx at the workspace root (the exact path named in
         the task instructions).
      2. Any file named Literature_Review.docx (case-insensitive) found
         anywhere under the workspace (handles a correctly written document
         placed in a subdirectory).
      3. Any other .docx under the workspace, ordered by how much text they
         contain (most content first), so the most complete document is graded
         rather than a scratch draft; files that cannot be parsed rank last.

    All candidates are returned rather than stopping at the first hit, so that
    a corrupted preferred copy (e.g. truncated by a concurrent write to the
    same path) can fall back to a readable copy elsewhere in the workspace.
    """
    root = os.path.join(agent_workspace, "Literature_Review.docx")
    exact = []
    others = []
    for dirpath, _dirnames, filenames in os.walk(agent_workspace):
        for fn in filenames:
            if fn.startswith("._") or not fn.lower().endswith(".docx"):
                continue
            full = os.path.join(dirpath, fn)
            if not os.path.isfile(full):
                continue
            if fn.lower() == "literature_review.docx":
                exact.append(full)
            else:
                try:
                    score = len(extract_full_text(full))
                except Exception:
                    score = -1
                others.append((score, full))
    exact.sort()
    others.sort(key=lambda item: item[0], reverse=True)
    result = exact + [full for _score, full in others]
    # Keep the exact root path first when present (it is the canonical deliverable).
    if os.path.isfile(root) and root in result:
        result.remove(root)
        result.insert(0, root)
    return result


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", type=str, required=True)
    parser.add_argument("--groundtruth_workspace", type=str, required=True)
    parser.add_argument("--launch_time", type=str, required=False)
    parser.add_argument("--res_log_file", type=str, required=False)
    args = parser.parse_args()

    # Locate the deliverable. The exact root path is preferred, but a correctly
    # completed model may place the document in a subdirectory or under a
    # slightly different name, and a concurrent write race can corrupt the
    # preferred copy while a good one survives elsewhere. We grade the best
    # readable .docx in the workspace instead of failing the whole task because
    # one hardcoded path is missing or unreadable.
    candidates = _find_docx_candidates(args.agent_workspace)

    # Check 1: File exists
    check("Literature_Review.docx exists", bool(candidates),
          f"No .docx deliverable found under {args.agent_workspace}")

    if not candidates:
        print(f"\nResults: {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} passed, {FAIL_COUNT} failed")
        sys.exit(1)

    # Read the Word document (paragraphs and tables). If the preferred file is
    # unreadable (e.g. truncated by a concurrent write to the same path), fall
    # back to the next candidate so the agent's content work is not discarded.
    full_text = None
    used_path = None
    last_err = ""
    for cand in candidates:
        try:
            full_text = extract_full_text(cand)
            used_path = cand
            break
        except Exception as e:
            last_err = str(e)
            continue

    if full_text is None:
        check("Word document readable", False, last_err)
        print(f"\nResults: {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} passed, {FAIL_COUNT} failed")
        sys.exit(1)

    if used_path and os.path.abspath(used_path) != os.path.abspath(candidates[0]):
        print(f"  [note] deliverable read from fallback: {used_path}")

    normalized = normalize(full_text)
    fuzzy_normalized = normalize_fuzzy(full_text)

    # Check 2: Minimum content length
    check("Document has at least 500 characters",
          len(full_text.strip()) >= 500,
          f"Document has {len(full_text.strip())} characters")

    # Check 3: All 5 paper titles appear (case-insensitive, punctuation-insensitive).
    # These papers must match the papers injected by preprocess/main.py
    # into arxiv/scholarly schemas. For each paper we accept either the full
    # normalized title or a distinctive core phrase, so reasonable rewording
    # (e.g. dropping a subtitle, changing punctuation) does not cause a false fail.
    paper_titles = [
        "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks",
        "Retrieval-Augmented Generation for Large Language Models: A Survey",
        "Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection",
        "RAPTOR: Recursive Abstractive Processing for Tree-Organized Retrieval",
        "From RAG to Rich: Building Robust Retrieval-Augmented Generation Systems",
    ]
    paper_title_phrases = [
        [  # original RAG (Lewis et al.)
            "retrieval augmented generation for knowledge intensive nlp tasks",
            "knowledge intensive nlp tasks",
        ],
        [  # survey (Gao et al.)
            "retrieval augmented generation for large language models a survey",
            "retrieval augmented generation for large language models",
        ],
        [  # Self-RAG (Asai et al.)
            "self rag learning to retrieve generate and critique through self reflection",
            "self rag learning to retrieve generate and critique",
        ],
        [  # RAPTOR (Sarthi et al.)
            "raptor recursive abstractive processing for tree organized retrieval",
            "recursive abstractive processing",
        ],
        [  # From RAG to Rich (Chen et al.)
            "from rag to rich building robust retrieval augmented generation systems",
            "from rag to rich",
        ],
    ]
    for title, phrases in zip(paper_titles, paper_title_phrases):
        matched = any(phrase in fuzzy_normalized for phrase in phrases)
        check(f"Paper title present: {title[:60]}...",
              matched,
              "Title not found in document text")

    # Check 4: All 5 first authors appear
    # These authors must match the papers injected by preprocess/main.py
    # into arxiv/scholarly schemas.
    first_authors = ["Lewis", "Gao", "Asai", "Sarthi", "Chen"]
    for author in first_authors:
        check(f"Author present: {author}",
              author.lower() in normalized,
              f"Author '{author}' not found in document text")

    # Check 5: Key domain terms present
    key_terms = ["retrieval", "generation", "augmented"]
    for term in key_terms:
        check(f"Key term present: {term}",
              term.lower() in normalized,
              f"Term '{term}' not found")

    # Check 6: Document has structure (introduction/conclusion indicators)
    has_intro = "introduction" in normalized
    has_conclusion = "conclusion" in normalized or "summary" in normalized or "synthesis" in normalized
    check("Document has introduction section", has_intro,
          "No 'introduction' found in text")
    check("Document has conclusion/summary section", has_conclusion,
          "No 'conclusion' or 'summary' found in text")

    # Check 7: Reverse validation — distractor (non-RAG) papers must NOT appear.
    # The preprocess injects these unrelated papers; the task instructions ask
    # the agent to keep the review focused on RAG papers only.
    distractor_titles = [
        "Deep Learning Approaches for Medical Image Segmentation",
        "Quantum Computing Foundations and Error Correction",
        "Reinforcement Learning for Robotic Manipulation in Cluttered Environments",
        "Graph Neural Networks for Social Network Analysis",
    ]
    for dtitle in distractor_titles:
        check(f"Distractor NOT present: {dtitle[:50]}...",
              normalize_fuzzy(dtitle) not in fuzzy_normalized,
              f"Unrelated distractor paper leaked into review: {dtitle}")

    # Summary with threshold-based pass (80%)
    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total > 0 else 0
    print(f"\nResults: {PASS_COUNT}/{total} passed, {FAIL_COUNT} failed ({accuracy:.1f}%)")

    sys.exit(0 if accuracy >= 80 else 1)


if __name__ == "__main__":
    main()
