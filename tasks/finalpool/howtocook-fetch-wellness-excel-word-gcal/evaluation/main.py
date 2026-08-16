"""Evaluation script for howtocook-fetch-wellness-excel-word-gcal."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Corporate_Wellness_Plan.xlsx
    excel_path = os.path.join(agent_workspace, "Corporate_Wellness_Plan.xlsx")
    check("Corporate_Wellness_Plan.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Corporate_Wellness_Plan.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    gt_rows = [r for r in gt_rows if r and r[0] is not None]
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    data_rows = [r for r in data_rows if r and r[0] is not None]
                    # Require at least len(gt_rows) data rows for Weekly_Plan (5 days),
                    # Nutritional_Summary (3 slots), Program_Metrics (5 metrics).
                    # Recipe_Evaluation is discretionary (task says "at least 10 recipes").
                    if sheet_name == "Recipe_Evaluation":
                        min_rows = max(3, len(gt_rows))
                    else:
                        min_rows = len(gt_rows)
                    check(f"{sheet_name} has >= {min_rows} data rows", len(data_rows) >= min_rows, f"got {len(data_rows)}")

                    # Cell value comparison against groundtruth
                    header_map = {h: i for i, h in enumerate(headers)}
                    # Build agent row lookup by first-column key to compare keyed rows (not positional).
                    agent_lookup = {}
                    for ar in data_rows:
                        if ar and ar[0] is not None:
                            agent_lookup[str(ar[0]).strip().lower()] = ar
                    # Compare ALL GT rows (not just first 3) matched by first-column key.
                    for ri, gt_row in enumerate(gt_rows):
                        key = str(gt_row[0]).strip().lower() if gt_row[0] is not None else ""
                        agent_row = agent_lookup.get(key)
                        if agent_row is None:
                            # Some sheets (Recipe_Evaluation) may have different row order; fall back to positional.
                            if ri < len(data_rows):
                                agent_row = data_rows[ri]
                            else:
                                continue
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                # Tighter tolerance: absolute 0.5, relative 10% (Wellness_Score/Compliance are subjective)
                                tol = max(0.5, abs(gf) * 0.10)
                                check(f"{sheet_name} R{ri+2} {gt_h} ~{gf:.1f}",
                                      abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    # Exact case-insensitive match; bidirectional substring removed.
                                    check(f"{sheet_name} R{ri+2} {gt_h} text",
                                          gs == avs,
                                          f"expected {gs[:50]}, got {avs[:50]}")

    # Check Wellness_Program_Guide.docx
    docx_path = os.path.join(agent_workspace, "Wellness_Program_Guide.docx")
    check("Wellness_Program_Guide.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Wellness_Program_Guide.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        gt_doc_path = os.path.join(groundtruth_workspace, "Wellness_Program_Guide.docx")
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if p.style.name.startswith("Heading")]
            for gh in gt_headings:
                if gh:
                    # Require exact or agent heading containing GT heading (GT is canonical).
                    found = any(h == gh or gh in h for h in headings)
                    check(f"Wellness_Program_Guide.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("Wellness_Program_Guide.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Track file-level failures (blocking) separately from runtime checks.
    file_fail_before_runtime = FAIL_COUNT

    # Check Python script exists (terminal usage) - only when agent_workspace differs
    # from groundtruth (skip in GT-self-test where no script is expected).
    try:
        gt_canon = os.path.realpath(groundtruth_workspace)
        ag_canon = os.path.realpath(agent_workspace)
    except Exception:
        gt_canon, ag_canon = groundtruth_workspace, agent_workspace
    if gt_canon != ag_canon:
        py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
        check("Python analysis script exists", len(py_files) >= 1, f"found: {py_files}")

    # Database checks (runtime-only; non-blocking when agent has not yet run)
    try:
        conn = get_conn()
        cur = conn.cursor()
        cur.execute("SELECT summary, start_datetime FROM gcal.events WHERE summary ILIKE '%wellness%'")
        event_row = cur.fetchone()
        check("Calendar event with correct summary", event_row is not None, "no matching event found")
        # Reverse verification: noise events should not match task keyword
        cur.execute("SELECT COUNT(*) FROM gcal.events WHERE summary ILIKE '%standup%' OR summary ILIKE '%lunch%'")
        noise_events = cur.fetchone()[0]
        check("Noise events exist (not deleted by agent)", noise_events >= 1, f"noise events: {noise_events}")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    # Local file checks must be zero; runtime checks allowed to fail at 85% accuracy.
    total = PASS_COUNT + FAIL_COUNT
    accuracy = (PASS_COUNT / total * 100) if total else 0
    file_ok = file_fail_before_runtime == 0
    return (file_ok and accuracy >= 85), f"Passed {PASS_COUNT}/{total} checks (accuracy={accuracy:.1f}%, file_fail={file_fail_before_runtime})"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()