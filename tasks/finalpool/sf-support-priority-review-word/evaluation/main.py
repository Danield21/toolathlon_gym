"""
Evaluation script for sf-support-priority-review-word task.

Checks:
1. Word file (Priority_Handling_Report.docx) - title, date 2026-03-06, per-priority numeric values, SLA targets
2. Google Sheet 'Priority Metrics Dashboard' / 'By Priority' with rows per priority
3. Email from support-analytics@company.com to support-director@company.com,
   subject 'Priority Handling Performance Report', body contains response time numbers
"""

import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PRIORITIES = ["High", "Medium", "Low"]

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def get_expected():
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute('SELECT "PRIORITY", COUNT(*), ROUND(AVG("RESPONSE_TIME_HOURS")::numeric, 2), ROUND(AVG("CUSTOMER_SATISFACTION")::numeric, 2) FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS" GROUP BY "PRIORITY"')
    p = {r[0]: {"count": int(r[1]), "avg_resp": float(r[2]), "avg_sat": float(r[3])} for r in cur.fetchall()}
    cur.execute('SELECT "PRIORITY", "RESPONSE_TARGET_HOURS" FROM sf_data."SUPPORT_CENTER__PUBLIC__SLA_POLICIES"')
    sla = {r[0]: float(r[1]) for r in cur.fetchall()}
    for prio, info in p.items():
        info["sla_target"] = sla.get(prio)
    cur.execute('SELECT COUNT(*) FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"')
    total = cur.fetchone()[0]
    conn.close()
    return p, total


def num_in(text_lower, value, tol=0.05):
    """Check if a numeric value (rounded similarly) appears in text."""
    if value is None:
        return False
    candidates = []
    if isinstance(value, int) or float(value).is_integer():
        candidates.append(str(int(value)))
    candidates.append(f"{float(value):.2f}")
    candidates.append(f"{float(value):.1f}")
    return any(c in text_lower for c in candidates)


def check_word(agent_workspace, expected, total_tickets):
    print("\n=== Checking Word Output ===")
    docx_path = os.path.join(agent_workspace, "Priority_Handling_Report.docx")
    check("Word file exists", os.path.isfile(docx_path), f"Expected {docx_path}")
    if not os.path.isfile(docx_path):
        return False

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        check("Word file readable", False, str(e))
        return False

    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += cell.text + "\n"
    text_lower = full_text.lower()

    check("Word title 'Support Ticket Priority Analysis' present",
          "support ticket priority analysis" in text_lower,
          f"text sample: {full_text[:200]}")
    check("Word contains date '2026-03-06'", "2026-03-06" in full_text,
          f"text sample: {full_text[:300]}")

    for prio in PRIORITIES:
        info = expected.get(prio)
        check(f"Word contains '{prio} priority' section header", prio.lower() in text_lower)
        if info is None:
            continue
        check(f"Word: {prio}.Number of Tickets == {info['count']}",
              str(info["count"]) in full_text,
              f"missing {info['count']}")
        check(f"Word: {prio}.Avg Response Time ≈ {info['avg_resp']}",
              num_in(text_lower, info["avg_resp"]),
              f"missing ~{info['avg_resp']}")
        check(f"Word: {prio}.Avg Satisfaction ≈ {info['avg_sat']}",
              num_in(text_lower, info["avg_sat"]),
              f"missing ~{info['avg_sat']}")
        if info.get("sla_target") is not None:
            check(f"Word: {prio}.SLA Target {info['sla_target']} hrs",
                  num_in(text_lower, info["sla_target"]),
                  f"missing SLA {info['sla_target']}")

    check("Word contains overall summary section",
          "overall" in text_lower or "summary" in text_lower,
          "no 'overall'/'summary' keyword")
    check(f"Word contains total tickets {total_tickets}",
          str(total_tickets) in full_text,
          f"missing total {total_tickets}")

    return True


def check_gsheet(expected):
    print("\n=== Checking Google Sheet ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
    except psycopg2.Error as e:
        check("DB connect", False, str(e))
        return
    cur.execute("SELECT id, title FROM gsheet.spreadsheets WHERE LOWER(title) LIKE '%%priority%%metrics%%' OR LOWER(title) LIKE '%%priority metrics dashboard%%'")
    sheets = cur.fetchall()
    check("Spreadsheet 'Priority Metrics Dashboard' exists", len(sheets) >= 1, f"Found {len(sheets)}")
    if not sheets:
        conn.close()
        return
    sid = sheets[0][0]
    cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id=%s", (sid,))
    sub = cur.fetchall()
    sheet_id = next((s_id for s_id, t in sub if t and t.strip().lower() == "by priority"), None)
    check("Sheet 'By Priority' present", sheet_id is not None, f"got {[t for _, t in sub]}")
    if sheet_id is None:
        conn.close()
        return
    cur.execute("SELECT row_index, col_index, value FROM gsheet.cells WHERE sheet_id=%s ORDER BY row_index, col_index", (sheet_id,))
    cells = cur.fetchall()
    grid = {}
    for r, c, v in cells:
        grid.setdefault(r, {})[c] = str(v).strip() if v is not None else ""
    # Find header row
    header_row = None
    for r in sorted(grid.keys()):
        joined = " ".join(grid[r].values()).lower()
        if "priority" in joined and "ticket" in joined and "response" in joined:
            header_row = r
            break
    check("GSheet header row present", header_row is not None)
    if header_row is None:
        conn.close()
        return
    col_map = {}
    for col, val in grid[header_row].items():
        v = val.strip().lower()
        if v == "priority":
            col_map["priority"] = col
        elif v == "ticket_count":
            col_map["count"] = col
        elif v == "avg_response_hours":
            col_map["resp"] = col
        elif v == "avg_satisfaction":
            col_map["sat"] = col
        elif v == "sla_target":
            col_map["sla"] = col

    data_rows = [grid[r] for r in sorted(grid.keys()) if r > header_row]
    by_prio = {}
    for row in data_rows:
        p = row.get(col_map.get("priority", -1), "").strip()
        if p:
            by_prio[p.lower()] = row
    for prio in PRIORITIES:
        info = expected.get(prio)
        if info is None:
            continue
        row = by_prio.get(prio.lower())
        check(f"GSheet row for {prio} present", row is not None)
        if row is None:
            continue
        try:
            cnt = int(float(row.get(col_map.get("count", -1), "0")))
        except (TypeError, ValueError):
            cnt = None
        try:
            resp = float(row.get(col_map.get("resp", -1), "0"))
        except (TypeError, ValueError):
            resp = None
        try:
            sat = float(row.get(col_map.get("sat", -1), "0"))
        except (TypeError, ValueError):
            sat = None
        try:
            sla = float(row.get(col_map.get("sla", -1), "0"))
        except (TypeError, ValueError):
            sla = None
        check(f"GSheet {prio}.Ticket_Count == {info['count']}", cnt == info["count"], f"got {cnt}")
        check(f"GSheet {prio}.Avg_Response_Hours ≈ {info['avg_resp']}",
              resp is not None and abs(resp - info["avg_resp"]) <= 0.05, f"got {resp}")
        check(f"GSheet {prio}.Avg_Satisfaction ≈ {info['avg_sat']}",
              sat is not None and abs(sat - info["avg_sat"]) <= 0.05, f"got {sat}")
        check(f"GSheet {prio}.SLA_Target == {info['sla_target']}",
              sla is not None and abs(sla - info["sla_target"]) <= 0.001, f"got {sla}")

    conn.close()


def check_email(expected):
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
    except psycopg2.Error as e:
        check("DB connect", False, str(e))
        return
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    print(f"  Found {len(all_emails)} emails.")
    found = None
    for subject, from_addr, to_addr, body in all_emails:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                if isinstance(parsed, list):
                    to_str = " ".join(str(r).lower() for r in parsed)
                else:
                    to_str = str(to_addr).lower()
            except (json.JSONDecodeError, TypeError):
                to_str = str(to_addr).lower()
        if "support-director@company.com" in to_str:
            found = (subject, from_addr, to_addr, body)
            break
    check("Email sent to support-director@company.com", found is not None)
    if found is None:
        return
    subject, from_addr, to_addr, body = found
    check("Email subject == 'Priority Handling Performance Report'",
          (subject or "").strip().lower() == "priority handling performance report",
          f"got: {subject}")
    check("Email from_addr == 'support-analytics@company.com'",
          str(from_addr or "").strip().lower() == "support-analytics@company.com",
          f"got: {from_addr}")
    body_lower = (body or "").lower()
    body_text = body or ""
    for prio in PRIORITIES:
        info = expected.get(prio)
        if info is None:
            continue
        check(f"Email body mentions '{prio}' priority", prio.lower() in body_lower,
              f"missing {prio}")
        check(f"Email body mentions response time {info['avg_resp']:.2f} hrs",
              num_in(body_lower, info["avg_resp"]),
              f"missing ~{info['avg_resp']}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    expected, total_tickets = get_expected()
    check_word(args.agent_workspace, expected, total_tickets)
    check_gsheet(expected)
    check_email(expected)

    all_ok = FAIL_COUNT == 0
    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": all_ok}, f, indent=2)
    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
