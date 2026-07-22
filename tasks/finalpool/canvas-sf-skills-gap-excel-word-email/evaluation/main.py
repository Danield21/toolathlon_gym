"""Evaluation script for canvas-sf-skills-gap-excel-word-email."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Skills_Gap_Report.xlsx
    excel_path = os.path.join(agent_workspace, "Skills_Gap_Report.xlsx")
    check("Skills_Gap_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            # Define key column for each sheet (for row-key matching)
            # Subjective free-text columns are skipped from strict equality
            sheet_key_col = {
                "department_skills": "department",
                "course_utilization": "course_code",
                "gap_details": "department",  # composite below for multi-row per dept
                "summary": "metric",
            }
            # Sheets where multiple rows can share the same primary key value;
            # use composite (key + secondary) to avoid agent_by_key overwrites
            sheet_secondary_col = {
                "gap_details": "skill_needed",
            }
            # Columns to skip (subjective free text, hand-written)
            skip_text_cols = {"recommended_action", "skill_needed", "current_coverage", "priority", "course_name", "target_department"}

            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count - allow agent to produce >= GT rows (Course_Utilization may have 22 vs GT's 3)
                    gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True)
                               if any(c is not None for c in r)]
                    data_rows = [r for r in ws.iter_rows(min_row=2, values_only=True)
                                 if any(c is not None for c in r)]
                    check(f"{sheet_name} has at least {len(gt_rows)} data rows",
                          len(data_rows) >= len(gt_rows), f"got {len(data_rows)}")

                    # Build row-key index for agent data
                    header_map = {h: i for i, h in enumerate(headers)}
                    gt_header_map = {h: i for i, h in enumerate(gt_headers)}
                    key_col = sheet_key_col.get(sheet_name.lower())

                    # Determine secondary key for sheets with multi-row per primary key
                    sec_col = sheet_secondary_col.get(sheet_name.lower())

                    def _row_key(row, hmap, kc, sc):
                        parts = []
                        for c in (kc, sc):
                            if c and c in hmap and hmap[c] < len(row) and row[hmap[c]] is not None:
                                parts.append(str(row[hmap[c]]).strip().lower())
                        return "|".join(parts) if parts else None

                    # Map agent rows by composite key value (lowercase normalized)
                    agent_by_key = {}
                    if key_col and key_col in header_map:
                        for ar in data_rows:
                            k = _row_key(ar, header_map, key_col, sec_col)
                            if k:
                                agent_by_key[k] = ar

                    # For each GT row, find the matching agent row by key
                    gt_ki = gt_header_map.get(key_col) if key_col else None
                    for gi, gt_row in enumerate(gt_rows):
                        # Determine key for GT row
                        agent_row = None
                        gk = _row_key(gt_row, gt_header_map, key_col, sec_col)
                        if gk:
                            agent_row = agent_by_key.get(gk)
                        # Fallback: positional
                        if agent_row is None and gi < len(data_rows):
                            agent_row = data_rows[gi]
                        if agent_row is None:
                            continue
                        # Determine display key for messages
                        display_key = ""
                        if gt_ki is not None and gt_ki < len(gt_row) and gt_row[gt_ki] is not None:
                            display_key = str(gt_row[gt_ki])

                        # Canonical (non-free-text) fields require EXACT match
                        # rather than substring (Department / Course_Code / Metric).
                        canonical_text_cols = {"department", "course_code", "metric"}
                        # Count metrics in Summary sheet require exact equality.
                        count_metrics = {
                            "total_departments", "total_courses",
                            "critical_gaps_count",
                        }
                        sn_lower2 = sheet_name.strip().lower()

                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            # Skip subjective free-text columns
                            if gt_h in skip_text_cols:
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                # Special case: Summary count metrics require
                                # EXACT equality (not 30% tolerance).
                                metric_label = ""
                                if "summary" in sn_lower2:
                                    # In Summary, key column "metric" identifies row
                                    metric_idx = gt_header_map.get("metric", -1)
                                    if metric_idx >= 0 and metric_idx < len(gt_row) and gt_row[metric_idx] is not None:
                                        metric_label = str(gt_row[metric_idx]).strip().lower()
                                if metric_label in count_metrics:
                                    check(
                                        f"{sheet_name}[{display_key or gi}] {gt_h} (exact count) = {gf}",
                                        abs(gf - af) <= 0.5,
                                        f"got {af}",
                                    )
                                else:
                                    # Wider numeric tolerance for hand-fabricated GT values
                                    tol = max(1.0, abs(gf) * 0.30)
                                    check(f"{sheet_name}[{display_key or gi}] {gt_h} ~{gf:.1f}",
                                          abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    # Canonical fields: exact equality (lower-trimmed).
                                    if gt_h in canonical_text_cols:
                                        check(
                                            f"{sheet_name}[{display_key or gi}] {gt_h} (exact) = {gs!r}",
                                            gs == avs,
                                            f"got {avs[:50]}",
                                        )
                                    else:
                                        # Other retained text columns: substring match (lenient).
                                        check(f"{sheet_name}[{display_key or gi}] {gt_h} text",
                                              gs == avs or gs in avs or avs in gs,
                                              f"expected {gs[:50]}, got {avs[:50]}")

    # Check Skills_Gap_Brief.docx
    docx_path = os.path.join(agent_workspace, "Skills_Gap_Brief.docx")
    check("Skills_Gap_Brief.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Skills_Gap_Brief.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth
        headings = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
        gt_doc_path = os.path.join(groundtruth_workspace, "Skills_Gap_Brief.docx")
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if p.style.name.startswith("Heading")]
            for gh in gt_headings:
                if gh:
                    # One-directional substring: GT heading must appear in
                    # an agent heading (e.g., GT 'Methodology' allowed inside
                    # 'Methodology Overview'). Reject the inverse to prevent
                    # micro-headings like 'Gaps' passing for 'Critical Gaps Identified'.
                    found = any(gh in h for h in headings)
                    check(f"Skills_Gap_Brief.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("Skills_Gap_Brief.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Check Python script exists by name (skills_gap_analyzer.py per task.md)
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("skills_gap_analyzer.py exists", "skills_gap_analyzer.py" in py_files, f"found: {py_files}")

    # Database checks
    try:
        conn = get_conn()
        cur = conn.cursor()
        # Look for email with correct subject and recipient hr-director@company.com
        # Don't depend on folder placement - check by recipient and subject
        cur.execute("""
            SELECT subject, to_addr FROM email.messages
            WHERE to_addr::text ILIKE '%hr-director@company.com%'
              AND subject ILIKE '%skill%gap%analysis%complete%'
        """)
        email_row = cur.fetchone()
        check("Email to hr-director@company.com with subject 'Skills Gap Analysis Complete' sent",
              email_row is not None, "no matching email found")
        # Check body mentions top 3 departments dynamically derived from agent's own Excel
        # Read agent's Department_Skills sheet and pick top 3 by Gap_Score
        cur.execute("""
            SELECT body_text FROM email.messages
            WHERE to_addr::text ILIKE '%hr-director@company.com%'
              AND subject ILIKE '%skill%'
            ORDER BY id DESC LIMIT 1
        """)
        body_row = cur.fetchone()
        if body_row and body_row[0]:
            body_lower = body_row[0].lower()
            # Dynamically derive expected top departments from agent's own Excel
            top_depts_from_agent = []
            try:
                if os.path.exists(excel_path):
                    a_wb = openpyxl.load_workbook(excel_path)
                    # Find department sheet
                    dept_sheet = None
                    for sn in a_wb.sheetnames:
                        if "department" in sn.lower() and "skill" in sn.lower():
                            dept_sheet = a_wb[sn]
                            break
                    if dept_sheet is not None:
                        a_headers = [str(c.value).strip().lower() if c.value else "" for c in dept_sheet[1]]
                        dept_idx = None
                        gap_idx = None
                        for hi, h in enumerate(a_headers):
                            if "department" in h:
                                dept_idx = hi
                            if "gap" in h and "score" in h:
                                gap_idx = hi
                        if dept_idx is not None and gap_idx is not None:
                            rows = []
                            for r in dept_sheet.iter_rows(min_row=2, values_only=True):
                                if r[dept_idx] is None:
                                    continue
                                gv = safe_float(r[gap_idx])
                                if gv is None:
                                    continue
                                rows.append((str(r[dept_idx]).strip(), gv))
                            rows.sort(key=lambda x: -x[1])
                            top_depts_from_agent = [d for d, _ in rows[:3]]
            except Exception:
                pass
            # Fallback: if agent's data unparsable, use GT-derived list
            if not top_depts_from_agent:
                top_depts_from_agent = ["Support", "Operations", "Sales"]
            for dept in top_depts_from_agent:
                check(f"Email body mentions top dept: {dept}",
                      dept.lower() in body_lower, body_lower[:200])
        else:
            check("Email body mentions top departments", False, "no body found")
        conn.close()
    except Exception as e:
        check("DB checks", False, str(e))

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()