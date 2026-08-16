"""Evaluation for canvas-late-submission-word-gcal."""
import argparse
import json
import os
import sys

import psycopg2
from zoneinfo import ZoneInfo

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# Source of truth: docs/task.md. The Calendar MCP writes events as UTC
# instants into the DB (`TIMESTAMPTZ`), so all comparisons anchor to the
# timezone that task.md declares for the event's wall-clock semantics.
# Never use bare `sd.hour` / `sd.date()` on rows read from `gcal.events`:
# psycopg2 returns them in the PG session `TimeZone` (case-study: compute
# node default was Asia/Shanghai). Use `utils.evaluation.gcal_helpers`
# instead (session-tz-independent).
# ──────────────────────────────────────────────────────────────────────────
# task.md: "Schedule them at 10:00, 11:00, and 14:00 respectively, each
# lasting one hour" on 2026-04-01. Canvas default timezone is ET
# (America/New_York), so the wall-clock hours anchor to America/New_York.
EXPECTED_TIMEZONE = ZoneInfo("America/New_York")

from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def _get_top3_courses_from_db():
    """Compute the top-3 courses by late submission rate from canvas DB.
    Falls back to known names if DB unreachable."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT c.name,
                   COUNT(*) AS total_subs,
                   SUM(CASE WHEN s.late THEN 1 ELSE 0 END) AS late_subs,
                   ROUND(100.0 * SUM(CASE WHEN s.late THEN 1 ELSE 0 END)
                         / NULLIF(COUNT(*), 0), 1) AS rate_pct
            FROM canvas.courses c
            JOIN canvas.assignments a ON a.course_id = c.id
            JOIN canvas.submissions s ON s.assignment_id = a.id
            GROUP BY c.name
            HAVING COUNT(*) > 0
            ORDER BY rate_pct DESC, c.name ASC
            LIMIT 3
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        if rows:
            return [r[0] for r in rows], rows
    except (psycopg2.OperationalError, Exception):
        pass
    fallback = [
        "Creative Computing & Culture (Spring 2014)",
        "Creative Computing & Culture (Fall 2014)",
        "Data-Driven Design (Spring 2013)",
    ]
    return fallback, []


TOP3_COURSES, _TOP3_ROWS = _get_top3_courses_from_db()

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Late_Submission_Report.docx")
    if not os.path.isfile(docx_path):
        check("Late_Submission_Report.docx exists", False, f"Not found: {docx_path}")
        return
    check("Late_Submission_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        para_text = " ".join(p.text for p in doc.paragraphs).lower()
        cell_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text += " " + cell.text.lower()
        all_text = para_text + " " + cell_text

        check("Word doc has meaningful content (>= 100 chars)",
              len(all_text.strip()) >= 100,
              f"Content length: {len(all_text)}")
        check("Word doc contains 'late submission'",
              "late submission" in all_text,
              f"Sample: {all_text[:300]}")

        # Require a table with the four required columns
        required_headers = ["course_name", "total_submissions", "late_submissions",
                            "late_rate_pct"]
        found_table = False
        target_table = None
        for table in doc.tables:
            if not table.rows:
                continue
            header_cells = [str(c.text or "").strip().lower().replace(" ", "_")
                            for c in table.rows[0].cells]
            joined = "|".join(header_cells)
            if all(h.replace("_", "") in joined.replace("_", "") for h in required_headers):
                found_table = True
                target_table = table
                break
        check("Word doc has table with Course_Name/Total_Submissions/Late_Submissions/Late_Rate_Pct headers",
              found_table,
              "Required table not found")

        # Validate top-3 courses appear in table rows with correct course names
        top3_in_table_count = 0
        if target_table:
            table_rows = []
            for row in list(target_table.rows)[1:]:
                table_rows.append([str(c.text or "").strip() for c in row.cells])
            for course in TOP3_COURSES:
                if any(course.lower() in " ".join(r).lower() for r in table_rows):
                    top3_in_table_count += 1
        check(f"Word doc table includes all top-3 highest-late-rate courses",
              top3_in_table_count == 3,
              f"Found {top3_in_table_count}/3 top courses in table")

        # Validate Late_Rate values for top-3 (within 1.0pp tolerance)
        if _TOP3_ROWS and target_table:
            rate_col = -1
            header_cells = [str(c.text or "").strip().lower().replace(" ", "_")
                            for c in target_table.rows[0].cells]
            for i, h in enumerate(header_cells):
                if "rate" in h:
                    rate_col = i
                    break
            if rate_col >= 0:
                course_to_gt_rate = {r[0]: float(r[3]) for r in _TOP3_ROWS}
                for row in list(target_table.rows)[1:]:
                    cells = [str(c.text or "").strip() for c in row.cells]
                    name = cells[0] if cells else ""
                    if name in course_to_gt_rate:
                        try:
                            v = float(str(cells[rate_col]).rstrip("%").strip())
                            gt = course_to_gt_rate[name]
                            check(f"Late_Rate_Pct for '{name}' matches GT ({gt})",
                                  abs(v - gt) <= 1.0,
                                  f"got {v}, expected ~{gt}")
                        except (ValueError, IndexError):
                            check(f"Late_Rate_Pct for '{name}' parseable", False,
                                  f"row={cells}")

        check("Word doc has recommendations section",
              "recommend" in all_text or "policy" in all_text or "intervention" in all_text,
              f"Sample: {all_text[:300]}")
    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, description, start_datetime, end_datetime
            FROM gcal.events
            WHERE LOWER(summary) LIKE '%%late submission review%%'
        """)
        events = cur.fetchall()
        check("At least 3 'Late Submission Review' events scheduled",
              len(events) >= 3,
              f"Found {len(events)} events")

        # Filter to events on 2026-04-01 (ET). gcal.events.start_datetime is
        # TIMESTAMPTZ (UTC instant); extract the wall-clock date in
        # EXPECTED_TIMEZONE via get_zone_components (session-tz-independent).
        on_target_date = []
        for e in events:
            if e[2] is None:
                continue
            ev_date, _, _ = get_zone_components(e[2], EXPECTED_TIMEZONE)
            if ev_date is not None and str(ev_date) == "2026-04-01":
                on_target_date.append(e)
        check("Events scheduled on 2026-04-01",
              len(on_target_date) >= 3,
              f"{len(on_target_date)} events on 2026-04-01 out of {len(events)} total")

        # Required slot times (hours, ET): 10, 11, 14
        required_hours = {10, 11, 14}
        observed_hours = set()
        for e in on_target_date:
            try:
                # session-tz-independent hour in EXPECTED_TIMEZONE (ET).
                _, ev_hour, _ = get_zone_components(e[2], EXPECTED_TIMEZONE)
                if ev_hour is not None:
                    observed_hours.add(ev_hour)
            except (AttributeError, TypeError):
                # may be a string; try parse
                s = str(e[2])
                if "10:00" in s:
                    observed_hours.add(10)
                elif "11:00" in s:
                    observed_hours.add(11)
                elif "14:00" in s:
                    observed_hours.add(14)
        missing = required_hours - observed_hours
        check("Events scheduled at 10:00, 11:00, and 14:00 on 2026-04-01",
              not missing,
              f"Missing hour slots: {sorted(missing)}")

        # Each event should last 1 hour
        if on_target_date:
            for e in on_target_date:
                if e[3] is None or e[2] is None:
                    continue
                try:
                    delta = (e[3] - e[2]).total_seconds() / 60.0
                    check(f"Event '{e[0]}' duration is 60 minutes",
                          abs(delta - 60) < 5,
                          f"got {delta} min")
                except Exception:
                    pass

        # Each top-3 course should appear in some event description
        for course in TOP3_COURSES:
            found = any(
                course.lower() in (str(e[1] or "") + " " + str(e[0] or "")).lower()
                for e in on_target_date
            )
            check(f"Calendar event mentions course '{course}'",
                  found,
                  f"Course '{course}' missing from event summaries/descriptions")

        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar check", False, str(e))


def check_email():
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # AND on recipient and subject (no OR)
        cur.execute("""
            SELECT id, subject, to_addr, body_text
            FROM email.messages
            WHERE to_addr::text ILIKE '%%registrar@university.edu%%'
              AND LOWER(subject) = LOWER(%s)
        """, ("Late Submission Analysis Report",))
        emails = cur.fetchall()
        check("Email to registrar@university.edu with subject 'Late Submission Analysis Report'",
              len(emails) >= 1,
              "No exact-subject email to registrar found")
        if emails:
            email = emails[0]
            body = str(email[3] or "").lower()
            check("Email body has substantial content (>= 30 chars)",
                  len(body) > 30,
                  f"Body length: {len(body)}")
            # Body should mention top late submission offenders (course names)
            mentions = sum(1 for c in TOP3_COURSES if c.lower() in body)
            check("Email body mentions top late-submission offender courses (>=2 of 3)",
                  mentions >= 2,
                  f"Mentioned {mentions}/3 top courses in body")
        cur.close()
        conn.close()
    except Exception as e:
        check("Email check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_calendar()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ===")
    if FAIL_COUNT > 0:
        print(f"{FAIL_COUNT} checks failed")
        sys.exit(1)
    else:
        print("All checks passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
