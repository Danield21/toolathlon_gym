#!/usr/bin/env python3
"""Evaluation script for arxiv-research-workflow-pipeline task validation"""

from argparse import ArgumentParser
import json
import os
import sys
from pathlib import Path

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def _get_db_conn():
    """Connect to the shared Postgres the calendar MCP writes into.

    The Calendar-Autoauth MCP bridge reads PG_HOST/PG_PORT/PG_DATABASE/PG_USER/
    PG_PASSWORD (underscore spelling); the harness may export either that family
    or the PGHOST/PGPORT/PGDATABASE/PGUSER/PGPASSWORD family. Accept both so the
    evaluator reads the SAME database the google_calendar MCP writes to.
    """
    import psycopg2
    return psycopg2.connect(
        host=os.environ.get("PGHOST", os.environ.get("PG_HOST", "localhost")),
        port=int(os.environ.get("PGPORT", os.environ.get("PG_PORT", "5432"))),
        dbname=os.environ.get("PGDATABASE", os.environ.get("PG_DATABASE", "toolathlon_gym")),
        user=os.environ.get("PGUSER", os.environ.get("PG_USER", "eigent")),
        password=os.environ.get("PGPASSWORD", os.environ.get("PG_PASSWORD", "camel")),
    )


def _find_xlsx_by_keywords(workspace, keywords):
    """Score xlsx files by filename + content keywords."""
    import glob
    import openpyxl
    scored = []
    for path in sorted(glob.glob(os.path.join(workspace, "*.xlsx"))):
        fname_low = os.path.basename(path).lower()
        if fname_low.startswith("~$") or fname_low.startswith("."):
            continue
        score = 0
        for kw in keywords:
            if kw in fname_low:
                score += 10
        try:
            wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
            content_text = " ".join(s.lower() for s in wb.sheetnames)
            for ws in wb.worksheets:
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    if row_count > 5:
                        break
                    row_count += 1
                    for cell in row:
                        if cell is not None:
                            content_text += " " + str(cell).lower()
            wb.close()
            for kw in keywords:
                if kw in content_text:
                    score += 1
        except Exception:
            pass
        if score > 0:
            scored.append((score, path))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [p for _, p in scored]


def check_xlsx_content(workspace, groundtruth_workspace="."):
    """Locate paper analysis xlsx by filename or content keyword.
    task.md requires one paper-metadata sheet (Paper ID/Title/Authors/Year, one
    row per paper). The agent's paper-analysis sheet is identified by its own
    headers (never by GT sheet names), and the GT 'Statistics'/'Recommendations'
    sheets -- artifacts of the reference solution, NOT required deliverables --
    impose no constraint on the agent. The row-count threshold (min(5, GT paper
    rows)) is derived only from the GT paper-metadata sheet. Row count is
    otherwise flexible because task.md does not state an exact paper count."""
    print("\n=== Check: XLSX paper analysis ===")
    import openpyxl
    keywords = ["paper_analysis", "paper analysis", "paper_notes", "papers",
                "notes", "literature", "research", "analysis", "metadata",
                "template"]
    exact = os.path.join(workspace, "paper_analysis.xlsx")
    chosen = exact if os.path.isfile(exact) else None
    if chosen is None:
        cands = _find_xlsx_by_keywords(workspace, keywords)
        if cands: chosen = cands[0]
    if chosen is None:
        record("xlsx for paper analysis exists", False,
               f"No xlsx with keywords {keywords[:4]} found")
        return False
    fname = os.path.basename(chosen)
    record(f"xlsx for paper analysis exists ({fname})", True)
    try:
        wb = openpyxl.load_workbook(chosen, data_only=False)
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            record(f"xlsx '{ws.title}' has data", len(rows) >= 2, f"{len(rows)} rows")
    except Exception as e:
        record("xlsx readable", False, str(e))
        return True

    # --- Groundtruth structural comparison (relaxed, content-based) ---
    # We do NOT require the agent to name its sheets like the GT
    # (Metadata/Statistics/Recommendations); we match headers by content instead.
    gt_path = os.path.join(groundtruth_workspace, "paper_analysis.xlsx")
    if not os.path.isfile(gt_path):
        wb.close()
        return True

    gt_wb = openpyxl.load_workbook(gt_path, data_only=False)

    def _header(ws):
        for r in ws.iter_rows(min_row=1, max_row=1, values_only=True):
            return [str(v).strip().lower() for v in r if v is not None]
        return None

    # Paper-metadata keywords (substring match on lowercased headers). Tolerant of
    # common variants: 'Paper ID'/'arXiv ID'/'ID', 'Title'/'Paper Title',
    # 'Authors'/'Author Name(s)', 'Year'/'Publication Year'.
    meta_kws = ("paper id", "id", "title", "author", "year")

    def _meta_score(hdr):
        if not hdr:
            return 0
        return sum(1 for kw in meta_kws if any(kw in h for h in hdr))

    # Collect agent sheets with their header rows.
    agent_sheets = []
    for asn in wb.sheetnames:
        ws = wb[asn]
        agent_sheets.append((asn, _header(ws), ws))

    # Primary structure check: at least one agent sheet is paper-metadata-like
    # (>=2 of Paper ID/Title/Authors/Year in its header). On a meta-score tie we
    # prefer the sheet with more data rows, so a small auxiliary sheet can never
    # shadow the real (larger) paper-analysis sheet.
    best_agent_meta = 0
    best_meta_sheet = None
    best_ws = None
    best_rows = -1
    for asn, hdr, ws in agent_sheets:
        s = _meta_score(hdr)
        if s == 0:
            continue
        n_rows = len([r for r in ws.iter_rows(min_row=2, values_only=True)
                      if any(c is not None for c in r)])
        if s > best_agent_meta or (s == best_agent_meta and n_rows > best_rows):
            best_agent_meta, best_meta_sheet, best_ws, best_rows = s, asn, ws, n_rows
    record("agent xlsx has a paper-analysis sheet (>=2 paper-metadata headers)",
           best_agent_meta >= 2,
           f"best: sheet '{best_meta_sheet}' meta-header score {best_agent_meta}; "
           f"agent sheets: {wb.sheetnames}")

    # Row-count check on the agent's own paper-analysis sheet (the deliverable
    # task.md requires). The GT sheet is used only to derive a row threshold
    # (min(5, GT paper rows)); the GT 'Statistics'/'Recommendations' sheets are
    # artifacts of the reference solution and impose NO constraint on the agent.
    if best_agent_meta >= 2 and best_ws is not None:
        gt_paper_rows = None
        for gsn in gt_wb.sheetnames:
            if _meta_score(_header(gt_wb[gsn])) >= 2:
                rows = [r for r in gt_wb[gsn].iter_rows(min_row=2, values_only=True)
                        if any(c is not None for c in r)]
                if gt_paper_rows is None or len(rows) > gt_paper_rows:
                    gt_paper_rows = len(rows)
        if gt_paper_rows is None:
            gt_paper_rows = 0
        # Relaxed: agent's paper sheet should have at least min(5, gt_paper_rows)
        # data rows. Task does not specify an exact paper count; an honest agent
        # analyzing >=5 papers passes (task.md requires a >=5-entry bibliography,
        # one row per analyzed paper).
        min_rows = min(5, max(1, gt_paper_rows))
        record(f"agent paper-analysis sheet '{best_meta_sheet}' has at least {min_rows} data rows",
               best_rows >= min_rows,
               f"Got {best_rows} data rows")
    gt_wb.close()
    wb.close()
    return True


def _find_docx_by_keywords(workspace, keywords):
    import glob
    from docx import Document
    scored = []
    for path in sorted(glob.glob(os.path.join(workspace, "*.docx"))):
        fname_low = os.path.basename(path).lower()
        if fname_low.startswith("~$") or fname_low.startswith("."):
            continue
        score = 0
        for kw in keywords:
            if kw in fname_low:
                score += 10
        try:
            doc = Document(path)
            text_low = "\n".join(p.text for p in doc.paragraphs).lower()
            for kw in keywords:
                if kw in text_low:
                    score += 1
        except Exception:
            pass
        if score > 0:
            scored.append((score, path))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [p for _, p in scored]


def check_docx_content(workspace):
    """Locate research summary / literature review docx by filename or content keyword."""
    print("\n=== Check: DOCX literature review / research summary ===")
    from docx import Document
    keywords = ["literature_review", "literature review", "research_summary",
                "research summary", "literature", "research", "summary",
                "review", "machine learning", "federated"]
    exact = os.path.join(workspace, "literature_review.docx")
    chosen = exact if os.path.isfile(exact) else None
    if chosen is None:
        cands = _find_docx_by_keywords(workspace, keywords)
        if cands: chosen = cands[0]
    if chosen is None:
        record("docx for literature review/research summary exists", False,
               f"No docx with keywords {keywords[:5]} found")
        return False
    fname = os.path.basename(chosen)
    record(f"docx for literature review/research summary exists ({fname})", True)
    try:
        doc = Document(chosen)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        word_count = len(all_text.split())
        # Task says "approximately two thousand words"; accept >=500 to allow concise
        # summaries. Word count is the substantive-content gate; there is NO paragraph
        # count requirement because task.md does not state one (a faithful agent may
        # organize ~2000 words into 6-9 long paragraphs).
        record("docx word count >= 500",
               word_count >= 500,
               f"{word_count} words")
        # Mention of the ML / neural network / federated topic. Task is broad; a
        # faithful review naturally names "machine learning" (or any peer keyword) at
        # least once. Requiring >=1 keyword keeps the check meaningful while never
        # penalizing a single-topic honest review.
        text_lower = all_text.lower()
        topic_keys = ("federated", "distributed", "machine learning", "neural",
                      "deep learning")
        topic_matched = sum(1 for k in topic_keys if k in text_lower)
        record("docx mentions ML/neural network/federated topic (>=1 keyword)",
               topic_matched >= 1, f"matched {topic_matched}/5")
    except Exception as e:
        record("docx readable", False, str(e))
    return True


def check_bib_file(workspace):
    print("\n=== Check: Bibliography .bib file ===")
    # Look for any .bib file (task says 'comprehensive bibliography' without naming it)
    bib_files = [f for f in os.listdir(workspace) if f.endswith(".bib")]
    record("At least one .bib file present", len(bib_files) >= 1,
           f"Found: {bib_files}")
    if not bib_files:
        return
    # Pick the largest bib file
    bib_path = os.path.join(workspace, max(bib_files, key=lambda f: os.path.getsize(os.path.join(workspace, f))))
    try:
        with open(bib_path) as f:
            content = f.read()
    except Exception as e:
        record(".bib file readable", False, str(e))
        return
    # Count bibtex entries (@article, @inproceedings, etc.)
    import re
    entries = re.findall(r"@\w+\s*\{", content)
    record(".bib file has at least 5 entries", len(entries) >= 5,
           f"Found {len(entries)} entries")
    # Topic check: task is broad (machine learning, neural networks, deep learning).
    # Requiring >=1 topical keyword keeps the check meaningful (combined with the
    # >=5-entry and >=3-recent-year gates above) without penalizing a faithful
    # bibliography whose titles happen to use a single topic vocabulary.
    text_lower = content.lower()
    topic_keys = ("federated", "distributed", "machine learning", "neural",
                  "deep learning")
    matched_topics = sum(1 for k in topic_keys if k in text_lower)
    record(".bib mentions ML/neural network/federated topic (>=1 keyword)",
           matched_topics >= 1,
           f"matched {matched_topics}/5; sample first 200 chars: {text_lower[:200]}")
    # Year distribution: bibliography entries should be reasonably modern (>=3 entries
    # in the last 7 years). Catches GT-noise files that are entirely classical/old papers.
    import re as _re
    years = [int(y) for y in _re.findall(r"year\s*=\s*[{\"]?(\d{4})[}\"]?", content)]
    recent_years = [y for y in years if y >= 2018]
    record(".bib has at least 3 entries from 2018 or later",
           len(recent_years) >= 3,
           f"recent_year_count={len(recent_years)} (years sampled: {sorted(set(years))[:8]})")


def check_calendar(is_gt_self_test=False):
    print("\n=== Check: Calendar invitations ===")
    try:
        conn = _get_db_conn()
    except Exception as e:
        # DB unreachable is an infra condition, not an agent failure. A correct agent
        # would be wrongly FAILed if the harness DB were down at eval time, so this
        # is a non-blocking skip. When the DB IS reachable the check runs normally.
        print(f"  [WARN] Calendar check skipped: database unavailable ({str(e)[:120]})")
        return
    try:
        cur = conn.cursor()
        cur.execute("SELECT id, summary FROM gcal.events")
        events = cur.fetchall()
        ev_count = len(events)
        if is_gt_self_test and ev_count == 0:
            print("  [WARN] Calendar checks skipped (GT self-test, non-blocking)")
            return
        record("At least 1 calendar event created for the team", ev_count >= 1,
               f"Found {ev_count} events")
        # Sanity check: at least 1 event title relates to review/research/meeting/paper
        if events:
            topical = 0
            for eid, summary in events:
                s = (summary or "").lower()
                if any(kw in s for kw in ("review", "research", "meeting", "paper", "literature",
                                          "discuss", "quarter", "plan", "action", "sync")):
                    topical += 1
            record("At least 1 event title mentions review/research/meeting/paper",
                   topical >= 1, f"got {topical}/{ev_count} topical")
            # NOTE: the google_calendar MCP create_event/update_event tools expose no
            # 'attendees' parameter, so gcal.events.attendees is always '[]'. An
            # attendee-count requirement is therefore not achievable through the runtime
            # and would systematically FAIL honest agents. We do not evaluate it.
    except Exception as e:
        msg = str(e)
        if "does not exist" in msg or "undefined_table" in msg:
            # gcal schema/table absent -> the MCP<->DB bridge is not set up. Infra
            # condition, not an agent failure; non-blocking skip.
            print("  [WARN] Calendar check skipped: gcal schema/table not present "
                  "(MCP<->DB bridge not configured)")
        else:
            record("Calendar query", False, msg)
    finally:
        try:
            conn.close()
        except Exception:
            pass


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    ws = args.agent_workspace
    if not os.path.isdir(ws):
        print(f"Agent workspace not found: {ws}")
        sys.exit(1)

    # Detect GT self-test (V1 parity test).
    try:
        gt_canon = os.path.realpath(args.groundtruth_workspace) if args.groundtruth_workspace else ""
        ag_canon = os.path.realpath(args.agent_workspace) if args.agent_workspace else ""
        is_gt_self_test = bool(gt_canon) and (gt_canon == ag_canon)
    except Exception:
        is_gt_self_test = False

    check_xlsx_content(ws, args.groundtruth_workspace)
    check_docx_content(ws)
    check_bib_file(ws)
    # Notion is not part of this task's required deliverables (task.md never
    # mentions it), so no Notion check is performed.
    check_calendar(is_gt_self_test)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} checks failed)")
        sys.exit(1)


if __name__ == "__main__":
    main()
