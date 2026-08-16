"""Evaluation for terminal-yf-sf-gsheet-word-gcal.
Checks:
1. Compensation_Review_Memo.docx content
2. Google Sheet "Sales Compensation Analysis" with 3 sheets
3. Google Calendar event "Q4 Compensation Review Meeting"
4. Script files exist (compute_bonuses.py, market_adjustment.py, validate_bonuses.py)
5. JSON output files with correct values
"""
import argparse
import json
import os
import sys
from zoneinfo import ZoneInfo

import psycopg2

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# Source of truth: docs/task.md. The Calendar MCP writes events as UTC
# instants into the DB (`TIMESTAMPTZ`), so all comparisons anchor to the
# timezone that task.md declares for the event's wall-clock semantics.
# Never use bare `start_dt.day` / `start_dt.month` on rows read from
# `gcal.events`: psycopg2 returns them in the PG session `TimeZone`
# (case-study: compute node default was Asia/Shanghai, shifting the
# wall-clock date). Use `utils.evaluation.gcal_helpers` instead
# (session-tz-independent).
# ──────────────────────────────────────────────────────────────────────────
# task.md line 13: "the week of March 9 to March 13, 2026 between 9 AM and
# 5 PM Eastern Time" (yf financial, ET).
EXPECTED_TIMEZONE = ZoneInfo("America/New_York")

# Evaluator runs as `python -m tasks.finalpool.<task>.evaluation.main` with
# cwd = /workspace (toolathlon_gym root), so `utils/` is importable directly.
from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def num_close(a, b, tol=2.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except Exception:
        return False


def pct_close(a, b, tol=1.0):
    """Check if two percentage values are close."""
    try:
        return abs(float(a) - float(b)) <= tol
    except Exception:
        return False


def check_word(workspace):
    print("\n=== Check 1: Compensation_Review_Memo.docx ===")
    path = os.path.join(workspace, "Compensation_Review_Memo.docx")
    if not os.path.exists(path):
        check("Word file exists", False, f"Not found at {path}")
        return
    check("Word file exists", True)

    try:
        from docx import Document
        doc = Document(path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()

        # Collect heading-style paragraphs (style.name starting with 'Heading')
        heading_texts = []
        for p in doc.paragraphs:
            sname = (p.style.name if p.style else "") or ""
            if sname.lower().startswith("heading"):
                heading_texts.append(p.text.strip().lower())
        headings_blob = " | ".join(heading_texts)

        check("Has title 'Q4 Compensation Review'",
              "q4 compensation review" in all_text or "compensation review memo" in all_text,
              f"Text snippet: {all_text[:100]}")
        # Tighten: section detection requires a Heading-styled paragraph containing the keyword
        check("Has Background heading",
              any("background" in h for h in heading_texts),
              f"Headings: {headings_blob[:200]}")
        check("Has Methodology heading",
              any("methodology" in h for h in heading_texts),
              f"Headings: {headings_blob[:200]}")
        check("Has Market Analysis heading",
              any("market" in h for h in heading_texts),
              f"Headings: {headings_blob[:200]}")
        check("Has Regional Performance heading",
              any(("regional" in h) or ("region" in h) for h in heading_texts),
              f"Headings: {headings_blob[:200]}")
        check("Has Budget Impact heading",
              any("budget" in h for h in heading_texts),
              f"Headings: {headings_blob[:200]}")
        check("Has Recommendations heading",
              any("recommend" in h for h in heading_texts),
              f"Headings: {headings_blob[:200]}")
        check("Mentions DJI or Dow Jones",
              "dji" in all_text or "dow jones" in all_text or "dow" in all_text,
              "No DJI/Dow Jones reference")
        check("Mentions adjustment factor or 0.9",
              "0.9" in all_text or "adjustment" in all_text,
              "No adjustment factor reference")
        check("Mentions budget cap",
              "budget cap" in all_text or "30,000,000" in all_text or "30000000" in all_text or "30 million" in all_text,
              "No budget cap reference")
    except Exception as e:
        check("Word readable", False, str(e))


def check_gsheet():
    print("\n=== Check 2: Google Sheet 'Sales Compensation Analysis' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        cur.execute("SELECT id, title FROM gsheet.spreadsheets WHERE lower(title) LIKE '%sales compensation%' OR lower(title) LIKE '%compensation analysis%'")
        rows = cur.fetchall()
        check("Spreadsheet exists", len(rows) >= 1, f"Found {len(rows)} matching spreadsheets")
        if not rows:
            return

        ss_id = rows[0][0]

        # Check sheets
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s ORDER BY index", (ss_id,))
        sheets = cur.fetchall()
        sheet_titles = [s[1].lower() for s in sheets]
        check("Has at least 3 sheets", len(sheets) >= 3, f"Found {len(sheets)} sheets: {sheet_titles}")

        # Tighten: require sheet names exactly matching Rep_Performance, Market_Adjustment, Adjusted_Bonuses
        # (case-insensitive, underscore/space tolerant)
        def _norm(s):
            return (s or "").lower().replace(" ", "_").replace("-", "_")
        norm_titles = [_norm(t) for t in sheet_titles]
        has_rep = any(t == "rep_performance" for t in norm_titles)
        has_market = any(t == "market_adjustment" for t in norm_titles)
        has_adjusted = any(t == "adjusted_bonuses" for t in norm_titles)
        check("Has Rep_Performance sheet (exact name)", has_rep, f"Sheets: {sheet_titles}")
        check("Has Market_Adjustment sheet (exact name)", has_market, f"Sheets: {sheet_titles}")
        check("Has Adjusted_Bonuses sheet (exact name)", has_adjusted, f"Sheets: {sheet_titles}")

        # Check Market_Adjustment sheet has DJI data
        for sid, title in sheets:
            if _norm(title) == "market_adjustment":
                cur.execute("""SELECT value FROM gsheet.cells
                    WHERE spreadsheet_id = %s AND sheet_id = %s""", (ss_id, sid))
                values = [r[0].lower() if r[0] else "" for r in cur.fetchall()]
                all_vals = " ".join(values)
                check("Market sheet has DJI", "dji" in all_vals or "^dji" in all_vals,
                      f"Values: {all_vals[:200]}")
                check("Market sheet has AMZN", "amzn" in all_vals,
                      f"Values: {all_vals[:200]}")
                check("Market sheet has XOM", "xom" in all_vals,
                      f"Values: {all_vals[:200]}")
                break

        # Check Rep_Performance sheet has data rows
        for sid, title in sheets:
            if _norm(title) == "rep_performance":
                cur.execute("""SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                    WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0""", (ss_id, sid))
                data_rows = cur.fetchone()[0]
                check("Rep sheet has data rows", data_rows >= 1, f"Found {data_rows} data rows")
                break

    except Exception as e:
        check("GSheet query", False, str(e))
    finally:
        cur.close()
        conn.close()


def check_gcal():
    print("\n=== Check 3: Calendar Event ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        cur.execute("""SELECT summary, description, start_datetime, end_datetime
            FROM gcal.events
            WHERE lower(summary) LIKE '%compensation%' OR lower(summary) LIKE '%q4%bonus%'""")
        events = cur.fetchall()
        check("Compensation meeting scheduled", len(events) >= 1, f"Found {len(events)} matching events")

        if events:
            evt = events[0]
            t = evt[0].lower()
            # Tighten: title must be exactly "Q4 Compensation Review Meeting"
            # (case-insensitive, whitespace-normalized) per task.md.
            t_norm = " ".join(t.split())
            check("Event title is exactly 'Q4 Compensation Review Meeting'",
                  t_norm == "q4 compensation review meeting",
                  f"Title: {evt[0]}")

            # Check it's in the March 9-13 week
            start_dt = evt[2]
            if start_dt:
                # gcal.events.start_datetime is TIMESTAMPTZ (UTC instant).
                # Extract components in EXPECTED_TIMEZONE so day/month are
                # session-tz-independent (case-study 2026-08-13). Bare
                # start_dt.day / start_dt.month silently compares against
                # the PG session tz.
                ev_date, _ev_hour, _ev_minute = get_zone_components(
                    start_dt, EXPECTED_TIMEZONE)
                if ev_date is not None:
                    check("Event in March 9-13 week",
                          ev_date.month == 3 and 9 <= ev_date.day <= 13,
                          f"Start: {start_dt}")

                # Check 90-minute duration
                if evt[3]:
                    duration = (evt[3] - start_dt).total_seconds() / 60
                    check("Event is 90 minutes", 85 <= duration <= 95, f"Duration: {duration} min")

                # Check no conflict with existing events
                try:
                    cur.execute("""SELECT summary, start_datetime, end_datetime FROM gcal.events
                        WHERE lower(summary) NOT LIKE '%%compensation%%'
                        AND start_datetime < %s AND end_datetime > %s""",
                        (evt[3], evt[2]))
                    conflicts = cur.fetchall()
                    check("No conflicts with existing events", len(conflicts) == 0,
                          f"Conflicts: {[(c[0], str(c[1])) for c in conflicts]}")
                except Exception as e2:
                    check("No conflicts with existing events", False, str(e2))
    except Exception as e:
        check("GCal query", False, str(e))
    finally:
        cur.close()
        conn.close()


def check_scripts(workspace):
    print("\n=== Check 4: Script Files ===")
    for script in ["compute_bonuses.py", "market_adjustment.py", "validate_bonuses.py"]:
        path = os.path.join(workspace, script)
        check(f"{script} exists", os.path.exists(path))


def _get_sales_count():
    """Pull number of Sales department employees from HR table."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        with conn.cursor() as cur:
            cur.execute(
                'SELECT COUNT(*) FROM sf_data."HR_ANALYTICS__PUBLIC__EMPLOYEES" '
                'WHERE "DEPARTMENT" = %s',
                ('Sales',),
            )
            n = cur.fetchone()[0]
        conn.close()
        return int(n)
    except Exception:
        return None


def check_json_outputs(workspace):
    print("\n=== Check 5: JSON Output Files ===")
    sales_n = _get_sales_count()
    # Allow ±2 tolerance for off-by-one boundaries; otherwise require exact count.
    expected_min = (sales_n - 2) if sales_n is not None else 100

    # current_bonuses.json
    cb_path = os.path.join(workspace, "current_bonuses.json")
    cb_data = None
    if not os.path.exists(cb_path):
        check("current_bonuses.json exists", False)
    else:
        check("current_bonuses.json exists", True)
        try:
            with open(cb_path) as f:
                cb_data = json.load(f)
            if isinstance(cb_data, list):
                check(
                    f"current_bonuses entry count >= Sales HR count ({expected_min})",
                    len(cb_data) >= expected_min,
                    f"Found {len(cb_data)} entries (expected >= {expected_min})",
                )
                # Check structure of first entry: tighten to use dict.keys()
                if cb_data and isinstance(cb_data[0], dict):
                    first = cb_data[0]
                    keys_lower = {k.lower() for k in first.keys()}
                    has_name = any("name" in k for k in keys_lower)
                    has_region = any("region" in k for k in keys_lower)
                    has_salary = any("salary" in k or "compensation" in k for k in keys_lower)
                    has_bonus = any("bonus" in k for k in keys_lower)
                    check("current_bonuses has 'name' field", has_name, f"Keys: {list(first.keys())}")
                    check("current_bonuses has 'region' field", has_region, f"Keys: {list(first.keys())}")
                    check("current_bonuses has 'salary' field", has_salary, f"Keys: {list(first.keys())}")
                    check("current_bonuses has 'bonus' field", has_bonus, f"Keys: {list(first.keys())}")
            elif isinstance(cb_data, dict):
                check("current_bonuses has entries", len(cb_data) > 0, f"Found dict with {len(cb_data)} keys")
        except Exception as e:
            check("current_bonuses.json valid JSON", False, str(e))

    # market_adjusted_bonuses.json
    mab_path = os.path.join(workspace, "market_adjusted_bonuses.json")
    mab_data = None
    if not os.path.exists(mab_path):
        check("market_adjusted_bonuses.json exists", False)
    else:
        check("market_adjusted_bonuses.json exists", True)
        try:
            with open(mab_path) as f:
                mab_data = json.load(f)
            if isinstance(mab_data, list):
                check(
                    f"adjusted_bonuses entry count >= Sales HR count ({expected_min})",
                    len(mab_data) >= expected_min,
                    f"Found {len(mab_data)} entries (expected >= {expected_min})",
                )
                if mab_data and isinstance(mab_data[0], dict):
                    keys_lower = {k.lower() for k in mab_data[0].keys()}
                    has_adjusted = any("adjust" in k for k in keys_lower)
                    check("adjusted_bonuses has adjusted field", has_adjusted,
                          f"Keys: {list(mab_data[0].keys())}")
            elif isinstance(mab_data, dict):
                check("adjusted_bonuses has data", len(mab_data) > 0, f"Found dict with {len(mab_data)} keys")
        except Exception as e:
            check("market_adjusted_bonuses.json valid JSON", False, str(e))

    # Validate budget cap and 20% constraint
    if isinstance(mab_data, list) and mab_data and isinstance(mab_data[0], dict):
        # Try common bonus field names
        def _get_val(d, keys):
            for k in d.keys():
                if k.lower() in keys or any(target in k.lower() for target in keys):
                    try:
                        return float(d[k])
                    except (TypeError, ValueError):
                        continue
            return None

        total_bonus = 0.0
        violations_20 = 0
        rows_with_bonus = 0
        for entry in mab_data:
            # Restrict to specific adjusted-bonus key names (drop generic 'bonus' fallback)
            adj = _get_val(entry, ["adjusted_bonus", "adjusted", "final_bonus", "scaled_bonus"])
            sal = _get_val(entry, ["salary", "annual_salary", "base_salary", "compensation"])
            if adj is not None:
                total_bonus += adj
                rows_with_bonus += 1
            if adj is not None and sal is not None and sal > 0 and adj > sal * 0.20 + 0.01:
                violations_20 += 1

        if rows_with_bonus > 0:
            # Budget cap: 30,000,000
            check("Total adjusted bonuses <= budget cap (30,000,000)",
                  total_bonus <= 30000000 + 1.0,
                  f"Total bonus: {total_bonus:.2f}")
            check("No bonus exceeds 20% of salary",
                  violations_20 == 0,
                  f"Found {violations_20} violations out of {rows_with_bonus} entries")


def check_reverse_validation():
    """Check no duplicate gcal events or wrong event types."""
    print("\n=== Reverse Validation ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        # Check no duplicate compensation review events
        cur.execute("""
            SELECT summary, start_datetime FROM gcal.events
            WHERE lower(summary) LIKE '%compensation%review%'
        """)
        comp_events = cur.fetchall()
        check("No duplicate Compensation Review events",
              len(comp_events) <= 1,
              f"Found {len(comp_events)} compensation review events: {[e[0] for e in comp_events]}")

        # Check no non-meeting event types were created (e.g., no 'bonus payout' or 'salary update' calendar events)
        cur.execute("""
            SELECT summary FROM gcal.events
            WHERE lower(summary) LIKE '%bonus%payout%'
               OR lower(summary) LIKE '%salary%update%'
               OR lower(summary) LIKE '%pay%raise%'
        """)
        wrong_events = cur.fetchall()
        check("No wrong event types (bonus payout, salary update, pay raise)",
              len(wrong_events) == 0,
              f"Found unexpected events: {[e[0] for e in wrong_events]}")

        # Check pre-existing events were not deleted
        cur.execute("""
            SELECT COUNT(*) FROM gcal.events
            WHERE lower(summary) NOT LIKE '%compensation%'
              AND lower(summary) NOT LIKE '%q4%bonus%'
        """)
        other_count = cur.fetchone()[0]
        check("Pre-existing calendar events preserved (>= 5)",
              other_count >= 5,
              f"Found {other_count} non-compensation events (expected >= 5 from original 11)")

    except Exception as e:
        check("Reverse validation (gcal noise)", False, str(e))
    finally:
        cur.close()
        conn.close()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_gcal()
    check_scripts(args.agent_workspace)
    check_json_outputs(args.agent_workspace)
    check_reverse_validation()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # Use FAIL_COUNT == 0 (replaces 70% accuracy threshold which masked critical fails)
    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
