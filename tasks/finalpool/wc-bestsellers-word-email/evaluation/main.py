"""
Evaluation for wc-bestsellers-word-email task.

Checks:
1. Word document Bestsellers_Report.docx with correct top 10 products
2. Email sent to sales-leads@store.com ( DB check)
"""
import argparse
import sys
import os
from pathlib import Path

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def get_expected_top10():
    """Query PostgreSQL for top 10 products by total_sales."""
    import psycopg2

    conn = psycopg2.connect(
        host=os.environ.get("PGHOST", "localhost"), port=int(os.environ.get("PGPORT", "5432")), dbname="toolathlon_gym",
        user="eigent", password="camel"
    )
    cur = conn.cursor()
    cur.execute("""
        SELECT name, total_sales::int as total_sold, price::float as price
        FROM wc.products
        ORDER BY total_sales::int DESC
        LIMIT 10
    """)
    products = cur.fetchall()
    conn.close()
    return products


def check_word(workspace, expected):
    """Check Bestsellers_Report.docx for correctness."""
    from docx import Document

    print("\n=== Checking Word Document ===")
    doc_path = Path(workspace) / "Bestsellers_Report.docx"

    if not doc_path.exists():
        check("Word file exists", False, f"Not found: {doc_path}")
        return
    check("Word file exists", True)

    doc = Document(str(doc_path))

    # Check for title heading
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()

    check("Has heading with 'best' or 'bestsell'",
          any("best" in h.lower() for h in headings),
          f"Headings: {[h[:50] for h in headings]}")

    # Check tables
    if len(doc.tables) == 0:
        check("Has a table", False, "No tables found")
        return
    check("Has a table", True)

    table = doc.tables[0]
    rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Table has 10 data rows", len(rows) == 10, f"Got {len(rows)} rows")

    # Check all 10 product names match (by rank)
    mismatches = 0
    for i, (exp_name, exp_sold, exp_price) in enumerate(expected):
        if i >= len(rows):
            break
        row = rows[i]
        if len(row) > 1:
            actual_name = row[1]
            if actual_name[:30].lower() != exp_name[:30].lower():
                mismatches += 1
    check("All 10 product names match (within first 30 chars)",
          mismatches == 0,
          f"{mismatches} of {len(expected)} product names mismatch")

    # Legacy single-product check for backward compat
    if len(rows) >= 1:
        top_name_in_doc = rows[0][1] if len(rows[0]) > 1 else ""
        exp_top_name = expected[0][0]
        check("Top product name matches",
              top_name_in_doc[:30].lower() == exp_top_name[:30].lower(),
              f"Expected '{exp_top_name[:50]}', got '{top_name_in_doc[:50]}'")

    # Check total_sold and price values
    import re
    for i, (exp_name, exp_sold, exp_price) in enumerate(expected):
        if i >= len(rows):
            break
        row = rows[i]
        if len(row) >= 3:
            try:
                actual_sold = int(row[2])
                check(f"Rank {i+1} Total_Sold",
                      actual_sold == exp_sold,
                      f"Expected {exp_sold}, got {actual_sold}")
            except ValueError:
                check(f"Rank {i+1} Total_Sold", False, f"Cannot parse: {row[2]}")
        # Price (col 3 typically)
        if len(row) >= 4:
            price_str = str(row[3])
            # Extract first numeric value from the cell
            m = re.search(r"(-?\d+(?:\.\d+)?)", price_str.replace(",", ""))
            if m:
                try:
                    actual_price = float(m.group(1))
                    if abs(actual_price - exp_price) > 0.02:
                        check(f"Rank {i+1} Price", False,
                              f"Expected {exp_price:.2f}, got {actual_price:.2f}")
                    else:
                        check(f"Rank {i+1} Price", True)
                except ValueError:
                    check(f"Rank {i+1} Price", False, f"Cannot parse: {price_str}")

    # Check that rows are sorted by Total_Sold descending
    sold_vals = []
    for row in rows:
        if len(row) >= 3:
            try:
                sold_vals.append(int(row[2]))
            except ValueError:
                pass
    check("Rows sorted by Total_Sold descending",
          all(sold_vals[i] >= sold_vals[i+1] for i in range(len(sold_vals)-1)),
          f"Sold vals: {sold_vals}")

    # Check summary section: total + avg price
    total_sold_all = sum(p[1] for p in expected)
    check("Summary mentions total units",
          str(total_sold_all) in full_text,
          f"Expected {total_sold_all} in text")
    # Avg price across top 10
    avg_price = sum(p[2] for p in expected) / len(expected)
    # Accept avg price with 1 decimal precision
    avg_price_str_2dec = f"{avg_price:.2f}"
    avg_price_str_1dec = f"{avg_price:.1f}"
    check("Summary mentions avg price of top 10",
          avg_price_str_2dec in full_text or avg_price_str_1dec in full_text,
          f"Expected ~{avg_price_str_2dec} in text")


def check_email(expected):
    """Check email was sent ( DB check)."""
    import psycopg2

    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"), port=int(os.environ.get("PGPORT", "5432")), dbname="toolathlon_gym",
            user="eigent", password="camel"
        )
        cur = conn.cursor()
    except Exception as e:
        check("DB connection", False, str(e))
        return

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    target = "sales-leads@store.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        to_str = str(to_addr or "").lower()
        if target in to_str:
            found = (subj, from_addr, to_addr, body)
            break

    check(f"Email sent to {target}", found is not None,
          f"Found {len(all_emails)} total emails")

    if found:
        subj, _, _, body = found
        body_lower = (body or "").lower()
        subj_lower = (subj or "").lower()

        check("Email subject mentions bestsellers/top",
              "bestseller" in subj_lower or "top" in subj_lower or "best" in subj_lower,
              f"Subject: {(subj or '')[:100]}")

        top_name_short = expected[0][0][:20].lower()
        check("Email body mentions top product",
              top_name_short in body_lower,
              f"Expected '{top_name_short}' in body")

        check("Email body mentions top product units",
              str(expected[0][1]) in (body or ""),
              f"Expected '{expected[0][1]}' in body")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("Fetching expected data...")
    expected = get_expected_top10()
    print(f"  Top 10 products loaded")

    check_word(args.agent_workspace, expected)
    check_email(expected)

    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if all_ok:
        print("\nPass all tests!")
        sys.exit(0)
    else:
        print("\nSome checks failed.")
        sys.exit(1)
