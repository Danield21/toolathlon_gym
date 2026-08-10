"""Evaluation script for yf-sector-scholarly-excel-word.

Pure file-based evaluation (no DB access). Groundtruth values that are
open-ended financial metrics (Avg_Price, Total_Market_Value, Volatility_Score)
have no unique correct answer because the task does not fix a calculation
window/methodology, so they are validated structurally (present, non-empty,
numeric/formula) rather than compared cell-by-cell against the GT. Research
paper mapping is checked against the deterministic paper set injected by
preprocess/main.py, and Investment_Thesis / docx headings are validated for
structure, not against GT wording.
"""
import argparse
import json
import os
import re
import sys

import openpyxl

PASS_COUNT = 0
FAIL_COUNT = 0

# Papers injected by preprocess/main.py into scholarly.arxiv_papers (the only
# relevant papers a faithful scholarly search can find). Used to verify the
# agent actually searched and mapped real papers.
RELEVANT_PAPER_TITLES = [
    "sector rotation patterns",
    "industry momentum",
    "market cycles and defensive allocation",
    "industry analysis framework",
]

SECTOR_OUTLOOKS = {"bullish", "neutral", "bearish", "positive", "negative", "mixed"}
VALIDATION_STATUSES = {"confirmed", "partial", "rejected", "inconclusive",
                       "pending", "unverified"}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")


_CURRENCY_SYMS = ("$", "€", "£", "¥", "₹", "₩", "₽", "¢", "₺", "₿", "₪",
                  "%", "‰")


def safe_float(val):
    """Robustly parse a cell value to float, or None.

    Accepts numeric literals and strings carrying currency symbols (single
    letters and multi-char ones), currency-code prefixes ("US$", "USD",
    "CAD", "EUR", ...), percent signs, thousand separators or spaces.
    Formula cells (starting with '=') are not parsed to a number (they are
    handled via is_formula).  A last-resort regex extracts the first
    well-formed decimal number from the remaining string, so faithful
    models that format a metric with a currency prefix/suffix still parse.
    """
    if val is None:
        return None
    if isinstance(val, bool):
        return None
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if s == "" or s.startswith("="):
        return None
    s2 = s.replace(",", "").replace(" ", "")
    for sym in _CURRENCY_SYMS:
        s2 = s2.replace(sym, "")
    # Drop a leading currency-code prefix, optionally followed by a '$'
    # (e.g. "US$ 2,350,303,674,368", "USD 3.6e12", "CAD 1,234").
    s2 = re.sub(r"^[A-Za-z]{1,5}\$?", "", s2)
    try:
        return float(s2)
    except (ValueError, TypeError):
        pass
    # Last resort: extract the first well-formed decimal number in the string
    # (handles e.g. "2.5 million", trailing unit letters, embedded codes).
    m = re.search(r"[-+]?\d+(?:[.,]\d+)*(?:[eE][-+]?\d+)?", s)
    if not m:
        return None
    try:
        return float(m.group(0).replace(",", ""))
    except ValueError:
        return None


def cell_text(val):
    if val is None:
        return ""
    return str(val).strip()


def _norm(s):
    """Normalise a header/sheet name for lenient comparison: lowercase and
    treat '_', '-' and whitespace runs equivalently (e.g. 'Stock_Count' and
    'Stock Count' both become 'stock count')."""
    return re.sub(r"[_\-\s]+", " ", s.strip().lower())


def is_formula(val):
    return isinstance(val, str) and val.strip().startswith("=")


def numeric_ok(val):
    """True when a metric cell is non-empty and either a non-negative number
    or a (possibly un-recalculated) Excel formula."""
    if val is None:
        return False
    if is_formula(val):
        return True
    f = safe_float(val)
    return f is not None and f >= 0


def _vals_close(agent_val, gt_val, abs_tol, rel_tol):
    """Compare an agent cell against the GT cell.

    Returns True/False when both sides parse to numbers; returns None when
    either side is unparseable so the caller can skip (the column-population
    check already flags a genuinely blank cell, so we do not double-penalize
    an unresolvable formula here)."""
    af = safe_float(agent_val)
    bf = safe_float(gt_val)
    if af is None or bf is None:
        return None
    return abs(af - bf) <= max(abs_tol, abs(bf) * rel_tol)


def _sig_tokens(s):
    stop = {"this", "that", "with", "from", "have", "into", "the", "and",
            "for", "are", "all", "section", "based", "data", "analysis",
            "detailed", "findings", "covers", "collected"}
    return [w for w in re.findall(r"[a-z0-9]+", s.lower()) if len(w) >= 4 and w not in stop]


def _tok_overlap(a, b):
    """Any significant token overlaps (including one being a prefix of the
    other, e.g. 'financial' vs 'financials')."""
    for x in a:
        for y in b:
            if x == y or x.startswith(y) or y.startswith(x):
                return True
    return False


def lenient_match(req, cand):
    """Loose match between two short text strings: exact, substring either
    way, or significant-token overlap."""
    r = req.lower().strip()
    c = cand.lower().strip()
    if not r or not c:
        return False
    if r == c or r in c or c in r:
        return True
    rt, ct = _sig_tokens(r), _sig_tokens(c)
    return bool(rt) and bool(ct) and _tok_overlap(rt, ct)


# Generic words shared across GICS sector names that must NOT count as a
# matching token (e.g. 'Financial Services' vs 'Communication Services').
_SECTOR_STOP = {"services", "consumer", "sector", "industry", "industries",
                "group", "division", "and", "the", "of", "technology"}


def sector_match(req, cand):
    """Lenient sector-name match that tolerates spelling/format variants
    ('Health Care' vs 'Healthcare', 'Financials' vs 'Financial Services',
    'Technology / Communication Services' vs 'Communication Services') but
    does NOT treat a shared generic word ('Services', 'Consumer', ...) as a
    match, so two genuinely different sectors never collide."""
    r = req.lower().strip()
    c = cand.lower().strip()
    if not r or not c:
        return False
    if r == c or r in c or c in r:
        return True
    rt = [w for w in re.findall(r"[a-z]+", r) if len(w) >= 4 and w not in _SECTOR_STOP]
    ct = [w for w in re.findall(r"[a-z]+", c) if len(w) >= 4 and w not in _SECTOR_STOP]
    return bool(rt) and bool(ct) and any(
        x == y or x.startswith(y) or y.startswith(x) for x in rt for y in ct)


def _header_index(ws):
    headers = [_norm(cell_text(c.value)) for c in ws[1]]
    return {h: i for i, h in enumerate(headers) if h}, headers


def _data_rows(ws):
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    return [r for r in rows if r and any(v is not None for v in r)]


def _check_sector_performance(ws, gt_ws):
    idx, headers = _header_index(ws)
    if "sector" not in idx:
        return
    rows = [r for r in _data_rows(ws) if cell_text(r[idx["sector"]])]
    agent_sectors = [cell_text(r[idx["sector"]]) for r in rows]

    # The seed universe is one stock per GICS sector. The task asks for a
    # cross-sector analysis ("stocks across different sectors") without
    # enumerating which sectors, so a model that covers a clear majority of
    # the available sectors passes (avoids false FAIL on partial pulls),
    # while a 1-2-sector sheet still fails (not a cross-sector analysis).
    gt_rows = _data_rows(gt_ws)
    expected_sectors = [cell_text(r[0]) for r in gt_rows if r[0] is not None]
    covered = [es for es in expected_sectors
               if any(sector_match(es, a) for a in agent_sectors)]
    need = max(1, (len(expected_sectors) + 1) // 2)  # majority
    missing = [x for x in expected_sectors if x not in covered]
    check(f"Sector_Performance covers >= {need}/{len(expected_sectors)} expected sectors",
          len(covered) >= need,
          f"covered {len(covered)}/{len(expected_sectors)}; missing: {missing}; "
          f"agent sectors: {agent_sectors[:10]}")

    # Metric columns: present, populated, and holding numbers/formulas. Exact
    # values are not compared because the task fixes no calc window/口径.
    for col_key, col_name in (("stock count", "Stock_Count"),
                              ("avg price", "Avg_Price"),
                              ("total market value", "Total_Market_Value"),
                              ("volatility score", "Volatility_Score")):
        ci = idx.get(col_key)
        if ci is None:
            continue
        vals = [r[ci] for r in rows if len(r) > ci]
        ok = all(numeric_ok(v) for v in vals) if vals else False
        check(f"Sector_Performance {col_name} values are numeric/populated", ok,
              f"samples: {[cell_text(v)[:14] for v in vals[:6]]}")

    # Deterministic value checks against the GT for the three well-defined
    # metrics. The seed holds exactly one stock per GICS sector, so any sector
    # the agent did analyze has a unique correct Stock_Count / Avg_Price /
    # Total_Market_Value. We compare only on sectors present in BOTH the agent
    # sheet and the GT (a valid cross-sector analysis need not cover the exact
    # same subset), so partial coverage is never punished here and a fabricated
    # number on a covered sector is. Volatility_Score is deliberately excluded:
    # task.md leaves its window/method open ("a volatility measure (e.g. the
    # standard deviation of daily returns); use one consistent time window"),
    # so two correct solvers can legitimately differ. The GT self-test always
    # passes because the agent workbook IS the GT workbook in that mode.
    gt_idx, _ = _header_index(gt_ws)
    gt_sec_i = gt_idx.get("sector")
    gt_stock_i = gt_idx.get("stock count")
    gt_avg_i = gt_idx.get("avg price")
    gt_tmv_i = gt_idx.get("total market value")
    gt_by_sec = {}
    if gt_sec_i is not None:
        for gr in _data_rows(gt_ws):
            gname = cell_text(gr[gt_sec_i]) if gt_sec_i < len(gr) else ""
            if gname:
                gt_by_sec[gname] = gr
    a_stock_i = idx.get("stock count")
    a_avg_i = idx.get("avg price")
    a_tmv_i = idx.get("total market value")
    for r in rows:
        aname = cell_text(r[idx["sector"]])
        gr = next((grow for gname, grow in gt_by_sec.items()
                   if sector_match(aname, gname)), None)
        if gr is None:
            continue  # agent sector not represented in GT -> coverage handles it
        # Stock_Count is an exact integer (number of stocks in the sector).
        if a_stock_i is not None and gt_stock_i is not None:
            av = r[a_stock_i] if a_stock_i < len(r) else None
            gv = gr[gt_stock_i] if gt_stock_i < len(gr) else None
            res = _vals_close(av, gv, 0, 0)
            if res is not None:
                check(f"Sector_Performance '{aname}' Stock_Count matches GT",
                      res, f"agent={cell_text(av)}, gt={cell_text(gv)}")
        # Avg_Price = average closing price; tolerate ~3% / $1 (rounding).
        if a_avg_i is not None and gt_avg_i is not None:
            av = r[a_avg_i] if a_avg_i < len(r) else None
            gv = gr[gt_avg_i] if gt_avg_i < len(gr) else None
            res = _vals_close(av, gv, 1.0, 0.03)
            if res is not None:
                check(f"Sector_Performance '{aname}' Avg_Price matches GT",
                      res, f"agent={cell_text(av)}, gt={cell_text(gv)}")
        # Total_Market_Value = combined market cap; tolerate ~3% (huge numbers).
        if a_tmv_i is not None and gt_tmv_i is not None:
            av = r[a_tmv_i] if a_tmv_i < len(r) else None
            gv = gr[gt_tmv_i] if gt_tmv_i < len(gr) else None
            res = _vals_close(av, gv, 0.0, 0.03)
            if res is not None:
                check(f"Sector_Performance '{aname}' Total_Market_Value matches GT",
                      res, f"agent={cell_text(av)[:18]}, gt={cell_text(gv)[:18]}")


def _check_research_mapping(ws):
    idx, headers = _header_index(ws)
    if "paper title" not in idx:
        return
    rows = [r for r in _data_rows(ws) if cell_text(r[idx["paper title"]])]
    for ri, r in enumerate(rows):
        title = cell_text(r[idx["paper title"]])
        finding = cell_text(r[idx["key finding"]]) if idx.get("key finding") is not None and len(r) > idx["key finding"] else ""
        sector = cell_text(r[idx["applicable sector"]]) if idx.get("applicable sector") is not None and len(r) > idx["applicable sector"] else ""
        status = cell_text(r[idx["validation status"]]).lower() if idx.get("validation status") is not None and len(r) > idx["validation status"] else ""
        check(f"Research_Mapping R{ri+2} has key finding", finding != "",
              f"row: {title[:40]}")
        check(f"Research_Mapping R{ri+2} has applicable sector", sector != "",
              f"row: {title[:40]}")
        check(f"Research_Mapping R{ri+2} valid Validation_Status",
              status in VALIDATION_STATUSES, f"status: {status!r}")

    # At least one mapped paper must be a real relevant paper from the seed.
    hits = 0
    for r in rows:
        t = cell_text(r[idx["paper title"]]).lower()
        if any(kt in t or t in kt for kt in RELEVANT_PAPER_TITLES):
            hits += 1
    check("Research_Mapping references relevant papers", hits >= 1,
          f"known-paper hits: {hits}/{len(rows)}")


def _check_investment_thesis(ws):
    idx, headers = _header_index(ws)
    if "sector" not in idx:
        return
    rows = [r for r in _data_rows(ws) if cell_text(r[idx["sector"]])]
    for ri, r in enumerate(rows):
        outlook = cell_text(r[idx["outlook"]]).lower() if idx.get("outlook") is not None and len(r) > idx["outlook"] else ""
        evidence = cell_text(r[idx["supporting evidence"]]) if idx.get("supporting evidence") is not None and len(r) > idx["supporting evidence"] else ""
        risk = cell_text(r[idx["risk factor"]]) if idx.get("risk factor") is not None and len(r) > idx["risk factor"] else ""
        check(f"Investment_Thesis R{ri+2} valid Outlook", outlook in SECTOR_OUTLOOKS,
              f"outlook: {outlook!r}")
        check(f"Investment_Thesis R{ri+2} has Supporting_Evidence", evidence != "",
              f"row: {outlook}")
        check(f"Investment_Thesis R{ri+2} has Risk_Factor", risk != "",
              f"row: {outlook}")


def _check_sheet(ws, gt_ws, sheet_name):
    idx, headers = _header_index(ws)
    gt_rows = _data_rows(gt_ws)

    # Headers from the GT must be present in the agent sheet.
    for h in [_norm(cell_text(c.value)) for c in gt_ws[1]]:
        if h:
            check(f"{sheet_name} has {h} column", h in headers,
                  f"headers: {headers[:10]}")

    # Row-count floor (GT-derived; faithful agents exceed it).
    min_rows = max(1, len(gt_rows) - 2)
    agent_data = _data_rows(ws)
    check(f"{sheet_name} has >= {min_rows} data rows",
          len(agent_data) >= min_rows, f"got {len(agent_data)}")

    if sheet_name == "Sector_Performance":
        _check_sector_performance(ws, gt_ws)
    elif sheet_name == "Research_Mapping":
        _check_research_mapping(ws)
    elif sheet_name == "Investment_Thesis":
        _check_investment_thesis(ws)


def check_word_doc(agent_workspace, groundtruth_workspace):
    docx_path = os.path.join(agent_workspace, "Sector_Research_Brief.docx")
    check("Sector_Research_Brief.docx exists", os.path.exists(docx_path))
    if not os.path.exists(docx_path):
        return
    from docx import Document
    doc = Document(docx_path)
    text = " ".join(p.text for p in doc.paragraphs)
    check("Sector_Research_Brief.docx has content", len(text) > 50,
          f"text length: {len(text)}")

    # Candidate headings: heading-styled paragraphs, plus short title-like
    # paragraphs (some models write headings as bold/plain text). Leading/
    # trailing periods are ignored so a title written as a plain paragraph
    # ending in a full stop (e.g. "Financial Performance Review.") still
    # counts as the heading the task asks for.
    candidates = []
    for p in doc.paragraphs:
        t = p.text.strip().strip(".")
        if not t:
            continue
        if p.style.name.lower().startswith("heading"):
            candidates.append(t)
        elif len(t) <= 70:
            candidates.append(t)

    # Required headings come from the GT doc (which mirrors task.md).
    gt_path = os.path.join(groundtruth_workspace, "Sector_Research_Brief.docx")
    required = []
    if os.path.exists(gt_path):
        gt_doc = Document(gt_path)
        for p in gt_doc.paragraphs:
            st = p.style.name.lower()
            t = p.text.strip().strip(".")
            if t and (st.startswith("heading") or st == "title"):
                required.append(t)
    # Fall back to the task-specified heading text if GT is unavailable.
    if not required:
        required = ["Cross-Disciplinary Sector Analysis",
                    "Financial Performance Review",
                    "Academic Research Insights",
                    "Theory vs Practice Comparison",
                    "Investment Implications"]
    for req in required:
        if not req:
            continue
        found = any(lenient_match(req, c) for c in candidates)
        check(f"Sector_Research_Brief.docx has \"{req[:40]}\"", found,
              f"agent headings: {candidates[:8]}")


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # ---- Sector_Analysis_Report.xlsx ----
    excel_path = os.path.join(agent_workspace, "Sector_Analysis_Report.xlsx")
    check("Sector_Analysis_Report.xlsx exists", os.path.exists(excel_path))
    gt_path = os.path.join(groundtruth_workspace, "Sector_Analysis_Report.xlsx")
    gt_wb = openpyxl.load_workbook(gt_path, data_only=False) if os.path.exists(gt_path) else None

    if os.path.exists(excel_path):
        # Read with data_only=False so formula cells keep their formula
        # strings instead of becoming None; numeric_ok handles them (R2/R3).
        wb = openpyxl.load_workbook(excel_path, data_only=False)
        if gt_wb:
            agent_sheet_map = {_norm(s): s for s in wb.sheetnames}
            for gt_sheet in gt_wb.sheetnames:
                agent_sheet = agent_sheet_map.get(_norm(gt_sheet))
                check(f"{gt_sheet} sheet exists", agent_sheet is not None,
                      f"agent sheets: {wb.sheetnames}")
                if agent_sheet is not None:
                    _check_sheet(wb[agent_sheet], gt_wb[gt_sheet], gt_sheet)
        wb.close()

    # ---- Sector_Research_Brief.docx ----
    check_word_doc(agent_workspace, groundtruth_workspace)

    # ---- sector_analyst.py (terminal usage) ----
    sa_path = os.path.join(agent_workspace, "sector_analyst.py")
    check("sector_analyst.py exists", os.path.isfile(sa_path), f"missing: {sa_path}")
    if os.path.isfile(sa_path):
        with open(sa_path) as _f:
            _src = _f.read()
        check("sector_analyst.py references financial_data.json",
              "financial_data.json" in _src,
              "missing reference to financial_data.json")
        check("sector_analyst.py references research_findings.json",
              "research_findings.json" in _src,
              "missing reference to research_findings.json")
        check("sector_analyst.py references sector_analysis.json",
              "sector_analysis.json" in _src,
              "missing reference to sector_analysis.json")

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0.0
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"total_passed": PASS_COUNT, "total_checks": total,
                       "accuracy": accuracy}, f, indent=2)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
