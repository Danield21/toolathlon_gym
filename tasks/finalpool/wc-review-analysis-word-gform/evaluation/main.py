"""Evaluation for wc-review-analysis-word-gform."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=0.5):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Review_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "review" in p.text.lower() and "analysis" in p.text.lower():
            has_heading = True
            break
    check("Document has review analysis heading", has_heading)

    # Check tables
    check("Document has at least 3 tables", len(doc.tables) >= 3,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 3:
        return False

    # Load groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Review_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return False

    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check category table (table 0)
    gt_cats = list(gt_wb["By Category"].iter_rows(min_row=2, values_only=True))
    table1 = doc.tables[0]
    cat_rows = []
    for row in table1.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        cat_rows.append(cells)

    expected_cat_count = len(gt_cats)
    check(f"Category table has {expected_cat_count} rows",
          len(cat_rows) == expected_cat_count,
          f"Got {len(cat_rows)} rows, expected {expected_cat_count}")

    # Validate per-category Avg_Rating, Review_Count, Positive_Count, Negative_Count
    for gt_row in gt_cats:
        cat_name, review_count, avg_rating, pos_count, neg_count = gt_row
        matched = None
        for r in cat_rows:
            if r and r[0] and cat_name.lower() in r[0].lower():
                matched = r
                break
        if matched and len(matched) >= 5:
            # Header order: Category, Review_Count, Avg_Rating, Positive_Count, Negative_Count
            # Avg_Rating
            try:
                actual_avg = float(matched[2].replace(",", ""))
                check(f"Category {cat_name} avg rating", num_close(actual_avg, avg_rating, 0.05),
                      f"Expected {avg_rating}, got {actual_avg}")
            except (ValueError, AttributeError):
                check(f"Category {cat_name} avg rating parseable", False, f"got '{matched[2]}'")
            # Review_Count
            try:
                actual_rc = int(matched[1].replace(",", ""))
                check(f"Category {cat_name} review count",
                      actual_rc == review_count,
                      f"Expected {review_count}, got {actual_rc}")
            except (ValueError, AttributeError):
                pass
            # Positive_Count
            try:
                actual_pos = int(matched[3].replace(",", ""))
                check(f"Category {cat_name} positive count",
                      actual_pos == pos_count,
                      f"Expected {pos_count}, got {actual_pos}")
            except (ValueError, AttributeError):
                pass
            # Negative_Count
            try:
                actual_neg = int(matched[4].replace(",", ""))
                check(f"Category {cat_name} negative count",
                      actual_neg == neg_count,
                      f"Expected {neg_count}, got {actual_neg}")
            except (ValueError, AttributeError):
                pass
        else:
            check(f"Category {cat_name} found in table", False)

    # Verify category sort: avg_rating desc
    if cat_rows:
        try:
            cat_ratings = []
            for r in cat_rows:
                if r and len(r) >= 3 and r[2]:
                    try:
                        cat_ratings.append(float(r[2].replace(",", "")))
                    except ValueError:
                        pass
            check("Category table sorted by Avg_Rating desc",
                  cat_ratings == sorted(cat_ratings, reverse=True),
                  f"Got {cat_ratings}")
        except Exception:
            pass

    # Check Top Rated Products table (table 1)
    if len(doc.tables) >= 2:
        gt_top = list(gt_wb["Best Products"].iter_rows(min_row=2, values_only=True))
        table2 = doc.tables[1]
        top_rows = []
        for row in table2.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            top_rows.append(cells)
        check(f"Top Rated table has {len(gt_top)} rows",
              len(top_rows) == len(gt_top),
              f"Got {len(top_rows)} rows")
        # Match by product name (substring)
        for gt_row in gt_top:
            pname, pcat, prev, prating = gt_row
            # truncate name for matching (40 chars)
            pname_key = (pname or "")[:40].lower()
            matched = None
            for r in top_rows:
                if r and r[0] and pname_key in r[0].lower():
                    matched = r
                    break
            if matched and len(matched) >= 4:
                try:
                    actual_avg = float(matched[3].replace(",", ""))
                    check(f"Top product '{pname[:40]}' avg rating",
                          num_close(actual_avg, prating, 0.05),
                          f"Expected {prating}, got {actual_avg}")
                except (ValueError, AttributeError):
                    pass
            else:
                check(f"Top product '{pname[:40]}' found in table", False)

    # Check Lowest Rated Products table (table 2)
    if len(doc.tables) >= 3:
        gt_worst = list(gt_wb["Worst Products"].iter_rows(min_row=2, values_only=True))
        table3 = doc.tables[2]
        worst_rows = []
        for row in table3.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            worst_rows.append(cells)
        check(f"Lowest Rated table has {len(gt_worst)} rows",
              len(worst_rows) == len(gt_worst),
              f"Got {len(worst_rows)} rows")
        for gt_row in gt_worst:
            pname, pcat, prev, prating = gt_row
            pname_key = (pname or "")[:40].lower()
            matched = None
            for r in worst_rows:
                if r and r[0] and pname_key in r[0].lower():
                    matched = r
                    break
            if matched and len(matched) >= 4:
                try:
                    actual_avg = float(matched[3].replace(",", ""))
                    check(f"Worst product '{pname[:40]}' avg rating",
                          num_close(actual_avg, prating, 0.05),
                          f"Expected {prating}, got {actual_avg}")
                except (ValueError, AttributeError):
                    pass
            else:
                check(f"Worst product '{pname[:40]}' found in table", False)

    # Check summary paragraph contains key info
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    # Total reviews number
    total_reviews = sum(r[1] for r in gt_cats)
    check(f"Summary mentions total review count {total_reviews}",
          str(total_reviews) in full_text,
          f"Missing total {total_reviews}")
    # Number of categories
    check(f"Summary mentions {expected_cat_count} categories",
          str(expected_cat_count) in full_text,
          f"Missing category count {expected_cat_count}")
    # Highest category
    highest_cat = gt_cats[0][0]
    check(f"Summary mentions highest category '{highest_cat}'",
          highest_cat.lower() in full_text,
          "missing highest category name")
    # Lowest category
    lowest_cat = gt_cats[-1][0]
    check(f"Summary mentions lowest category '{lowest_cat}'",
          lowest_cat.lower() in full_text,
          "missing lowest category name")

    # Check top/worst products sections
    has_top = False
    has_worst = False
    for p in doc.paragraphs:
        if "top" in p.text.lower() and "rated" in p.text.lower():
            has_top = True
        if "lowest" in p.text.lower() or "worst" in p.text.lower():
            has_worst = True
    check("Has top rated products section", has_top)
    check("Has lowest rated products section", has_worst)

    return True


def check_gform():
    """Check Google Form."""
    print("\n=== Checking Google Form ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SELECT id, title FROM gform.forms")
    forms = cur.fetchall()
    cur.execute("SELECT form_id, title, question_type FROM gform.questions")
    questions = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 1 form created", len(forms) >= 1,
          f"Found {len(forms)}")

    target_form_id = None
    if forms:
        # Title exact match: 'Product Improvement Feedback'
        for f in forms:
            t = (f[1] or "").lower().strip()
            if "product improvement feedback" in t:
                target_form_id = f[0]
                break
        check("Form titled 'Product Improvement Feedback' exists",
              target_form_id is not None,
              f"Forms: {[f[1] for f in forms]}")

    # Filter questions for target form
    if target_form_id:
        target_questions = [q for q in questions if q[0] == target_form_id]
    else:
        target_questions = questions
    check("At least 3 questions in form", len(target_questions) >= 3,
          f"Found {len(target_questions)}")

    if target_questions:
        q_lookup = {(q[1] or "").lower(): q for q in target_questions}
        types_present = [(q[2] or "").lower() for q in target_questions]

        # Q1: category multiple choice
        cat_q = None
        for txt, q in q_lookup.items():
            if "category" in txt and ("which" in txt or "primarily" in txt):
                cat_q = q
                break
        if cat_q:
            t = (cat_q[2] or "").lower()
            check("Category question is multiple choice",
                  any(x in t for x in ["multiple", "choice", "radio"]),
                  f"type='{cat_q[2]}'")
        else:
            check("Has category multiple-choice question", False, "not found")

        # Q2: linear scale 1-5
        rate_q = None
        for txt, q in q_lookup.items():
            if "experience" in txt or "shopping experience" in txt or "overall shopping" in txt:
                rate_q = q
                break
        if rate_q:
            t = (rate_q[2] or "").lower()
            check("Rating question is linear scale",
                  any(x in t for x in ["linear", "scale", "rating"]),
                  f"type='{rate_q[2]}'")
        else:
            check("Has linear scale rating question", False, "not found")

        # Q3: paragraph
        sug_q = None
        for txt, q in q_lookup.items():
            if "improve" in txt or "suggest" in txt:
                sug_q = q
                break
        if sug_q:
            t = (sug_q[2] or "").lower()
            check("Suggestion question is paragraph",
                  any(x in t for x in ["paragraph", "text", "long"]),
                  f"type='{sug_q[2]}'")
        else:
            check("Has paragraph suggestion question", False, "not found")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("WC REVIEW ANALYSIS WORD GFORM - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_gform()

    total_pass = PASS_COUNT
    total_fail = FAIL_COUNT
    all_ok = FAIL_COUNT == 0

    print(f"\n=== SUMMARY ===")
    print(f"  Total checks - Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": total_pass, "failed": total_fail, "success": all_ok}, f, indent=2)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
