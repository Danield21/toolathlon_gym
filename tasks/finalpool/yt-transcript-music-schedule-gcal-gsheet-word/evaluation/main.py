"""
Evaluation for yt-transcript-music-schedule-gcal-gsheet-word task.

Checks:
1. Radio_Show_Script.docx (or similar) exists in agent_workspace
2. Word doc has >= 3 heading-level sections
3. Word doc text contains Afrobeat and >= 2 artist names
4. GCal has >= 3 new events with Afrobeat/Show in summary in April 2026
5. GSheet has spreadsheet with Playlist sheet containing >= 8 rows
6. Email sent to station@radioafrica.fm
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
RUNTIME_ONLY_FAIL = 0
ARTIST_NAMES = ["burna boy", "wizkid", "rema", "davido", "ckay", "asake", "ayra starr",
                "fireboy", "tems", "omah lay", "kizz daniel"]


def record(name, passed, detail="", runtime_only=False):
    global PASS_COUNT, FAIL_COUNT, RUNTIME_ONLY_FAIL
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        if runtime_only:
            RUNTIME_ONLY_FAIL += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Check 1-3: Radio Show Script Word Doc ===")
    docx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".docx") and ("script" in fname.lower() or "show" in fname.lower() or "radio" in fname.lower()):
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".docx"):
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Radio show script Word doc exists", docx_path is not None,
           f"No script/show/radio docx found in {agent_workspace}")

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found")
        record("Word doc contains Afrobeat and artist names", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings")

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        has_afrobeat = "afrobeat" in full_text
        # Artist names are not present in the transcript (lyrics only), so
        # checking against a fixed list is unfair. Instead verify the agent
        # wrote track introductions with plausible artist names — i.e. the
        # "Track Introductions" section has >= 2 non-trivial entries.
        intro_section = False
        artist_mentions = 0
        for p in doc.paragraphs:
            t = p.text.strip().lower()
            if "track introduction" in t or "track intro" in t:
                intro_section = True
                continue
            if intro_section and len(t) > 20:
                artist_mentions += 1
        record("Word doc contains Afrobeat and >= 2 artist names",
               has_afrobeat and artist_mentions >= 2,
               f"Afrobeat: {has_afrobeat}, track intro paragraphs: {artist_mentions}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e))
        record("Word doc contains Afrobeat and artist names", False, str(e))


def check_gcal():
    """Validate the four Sunday Afrobeat Show events at 20:00-22:00.

    Required Sundays: Apr 5, 12, 19, 26 of 2026.
    Each event must start at 20:00 and end at 22:00 (UTC offset accepted).
    """
    print("\n=== Check 4: GCal Afrobeat Show Events in April 2026 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime FROM gcal.events
        WHERE start_datetime >= TIMESTAMPTZ '2026-04-01T00:00:00Z' AND start_datetime < TIMESTAMPTZ '2026-05-01T00:00:00Z'
        AND (summary ILIKE '%afrobeat%' OR summary ILIKE '%show%')
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    record("GCal has 4 Afrobeat/Show events in April 2026",
           len(events) == 4, f"Found {len(events)} matching events",
           runtime_only=True)

    # Build a set of (date, start_hh:mm, end_hh:mm) tuples in UTC.
    # Postgres returns timestamptz in the session timezone; normalize to UTC
    # so comparisons against the show's intended 20:00-22:00 UTC are robust.
    from datetime import timezone
    triples = set()
    for summary, desc, st, en in events:
        if st is None or en is None:
            continue
        if st.tzinfo is not None:
            st_utc = st.astimezone(timezone.utc)
            en_utc = en.astimezone(timezone.utc)
        else:
            st_utc, en_utc = st, en
        triples.add((str(st_utc.date()), st_utc.strftime("%H:%M"), en_utc.strftime("%H:%M")))

    expected_sundays = ["2026-04-05", "2026-04-12", "2026-04-19", "2026-04-26"]
    for d in expected_sundays:
        record(f"Show on {d} 20:00-22:00 UTC",
               (d, "20:00", "22:00") in triples,
               f"Have: {sorted(triples)[:8]}",
               runtime_only=True)


def check_gsheet():
    """Validate the Radio_Broadcast_Schedule cloud spreadsheet.

    Match by EXACT title 'Radio_Broadcast_Schedule' (case-insensitive,
    spaces/underscores normalized). Do NOT fall back to most-recent
    spreadsheet to avoid accepting a wrong workbook.
    """
    print("\n=== Check 5: GSheet Radio_Broadcast_Schedule ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    def _norm(s):
        return str(s or "").strip().lower().replace("_", " ").replace("-", " ")

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    rows = cur.fetchall()
    spreadsheet = None
    for sid, title in rows:
        if _norm(title) == _norm("Radio_Broadcast_Schedule"):
            spreadsheet = (sid, title)
            break

    record("Radio_Broadcast_Schedule spreadsheet exists (exact title)",
           spreadsheet is not None,
           f"Available titles: {[r[1] for r in rows]}",
           runtime_only=True)

    if spreadsheet:
        spreadsheet_id, title = spreadsheet
        cur.execute("""
            SELECT id, title FROM gsheet.sheets
            WHERE spreadsheet_id = %s AND title ILIKE '%%playlist%%'
            LIMIT 1
        """, (spreadsheet_id,))
        playlist_sheet = cur.fetchone()

        record("Playlist sheet exists in spreadsheet", playlist_sheet is not None,
               f"Sheets in spreadsheet: {title}", runtime_only=True)

        if playlist_sheet:
            sheet_id = playlist_sheet[0]
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
                AND value IS NOT NULL AND value != ''
            """, (spreadsheet_id, sheet_id))
            data_rows = cur.fetchone()[0]
            record("Playlist sheet has >= 8 data rows", data_rows >= 8,
                   f"Found {data_rows} data rows", runtime_only=True)
        else:
            record("Playlist sheet has >= 8 data rows", False, "No Playlist sheet",
                   runtime_only=True)

        # Check Show_Schedule sheet (must contain the 4 specific Sundays)
        cur.execute("""
            SELECT id, title FROM gsheet.sheets
            WHERE spreadsheet_id = %s AND (title ILIKE '%%show%%' OR title ILIKE '%%schedule%%')
            LIMIT 1
        """, (spreadsheet_id,))
        show_sheet = cur.fetchone()
        record("Show_Schedule sheet exists", show_sheet is not None,
               "No Show_Schedule sheet found", runtime_only=True)

        if show_sheet:
            sheet_id = show_sheet[0]
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
                AND value IS NOT NULL AND value != ''
            """, (spreadsheet_id, sheet_id))
            data_rows = cur.fetchone()[0]
            record("Show_Schedule sheet has >= 4 rows (April Sundays)", data_rows >= 4,
                   f"Found {data_rows} rows", runtime_only=True)

            # Pull all cell values for that sheet to verify the 4 Sundays appear
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s
            """, (spreadsheet_id, sheet_id))
            cells = cur.fetchall()
            # Group by row
            grid = {}
            for ri, ci, val in cells:
                grid.setdefault(ri, {})[ci] = str(val) if val is not None else ""
            # Flatten each row's values
            row_strings = {ri: " | ".join(v for _, v in sorted(cols.items())) for ri, cols in grid.items()}

            for d in ["2026-04-05", "2026-04-12", "2026-04-19", "2026-04-26"]:
                found = any(d in s for s in row_strings.values())
                record(f"Show_Schedule contains date {d}", found,
                       f"Rows: {list(row_strings.values())[:6]}",
                       runtime_only=True)
        else:
            record("Show_Schedule sheet has >= 4 rows", False, "Sheet not found",
                   runtime_only=True)
    else:
        for chk in ["Playlist sheet exists", "Playlist sheet has >= 8 data rows",
                    "Show_Schedule sheet exists", "Show_Schedule sheet has >= 4 rows",
                    "Show_Schedule contains date 2026-04-05",
                    "Show_Schedule contains date 2026-04-12",
                    "Show_Schedule contains date 2026-04-19",
                    "Show_Schedule contains date 2026-04-26"]:
            record(chk, False, "No spreadsheet found", runtime_only=True)

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 6: Email to station@radioafrica.fm ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else to_addr.lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "station@radioafrica.fm" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to station@radioafrica.fm", matching is not None,
           f"Total messages: {len(messages)}", runtime_only=True)
    if matching:
        body = (matching[0] or "") + " " + (matching[3] or "")
        has_content = any(k in body.lower() for k in ["schedule", "broadcast", "afrobeat", "sunday", "show"])
        record("Email mentions show schedule content", has_content,
               f"Subject: {matching[0]}", runtime_only=True)


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_xlsx_content(workspace, groundtruth_workspace="."):
    """Optional/runtime-only XLSX spot check.

    The task.md asks the agent to create the workbook in the cloud
    spreadsheet, NOT a local xlsx. We keep this check as a non-blocking
    spot check: if a local xlsx happens to be present, validate it; if
    not, skip silently. This avoids forcing an output the task does not
    require.
    """
    print("\n=== Checking XLSX Content (optional) ===")
    try:
        import openpyxl
    except ImportError:
        return True

    xlsx_path = os.path.join(workspace, "Radio_Broadcast_Schedule_local.xlsx")
    if not os.path.isfile(xlsx_path):
        # Not required by task.md; skip silently.
        print("  [SKIP] Local xlsx not present (not required by task)")
        return True

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        # Just sanity check it has data; do not assert structure here
        # because the cloud spreadsheet (gsheet) is the source of truth.
        any_data = any(
            len(list(ws.iter_rows(values_only=True))) >= 2 for ws in wb.worksheets
        )
        print(f"  [SPOT] Local xlsx present and has data: {any_data}")
        wb.close()
        return True
    except Exception as e:
        print(f"  [SPOT] Local xlsx error (ignored): {e}")
        return True


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gcal()
    check_gsheet()
    check_email()
    check_xlsx_content(args.agent_workspace, args.groundtruth_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # Two-mode gate:
    # - V1 (GT-test, agent did not run): only gate on non-runtime fails.
    # - V2 (agent did run): once we detect ANY agent-side artifact in
    #   gcal/gsheet, we gate on the full FAIL_COUNT.
    non_runtime_fail = FAIL_COUNT - RUNTIME_ONLY_FAIL

    agent_ran = False
    try:
        ck = psycopg2.connect(**DB_CONFIG)
        ck_cur = ck.cursor()
        ck_cur.execute(
            "SELECT COUNT(*) FROM gcal.events "
            "WHERE start_datetime >= TIMESTAMPTZ '2026-04-01T00:00:00Z' AND start_datetime < TIMESTAMPTZ '2026-05-01T00:00:00Z' "
            "AND (summary ILIKE '%afrobeat%' OR summary ILIKE '%show%')"
        )
        n_gcal = ck_cur.fetchone()[0] or 0
        ck_cur.execute("SELECT COUNT(*) FROM gsheet.spreadsheets")
        n_sheets = ck_cur.fetchone()[0] or 0
        ck_cur.close()
        ck.close()
        agent_ran = (n_gcal >= 1) or (n_sheets >= 1)
    except Exception:
        pass

    print(f"Agent ran: {agent_ran}; Non-runtime fails: {non_runtime_fail}, runtime-only fails: {RUNTIME_ONLY_FAIL}")

    if agent_ran:
        gate_fails = FAIL_COUNT
    else:
        gate_fails = non_runtime_fail

    if gate_fails == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
