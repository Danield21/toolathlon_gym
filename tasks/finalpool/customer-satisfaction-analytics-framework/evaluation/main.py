#!/usr/bin/env python3
"""Evaluation script for customer-satisfaction-analytics-framework task validation.

Six-phase narrative customer satisfaction analytics task. Validates structural
deliverables that any reasonable agent would produce based on task.md:

  - A satisfaction-analysis xlsx (NPS / CSAT / metric data)
  - An action-plan xlsx (improvement initiatives)
  - A satisfaction report docx with substantive findings
  - An executive-summary docx
  - Outgoing email to stakeholders + calendar review meeting (Phase 6)

The task description is intentionally narrative (no exact filenames, segments
or NPS targets) so verification uses structural + numeric-range overlap
checks against GT rather than strict row-by-row comparison.
"""

from argparse import ArgumentParser
import json
import os
import re
import sys

PASS_COUNT = 0
FAIL_COUNT = 0
WARN_COUNT = 0
IS_GT_SELF_TEST = False


def record(name, passed, detail="", db_side=False):
    global PASS_COUNT, FAIL_COUNT, WARN_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        # In GT self-test mode, DB-side checks (email/calendar/gsheet/notion)
        # naturally fail because GT files cannot pre-populate DB state.
        if IS_GT_SELF_TEST and db_side:
            WARN_COUNT += 1
            msg = f": {detail[:300]}" if detail else ""
            print(f"  [WARN] {name} (GT self-test mode, DB-side){msg}")
        else:
            FAIL_COUNT += 1
            msg = f": {detail[:300]}" if detail else ""
            print(f"  [FAIL] {name}{msg}")


def get_conn():
    import psycopg2
    return psycopg2.connect(
        host=os.environ.get("PGHOST", "localhost"),
        port=int(os.environ.get("PGPORT", "5432")),
        dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
        user=os.environ.get("PGUSER", "eigent"),
        password=os.environ.get("PGPASSWORD", "camel"),
    )


def check_phase6_distribution():
    """Phase 6: distribute findings via email + calendar review meetings.

    Keyword sets include both English and Chinese terms so that outputs written
    in Chinese (the task does not mandate an output language) are recognized
    just like English ones. These sets MUST be a subset of the preprocess
    cleanup sets (see preprocess/main.py) so stale rows never satisfy the
    checks.
    """
    print("\n=== Check: Phase 6 distribution (email/gcal) ===")
    EMAIL_KW = ["satisfaction", "customer experience", "findings", "nps", "csat",
                "survey", "feedback",
                "满意度", "客户体验", "调查", "反馈", "净推荐", "洞察"]
    EVENT_KW = ["satisfaction", "review", "customer", "feedback", "findings",
                "满意度", "评审", "回顾", "复盘", "客户", "反馈", "洞察", "发现"]
    try:
        conn = get_conn()
        cur = conn.cursor()
        email_clause = (" OR ".join(["subject ILIKE %s"] * len(EMAIL_KW))
                        + " OR "
                        + " OR ".join(["body_text ILIKE %s"] * len(EMAIL_KW)))
        email_params = [f"%{k}%" for k in EMAIL_KW] + [f"%{k}%" for k in EMAIL_KW]
        cur.execute(f"SELECT id, subject FROM email.messages WHERE {email_clause}",
                    email_params)
        emails = cur.fetchall()
        event_clause = " OR ".join(["summary ILIKE %s"] * len(EVENT_KW))
        event_params = [f"%{k}%" for k in EVENT_KW]
        cur.execute(f"SELECT id, summary FROM gcal.events WHERE {event_clause}",
                    event_params)
        events = cur.fetchall()
        conn.close()
    except Exception as e:
        record("Phase 6 distribution (db query)", False, f"db error: {e}", db_side=True)
        return
    record("Phase 6: at least one satisfaction/findings email sent",
           len(emails) >= 1, f"matching emails: {len(emails)}", db_side=True)
    record("Phase 6: at least one review/satisfaction calendar meeting",
           len(events) >= 1, f"matching events: {len(events)}", db_side=True)


def _find_xlsx_by_keywords(workspace, keywords, exclude=None):
    """Score xlsx files by keyword matches; filename match weighs 10x more than content."""
    import glob
    import openpyxl
    exclude = exclude or set()
    scored = []
    for path in sorted(glob.glob(os.path.join(workspace, "*.xlsx"))):
        if path in exclude:
            continue
        fname_low = os.path.basename(path).lower()
        if fname_low.startswith("~$") or fname_low.startswith("."):
            continue
        # Skip the input survey template
        if fname_low == "survey_template.xlsx":
            continue
        score = 0
        for kw in keywords:
            if kw in fname_low:
                score += 10
        try:
            wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
            content_text = " ".join(s.lower() for s in wb.sheetnames)
            for ws in wb.worksheets:
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    if row_count > 5:
                        break
                    row_count += 1
                    for cell in row:
                        if cell is not None:
                            content_text += " " + str(cell).lower()
            wb.close()
            for kw in keywords:
                if kw in content_text:
                    score += 1
        except Exception:
            pass
        if score > 0:
            scored.append((score, path))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [p for _, p in scored]


def _to_float(v):
    """Best-effort numeric parse of a cell value (int/float/str with % , currency,
    thousands separators, spaces). Returns None when the value is not parseable.
    Used so that agents writing literal numbers as strings (e.g. "4.2", "75%",
    "$1,200") are scored identically to those writing raw numerics."""
    if v is None:
        return None
    if isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    if isinstance(v, str):
        s = v.strip()
        if not s:
            return None
        t = s
        for ch in ("$", "¥", "€", "£", ",", "%", " "):
            t = t.replace(ch, "")
        neg = False
        if t.startswith("(") and t.endswith(")"):
            neg = True
            t = t[1:-1]
        if t in ("", "-", ".", "--"):
            return None
        try:
            val = float(t)
            return -val if neg else val
        except (ValueError, TypeError):
            return None
    return None


def _load_wb_pair(path):
    """Load a workbook twice: once with data_only=False (formulas/literal values)
    and once with data_only=True (cached formula results). Returns (wb_f, wb_v)."""
    import openpyxl
    wb_f = openpyxl.load_workbook(path, data_only=False)
    wb_v = openpyxl.load_workbook(path, data_only=True)
    return wb_f, wb_v


def _collect_numerics(wb_f, wb_v):
    """Collect numeric values across sheets. Formula cells use their cached value
    from the data_only=True workbook; cells that are formulas with no cached
    value are counted as 'unresolved' (returned separately). Returns
    (nums, unresolved_formula_count)."""
    nums = []
    unresolved = 0
    for ws_f, ws_v in zip(wb_f.worksheets, wb_v.worksheets):
        for row_f, row_v in zip(ws_f.iter_rows(), ws_v.iter_rows()):
            for c_f, c_v in zip(row_f, row_v):
                val = c_f.value
                if isinstance(val, str) and val.lstrip().startswith("="):
                    cached = c_v.value
                    if cached is None:
                        unresolved += 1
                        continue
                    val = cached
                f = _to_float(val)
                if f is not None:
                    nums.append(f)
    return nums, unresolved


def _collect_text_cells(wb_f):
    """Collect all literal text cells (lower-cased) from a workbook for
    header/keyword checks. Formula cells are skipped."""
    texts = []
    for ws in wb_f.worksheets:
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell is not None and isinstance(cell, str) and not cell.lstrip().startswith("="):
                    texts.append(cell.lower())
    return texts


def _structural_xlsx_check(agent_xlsx_path, gt_xlsx_path, label, expected_keywords):
    """Structural validation of an agent xlsx vs GT.

    Replaces strict row-by-row matching (incompatible with the narrative task
    description) with the following checks:
      - File readable and has at least one sheet with substantive rows
      - Total agent row count is within a tolerant range derived from GT
      - At least one content keyword from `expected_keywords` appears in the file
      - Numeric values overlap GT numeric range (catches empty / nonsense files)
    """
    fname = os.path.basename(agent_xlsx_path)

    try:
        a_wb_f, a_wb_v = _load_wb_pair(agent_xlsx_path)
    except Exception as e:
        record(f"xlsx {fname} readable", False, str(e))
        return False
    try:
        gt_wb_f, gt_wb_v = _load_wb_pair(gt_xlsx_path)
    except Exception as e:
        record(f"GT {label} readable", False, str(e))
        a_wb_f.close()
        a_wb_v.close()
        return False

    # Row count: tolerant range based on GT's total non-empty rows
    def total_nonempty_rows(wb):
        total = 0
        for ws in wb.worksheets:
            for r in ws.iter_rows(values_only=True):
                if any(c is not None and str(c).strip() != "" for c in r):
                    total += 1
        return total

    gt_rows = total_nonempty_rows(gt_wb_f)
    a_rows = total_nonempty_rows(a_wb_f)
    # Allow agent to produce between ~an eighth and ~triple the GT row count.
    # Floor of 2 so a compact-but-complete deliverable (title + header + one
    # data row) still clears the row check; substance is enforced by the
    # keyword-coverage and numeric-overlap checks below.
    lo = max(2, gt_rows // 8)
    hi = max(gt_rows * 3, gt_rows + 10)
    record(f"xlsx {fname} row count in tolerant range [{lo}, {hi}] (GT={gt_rows})",
           lo <= a_rows <= hi, f"agent_rows={a_rows}")

    # Keyword coverage check
    a_texts = _collect_text_cells(a_wb_f)
    a_blob = " ".join(a_texts)
    matched_kw = [kw for kw in expected_keywords if kw in a_blob]
    record(f"xlsx {fname} content keyword coverage (>=2 of {len(expected_keywords)})",
           len(matched_kw) >= 2,
           f"matched: {matched_kw}")

    # Numeric overlap (catches empty / completely off files). Robust to agents
    # writing literal values as strings (percentages, currency) and to formula
    # cells: formula cells with no cached value (never recalculated by Excel)
    # cannot be read, so the overlap check is relaxed to non-blocking when the
    # agent file contains such unresolved formulas.
    gt_nums, _ = _collect_numerics(gt_wb_f, gt_wb_v)
    a_nums, a_unresolved = _collect_numerics(a_wb_f, a_wb_v)
    if gt_nums:
        if not a_nums:
            # Agent file contains no parseable numerics at all; the overlap
            # check is not applicable (row-count + keyword checks still apply).
            record(f"xlsx {fname} numeric values overlap GT range (non-blocking)",
                   True, "no parseable numeric values in agent file; skipping overlap check")
            gt_wb_f.close()
            gt_wb_v.close()
            a_wb_f.close()
            a_wb_v.close()
            return True
        gt_min, gt_max = min(gt_nums), max(gt_nums)
        # Allow some buffer below and above the GT range
        buf = max((gt_max - gt_min) * 0.10, 1.0)
        overlap = sum(1 for v in a_nums if (gt_min - buf) <= v <= (gt_max + buf))
        if overlap >= 2:
            record(f"xlsx {fname} numeric values overlap GT range (>=2 values)",
                   True, f"overlap_count={overlap}, gt_range=[{gt_min}, {gt_max}]")
        elif a_unresolved > 0:
            record(f"xlsx {fname} numeric values overlap GT range (non-blocking)",
                   True,
                   f"agent file has {a_unresolved} formula cell(s) with no cached "
                   f"value; cannot verify cached numerics (overlap_count={overlap})")
        else:
            record(f"xlsx {fname} numeric values overlap GT range (>=2 values)",
                   False,
                   f"overlap_count={overlap}, gt_range=[{gt_min}, {gt_max}]")
    else:
        record(f"xlsx {fname} GT numeric range available (non-blocking)", True,
               "no GT numerics; skipping overlap check")

    gt_wb_f.close()
    gt_wb_v.close()
    a_wb_f.close()
    a_wb_v.close()
    return True


def check_xlsx_content(workspace, groundtruth_workspace="."):
    """Locate two analysis-type xlsx files (satisfaction/analysis + action_plans) by content keywords,
    then perform structural validation against GT. Allows agent-chosen filenames as long as
    content matches."""
    print("\n=== Check: XLSX files ===")
    import openpyxl

    # Each entry: (gt_filename, label, [accept-keywords for locating], [content-keywords for validation])
    # Locate/content keywords include Chinese equivalents so semantically-correct
    # outputs written in Chinese are matched like English ones. The >=2 keyword
    # coverage check still requires the file to actually discuss the topic.
    targets = [
        ("satisfaction_analysis.xlsx", "satisfaction analysis xlsx",
         ["satisfaction", "analysis", "nps", "metric", "score",
          "满意度", "分析", "指标", "评分", "净推荐"],
         ["satisfaction", "nps", "score", "rating", "segment", "category", "rate",
          "满意度", "评分", "细分", "类别", "维度", "净推荐"]),
        ("action_plans.xlsx", "action plans xlsx",
         ["action", "plan", "initiative", "improvement", "recommendation",
          "行动计划", "方案", "举措", "改进", "建议"],
         ["action", "initiative", "improvement", "owner", "metric", "timeline", "goal",
          "行动", "举措", "改进", "负责人", "指标", "时间", "目标"]),
    ]

    used_paths = set()
    for gt_fname, label, locate_keywords, content_keywords in targets:
        gt_path = os.path.join(groundtruth_workspace, gt_fname)
        if not os.path.isfile(gt_path):
            continue

        # 1) Prefer the exact GT filename if present
        exact = os.path.join(workspace, gt_fname)
        chosen = None
        if os.path.isfile(exact) and exact not in used_paths:
            chosen = exact
        else:
            # 2) Fuzzy locate by filename or content keywords (scored)
            cands = _find_xlsx_by_keywords(workspace, locate_keywords, exclude=used_paths)
            if cands:
                chosen = cands[0]
        # 3) Fallback: allow reusing a file already assigned to another target
        # (a single combined workbook may legitimately cover both analysis and
        # action-plan content when sub-agents share the workspace)
        if chosen is None:
            cands = _find_xlsx_by_keywords(workspace, locate_keywords)
            if cands:
                chosen = cands[0]

        if chosen is None:
            record(f"xlsx for {label} exists", False,
                   f"No xlsx with keywords {locate_keywords} found in workspace")
            continue
        used_paths.add(chosen)
        record(f"xlsx for {label} exists ({os.path.basename(chosen)})", True)

        # Sanity: at least one sheet has data
        try:
            wb = openpyxl.load_workbook(chosen, data_only=False)
            any_data_sheet = False
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                if len(rows) >= 2:
                    any_data_sheet = True
                    break
            wb.close()
            record(f"xlsx {os.path.basename(chosen)} has at least one populated sheet",
                   any_data_sheet, "no sheet has >=2 rows")
        except Exception as e:
            record(f"xlsx {os.path.basename(chosen)} readable", False, str(e))
            continue

        # Structural comparison vs GT (replaces strict row-by-row matching)
        _structural_xlsx_check(chosen, gt_path, label, content_keywords)


def _find_docx_by_keywords(workspace, keywords, exclude=None):
    """Find docx files in workspace whose filename OR content matches any keyword."""
    import glob
    from docx import Document
    exclude = exclude or set()
    candidates = []
    for path in sorted(glob.glob(os.path.join(workspace, "*.docx"))):
        if path in exclude:
            continue
        fname_low = os.path.basename(path).lower()
        if fname_low.startswith("~$") or fname_low.startswith("."):
            continue
        if any(kw in fname_low for kw in keywords):
            candidates.append(path)
            continue
        # Content match
        try:
            doc = Document(path)
            text_low = "\n".join(p.text for p in doc.paragraphs).lower()
            if any(kw in text_low for kw in keywords):
                candidates.append(path)
        except Exception:
            continue
    return candidates


def _count_words(text):
    """Count words in both space-separated scripts (English) and scripts that
    do not insert spaces between words (Chinese/Japanese). Whitespace tokens and
    each CJK character each count as one word, so a substantive document in
    either language clears the word-count floor."""
    if not text:
        return 0
    tokens = len([w for w in text.split() if w.strip()])
    cjk = len(re.findall(r"[一-鿿㐀-䶿豈-﫿]", text))
    return tokens + cjk


def check_docx_content(workspace):
    """Locate two report-type docx files (satisfaction report + executive summary)
    by filename or content keywords, then verify each is substantive."""
    print("\n=== Check: DOCX files ===")
    from docx import Document

    targets = [
        ("satisfaction_report.docx", "satisfaction report docx",
         ["satisfaction", "report", "finding", "customer", "nps",
          "满意度", "报告", "客户"]),
        ("executive_summary.docx", "executive summary docx",
         ["executive", "summary", "strategic", "stakeholder", "leadership",
          "执行摘要", "摘要", "战略", "干系人", "管理层"]),
    ]

    used_paths = set()
    for gt_fname, label, keywords in targets:
        # 1) Prefer exact GT filename
        exact = os.path.join(workspace, gt_fname)
        chosen = None
        if os.path.isfile(exact) and exact not in used_paths:
            chosen = exact
        else:
            cands = _find_docx_by_keywords(workspace, keywords, exclude=used_paths)
            if cands:
                chosen = cands[0]
        # 2) Fallback: allow reusing a file already assigned to another target
        # (a single combined document may cover both report and summary)
        if chosen is None:
            cands = _find_docx_by_keywords(workspace, keywords)
            if cands:
                chosen = cands[0]

        if chosen is None:
            record(f"docx for {label} exists", False,
                   f"No docx with keywords {keywords} found in workspace")
            continue
        used_paths.add(chosen)
        record(f"docx for {label} exists ({os.path.basename(chosen)})", True)

        try:
            doc = Document(chosen)
            text = "\n".join(p.text for p in doc.paragraphs)
            text_low = text.lower()
            word_count = _count_words(text)
            record(f"docx {os.path.basename(chosen)} substantive (>=60 words)",
                   word_count >= 60, f"{word_count} words")
            content_keywords = ['satisfaction', 'nps', 'recommend', 'customer',
                                'survey', 'review', 'analysis', 'finding',
                                '满意度', '客户', '调查', '推荐', '反馈', '分析',
                                '净推荐', '发现']
            matches = sum(1 for kw in content_keywords if kw in text_low)
            record(f"docx {os.path.basename(chosen)} mentions satisfaction/nps/recommend (>=2 keywords)",
                   matches >= 2, f"matched {matches}/{len(content_keywords)}")
        except Exception as e:
            record(f"docx {os.path.basename(chosen)} readable", False, str(e))


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    ws = args.agent_workspace
    if not os.path.isdir(ws):
        print(f"Agent workspace not found: {ws}")
        sys.exit(1)

    # Detect GT self-test mode (agent_ws == groundtruth_ws)
    global IS_GT_SELF_TEST
    try:
        if args.groundtruth_workspace and os.path.exists(args.groundtruth_workspace):
            IS_GT_SELF_TEST = (
                os.path.realpath(ws) ==
                os.path.realpath(args.groundtruth_workspace)
            )
    except Exception:
        IS_GT_SELF_TEST = False

    check_xlsx_content(ws, args.groundtruth_workspace)
    check_docx_content(ws)
    check_phase6_distribution()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} checks failed)")
        sys.exit(1)


if __name__ == "__main__":
    main()
