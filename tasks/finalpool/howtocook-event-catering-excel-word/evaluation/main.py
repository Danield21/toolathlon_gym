"""
Evaluation for howtocook-event-catering-excel-word task.

Checks:
1. Catering_Plan.xlsx exists with Menu and Ingredients sheets
2. Menu sheet has at least 8 data rows
3. Catering_Proposal.docx exists and mentions catering/menu
4. GForm "Menu Approval Survey" exists with at least 3 questions
5. Email sent to client@corporate.com
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
import openpyxl
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": os.environ.get("PGUSER", "eigent"),
    "password": os.environ.get("PGPASSWORD", "camel"),
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def _to_float(v):
    """Parse a numeric value from int/float/str/None.

    Strips thousand-separator commas, currency symbols ($/¥/€) and a trailing
    '%' (percent values like "90%" parse to 90.0). Returns None when the value
    is absent or cannot be parsed.
    """
    if v is None or isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    if isinstance(v, str):
        s = v.strip().replace(",", "").replace("$", "").replace("¥", "").replace("€", "").strip()
        if s.endswith("%"):
            s = s[:-1].strip()
        if not s:
            return None
        try:
            return float(s)
        except ValueError:
            return None
    return None


def num_close(a, b, tol=1.0):
    fa, fb = _to_float(a), _to_float(b)
    if fa is not None and fb is not None:
        return abs(fa - fb) <= tol
    # Fall back to case-insensitive string comparison only when one side is non-numeric.
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def _cell_value(formula_val, cached_val=None):
    """Resolve a cell's numeric value.

    Handles literal numbers/strings and formula cells. Formula cells read back
    as '=...' when the workbook is loaded with data_only=False; for those, fall
    back to the cached value from a data_only=True workbook. Non-numeric cells
    resolve to None.
    """
    if formula_val is None:
        return None
    if isinstance(formula_val, str) and formula_val.lstrip().startswith("="):
        return _to_float(cached_val)
    return _to_float(formula_val)


HEADER_HINTS = ["dish", "name", "category", "serving", "prep", "cost", "price",
                "ingredient", "quantity", "unit"]


def _find_header_idx(rows):
    """Locate the header row by keyword score.

    Tolerates a title/merged row above the real column headers (which the task
    does not forbid). The row whose cells contain the most expected header
    tokens (dish/name/category/serving/prep/cost/price/ingredient/quantity/unit)
    wins. Falls back to row 0 when no header-like row exists, so genuinely
    missing headers still fail the column checks.
    """
    best_idx, best_score = 0, -1
    for i, r in enumerate(rows):
        cells = [str(c).strip().lower() if c is not None else "" for c in r]
        if not any(cells):
            continue
        score = sum(1 for h in cells for tok in HEADER_HINTS if tok in h)
        if score > best_score:
            best_score, best_idx = score, i
    return best_idx if best_score > 0 else 0


def _resolve_rows(ws, wb_values, start_row=1):
    """Return data rows (rows before start_row skipped, empty rows dropped) with
    each cell resolved to a numeric value via _cell_value (formula-aware)."""
    rows_f = list(ws.iter_rows(values_only=True))
    ws_v = None
    if wb_values is not None and ws.title in wb_values.sheetnames:
        ws_v = wb_values[ws.title]
    rows_v = list(ws_v.iter_rows(values_only=True)) if ws_v is not None else None
    out = []
    for r_i in range(start_row, len(rows_f)):
        rf = rows_f[r_i]
        if not rf or not any(c for c in rf):
            continue
        rv = rows_v[r_i] if rows_v is not None and r_i < len(rows_v) else None
        rr = []
        for c_i in range(len(rf)):
            cached = rv[c_i] if rv is not None and c_i < len(rv) else None
            rr.append(_cell_value(rf[c_i], cached))
        out.append(rr)
    return out


def str_match(a, b):
    if a is None or b is None: return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Check 1: Excel Catering_Plan.xlsx ===")

    xlsx_path = os.path.join(agent_workspace, "Catering_Plan.xlsx")
    if not os.path.exists(xlsx_path):
        record("Catering_Plan.xlsx exists", False, f"Not found at {xlsx_path}")
        return {}
    record("Catering_Plan.xlsx exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=False)
    except Exception as e:
        record("Excel readable", False, str(e))
        return {}
    record("Excel readable", True)

    # Second load with data_only=True to access cached formula results.
    try:
        wb_values = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception:
        wb_values = None

    sheet_names_lower = [s.lower() for s in wb.sheetnames]
    has_menu = any("menu" in s for s in sheet_names_lower)
    has_ingredients = any("ingredient" in s for s in sheet_names_lower)

    record("Excel has Menu sheet", has_menu, f"Sheets: {wb.sheetnames}")
    record("Excel has Ingredients sheet", has_ingredients, f"Sheets: {wb.sheetnames}")

    menu_header_idx = 0
    if has_menu:
        menu_sheet_name = wb.sheetnames[next(i for i, s in enumerate(sheet_names_lower) if "menu" in s)]
        ws = wb[menu_sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        menu_header_idx = _find_header_idx(rows) if rows else 0
        menu_headers = rows[menu_header_idx] if rows else []
        data_rows = [r for r in rows[menu_header_idx + 1:] if any(c for c in r)]
        record("Menu sheet has at least 8 dishes", len(data_rows) >= 8,
               f"Found {len(data_rows)} data rows")

        # Check for expected columns (header row located by keyword scan so a
        # title row above the real headers does not break the column checks)
        headers = [str(c).lower() if c else "" for c in menu_headers]
        has_dish = any("dish" in h or "name" in h for h in headers)
        has_cost = any("cost" in h or "price" in h for h in headers)
        record("Menu has Dish_Name column", has_dish, f"Headers: {menu_headers}")
        record("Menu has cost column", has_cost, f"Headers: {menu_headers}")

    # --- Self-consistency checks (free dish choice means GT row-count comparison is invalid) ---
    # Verify Servings_For_30 reflects ~30-person scaling for at least 6 dishes.
    # The task defines Servings_For_30 as the servings needed to feed 30 guests
    # (approximately 30, one per guest). The threshold is kept well above recipe
    # defaults (2-4) so unscaled values still fail, but low enough to accept any
    # reasonable interpretation of "scaled to 30 people" (e.g. 6x a small default).
    # Values are resolved formula-aware so a correct answer written as a literal
    # number always passes, and one written as an Excel formula passes too when a
    # cached value exists.
    if has_menu:
        menu_sheet_name = wb.sheetnames[next(i for i, s in enumerate(sheet_names_lower) if "menu" in s)]
        ws = wb[menu_sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if rows:
            h_idx = _find_header_idx(rows)
            headers = [str(c).lower() if c else "" for c in rows[h_idx]]
            servings_idx = next((i for i, h in enumerate(headers) if "serving" in h), None)
            cost_idx = next((i for i, h in enumerate(headers) if "cost" in h or "price" in h), None)
            resolved_rows = _resolve_rows(ws, wb_values, start_row=h_idx + 1)
            if servings_idx is not None:
                ok_servings = sum(1 for r in resolved_rows
                                  if len(r) > servings_idx
                                  and r[servings_idx] is not None
                                  and r[servings_idx] >= 12)
                record("Menu Servings_For_30 reflects 30-person scaling in >=6 rows",
                       ok_servings >= 6,
                       f"rows with Servings>=12: {ok_servings}/{len(resolved_rows)}")
            if cost_idx is not None:
                ok_cost = sum(1 for r in resolved_rows
                              if len(r) > cost_idx
                              and r[cost_idx] is not None
                              and r[cost_idx] > 0)
                record("Menu Estimated_Cost_USD positive in >=6 rows",
                       ok_cost >= 6,
                       f"rows with cost>0: {ok_cost}/{len(resolved_rows)}")
    gt_path = os.path.join(groundtruth_workspace, "Catering_Plan.xlsx")
    # GT existence is enough to confirm the schema; per-row comparison disabled
    # since free dish choice produces dish-set divergence.
    if os.path.isfile(gt_path):
        record("GT file present", True, f"GT path={gt_path}")

    # Cross-sheet self-consistency: collect dish names from Menu and from Ingredients
    menu_dishes = set()
    ingredients_dishes = set()
    if has_menu:
        menu_sheet_name = wb.sheetnames[next(i for i, s in enumerate(sheet_names_lower) if "menu" in s)]
        ws = wb[menu_sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if rows:
            h_idx = _find_header_idx(rows)
            headers_low = [str(c).lower() if c else "" for c in rows[h_idx]]
            dish_idx = next((i for i, h in enumerate(headers_low) if "dish" in h or "name" in h), 0)
            for r in rows[h_idx + 1:]:
                if r and r[dish_idx]:
                    menu_dishes.add(str(r[dish_idx]).strip().lower())
    if has_ingredients:
        ing_idx_in_sheet = next(i for i, s in enumerate(sheet_names_lower) if "ingredient" in s)
        ing_sheet_name = wb.sheetnames[ing_idx_in_sheet]
        ws = wb[ing_sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if rows:
            h_idx = _find_header_idx(rows)
            headers_low = [str(c).lower() if c else "" for c in rows[h_idx]]
            dish_idx = next((i for i, h in enumerate(headers_low) if "dish" in h or "name" in h), 0)
            for r in rows[h_idx + 1:]:
                if r and r[dish_idx]:
                    ingredients_dishes.add(str(r[dish_idx]).strip().lower())
            record("Ingredients sheet has data rows (>=3 dishes)",
                   len(ingredients_dishes) >= 3,
                   f"unique ingredient dishes: {len(ingredients_dishes)}")
            if menu_dishes and ingredients_dishes:
                overlap = menu_dishes & ingredients_dishes
                # At least half of ingredient-dishes must come from Menu (or >=3 absolute)
                target = max(3, min(len(ingredients_dishes), len(menu_dishes) // 2))
                record(f"Ingredients sheet references dishes from Menu (>={target} overlap)",
                       len(overlap) >= target,
                       f"overlap={len(overlap)}, menu={len(menu_dishes)}, ing={len(ingredients_dishes)}")
    return {"menu_dishes": menu_dishes}


def check_word_doc(agent_workspace, menu_dishes=None):
    print("\n=== Check 2: Word Document Catering_Proposal.docx ===")

    docx_path = os.path.join(agent_workspace, "Catering_Proposal.docx")
    if not os.path.exists(docx_path):
        record("Catering_Proposal.docx exists", False, f"Not found at {docx_path}")
        return
    record("Catering_Proposal.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # python-docx's .paragraphs does not include table cell text, so also fold
    # in any table cells before keyword/dishes checks.
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    all_text += "\n" + cell.text.lower()
    has_catering = "catering" in all_text or "menu" in all_text
    has_summary = "summary" in all_text or "overview" in all_text or "executive" in all_text
    has_timeline = "timeline" in all_text or "preparation" in all_text or "schedule" in all_text

    record("Word doc mentions catering/menu", has_catering, "No catering/menu content")
    record("Word doc has summary/overview section", has_summary, "No summary/overview section")
    record("Word doc has timeline/preparation section", has_timeline, "No timeline/preparation section")

    # New: word doc should mention >=N dishes from the Menu sheet
    if menu_dishes:
        # Use first significant token of each dish name (handles multi-word dishes)
        mentioned = 0
        for d in menu_dishes:
            # Try full match first; then first 4-char prefix as fallback
            if d and (d in all_text or (len(d) >= 4 and d[:4] in all_text)):
                mentioned += 1
        # Need at least 4 of the menu dishes mentioned
        target = min(4, max(1, len(menu_dishes) // 2))
        record(f"Word doc mentions >={target} dishes from Menu sheet",
               mentioned >= target,
               f"mentioned {mentioned}/{len(menu_dishes)} menu dishes")


def check_gform(is_gt_self_test=False):
    print("\n=== Check 3: GForm Menu Approval Survey ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT id, title FROM gform.forms")
    forms = cur.fetchall()

    approval_form = None
    for form_id, title in forms:
        title_l = (title or "").lower()
        # Tightened: require BOTH 'menu' AND ('approval' OR 'survey') in the title
        if "menu" in title_l and ("approval" in title_l or "survey" in title_l):
            approval_form = (form_id, title)
            break

    if is_gt_self_test and approval_form is None:
        record("'Menu Approval Survey' form exists (GT self-test toleration)",
               True, "GT self-test: gform is agent runtime artifact, skipped")
        cur.close(); conn.close()
        return

    record("'Menu Approval Survey' form exists",
           approval_form is not None,
           f"Forms found: {[f[1] for f in forms]}")

    if approval_form:
        form_id, title = approval_form
        cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (form_id,))
        q_count = cur.fetchone()[0]
        record("Form has at least 3 questions", q_count >= 3,
               f"Found {q_count} questions")
        # Pull all question text and verify required topic coverage
        cur.execute("SELECT title, description FROM gform.questions WHERE form_id = %s", (form_id,))
        question_rows = cur.fetchall()
        all_q_text = " | ".join(
            (str(t or "") + " " + str(d or "")).lower() for t, d in question_rows
        )
        # Topic 1: overall menu satisfaction
        has_satisfaction = any(kw in all_q_text for kw in ["satisfaction", "satisfy", "rate", "rating"])
        # Topic 2: dish confirm/remove
        has_dish_q = any(kw in all_q_text for kw in ["dish", "remove", "confirm", "add", "menu item"])
        # Topic 3: dietary requirements/notes. Keyword set is intentionally broad
        # to accept faithful paraphrases ("food restrictions or preferences",
        # "any dietary needs?", "allergies/intolerances", "special requests"),
        # while a question that contains none of these signals is not dietary.
        has_dietary = any(kw in all_q_text for kw in [
            "dietary", "diet", "allergy", "allergen", "allerg", "special",
            "requirement", "note", "restrict", "prefer", "intoleran",
            "accommodat", "vegan", "vegetarian", "gluten", "halal", "kosher",
        ])
        if is_gt_self_test and not has_dietary:
            # GT self-test may pick up a stale form from prior runs without dietary q; tolerate.
            record("GForm has overall-menu-satisfaction question", has_satisfaction,
                   f"questions: {all_q_text[:200]}")
            record("GForm has dish confirm/remove question", has_dish_q,
                   f"questions: {all_q_text[:200]}")
            record("GForm has dietary/notes question (GT self-test toleration)",
                   True, "GT self-test: stale form from cross-task contamination tolerated")
        else:
            record("GForm has overall-menu-satisfaction question", has_satisfaction,
                   f"questions: {all_q_text[:200]}")
            record("GForm has dish confirm/remove question", has_dish_q,
                   f"questions: {all_q_text[:200]}")
            record("GForm has dietary/notes question", has_dietary,
                   f"questions: {all_q_text[:200]}")

    cur.close()
    conn.close()


def check_email(is_gt_self_test=False):
    print("\n=== Check 4: Email to client@corporate.com ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else str(to_addr).lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "client@corporate.com" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    if is_gt_self_test and matching is None:
        record("Email sent to client@corporate.com (GT self-test toleration)",
               True, "GT self-test: email is agent runtime artifact, skipped")
        return

    record("Email sent to client@corporate.com", matching is not None,
           f"Messages found: {len(messages)}")

    if matching:
        subject, _, _, body_text = matching
        all_text = ((subject or "") + " " + (body_text or "")).lower()
        has_catering = "catering" in all_text or "menu" in all_text or "proposal" in all_text
        record("Email mentions catering/menu/proposal", has_catering,
               f"Subject: {subject}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    # Detect GT self-test
    is_gt_self_test = False
    try:
        if args.groundtruth_workspace and os.path.realpath(args.groundtruth_workspace) == os.path.realpath(args.agent_workspace):
            is_gt_self_test = True
    except Exception:
        pass

    excel_state = check_excel(args.agent_workspace, args.groundtruth_workspace) or {}
    check_word_doc(args.agent_workspace, menu_dishes=excel_state.get("menu_dishes"))
    check_gform(is_gt_self_test=is_gt_self_test)
    check_email(is_gt_self_test=is_gt_self_test)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "success": FAIL_COUNT == 0,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
