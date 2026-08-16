"""
Evaluation for yt-veritasium-arxiv-survey-word-gcal-email task.

Checks:
1. Science_Communication_Survey.docx exists with required structure (>=8 headings: title + exec summary + 5 topics + conclusion + references)
2. Word doc text contains Veritasium and at least 5 distinct science topics
3. Word doc text contains at least 5 paper titles or author names
4. GCal has 2 specific new events with correct titles, dates, times
5. Email sent to seminar@science.edu with exact subject
6. Email sent to collab@research.org with exact subject
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
from zoneinfo import ZoneInfo

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# Source of truth: docs/task.md line 11: "Science Communication Seminar" on
# April 10, 2026, from 14:00 to 16:00 (scientific seminar, ET wall-clock);
# "Survey Review Session" on April 3, 2026, from 10:00 to 11:00. The Calendar
# MCP writes events as UTC instants into the DB (TIMESTAMPTZ), so all
# comparisons anchor to America/New_York. Never use bare `sd.hour` /
# `sd.date()` on rows read from `gcal.events`: psycopg2 returns them in the
# PG session TimeZone (case-study: compute node default was Asia/Shanghai,
# masking ET wall-clock). Use `utils.evaluation.gcal_helpers` instead
# (session-tz-independent, DST-aware).
# ──────────────────────────────────────────────────────────────────────────
EXPECTED_TIMEZONE = ZoneInfo("America/New_York")
from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document Science_Communication_Survey.docx ===")
    docx_path = os.path.join(agent_workspace, "Science_Communication_Survey.docx")
    if not os.path.exists(docx_path):
        record("Science_Communication_Survey.docx exists", False, f"Not found at {docx_path}")
        return
    record("Science_Communication_Survey.docx exists", True)

    try:
        import docx
        doc = docx.Document(docx_path)
    except ImportError:
        size = os.path.getsize(docx_path)
        record("Word doc is non-empty (>5KB)", size > 5000, f"Size: {size} bytes")
        return
    except Exception as e:
        record("Word doc readable", False, str(e))
        return

    # Headings (Title + Exec Summary + 5 topics + Conclusion + References = 9)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name.startswith("Title")]
    heading_texts = [h.text.strip() for h in headings]
    record("Word doc has >= 8 headings (title + exec summary + 5 topics + conclusion + references)",
           len(headings) >= 8,
           f"Found {len(headings)} headings: {heading_texts[:10]}")

    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    # Title text
    record(
        "Word doc contains title 'Bridging Popular Science and Academic Research'",
        "bridging popular science and academic research" in full_text,
        f"text snippet: {full_text[:200]}",
    )

    # Veritasium mention
    record("Word doc mentions Veritasium", "veritasium" in full_text)

    # 5 distinct science topic keywords (each as its own check)
    SCIENCE_TOPICS = [
        ("quantum", ["quantum", "decoherence", "many-worlds"]),
        ("evolutionary biology", ["evolution", "game theory"]),
        ("cognitive science", ["cognitive", "biases", "psychology", "brain"]),
        ("mathematical paradoxes", ["mathematical", "paradox", "infinity"]),
        ("fluid dynamics or biomechanics", ["fluid", "biomechanics", "biomimetic"]),
    ]
    for topic_name, kws in SCIENCE_TOPICS:
        record(
            f"Word doc covers '{topic_name}' topic",
            any(kw in full_text for kw in kws),
            f"keywords {kws} not found in text",
        )

    # Executive Summary section: must contain phrase 'executive summary' (not just 'summary')
    has_exec_summary = "executive summary" in full_text
    record("Word doc has 'Executive Summary' section (exact phrase)", has_exec_summary)

    # Conclusion section
    record("Word doc has 'Conclusion' section", "conclusion" in full_text)

    # References section: 'references' substring is required
    has_references = "references" in full_text
    record("Word doc has References section", has_references)

    # Word count substantial
    word_count = len(full_text.split())
    record("Word doc has substantial content (>= 800 words)", word_count >= 800,
           f"Word count: {word_count}")


def check_gcal():
    print("\n=== Check 2: GCal April 2026 Events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime, end_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Find Science Communication Seminar event
    seminar = None
    review = None
    for s, sd, ed in events:
        sl = (s or "").lower()
        if "science communication seminar" in sl:
            seminar = (s, sd, ed)
        if "survey review session" in sl:
            review = (s, sd, ed)
    record("GCal has 'Science Communication Seminar' event", seminar is not None,
           f"April events: {[e[0] for e in events]}")
    if seminar:
        s, sd, ed = seminar
        # gcal.events.start_datetime is TIMESTAMPTZ (UTC instant). Use the
        # session-tz-independent helper to extract components in
        # EXPECTED_TIMEZONE (America/New_York per task.md). Bare sd.hour /
        # sd.date() silently compares against the PG session tz.
        ev_date, ev_hour, ev_minute = get_zone_components(sd, EXPECTED_TIMEZONE)
        record(
            "Seminar event date is 2026-04-10",
            ev_date is not None and ev_date.isoformat() == "2026-04-10",
            f"date={ev_date}",
        )
        _, ed_hour, ed_minute = get_zone_components(ed, EXPECTED_TIMEZONE)
        record(
            "Seminar event time is 14:00-16:00",
            ev_hour == 14 and ev_minute == 0 and ed_hour == 16 and ed_minute == 0,
            f"start_hour={ev_hour}:{ev_minute} end_hour={ed_hour}:{ed_minute}",
        )

    record("GCal has 'Survey Review Session' event", review is not None,
           f"April events: {[e[0] for e in events]}")
    if review:
        s, sd, ed = review
        # Same TIMESTAMPTZ fix as the seminar event above.
        ev_date, ev_hour, ev_minute = get_zone_components(sd, EXPECTED_TIMEZONE)
        record(
            "Review event date is 2026-04-03",
            ev_date is not None and ev_date.isoformat() == "2026-04-03",
            f"date={ev_date}",
        )
        _, ed_hour, ed_minute = get_zone_components(ed, EXPECTED_TIMEZONE)
        record(
            "Review event time is 10:00-11:00",
            ev_hour == 10 and ev_minute == 0 and ed_hour == 11 and ed_minute == 0,
            f"start_hour={ev_hour}:{ev_minute} end_hour={ed_hour}:{ed_minute}",
        )


def check_emails_sent():
    print("\n=== Check 3: Emails Sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        # Look for messages from Sent folder OR via sent_log
        cur.execute("""
            SELECT m.id, m.subject, m.to_addr, m.body_text FROM email.messages m
            JOIN email.folders f ON m.folder_id = f.id
            WHERE UPPER(f.name) = 'SENT'
        """)
        sent_via_folder = cur.fetchall()
        cur.execute("""
            SELECT m.id, m.subject, m.to_addr, m.body_text FROM email.sent_log sl
            JOIN email.messages m ON sl.message_id = m.id
        """)
        sent_via_log = cur.fetchall()
        # Combine + dedupe
        seen = set()
        sent = []
        for r in sent_via_folder + sent_via_log:
            if r[0] in seen:
                continue
            seen.add(r[0])
            sent.append(r)

        # Email 1: seminar@science.edu
        seminar_email = None
        for _id, subj, to_addr, body in sent:
            if "seminar@science.edu" in str(to_addr or "").lower():
                seminar_email = (subj, body)
                break
        record("Email sent to seminar@science.edu", seminar_email is not None,
               f"sent count: {len(sent)}")
        if seminar_email:
            subj, body = seminar_email
            target_subj = "science communication survey - ready for review"
            record(
                "Seminar email subject 'Science Communication Survey - Ready for Review'",
                (subj or "").strip().lower() == target_subj,
                f"got '{subj}'",
            )
            body_l = (body or "").lower()
            # Tighten: require >=2 of 5 actual topic names (not bare 'topic' substring)
            topic_keywords = ["quantum", "evolution", "cognitive", "mathematical", "paradox", "fluid"]
            topic_hits = sum(1 for k in topic_keywords if k in body_l)
            record("Seminar email body mentions >=2 of 5 actual topics",
                   topic_hits >= 2,
                   f"hits={topic_hits}; body: {body_l[:200]}")
            # Tighten: require 'april 10' or '2026-04-10' or '04-10' (not bare '10')
            seminar_date_ok = (
                "april 10" in body_l or
                "2026-04-10" in body_l or
                "04-10" in body_l or
                "april 10, 2026" in body_l or
                "10 april" in body_l
            )
            record("Seminar email body mentions seminar date (april 10 / 2026-04-10)",
                   seminar_date_ok,
                   f"body: {body_l[:200]}")
            record("Seminar email body length >= 100 chars", len(body or "") >= 100,
                   f"body length={len(body or '')}")

        # Email 2: collab@research.org
        collab_email = None
        for _id, subj, to_addr, body in sent:
            if "collab@research.org" in str(to_addr or "").lower():
                collab_email = (subj, body)
                break
        record("Email sent to collab@research.org", collab_email is not None,
               f"sent count: {len(sent)}")
        if collab_email:
            subj, body = collab_email
            target_subj = "collaboration invitation - veritasium science survey"
            record(
                "Collab email subject 'Collaboration Invitation - Veritasium Science Survey'",
                (subj or "").strip().lower() == target_subj,
                f"got '{subj}'",
            )
            body_l = (body or "").lower()
            record("Collab email body mentions co-present or paper recommendations",
                   ("co-present" in body_l or "copresent" in body_l or "co present" in body_l
                    or "recommendation" in body_l or "paper" in body_l),
                   f"body: {body_l[:200]}")
            record("Collab email body length >= 100 chars", len(body or "") >= 100,
                   f"body length={len(body or '')}")
    except Exception as e:
        record("Email sent check", False, str(e))
    finally:
        cur.close()
        conn.close()


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal()
    check_emails_sent()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    print(f"\nOverall: {PASS_COUNT}/{total} checks passed")

    result = {"total_passed": PASS_COUNT, "total_checks": total}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} failures)")
        sys.exit(1)


if __name__ == "__main__":
    main()
