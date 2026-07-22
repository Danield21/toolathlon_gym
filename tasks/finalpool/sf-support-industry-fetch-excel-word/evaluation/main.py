"""Evaluation script for sf-support-industry-fetch-excel-word."""
import os
import argparse, json, os, sys
import openpyxl

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None: return default
        return float(str(val).replace(",", "").replace("%", "").replace("$", "").strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    # Check Support_Benchmark_Report.xlsx
    excel_path = os.path.join(agent_workspace, "Support_Benchmark_Report.xlsx")
    check("Support_Benchmark_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Support_Benchmark_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        if gt_wb:
            for sheet_name in gt_wb.sheetnames:
                check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    gt_ws = gt_wb[sheet_name]
                    # Check headers
                    gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                    headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                    for h in gt_headers:
                        if h:
                            check(f"{sheet_name} has {h} column", h in headers, f"headers: {headers[:10]}")
                    # Check row count - exact match for fixed sheets; ">=GT" for Action_Plan (task says "at least 3")
                    gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                    # Filter empty trailing rows in GT
                    gt_rows = [r for r in gt_rows if r and any(c is not None and str(c).strip() != "" for c in r)]
                    data_rows = list(ws.iter_rows(min_row=2, values_only=True))
                    data_rows = [r for r in data_rows if r and any(c is not None and str(c).strip() != "" for c in r)]
                    if sheet_name.lower() == "action_plan":
                        check(f"{sheet_name} has at least {len(gt_rows)} data rows", len(data_rows) >= len(gt_rows), f"got {len(data_rows)}")
                    else:
                        check(f"{sheet_name} has {len(gt_rows)} data rows", len(data_rows) == len(gt_rows), f"got {len(data_rows)}")

                    # Cell value comparison against groundtruth - iterate ALL rows, key by first column
                    header_map = {h: i for i, h in enumerate(headers)}
                    # Build lookup by first-column key
                    a_lookup = {}
                    for r in data_rows:
                        if r and r[0] is not None:
                            a_lookup[str(r[0]).strip().lower()] = r
                    for ri, gt_row in enumerate(gt_rows):
                        if not gt_row or gt_row[0] is None:
                            continue
                        key = str(gt_row[0]).strip().lower()
                        agent_row = a_lookup.get(key)
                        if agent_row is None:
                            check(f"{sheet_name} row '{key}' present", False, "row missing")
                            continue
                        for ci, gt_h in enumerate(gt_headers):
                            if not gt_h or ci >= len(gt_row):
                                continue
                            gv = gt_row[ci]
                            agent_ci = header_map.get(gt_h)
                            if agent_ci is None or agent_ci >= len(agent_row):
                                continue
                            av = agent_row[agent_ci]
                            gf = safe_float(gv)
                            af = safe_float(av)
                            if gf is not None and af is not None:
                                # Tighter tolerance: max(1.0, 5% rel) - drops 15% rel
                                tol = max(1.0, abs(gf) * 0.05)
                                check(f"{sheet_name} {key} {gt_h} ~{gf:.1f}",
                                      abs(gf - af) <= tol, f"got {af}")
                            elif gv is not None and av is not None:
                                gs = str(gv).strip().lower()
                                avs = str(av).strip().lower()
                                if gs:
                                    # Exact match required (not bidirectional substring)
                                    check(f"{sheet_name} {key} {gt_h} text",
                                          gs == avs,
                                          f"expected {gs[:50]}, got {avs[:50]}")

    # Check Benchmark_Analysis.docx
    docx_path = os.path.join(agent_workspace, "Benchmark_Analysis.docx")
    check("Benchmark_Analysis.docx exists", os.path.exists(docx_path))
    if os.path.exists(docx_path):
        from docx import Document
        doc = Document(docx_path)
        text = " ".join([p.text for p in doc.paragraphs])
        check("Benchmark_Analysis.docx has content", len(text) > 50, f"text length: {len(text)}")
        # Check headings match groundtruth - accept Heading 1-6 OR Title style
        def _is_heading_style(s):
            return s.startswith("Heading") or s == "Title"
        headings = [p.text.strip().lower() for p in doc.paragraphs if _is_heading_style(p.style.name)]
        gt_doc_path = os.path.join(groundtruth_workspace, "Benchmark_Analysis.docx")
        if os.path.exists(gt_doc_path):
            gt_doc = Document(gt_doc_path)
            gt_headings = [p.text.strip().lower() for p in gt_doc.paragraphs if _is_heading_style(p.style.name)]
            for gh in gt_headings:
                if gh:
                    # Require exact match (preferred) OR agent's heading equals GT
                    # without trailing decoration. Reject "GT-substring-of-agent" style
                    # like 'Performance Overview Section' matching 'Performance Overview'
                    # since that would also match 'Performance Overview Survey Section'.
                    def _norm(s):
                        return s.strip().lower().rstrip(":.")
                    gh_n = _norm(gh)
                    found = any(_norm(h) == gh_n for h in headings)
                    check(f"Benchmark_Analysis.docx has heading \"{gh[:40]}\"", found, f"agent headings: {headings[:5]}")
        else:
            check("Benchmark_Analysis.docx has headings", len(headings) >= 2, f"found {len(headings)} headings")

    # Check Python script exists with substantive comparison logic
    py_files = [f for f in os.listdir(agent_workspace) if f.endswith(".py")]
    check("benchmark_analyzer.py exists", "benchmark_analyzer.py" in py_files, f"found: {py_files}")
    analyzer_path = os.path.join(agent_workspace, "benchmark_analyzer.py")
    if os.path.exists(analyzer_path):
        try:
            with open(analyzer_path, "r", encoding="utf-8", errors="ignore") as f:
                analyzer_src = f.read().lower()
        except Exception:
            analyzer_src = ""
        # Must reference both inputs and exercise comparison logic
        has_inputs = (
            "support_metrics" in analyzer_src
            and "industry_benchmarks" in analyzer_src
        )
        has_compare = any(
            kw in analyzer_src
            for kw in ("compare", "benchmark", "variance", "diff")
        )
        has_output = (
            "benchmark_comparison" in analyzer_src
            or "json.dump" in analyzer_src
            or "open(" in analyzer_src
        )
        check(
            "benchmark_analyzer.py reads support_metrics + industry_benchmarks",
            has_inputs,
            f"src has support_metrics={('support_metrics' in analyzer_src)} industry_benchmarks={('industry_benchmarks' in analyzer_src)}",
        )
        check(
            "benchmark_analyzer.py performs comparison/benchmark logic",
            has_compare,
            "no compare/benchmark/variance/diff keyword found",
        )
        check(
            "benchmark_analyzer.py writes output (benchmark_comparison.json or JSON dump)",
            has_output,
            "no benchmark_comparison reference / json.dump / open()",
        )

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()