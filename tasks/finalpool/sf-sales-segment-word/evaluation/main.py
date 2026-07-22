"""Evaluation for sf-sales-segment-word."""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def parse_targets_from_pdf(pdf_path):
    """Parse Q4 targets from the PDF; falls back to hardcoded values if parse fails.

    Keeps the hardcoded values in sync with Q4_Targets.pdf by re-reading at runtime.
    """
    fallback = {"Consumer": 210000, "Enterprise": 200000, "Government": 170000, "SMB": 180000}
    if not pdf_path or not os.path.isfile(pdf_path):
        return fallback
    try:
        import fitz
        doc = fitz.open(pdf_path)
        text = ""
        for p in doc:
            text += p.get_text()
        doc.close()
        import re
        out = {}
        for seg in ("Enterprise", "SMB", "Consumer", "Government"):
            # Match "Segment\n<number>" or "Segment <number>" (with commas)
            m = re.search(rf"{seg}\s*[:\n]?\s*\$?\s*([0-9][0-9,]*)", text, re.IGNORECASE)
            if m:
                out[seg] = int(m.group(1).replace(",", ""))
        # Only accept if we have all 4
        if len(out) == 4:
            return out
    except Exception as e:
        print(f"  (PDF parse failed: {e}; using fallback)")
    return fallback


def check_word_doc(agent_workspace, initial_workspace=None):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Q4_Segment_Report.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "q4 2025" in p.text.lower() and "segment" in p.text.lower():
            has_heading = True
            break
    check("Document has Q4 2025 segment heading", has_heading)

    # Check table exists
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    # Get expected data from DB
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT c."SEGMENT", COUNT(o."ORDER_ID"), ROUND(SUM(o."TOTAL_AMOUNT")::numeric, 2)
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        JOIN sf_data."SALES_DW__PUBLIC__CUSTOMERS" c ON o."CUSTOMER_ID" = c."CUSTOMER_ID"
        WHERE o."ORDER_DATE" >= '2025-10-01' AND o."ORDER_DATE" <= '2025-12-31'
        GROUP BY c."SEGMENT"
        ORDER BY c."SEGMENT"
    """)
    expected = cur.fetchall()
    conn.close()

    # Parse targets from PDF (runtime) to avoid hardcode/PDF drift
    pdf_dir = initial_workspace
    if not pdf_dir:
        pdf_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "initial_workspace")
    pdf_path = os.path.join(pdf_dir, "Q4_Targets.pdf") if pdf_dir else None
    targets = parse_targets_from_pdf(pdf_path)

    # Check table content
    table = doc.tables[0]
    rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Table has 4 segment rows", len(rows) == 4, f"Got {len(rows)} rows")

    for segment, count, actual in expected:
        actual_f = float(actual)
        target = targets[segment]
        exp_variance = actual_f - target
        exp_pct = round(actual_f / target * 100, 2) if target else 0
        matched = None
        for r in rows:
            if r and r[0].lower() == segment.lower():
                matched = r
                break
        if matched:
            # Parse all numeric cells
            nums = []
            for cell in matched[1:]:
                try:
                    nums.append(float(str(cell).replace(",", "").replace("$", "").replace("%", "")))
                except (ValueError, AttributeError):
                    continue
            found_target = any(num_close(v, target, 50.0) for v in nums)
            found_actual = any(num_close(v, actual_f, 50.0) for v in nums)
            found_variance = any(num_close(v, exp_variance, 100.0) for v in nums)
            found_pct = any(num_close(v, exp_pct, 1.0) for v in nums)
            check(f"Segment {segment} Q4_Target (~{target}) in table", found_target,
                  f"Expected target {target}, nums={nums}")
            check(f"Segment {segment} actual value in table", found_actual,
                  f"Expected ~{actual_f}, nums={nums}")
            check(f"Segment {segment} Variance (~{exp_variance:.0f}) in table", found_variance,
                  f"Expected ~{exp_variance}, nums={nums}")
            check(f"Segment {segment} Achievement_Pct (~{exp_pct}%) in table", found_pct,
                  f"Expected ~{exp_pct}%, nums={nums}")
        else:
            check(f"Segment {segment} found in table", False)

    # Check summary paragraph
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    has_summary = "total" in full_text or "overall" in full_text
    check("Document has summary text", has_summary)
    # Check summary mentions total target and total actual
    total_target = sum(targets.values())
    total_actual = sum(float(e[2]) for e in expected)
    import re
    nums_in_text = re.findall(r"\d[\d,]*(?:\.\d+)?", full_text.replace(",", ""))
    found_total_target = any(abs(float(n) - total_target) < 500 for n in nums_in_text if n)
    found_total_actual = any(abs(float(n) - total_actual) < 500 for n in nums_in_text if n)
    check(f"Summary mentions total target (~{total_target})", found_total_target)
    check(f"Summary mentions total actual (~{total_actual:.0f})", found_total_actual)

    return True


def check_excel_groundtruth(agent_workspace, groundtruth_workspace):
    """Cross-check with groundtruth Excel for numeric accuracy."""
    print("\n=== Cross-checking with Groundtruth Data ===")

    gt_file = os.path.join(groundtruth_workspace, "Q4_Segment_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return

    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)
    gt_rows = list(gt_wb["Segment Comparison"].iter_rows(min_row=2, values_only=True))

    # Verify groundtruth values against DB
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT c."SEGMENT", ROUND(SUM(o."TOTAL_AMOUNT")::numeric, 2)
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        JOIN sf_data."SALES_DW__PUBLIC__CUSTOMERS" c ON o."CUSTOMER_ID" = c."CUSTOMER_ID"
        WHERE o."ORDER_DATE" >= '2025-10-01' AND o."ORDER_DATE" <= '2025-12-31'
        GROUP BY c."SEGMENT"
    """)
    db_data = {r[0]: float(r[1]) for r in cur.fetchall()}
    conn.close()

    for segment, actual in db_data.items():
        gt_match = None
        for r in gt_rows:
            if str_match(r[0], segment):
                gt_match = r
                break
        if gt_match:
            check(f"Groundtruth {segment} actual matches DB",
                  num_close(gt_match[2], actual, 1.0),
                  f"GT={gt_match[2]}, DB={actual}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("SF SALES SEGMENT WORD - EVALUATION")
    print("=" * 70)

    initial_ws = os.path.join(task_root, "initial_workspace")
    check_word_doc(args.agent_workspace, initial_workspace=initial_ws)
    check_excel_groundtruth(args.agent_workspace, gt_dir)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
