"""Evaluation for terminal-canvas-arxiv-word-excel-notion."""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


EXPECTED_SUBJECTS = [
    "applied analytics", "biochemistry", "creative computing",
    "data-driven design", "environmental economics", "foundations of finance",
    "global governance"
]


def _norm_header(s):
    return str(s or "").strip().lower().replace(" ", "_")


def check_excel(agent_workspace):
    print("\n=== Checking Excel Output ===")
    fpath = os.path.join(agent_workspace, "Curriculum_Gap_Analysis.xlsx")
    if not os.path.isfile(fpath):
        record("Excel file exists", False, f"Not found: {fpath}")
        return False
    record("Excel file exists", True)

    wb = openpyxl.load_workbook(fpath, data_only=True)

    # Sheet 1: Course_Topics
    ct_sheet = None
    for name in wb.sheetnames:
        if "course" in name.lower() and "topic" in name.lower():
            ct_sheet = name
            break
    if not ct_sheet:
        record("Course_Topics sheet exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Course_Topics sheet exists", True)
        ws = wb[ct_sheet]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        # Header validation
        headers = [_norm_header(c) for c in next(ws.iter_rows(values_only=True), [])]
        for h in ["course_subject", "current_topics"]:
            record(f"Course_Topics has '{h}' header", h in headers, f"got {headers}")
        # Query dynamic unique subject count from Canvas DB
        try:
            conn = psycopg2.connect(**DB_CONFIG)
            cur = conn.cursor()
            cur.execute("SELECT COUNT(DISTINCT regexp_replace(name, ' \\(.*\\)$', '')) FROM canvas.courses")
            expected_course_count = cur.fetchone()[0]
            cur.close(); conn.close()
        except Exception:
            expected_course_count = 7
        record(f"Course_Topics has exactly {expected_course_count} rows",
               len(rows) == expected_course_count, f"Got {len(rows)}")

    # Sheet 2: Research_Frontiers
    rf_sheet = None
    for name in wb.sheetnames:
        if "research" in name.lower() or "frontier" in name.lower():
            rf_sheet = name
            break
    if not rf_sheet:
        record("Research_Frontiers sheet exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Research_Frontiers sheet exists", True)
        ws = wb[rf_sheet]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        # Header validation
        headers = [_norm_header(c) for c in next(ws.iter_rows(values_only=True), [])]
        for h in ["paper_title", "primary_category", "relevance_score"]:
            record(f"Research_Frontiers has '{h}' header", h in headers, f"got {headers}")
        record("Research_Frontiers has >= 5 rows", len(rows) >= 5, f"Got {len(rows)}")
        # Relevance scores must be in [1, 10]
        try:
            score_idx = headers.index("relevance_score")
        except ValueError:
            score_idx = None
        if score_idx is not None:
            bad = []
            for r in rows:
                if r and len(r) > score_idx:
                    try:
                        v = float(r[score_idx])
                        if v < 1 or v > 10:
                            bad.append(v)
                    except (TypeError, ValueError):
                        bad.append(r[score_idx])
            record("Research_Frontiers Relevance_Score in [1,10]", not bad, f"out-of-range: {bad[:5]}")

    # Sheet 3: Gap_Matrix
    gm_sheet = None
    for name in wb.sheetnames:
        if "gap" in name.lower() and "matrix" in name.lower():
            gm_sheet = name
            break
    if not gm_sheet:
        record("Gap_Matrix sheet exists", False, f"Sheets: {wb.sheetnames}")
    else:
        record("Gap_Matrix sheet exists", True)
        ws = wb[gm_sheet]
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        # Header validation
        headers = [_norm_header(c) for c in next(ws.iter_rows(values_only=True), [])]
        for h in ["course_subject", "paper_title", "relevance_score", "gap_type"]:
            record(f"Gap_Matrix has '{h}' header", h in headers, f"got {headers}")
        record("Gap_Matrix has >= 20 rows", len(rows) >= 20, f"Got {len(rows)}")

        # Check gap types: must include all 3 (Covered, Partial, Gap)
        gap_types = set()
        try:
            type_idx = headers.index("gap_type")
        except ValueError:
            type_idx = 3
        for row in rows:
            if row and len(row) > type_idx and row[type_idx]:
                gap_types.add(str(row[type_idx]).strip().lower())
        record("Gap_Matrix has Covered type", "covered" in gap_types, f"Found types: {gap_types}")
        record("Gap_Matrix has Partial type", "partial" in gap_types, f"Found types: {gap_types}")
        record("Gap_Matrix has Gap type", "gap" in gap_types, f"Found types: {gap_types}")
        # Verify gap_type matches the rule (>=7 covered, 4..6 partial, <4 gap)
        try:
            score_idx = headers.index("relevance_score")
        except ValueError:
            score_idx = 2
        bad_classifications = []
        for row in rows:
            if not row or len(row) <= max(score_idx, type_idx):
                continue
            try:
                score = float(row[score_idx])
            except (TypeError, ValueError):
                continue
            actual_type = str(row[type_idx] or "").strip().lower()
            if score >= 7:
                expected = "covered"
            elif score >= 4:
                expected = "partial"
            else:
                expected = "gap"
            if actual_type != expected:
                bad_classifications.append((score, actual_type, expected))
        record(
            "Gap_Matrix gap_type classifications match scoring rule",
            not bad_classifications,
            f"mismatches: {bad_classifications[:5]}",
        )

    wb.close()
    return True


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    fpath = os.path.join(agent_workspace, "Curriculum_Enhancement_Proposal.docx")
    if not os.path.isfile(fpath):
        record("Word document exists", False, f"Not found: {fpath}")
        return False
    record("Word document exists", True)

    from docx import Document
    doc = Document(fpath)
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name.startswith("Title")]
    headings_l = [h.lower() for h in headings]
    # Title presence
    record(
        "Document title 'Curriculum Enhancement Proposal'",
        any("curriculum enhancement proposal" in h for h in headings_l)
        or "curriculum enhancement proposal" in full_text,
        f"headings: {headings[:6]}",
    )
    # Required sections: introduction, conclusion — require either as a Heading style or paragraph starting with the word
    def _has_section(name):
        # Heading-style match
        for h in headings_l:
            if name in h:
                return True
        # Paragraph-prefix match (paragraph starts with the word, allowing punctuation/spacing)
        import re as _re
        for p in doc.paragraphs:
            t = (p.text or "").strip().lower()
            if _re.match(rf"^{name}\b", t):
                return True
        return False

    record("Document has introduction section (heading or paragraph prefix)", _has_section("introduction"))
    record("Document has conclusion section (heading or paragraph prefix)", _has_section("conclusion"))
    record("Document mentions gap", "gap" in full_text)
    # Must mention at least 5 of 7 expected subjects (more rigorous than 3)
    mentioned = sum(1 for s in EXPECTED_SUBJECTS if s in full_text)
    record("Document mentions >= 5 of 7 course subjects", mentioned >= 5, f"Found {mentioned}")
    return True


def check_terminal_output(agent_workspace):
    print("\n=== Checking Terminal Output ===")
    fpath = os.path.join(agent_workspace, "curriculum_gap_output.txt")
    if not os.path.isfile(fpath):
        record("curriculum_gap_output.txt exists", False)
        return False
    record("curriculum_gap_output.txt exists", True)
    with open(fpath) as f:
        content = f.read().lower()
    # Must mention at least relevance and gap (not OR)
    record("Output mentions 'relevance'", "relevance" in content)
    record("Output mentions 'gap'", "gap" in content)
    # Must mention at least 3 expected subjects (the agent must list per-subject results)
    matched = sum(1 for s in EXPECTED_SUBJECTS if s in content)
    record(
        "curriculum_gap_output.txt mentions >= 3 course subjects",
        matched >= 3,
        f"found {matched}",
    )
    return True


def check_notion():
    print("\n=== Checking Notion ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, title FROM notion.databases
            WHERE title::text ILIKE '%curriculum%review%'
        """)
        dbs = cur.fetchall()
        record(
            "Notion Curriculum Review Tracker database exists",
            len(dbs) >= 1,
            f"Found {len(dbs)}",
        )
        # Determine expected course-subject count from Canvas
        try:
            cur.execute("SELECT COUNT(DISTINCT regexp_replace(name, ' \\(.*\\)$', '')) FROM canvas.courses")
            expected_count = cur.fetchone()[0]
        except Exception:
            expected_count = 7
        # Pages must be in the specific Curriculum Review database (not entire notion.pages)
        if dbs:
            db_id = dbs[0][0]
            cur.execute(
                "SELECT COUNT(*) FROM notion.pages WHERE parent::jsonb->>'database_id' = %s AND archived = false",
                (db_id,),
            )
            page_count = cur.fetchone()[0]
            record(
                f"Curriculum Review DB has exactly {expected_count} entries",
                page_count == expected_count,
                f"Found {page_count}, expected {expected_count}",
            )
            # Inspect properties for expected select fields
            cur.execute(
                "SELECT properties FROM notion.pages WHERE parent::jsonb->>'database_id' = %s AND archived = false",
                (db_id,),
            )
            props_rows = cur.fetchall()
            seen_status = set()
            seen_priority = set()
            for (props,) in props_rows:
                ps = str(props or "").lower()
                for s in ["not started", "in progress", "completed"]:
                    if s in ps:
                        seen_status.add(s)
                for p in ["high", "medium", "low"]:
                    if p in ps:
                        seen_priority.add(p)
            record(
                "Notion entries reference 'Not Started' status",
                "not started" in seen_status,
                f"statuses seen: {seen_status}",
            )
        cur.close()
        conn.close()
        return len(dbs) >= 1
    except Exception as e:
        record("Notion check", False, str(e))
        return False


def check_reverse_validation(workspace):
    print("\n=== Reverse Validation ===")
    # Check no duplicate pages in Notion (each course should appear only once)
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(*) FROM notion.databases
            WHERE title::text ILIKE '%curriculum%review%'
        """)
        count = cur.fetchone()[0]
        if count >= 1:
            cur.execute("""
                SELECT p.properties::text FROM notion.pages p
                JOIN notion.databases d ON p.parent->>'database_id' = d.id
                WHERE d.title::text ILIKE '%%curriculum%%review%%'
            """)
            pages = cur.fetchall()
            page_texts = [str(p[0]).lower() for p in pages]
            # Check for duplicate entries (same text appearing twice)
            if len(page_texts) > len(set(page_texts)):
                record("No duplicate pages in Notion database", False,
                       f"Found {len(page_texts)} pages but only {len(set(page_texts))} unique")
            else:
                record("No duplicate pages in Notion database", True)
        cur.close(); conn.close()
    except Exception as e:
        record("Reverse validation", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace)
    check_word(args.agent_workspace)
    check_terminal_output(args.agent_workspace)
    check_notion()
    check_reverse_validation(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
