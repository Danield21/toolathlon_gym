"""Evaluation for terminal-howtocook-pw-nutrition-gsheet-word.
Checks:
1. Google Sheets "Nutrition Dashboard" with Recipe Comparison and Daily Plan sheets
2. Wellness_Diet_Plan.docx with required sections
3. nutrition_calculator.py script exists
"""
import argparse
import json
import os
import sys
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def norm(s):
    return str(s or "").strip().lower().replace("_", " ")


def check_gsheet():
    print("\n=== Check 1: Google Sheets Nutrition Dashboard ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
    except Exception as e:
        check("DB connect", False, str(e))
        return

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()
    dashboard = None
    for ss_id, title in spreadsheets:
        if title and norm(title) == norm("Nutrition Dashboard"):
            dashboard = (ss_id, title)
            break

    check("'Nutrition Dashboard' spreadsheet exists (exact)", dashboard is not None,
          f"Spreadsheets found: {[s[1] for s in spreadsheets]}")

    if dashboard:
        ss_id = dashboard[0]
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
        sheet_list = cur.fetchall()
        sheet_norm = {norm(s[1]): s for s in sheet_list}
        check("Has at least 2 sheets", len(sheet_list) >= 2, f"Found {len(sheet_list)}: {[s[1] for s in sheet_list]}")

        rc = sheet_norm.get(norm("Recipe Comparison"))
        dp = sheet_norm.get(norm("Daily Plan"))
        check("'Recipe Comparison' sheet exists (exact)", rc is not None, f"sheets: {list(sheet_norm.keys())}")
        check("'Daily Plan' sheet exists (exact)", dp is not None, f"sheets: {list(sheet_norm.keys())}")

        # Recipe Comparison: 9+ rows, 8 columns header
        if rc:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, rc[0]))
            row_count = cur.fetchone()[0]
            check("Recipe Comparison has >= 9 recipe rows", row_count >= 9,
                  f"Found {row_count} data rows")
            # Header check
            cur.execute("""
                SELECT value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index = 0
                ORDER BY col_index
            """, (ss_id, rc[0]))
            header_cells = [norm(r[0]) for r in cur.fetchall()]
            for h in ["recipe name", "meal type", "calories", "protein g", "carbs g", "fat g", "fiber g", "sodium mg"]:
                check(f"Recipe Comparison has '{h}' header",
                      h in header_cells, f"got: {header_cells}")

        # Daily Plan: 4+ rows (3 meals + totals), 9 column header
        if dp:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, dp[0]))
            row_count = cur.fetchone()[0]
            check("Daily Plan has >= 4 rows (3 meals + totals)", row_count >= 4,
                  f"Found {row_count} data rows")
            cur.execute("""
                SELECT value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index = 0
                ORDER BY col_index
            """, (ss_id, dp[0]))
            header_cells = [norm(r[0]) for r in cur.fetchall()]
            for h in ["meal", "selected recipe", "calories", "protein g", "carbs g", "fat g", "fiber g", "sodium mg", "pct daily calories"]:
                check(f"Daily Plan has '{h}' header",
                      h in header_cells, f"got: {header_cells}")
            # Check for breakfast/lunch/dinner meal column entries
            if "meal" in header_cells:
                col_idx = header_cells.index("meal")
                cur.execute("""
                    SELECT value FROM gsheet.cells
                    WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0 AND col_index = %s
                    ORDER BY row_index
                """, (ss_id, dp[0], col_idx))
                meal_vals = [norm(r[0]) for r in cur.fetchall()]
                joined = " ".join(meal_vals)
                check("Daily Plan mentions breakfast", "breakfast" in joined, f"meals: {meal_vals}")
                check("Daily Plan mentions lunch", "lunch" in joined, f"meals: {meal_vals}")
                check("Daily Plan mentions dinner", "dinner" in joined, f"meals: {meal_vals}")

            # Validate Calories column has numeric values > 0 in meal rows
            # (avoid placeholder zeros). Sum should be in a reasonable range.
            def _col_idx(name):
                return header_cells.index(name) if name in header_cells else None

            cal_idx = _col_idx("calories")
            prot_idx = _col_idx("protein g")
            pdc_idx = _col_idx("pct daily calories")
            if cal_idx is not None and prot_idx is not None:
                cur.execute("""
                    SELECT row_index, col_index, value FROM gsheet.cells
                    WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
                """, (ss_id, dp[0]))
                cells = cur.fetchall()
                # Build per-row dict
                rows = {}
                for ri, ci, v in cells:
                    rows.setdefault(ri, {})[ci] = v
                meal_rows = []
                meal_idx = _col_idx("meal")
                for ri, vals in rows.items():
                    meal_label = norm(vals.get(meal_idx, ""))
                    if meal_label in ("breakfast", "lunch", "dinner"):
                        meal_rows.append(vals)
                # Each meal row must have positive numeric Calories and Protein_g
                bad_calories = 0
                bad_protein = 0
                total_cal = 0.0
                for vals in meal_rows:
                    try:
                        c = float(vals.get(cal_idx) or 0)
                        if c <= 0:
                            bad_calories += 1
                        total_cal += c
                    except (TypeError, ValueError):
                        bad_calories += 1
                    try:
                        p = float(vals.get(prot_idx) or 0)
                        if p <= 0:
                            bad_protein += 1
                    except (TypeError, ValueError):
                        bad_protein += 1
                check("Daily Plan: each meal row has Calories > 0",
                      bad_calories == 0,
                      f"{bad_calories} meal rows with zero/missing calories")
                check("Daily Plan: each meal row has Protein_g > 0",
                      bad_protein == 0,
                      f"{bad_protein} meal rows with zero/missing protein")
                if meal_rows:
                    # Daily total calories sanity range (1500-2500 kcal typical adult).
                    check("Daily Plan: total meal calories in 1500-2500 range",
                          1500 <= total_cal <= 2500,
                          f"total={total_cal}")

    cur.close()
    conn.close()


def check_word(workspace):
    print("\n=== Check 2: Wellness_Diet_Plan.docx ===")
    path = os.path.join(workspace, "Wellness_Diet_Plan.docx")
    if not os.path.exists(path):
        check("Word document exists", False, f"Not found at {path}")
        return
    check("Word document exists", True)

    try:
        from docx import Document
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        full_text_lower = full_text.lower()

        # Collect heading-style paragraph texts (style starts with 'Heading' or is 'Title')
        heading_texts = []
        for p in doc.paragraphs:
            sname = (p.style.name if p.style else "") or ""
            if sname.lower().startswith("heading") or sname.lower() == "title":
                heading_texts.append(p.text.strip().lower())

        # Title contains specific phrase
        check("Doc contains title 'Corporate Wellness Diet Plan'",
              "corporate wellness diet plan" in full_text_lower,
              f"Text: {full_text[:200]}")
        # Required sections — must appear as heading/title style paragraphs
        check("Doc has 'Recommended Daily Nutritional Standards' heading",
              any(("recommended daily" in h) or ("nutritional standards" in h) for h in heading_texts),
              f"Headings: {heading_texts[:6]}")
        check("Doc has 'Recipes Considered' heading",
              any(("recipes considered" in h) or ("recipes" in h and "consider" in h) for h in heading_texts),
              f"Headings: {heading_texts[:6]}")
        check("Doc has 'Recommended Daily Meal Plan' heading",
              any("meal plan" in h for h in heading_texts),
              f"Headings: {heading_texts[:6]}")
        check("Doc has 'Nutritional Gap Analysis' heading",
              any(("nutritional gap" in h) or ("gap analysis" in h and "nutri" in h) for h in heading_texts),
              f"Headings: {heading_texts[:6]}")
        check("Doc mentions breakfast", "breakfast" in full_text_lower)
        check("Doc mentions lunch", "lunch" in full_text_lower)
        check("Doc mentions dinner", "dinner" in full_text_lower)
        # Specific nutrient values - 2000 calories or 2000 kcal
        check("Doc mentions 2000 cal target", "2000" in full_text_lower and ("cal" in full_text_lower or "kcal" in full_text_lower),
              f"Text: {full_text_lower[:300]}")
        # Mention of protein, carbs, fat, fiber, sodium
        for nutrient in ["protein", "carb", "fat", "fiber", "sodium"]:
            check(f"Doc mentions {nutrient}", nutrient in full_text_lower,
                  f"Text: {full_text_lower[:200]}")
    except Exception as e:
        check("Word readable", False, str(e))


def check_script(workspace):
    print("\n=== Check 3: nutrition_calculator.py ===")
    path = os.path.join(workspace, "nutrition_calculator.py")
    check("nutrition_calculator.py exists", os.path.exists(path))
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                code = f.read().lower()
            check("nutrition_calculator.py mentions calories",
                  "calori" in code or "kcal" in code, f"size: {len(code)}")
            # Stronger: percentage-or-pct AND a daily-target reference.
            has_pct = "percent" in code or "pct" in code or "%" in code
            has_target = ("2000" in code) or ("daily" in code and "target" in code) or ("daily target" in code)
            check("nutrition_calculator.py performs percent calc with daily-target reference",
                  has_pct and has_target,
                  f"has_pct={has_pct}, has_target={has_target}, size: {len(code)}")
        except Exception as e:
            check("Script readable", False, str(e))


def check_reverse_validation(workspace):
    """Verify things that should NOT exist in output."""
    print("\n=== Reverse Validation ===")
    path = os.path.join(workspace, "Wellness_Diet_Plan.docx")
    if os.path.isfile(path):
        try:
            from docx import Document
            doc = Document(path)
            full_text = " ".join(p.text for p in doc.paragraphs).lower()
            check("Word doc has no placeholder text",
                  "[insert" not in full_text and "todo" not in full_text and "xxx" not in full_text,
                  "Found placeholder text in document")
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_gsheet()
    check_word(args.agent_workspace)
    check_script(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
