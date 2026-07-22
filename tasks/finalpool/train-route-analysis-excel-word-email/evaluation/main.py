"""
Evaluation for train-route-analysis-excel-word-email task.

Checks:
1. Route_Analysis.xlsx exists with Routes and Summary sheets, exact column order, exact 3 rows in expected sort order, correct values
2. Summary sheet has exact label/value pairs matching DB-derived ground truth
3. Route_Analysis_Report.docx exists with 4 required sections and recommends correct train
4. Email to logistics@company.com with required body content
"""
import json
import os
import re
import sys
from argparse import ArgumentParser

import psycopg2
import openpyxl


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, rel_tol=0.05, abs_tol=0.5):
    try:
        return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)
    except Exception:
        return False


# -------------------- Build expected from DB (single source of truth) --------------------

def build_expected():
    """
    Returns dict with:
      routes: list of dicts in departure-time ascending order
        {train_code, destination, depart_time, arrive_time, duration_min, first_class_price, second_class_price, available_second_class_seats}
      summary: dict {label: value}
    Falls back to known constants if DB unavailable.
    """
    # Known fallback (from DB snapshot)
    fallback_routes = [
        {"train_code": "G11", "destination_keywords": ["shanghai"],
         "depart_time": "07:00", "arrive_time": "12:31", "duration_min": 331,
         "first_class_price": 553.0, "second_class_price": 349.0, "second_class_seats": "有"},
        {"train_code": "G105", "destination_keywords": ["guangzhou"],
         "depart_time": "08:00", "arrive_time": "15:48", "duration_min": 468,
         "first_class_price": 862.5, "second_class_price": 545.5, "second_class_seats": "有"},
        {"train_code": "G1", "destination_keywords": ["shanghai"],
         "depart_time": "09:00", "arrive_time": "13:28", "duration_min": 268,
         "first_class_price": 553.0, "second_class_price": 349.0, "second_class_seats": "有"},
    ]
    fallback = {
        "routes": fallback_routes,
        "summary": {
            "Total_Routes": 3,
            "Earliest_Departure": "07:00",
            "Latest_Departure": "09:00",
            "Cheapest_Second_Class": "G11",  # tie with G1 at 349; pick by depart order
            "Fastest_Route": "G1",
            "Most_Available_Seats": "G1",
        }
    }
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # Beijing south = VNP
        cur.execute("""
            SELECT t.id, t.station_train_code, t.from_station_telecode, t.to_station_telecode,
                   t.start_time, t.arrive_time
            FROM train.trains t
            WHERE t.depart_date = '2026-03-10' AND t.from_station_telecode = 'VNP'
            ORDER BY t.start_time
        """)
        train_rows = cur.fetchall()
        if not train_rows:
            return fallback

        # Get destination station name
        cur.execute("SELECT station_code, station_name FROM train.stations")
        station_map = {r[0]: r[1] for r in cur.fetchall()}

        # Get seat info per train
        cur.execute("""
            SELECT t.id, ts.seat_short, ts.num, ts.price
            FROM train.trains t JOIN train.train_seats ts ON ts.train_id = t.id
            WHERE t.depart_date = '2026-03-10' AND t.from_station_telecode = 'VNP'
        """)
        seats_map = {}
        for tid, sshort, num, price in cur.fetchall():
            seats_map.setdefault(tid, {})[sshort] = (num, price)

        cur.close()
        conn.close()

        routes = []
        for tid, code, from_code, to_code, depart, arrive in train_rows:
            seats = seats_map.get(tid, {})
            zy = seats.get("zy", (None, None))  # first class
            ze = seats.get("ze", (None, None))  # second class
            # duration in minutes
            try:
                d_h, d_m = map(int, depart.split(":"))
                a_h, a_m = map(int, arrive.split(":"))
                duration = (a_h * 60 + a_m) - (d_h * 60 + d_m)
                if duration < 0:
                    duration += 24 * 60
            except Exception:
                duration = None
            dest_name = station_map.get(to_code, "")
            # English keyword guess from station mapping
            dest_kw = []
            for kw in ["shanghai", "guangzhou", "shenzhen", "hongqiao", "south", "tianjin", "nanjing", "hangzhou", "wuhan"]:
                if kw in dest_name.lower():
                    dest_kw.append(kw)
            # also add station_pinyin
            if to_code == "SHH":
                dest_kw = ["shanghai"]
            elif to_code == "GGQ":
                dest_kw = ["guangzhou"]
            routes.append({
                "train_code": code,
                "destination_keywords": dest_kw or [dest_name.lower()],
                "depart_time": depart,
                "arrive_time": arrive,
                "duration_min": duration,
                "first_class_price": float(zy[1]) if zy[1] is not None else None,
                "second_class_price": float(ze[1]) if ze[1] is not None else None,
                "second_class_seats": ze[0],
            })

        # Compute summary
        n = len(routes)
        depart_sorted = sorted(routes, key=lambda r: r["depart_time"])
        cheap = min((r for r in routes if r["second_class_price"] is not None),
                    key=lambda r: (r["second_class_price"], r["depart_time"]))
        fast = min((r for r in routes if r["duration_min"] is not None),
                   key=lambda r: (r["duration_min"], r["depart_time"]))

        # most_available_seats: numeric > "有" > others; tie-break by largest numeric or "有" first depart
        def avail_rank(r):
            v = r["second_class_seats"]
            try:
                n = int(v)
                return (2, n)
            except Exception:
                if v == "有":
                    return (1, 0)
                return (0, 0)
        best = max(routes, key=lambda r: (avail_rank(r), -depart_sorted.index(r)))
        most_av = best["train_code"]

        summary = {
            "Total_Routes": n,
            "Earliest_Departure": depart_sorted[0]["depart_time"],
            "Latest_Departure": depart_sorted[-1]["depart_time"],
            "Cheapest_Second_Class": cheap["train_code"],
            "Fastest_Route": fast["train_code"],
            "Most_Available_Seats": most_av,
        }
        return {"routes": depart_sorted, "summary": summary}
    except Exception as e:
        print(f"[expected] DB fetch failed: {e}; using fallback")
        return fallback


EXPECTED = build_expected()
EXPECTED_TOTAL = 22  # number of checks below; used for missing-file failure attribution


# -------------------- Excel checks --------------------

EXPECTED_COLUMNS = [
    "Train_Code", "Destination", "Depart_Time", "Arrive_Time",
    "Duration_Min", "First_Class_Price", "Second_Class_Price", "Available_Second_Class_Seats",
]


def norm(s):
    return str(s or "").strip().lower().replace("_", " ")


def check_excel(agent_workspace):
    print("\n=== Check 1: Route_Analysis.xlsx ===")

    xlsx_path = os.path.join(agent_workspace, "Route_Analysis.xlsx")
    if not os.path.exists(xlsx_path):
        record("Route_Analysis.xlsx exists", False, f"Not found at {xlsx_path}")
        # Force-fail downstream checks
        for k in ["Excel readable", "Routes sheet exists", "Summary sheet exists",
                  "Routes column order", "Routes has exactly 3 data rows",
                  "Routes are sorted by depart_time ascending",
                  "Each train_code present in correct row",
                  "Each row train pricing matches DB",
                  "Each row duration matches DB",
                  "Summary Total_Routes correct", "Summary Earliest/Latest correct",
                  "Summary Cheapest_Second_Class correct", "Summary Fastest_Route correct",
                  "Summary Most_Available_Seats correct"]:
            record(k, False, "skipped (no excel)")
        return
    record("Route_Analysis.xlsx exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path)
    except Exception as e:
        record("Excel readable", False, str(e))
        return
    record("Excel readable", True)

    sheet_names_lower = [s.lower() for s in wb.sheetnames]

    # Routes sheet
    if "routes" not in sheet_names_lower:
        record("Routes sheet exists", False, f"Sheets: {wb.sheetnames}")
        for k in ["Routes column order", "Routes has exactly 3 data rows",
                  "Routes are sorted by depart_time ascending",
                  "Each train_code present in correct row",
                  "Each row train pricing matches DB",
                  "Each row duration matches DB"]:
            record(k, False, "skipped (no Routes sheet)")
    else:
        record("Routes sheet exists", True)
        ws = wb[wb.sheetnames[sheet_names_lower.index("routes")]]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            record("Routes column order", False, "empty sheet")
        else:
            header = [norm(c) for c in rows[0]]
            expected_norm = [norm(c) for c in EXPECTED_COLUMNS]
            # match first len(expected_norm) columns
            ok_cols = header[:len(expected_norm)] == expected_norm
            record("Routes column order", ok_cols,
                   f"Got header={rows[0]}; expected {EXPECTED_COLUMNS}")
        data_rows = [r for r in rows[1:] if any(c is not None and str(c).strip() != "" for c in r)]
        record("Routes has exactly 3 data rows", len(data_rows) == 3, f"Found {len(data_rows)}")

        # Build dict by train_code
        codes = []
        per_code = {}
        for r in data_rows:
            if r and r[0]:
                codes.append(str(r[0]).strip())
                per_code[str(r[0]).strip()] = r

        expected_codes_in_order = [er["train_code"] for er in EXPECTED["routes"]]
        record("Routes are sorted by depart_time ascending",
               codes == expected_codes_in_order,
               f"Got {codes}; expected {expected_codes_in_order}")

        all_present = all(c in per_code for c in expected_codes_in_order)
        record("Each train_code present in correct row", all_present,
               f"Codes: {codes}")

        # Per-row content checks
        prices_ok = True
        prices_detail = []
        for er in EXPECTED["routes"]:
            row = per_code.get(er["train_code"])
            if not row:
                prices_ok = False
                prices_detail.append(f"{er['train_code']} missing")
                continue
            # row layout: train_code, destination, depart, arrive, duration, first_price, second_price, second_seats
            try:
                first_p = row[5]
                second_p = row[6]
                if er["first_class_price"] is not None and not num_close(first_p, er["first_class_price"]):
                    prices_ok = False
                    prices_detail.append(f"{er['train_code']} first {first_p} != {er['first_class_price']}")
                if er["second_class_price"] is not None and not num_close(second_p, er["second_class_price"]):
                    prices_ok = False
                    prices_detail.append(f"{er['train_code']} second {second_p} != {er['second_class_price']}")
            except Exception as e:
                prices_ok = False
                prices_detail.append(f"{er['train_code']} parse error {e}")
        record("Each row train pricing matches DB", prices_ok, "; ".join(prices_detail)[:280])

        durations_ok = True
        durations_detail = []
        for er in EXPECTED["routes"]:
            row = per_code.get(er["train_code"])
            if not row:
                durations_ok = False
                durations_detail.append(f"{er['train_code']} missing")
                continue
            try:
                dur = row[4]
                # duration must be integer minutes
                dval = int(float(dur))
                if er["duration_min"] is not None and abs(dval - er["duration_min"]) > 5:
                    durations_ok = False
                    durations_detail.append(f"{er['train_code']} duration {dur} != {er['duration_min']}")
            except Exception as e:
                durations_ok = False
                durations_detail.append(f"{er['train_code']} parse error {e}")
        record("Each row duration matches DB", durations_ok, "; ".join(durations_detail)[:280])

    # Summary sheet
    if "summary" not in sheet_names_lower:
        record("Summary sheet exists", False, f"Sheets: {wb.sheetnames}")
        for k in ["Summary Total_Routes correct", "Summary Earliest/Latest correct",
                  "Summary Cheapest_Second_Class correct", "Summary Fastest_Route correct",
                  "Summary Most_Available_Seats correct"]:
            record(k, False, "skipped (no Summary sheet)")
    else:
        record("Summary sheet exists", True)
        ws_s = wb[wb.sheetnames[sheet_names_lower.index("summary")]]
        rows_s = list(ws_s.iter_rows(values_only=True))
        # build label->value pairs
        pairs = {}
        for r in rows_s:
            if not r or len(r) < 2:
                continue
            lab = norm(r[0])
            if not lab or lab == "label":
                continue
            pairs[lab] = r[1]

        # Total_Routes exact value (allow text or int 3)
        total_v = pairs.get(norm("Total_Routes"))
        try:
            ok_total = int(float(total_v)) == EXPECTED["summary"]["Total_Routes"]
        except Exception:
            ok_total = str(total_v).strip() == str(EXPECTED["summary"]["Total_Routes"])
        record("Summary Total_Routes correct", ok_total, f"got {total_v}")

        e_dep = pairs.get(norm("Earliest_Departure"))
        l_dep = pairs.get(norm("Latest_Departure"))

        def time_match(cell, expected):
            if cell is None:
                return False
            s = str(cell).strip()
            # accept HH:MM or HH:MM:SS or datetime/time obj
            try:
                from datetime import time, datetime
                if hasattr(cell, "strftime"):
                    s = cell.strftime("%H:%M")
            except Exception:
                pass
            # strip seconds
            m = re.match(r"^(\d{1,2}):(\d{2})", s)
            if m:
                hh = int(m.group(1))
                mm = int(m.group(2))
                s = f"{hh:02d}:{mm:02d}"
            return s == expected

        ok_e = time_match(e_dep, EXPECTED["summary"]["Earliest_Departure"])
        ok_l = time_match(l_dep, EXPECTED["summary"]["Latest_Departure"])
        record("Summary Earliest/Latest correct", ok_e and ok_l,
               f"earliest={e_dep} latest={l_dep}")

        cheap = pairs.get(norm("Cheapest_Second_Class"))
        # G11 and G1 tied at 349; accept either if multiple equal cheapest exist
        cheap_candidates = set()
        min_p = min(r["second_class_price"] for r in EXPECTED["routes"]
                    if r["second_class_price"] is not None)
        for r in EXPECTED["routes"]:
            if r["second_class_price"] is not None and abs(r["second_class_price"] - min_p) < 0.5:
                cheap_candidates.add(r["train_code"])
        ok_cheap = str(cheap or "").strip().upper() in {c.upper() for c in cheap_candidates}
        record("Summary Cheapest_Second_Class correct", ok_cheap,
               f"got {cheap}, allowed {cheap_candidates}")

        fast = pairs.get(norm("Fastest_Route"))
        ok_fast = str(fast or "").strip().upper() == EXPECTED["summary"]["Fastest_Route"].upper()
        record("Summary Fastest_Route correct", ok_fast,
               f"got {fast} expected {EXPECTED['summary']['Fastest_Route']}")

        most = pairs.get(norm("Most_Available_Seats"))
        # Accept any train code with same max numeric availability or "有" tie-allowance
        ok_most = str(most or "").strip().upper() in {r["train_code"].upper() for r in EXPECTED["routes"]}
        # but require it to be among trains with maximal availability
        # Use rank: numeric > "有"; numeric > 0
        def rank(v):
            try:
                return int(v)
            except Exception:
                return 0
        ranks = [(r["train_code"], rank(r["second_class_seats"])) for r in EXPECTED["routes"]]
        max_rank = max(rk for _, rk in ranks)
        if max_rank > 0:
            allowed = {c for c, rk in ranks if rk == max_rank}
        else:
            # all "有" - any code is acceptable as long as it's one of the 3
            allowed = {r["train_code"] for r in EXPECTED["routes"]}
        ok_most = str(most or "").strip().upper() in {c.upper() for c in allowed}
        record("Summary Most_Available_Seats correct", ok_most,
               f"got {most} allowed {allowed}")


# -------------------- Word checks --------------------

def check_word(agent_workspace):
    print("\n=== Check 2: Route_Analysis_Report.docx ===")

    docx_path = os.path.join(agent_workspace, "Route_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Route_Analysis_Report.docx exists", False, f"Not found at {docx_path}")
        for k in ["Has Executive Summary section", "Has Route Details section",
                  "Has Recommendation section", "Has Cost Comparison section",
                  "Recommends G1 for speed (word-boundary)"]:
            record(k, False, "skipped (no docx)")
        return
    record("Route_Analysis_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word file readable", False, str(e))
        return

    # Collect heading paragraphs
    heading_texts = []
    all_texts = []
    for p in doc.paragraphs:
        if p.style and "heading" in p.style.name.lower():
            heading_texts.append(p.text.strip().lower())
        all_texts.append(p.text.strip())

    full_lower = "\n".join(t.lower() for t in all_texts)
    headings_lower = " | ".join(heading_texts)

    record("Has Executive Summary section", "executive summary" in headings_lower or "executive summary" in full_lower)
    record("Has Route Details section", "route details" in headings_lower or "route detail" in full_lower)
    record("Has Recommendation section", "recommendation" in headings_lower or "recommendation" in full_lower)
    # Cost Comparison must be the heading wording, not just 'cost'
    record("Has Cost Comparison section", "cost comparison" in headings_lower or "cost comparison" in full_lower)

    # Recommendation: must match G1 with word boundary (NOT G11/G105)
    fastest = EXPECTED["summary"]["Fastest_Route"]
    pat = re.compile(rf"\b{re.escape(fastest)}\b", re.IGNORECASE)
    rec_section_text = []
    in_rec = False
    for p in doc.paragraphs:
        s = p.text.strip()
        if p.style and "heading" in p.style.name.lower():
            in_rec = "recommendation" in s.lower()
            continue
        if in_rec and s:
            rec_section_text.append(s)
    rec_text = "\n".join(rec_section_text).lower()
    has_word_boundary_g1 = bool(pat.search("\n".join(rec_section_text)))
    has_speed_kw = any(k in rec_text for k in ["fastest", "speed", "shortest", "duration"])
    record("Recommends G1 for speed (word-boundary)",
           has_word_boundary_g1 and has_speed_kw,
           f"rec_text head: {rec_text[:200]}")


# -------------------- Email check --------------------

def check_email():
    print("\n=== Check 3: Email to logistics@company.com ===")

    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT to_addr, subject, body_text, body_html FROM email.messages
        """)
        messages = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Email DB query", False, str(e))
        for k in ["Email sent to logistics@company.com",
                  "Email subject is exactly 'Beijing Train Route Analysis 2026-03-10'",
                  "Email body has all 3 second class prices",
                  "Email body has recommended train code G1 (word boundary)",
                  "Email body has recommended departure/arrival times"]:
            record(k, False, "skipped")
        return

    target_email = "logistics@company.com"
    expected_subject = "Beijing Train Route Analysis 2026-03-10"

    matched = []
    for to_addr, subj, btxt, bhtml in messages:
        to_str = str(to_addr or "").lower()
        if target_email in to_str:
            matched.append((to_addr, subj or "", (btxt or "") + "\n" + (bhtml or "")))

    record("Email sent to logistics@company.com", len(matched) >= 1,
           f"Found {len(matched)} matching emails")

    if not matched:
        for k in ["Email subject is exactly 'Beijing Train Route Analysis 2026-03-10'",
                  "Email body has all 3 second class prices",
                  "Email body has recommended train code G1 (word boundary)",
                  "Email body has recommended departure/arrival times"]:
            record(k, False, "skipped (no email)")
        return

    # Pick first matching
    _, subj, body = matched[0]
    body_lower = body.lower()
    # task.md gives exact subject: "Beijing Train Route Analysis 2026-03-10"
    # Require exact match (case-insensitive, whitespace-collapsed)
    subj_norm = re.sub(r"\s+", " ", subj.strip().lower())
    expected_subj_norm = re.sub(r"\s+", " ", expected_subject.strip().lower())
    ok_subj = (subj_norm == expected_subj_norm)
    record(f"Email subject is exactly '{expected_subject}'", ok_subj, f"got: {subj!r}")

    # Body must contain all 3 second class prices
    second_prices = [r["second_class_price"] for r in EXPECTED["routes"] if r["second_class_price"] is not None]
    # Extract numeric tokens from body
    nums_in_body = set()
    for m in re.finditer(r"(\d+(?:\.\d+)?)", body):
        try:
            nums_in_body.add(float(m.group(1)))
        except Exception:
            pass
    all_prices_present = True
    missing = []
    for sp in second_prices:
        ok = any(num_close(n, sp, rel_tol=0.01, abs_tol=0.5) for n in nums_in_body)
        if not ok:
            all_prices_present = False
            missing.append(sp)
    record("Email body has all 3 second class prices", all_prices_present,
           f"missing prices: {missing}; nums sample: {sorted(nums_in_body)[:20]}")

    fastest = EXPECTED["summary"]["Fastest_Route"]
    pat = re.compile(rf"\b{re.escape(fastest)}\b", re.IGNORECASE)
    record("Email body has recommended train code G1 (word boundary)",
           bool(pat.search(body)), f"body sample: {body[:200]}")

    # Recommended train depart/arrive
    rec_route = next((r for r in EXPECTED["routes"] if r["train_code"] == fastest), None)
    times_ok = False
    if rec_route:
        d = rec_route["depart_time"]
        a = rec_route["arrive_time"]
        # accept hh:mm format anywhere; allow seconds
        d_pat = re.compile(rf"\b{re.escape(d)}(?::\d{{2}})?\b")
        a_pat = re.compile(rf"\b{re.escape(a)}(?::\d{{2}})?\b")
        times_ok = bool(d_pat.search(body) and a_pat.search(body))
    record("Email body has recommended departure/arrival times", times_ok,
           f"need {rec_route and rec_route['depart_time']} and {rec_route and rec_route['arrive_time']}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print(f"[expected] routes order: {[r['train_code'] for r in EXPECTED['routes']]}")
    print(f"[expected] summary: {EXPECTED['summary']}")

    check_excel(args.agent_workspace)
    check_word(args.agent_workspace)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%) [FAIL_COUNT={FAIL_COUNT}]")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
        "fail_count": FAIL_COUNT,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
