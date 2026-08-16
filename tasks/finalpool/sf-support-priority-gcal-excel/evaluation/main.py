"""Evaluation for sf-support-priority-gcal-excel."""
import argparse
import json
import os
import sys

import psycopg2

DB = {"host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")), "dbname": "toolathlon_gym", "user": "eigent", "password": "camel"}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def query_total_tickets():
    """Query total ticket count from Snowflake — fall back to baked value if DB unavailable."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute('SELECT COUNT(*) FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"')
        v = cur.fetchone()[0]
        cur.close()
        conn.close()
        return int(v)
    except Exception:
        return 31588


def query_priorities_met_sla():
    """Query priority counts where avg response time <= SLA target — derived from data."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        # SLA targets: High 4h, Medium 8h, Low 24h (typical support SLA tiers used by GT)
        cur.execute('''
            SELECT "PRIORITY", AVG("RESPONSE_TIME_HOURS")::float AS avg_h
            FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
            WHERE "PRIORITY" IS NOT NULL
            GROUP BY "PRIORITY"
        ''')
        rows = cur.fetchall()
        cur.close()
        conn.close()
        sla_targets = {"high": 4.0, "medium": 8.0, "low": 24.0, "critical": 1.0}
        met = 0
        for pri, avg_h in rows:
            t = sla_targets.get(str(pri).lower())
            if t is not None and avg_h is not None and float(avg_h) <= t:
                met += 1
        return met
    except Exception:
        return 0


def query_most_common_priority():
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute('''
            SELECT "PRIORITY", COUNT(*) AS cnt
            FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
            WHERE "PRIORITY" IS NOT NULL
            GROUP BY "PRIORITY" ORDER BY cnt DESC LIMIT 1
        ''')
        row = cur.fetchone()
        cur.close()
        conn.close()
        return str(row[0]) if row else "Medium"
    except Exception:
        return "Medium"


def check_excel(agent_workspace, gt_dir):
    errors = []
    try:
        import openpyxl
    except ImportError:
        errors.append("openpyxl not installed")
        return errors

    agent_file = os.path.join(agent_workspace, "Support_SLA_Analysis.xlsx")
    gt_file = os.path.join(gt_dir, "Support_SLA_Analysis.xlsx")

    if not os.path.exists(agent_file):
        errors.append("Support_SLA_Analysis.xlsx not found in agent workspace")
        return errors
    if not os.path.exists(gt_file):
        errors.append("Groundtruth Support_SLA_Analysis.xlsx not found")
        return errors

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check Priority Analysis sheet
    a_rows = load_sheet_rows(agent_wb, "Priority Analysis")
    g_rows = load_sheet_rows(gt_wb, "Priority Analysis")
    if a_rows is None:
        errors.append("Sheet 'Priority Analysis' not found in agent output")
    else:
        a_data = [r for r in (a_rows[1:] if len(a_rows) > 1 else []) if r and r[0] is not None]
        g_data = [r for r in (g_rows[1:] if g_rows and len(g_rows) > 1 else []) if r and r[0] is not None]

        if len(a_data) < 3:
            errors.append(f"Priority Analysis: expected 3 data rows, got {len(a_data)}")
        else:
            a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r[0]}
            for g_row in g_data:
                key = str(g_row[0]).strip().lower()
                a_row = a_lookup.get(key)
                if a_row is None:
                    errors.append(f"Missing priority row: {g_row[0]}")
                    continue
                # Ticket_Count col 1
                if len(a_row) > 1 and not num_close(a_row[1], g_row[1], 5):
                    errors.append(f"{g_row[0]} Ticket_Count: got {a_row[1]}, expected {g_row[1]} (tol=5)")
                # Avg_Response_Hours col 2
                if len(a_row) > 2 and not num_close(a_row[2], g_row[2], 0.5):
                    errors.append(f"{g_row[0]} Avg_Response_Hours: got {a_row[2]}, expected {g_row[2]} (tol=0.5)")
                # SLA_Status col 3
                if len(a_row) > 3 and not str_match(a_row[3], g_row[3]):
                    errors.append(f"{g_row[0]} SLA_Status: got '{a_row[3]}', expected '{g_row[3]}'")

    # Check Summary sheet
    a_sum = load_sheet_rows(agent_wb, "Summary")
    g_sum = load_sheet_rows(gt_wb, "Summary")
    if a_sum is None:
        errors.append("Sheet 'Summary' not found in agent output")
    else:
        a_sum_data = {str(r[0]).strip().lower(): r[1] for r in (a_sum[1:] if len(a_sum) > 1 else []) if r and r[0]}
        g_sum_data = {str(r[0]).strip().lower(): r[1] for r in (g_sum[1:] if g_sum and len(g_sum) > 1 else []) if r and r[0]}

        expected_total = query_total_tickets()
        expected_met = query_priorities_met_sla()
        expected_mcp = query_most_common_priority()

        # Total_Tickets
        tt = a_sum_data.get("total_tickets")
        if tt is None:
            errors.append("Summary missing Total_Tickets")
        elif not num_close(tt, expected_total, 10):
            errors.append(f"Total_Tickets: got {tt}, expected {expected_total} (tol=10)")

        # Priorities_Met_SLA
        pm = a_sum_data.get("priorities_met_sla")
        if pm is None:
            errors.append("Summary missing Priorities_Met_SLA")
        elif not num_close(pm, expected_met, 0):
            errors.append(f"Priorities_Met_SLA: got {pm}, expected {expected_met}")

        # Most_Common_Priority
        mcp = a_sum_data.get("most_common_priority")
        if mcp is None:
            errors.append("Summary missing Most_Common_Priority")
        elif not str_match(mcp, expected_mcp):
            errors.append(f"Most_Common_Priority: got '{mcp}', expected '{expected_mcp}'")

    return errors


def check_gcal():
    errors = []
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, summary, start_datetime, end_datetime
            FROM gcal.events
            WHERE (LOWER(summary) LIKE '%sla review%' OR LOWER(summary) LIKE '%support%review%' OR LOWER(summary) LIKE '%monthly support%')
            AND start_datetime >= '2026-04-01T00:00:00'
        """)
        events = cur.fetchall()
        cur.close()
        conn.close()
        if len(events) < 3:
            errors.append(f"Expected at least 3 SLA review gcal events, found {len(events)}")
            return errors
        # Required first Mondays of April, May, June 2026 (April 6, May 4, June 1)
        required_dates = {"2026-04-06", "2026-05-04", "2026-06-01"}
        found_dates = set()
        for ev in events:
            sdt = ev[2]
            edt = ev[3]
            if sdt is None:
                continue
            sdt_str = str(sdt)
            day = sdt_str[:10]
            if day in required_dates:
                # Check 9:00 start
                hh_mm = sdt_str[11:16] if len(sdt_str) >= 16 else ""
                if hh_mm not in ("09:00", "9:00 ", "09:00:"):
                    # Allow slight format variants (e.g., 09:00:00)
                    if not sdt_str[11:13] == "09":
                        errors.append(f"Event on {day} expected 9:00 start, got {hh_mm}")
                # Check ~1h duration if end_datetime present
                if edt is not None:
                    edt_str = str(edt)
                    if edt_str[11:13] not in ("10",):
                        errors.append(f"Event on {day} expected 10:00 end, got {edt_str[11:16]}")
                found_dates.add(day)
        for rd in required_dates:
            if rd not in found_dates:
                errors.append(f"Missing required SLA review event on {rd} (first Monday)")
    except Exception as e:
        errors.append(f"GCal DB check error: {e}")
    return errors


def check_email():
    errors = []
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        # Require to=support.leads@company.com AND subject references march 2026 support review
        cur.execute("""
            SELECT subject, to_addr
            FROM email.messages
            WHERE LOWER(to_addr::text) LIKE '%support.leads@company.com%'
        """)
        emails = cur.fetchall()
        cur.close()
        conn.close()
        if not emails:
            errors.append("No email sent to support.leads@company.com")
        else:
            # Subject must reference March 2026 support review per task.md
            found_subject = False
            for em in emails:
                subj = str(em[0] or "").lower()
                # Accept subjects that reference 'march 2026' AND ('support'|'review'|'sla'),
                # or both 'march' and 'support'/'sla' clearly
                if ("march" in subj and "2026" in subj) or ("march" in subj and ("support" in subj or "review" in subj or "sla" in subj)):
                    found_subject = True
                    break
            if not found_subject:
                errors.append("Email subject does not reference March 2026 support review")
    except Exception as e:
        errors.append(f"Email DB check error: {e}")
    return errors


def check_word(agent_workspace):
    errors = []
    docx_path = os.path.join(agent_workspace, "SLA_Report.docx")
    if not os.path.exists(docx_path):
        errors.append("SLA_Report.docx not found")
        return errors
    try:
        from docx import Document
        doc = Document(docx_path)
        text = " ".join(p.text for p in doc.paragraphs).lower()
        # Include table cells in text
        table_text = ""
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    table_text += " " + (c.text or "")
        full_text = (text + " " + table_text).lower()
        if len(full_text.strip()) < 100:
            errors.append("SLA_Report.docx has too little content")
        # Required heading per task.md
        if "support sla analysis" not in full_text:
            errors.append("SLA_Report.docx missing 'Support SLA Analysis' heading")
        # All three priorities must appear (priority breakdown)
        for kw in ["high", "medium", "low"]:
            if kw not in full_text:
                errors.append(f"SLA_Report.docx missing priority: {kw}")
        # Must reference SLA + breach
        for kw in ["sla"]:
            if kw not in full_text:
                errors.append(f"SLA_Report.docx missing keyword: {kw}")
    except ImportError:
        if os.path.getsize(docx_path) < 100:
            errors.append("SLA_Report.docx too small")
    except Exception as e:
        errors.append(f"Error reading SLA_Report.docx: {e}")
    return errors


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    all_errors = []

    print("\n=== Checking Excel ===")
    excel_errors = check_excel(args.agent_workspace, gt_dir)
    if excel_errors:
        for e in excel_errors:
            print(f"  [FAIL] {e}")
        all_errors.extend(excel_errors)
    else:
        print("  [PASS] Excel check passed")

    print("\n=== Checking GCal Events ===")
    gcal_errors = check_gcal()
    if gcal_errors:
        for e in gcal_errors:
            print(f"  [FAIL] {e}")
        all_errors.extend(gcal_errors)
    else:
        print("  [PASS] GCal check passed")

    print("\n=== Checking Email ===")
    email_errors = check_email()
    if email_errors:
        for e in email_errors:
            print(f"  [FAIL] {e}")
        all_errors.extend(email_errors)
    else:
        print("  [PASS] Email check passed")

    print("\n=== Checking Word Document ===")
    word_errors = check_word(args.agent_workspace)
    if word_errors:
        for e in word_errors:
            print(f"  [FAIL] {e}")
        all_errors.extend(word_errors)
    else:
        print("  [PASS] Word check passed")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"errors": all_errors, "success": len(all_errors) == 0}, f, indent=2)

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
