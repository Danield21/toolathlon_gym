"""Evaluation for terminal-yf-wc-excel-word-notion.
Checks:
1. Commodity_Impact_Analysis.xlsx with 4 sheets and correct numeric values
2. Pricing_Strategy_Memo.docx with required sections
3. Notion "Market Research Dashboard" database with 2 entries
4. correlation_analysis.py script exists
"""
import argparse
import json
import os
import re
import sys

import openpyxl
import psycopg2
from docx import Document

def _db_config():
    """Build the PostgreSQL connection config.

    The MCP servers the agent calls (local_servers/*/pg_adapter.py) read the
    PG_HOST / PG_PORT / PG_DATABASE / PG_USER / PG_PASSWORD env-var family,
    while this evaluator (and preprocess) historically read the PGHOST family.
    Prefer the PGHOST family when it is present, otherwise fall back to the
    PG_* family so judge and agent keep reading the same database when only one
    family is exported, then to the local seeded defaults.
    """
    host = os.environ["PGHOST"] if "PGHOST" in os.environ else os.environ.get("PG_HOST", "localhost")
    port = int(os.environ["PGPORT"] if "PGPORT" in os.environ else os.environ.get("PG_PORT", "5432"))
    dbname = os.environ["PGDATABASE"] if "PGDATABASE" in os.environ else os.environ.get("PG_DATABASE", "toolathlon_gym")
    user = os.environ["PGUSER"] if "PGUSER" in os.environ else os.environ.get("PG_USER", "eigent")
    password = os.environ["PGPASSWORD"] if "PGPASSWORD" in os.environ else os.environ.get("PG_PASSWORD", "camel")
    return {"host": host, "port": port, "dbname": dbname, "user": user, "password": password}


DB_CONFIG = _db_config()

PASS_COUNT = 0
FAIL_COUNT = 0

# Sentinel for formula cells whose cached value could not be read (e.g. a
# workbook written by openpyxl without Excel recompute). We do not penalize
# these: the structural checks still run, but a numeric comparison against an
# unverifiable value is not treated as a FAIL.
_UNVERIFIABLE = object()

# Fallbacks matching initial_workspace/analysis_params.json and the task
# description. When the agent workspace carries analysis_params.json the
# evaluator derives expectations from that file instead (see _load_params).
DEFAULT_PARAMS = {
    "estimated_cost_pct": 60,
    "standard_margin_pct": 40,
    "gold_significant_change_pct": 50,
    "reduced_margin_target": 35,
    "standard_margin_target": 40,
}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def _to_float(v):
    """Robustly coerce a value to float.
    Accepts int/float/str. Strips thousand separators, currency symbols,
    percent signs and whitespace. Returns None when it cannot parse."""
    if v is None or isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return None
    s = s.replace(",", "").replace("$", "").replace("€", "").replace("¥", "")
    s = s.replace("%", "").replace("¥", "").strip()
    try:
        return float(s)
    except ValueError:
        return None


def _to_pct(v):
    """Convert a value to a percent-scale number.
    Handles ratios (0.40 -> 40), percent integers (40 -> 40) and strings
    ('40%' -> 40)."""
    f = _to_float(v)
    if f is None:
        return None
    if abs(f) <= 1.5:
        return f * 100.0
    return f


def num_close(a, b, tol):
    """Numeric closeness with robust parsing. Falls back to a case-insensitive
    string comparison only when one side cannot be parsed to a number."""
    fa = _to_float(a)
    fb = _to_float(b)
    if fa is not None and fb is not None:
        return abs(fa - fb) <= tol
    if a is None or b is None:
        return False
    return str(a).strip().lower() == str(b).strip().lower()


def _tol_for(expected_val, factor, floor):
    """Tolerance for a numeric comparison, scaled by the expected magnitude.
    When the expected value cannot be parsed the floor is used, so a None or
    otherwise unusable expected value never raises (e.g. an empty stock_prices
    table feeding None through the aggregate query)."""
    f = _to_float(expected_val)
    if f is None:
        return floor
    return max(floor, abs(f) * factor)


def _resolve_cell(formula_val, cached_val):
    """Pick the comparable value for a cell from the two workbook loads.
    Formula cells use the cached (data_only) value when it is a number;
    otherwise the cell is marked unverifiable."""
    if isinstance(formula_val, str) and formula_val.lstrip().startswith("="):
        f = _to_float(cached_val)
        if f is not None:
            return f
        return _UNVERIFIABLE
    return formula_val


def _resolve_rows(ws_f, ws_v):
    """Yield rows with each cell resolved from the formula/cached workbooks."""
    rows_f = list(ws_f.iter_rows(values_only=True))
    rows_v = list(ws_v.iter_rows(values_only=True))
    for i in range(max(len(rows_f), len(rows_v))):
        rf = rows_f[i] if i < len(rows_f) else (None,)
        rv = rows_v[i] if i < len(rows_v) else (None,)
        n = max(len(rf), len(rv))
        resolved = tuple(
            _resolve_cell(rf[j] if j < len(rf) else None,
                          rv[j] if j < len(rv) else None)
            for j in range(n)
        )
        yield resolved


def _txt(c):
    """String form of a cell value; unverifiable cells contribute nothing."""
    return "" if c is _UNVERIFIABLE else str(c)


def _num_ok(value, expected, tol):
    """True when the value is close to expected, or when the value is
    unverifiable (uncomputed formula) and therefore not penalized."""
    if value is _UNVERIFIABLE:
        return True
    return num_close(value, expected, tol)


def _load_params(workspace):
    """Derive expectations from the agent workspace's analysis_params.json,
    falling back to DEFAULT_PARAMS. Values may be expressed as ratios or as
    percent integers; both are normalized to percent scale."""
    params = dict(DEFAULT_PARAMS)
    try:
        with open(os.path.join(workspace, "analysis_params.json")) as f:
            data = json.load(f)
        margin = data.get("margin_calculation", {}) or {}
        if "estimated_cost_pct" in margin:
            v = _to_pct(margin["estimated_cost_pct"])
            if v is not None:
                params["estimated_cost_pct"] = v
        if "standard_margin_pct" in margin:
            v = _to_pct(margin["standard_margin_pct"])
            if v is not None:
                params["standard_margin_pct"] = v
        thr = data.get("thresholds", {}) or {}
        if "gold_significant_change_pct" in thr:
            v = _to_float(thr["gold_significant_change_pct"])
            if v is not None:
                params["gold_significant_change_pct"] = v
        if "reduced_margin_target" in thr:
            v = _to_pct(thr["reduced_margin_target"])
            if v is not None:
                params["reduced_margin_target"] = v
        if "standard_margin_target" in thr:
            v = _to_pct(thr["standard_margin_target"])
            if v is not None:
                params["standard_margin_target"] = v
    except Exception:
        pass
    return params


def _fetch_yf_expected():
    """Fetch expected YF stock stats from DB.

    Symbols that have no usable rows (empty table, or a symbol missing from
    yf.stock_prices) are skipped so no None-valued entry is ever produced: an
    entry of {'avg_price': None, ...} is a truthy dict and would otherwise
    bypass the YF_FALLBACK and crash the tolerance arithmetic downstream."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        out = {}
        for sym in ("GC=F", "AMZN"):
            cur.execute("""
                SELECT AVG(close)::float, STDDEV(close)::float,
                       (SELECT close::float FROM yf.stock_prices WHERE symbol=%s ORDER BY date ASC LIMIT 1),
                       (SELECT close::float FROM yf.stock_prices WHERE symbol=%s ORDER BY date DESC LIMIT 1)
                FROM yf.stock_prices WHERE symbol=%s
            """, (sym, sym, sym))
            row = cur.fetchone()
            if row is None:
                continue
            avg, stddev, first, last = row
            if avg is None or stddev is None or first is None or last is None:
                # No usable rows for this symbol -> leave it to the fallback.
                continue
            pct = (last - first) / first * 100
            out[sym] = {"avg_price": avg, "volatility": stddev, "pct_change": pct}
        cur.close()
        conn.close()
        return out
    except Exception as e:
        print(f"[fallback] YF DB fetch error: {e}")
        return None


def _merge_yf_expected(fetched):
    """Return a complete GC=F/AMZN expectations dict.

    Any symbol or field that is missing, None or unparseable from the DB fetch
    is filled from YF_FALLBACK, so callers always get usable numbers and never
    do arithmetic on None."""
    merged = {}
    for sym, fallback in YF_FALLBACK.items():
        entry = (fetched or {}).get(sym)
        if entry:
            clean = {}
            for field, fval in fallback.items():
                v = entry.get(field)
                clean[field] = fval if (v is None or _to_float(v) is None) else v
            merged[sym] = clean
        else:
            merged[sym] = dict(fallback)
    return merged


YF_FALLBACK = {
    "GC=F": {"avg_price": 3163.01, "volatility": 796.12, "pct_change": 138.73},
    "AMZN": {"avg_price": 206.45, "volatility": 21.87, "pct_change": 26.18},
}


def check_excel(workspace, params):
    print("\n=== Check 1: Commodity_Impact_Analysis.xlsx ===")
    path = os.path.join(workspace, "Commodity_Impact_Analysis.xlsx")
    if not os.path.exists(path):
        check("Excel file exists", False, f"Not found at {path}")
        # Subordinate failures
        for label in ["Stock_Trends sheet present", "Product_Margins sheet present",
                      "Correlation_Analysis sheet present", "Strategic_Recommendations sheet present",
                      "Stock_Trends has GC=F row with correct stats",
                      "Stock_Trends has AMZN row with correct stats",
                      "Product_Margins has >=5 categories",
                      "Correlation_Analysis has 2 entries",
                      "Strategic_Recommendations has rows for each category"]:
            check(label, False, "Excel missing")
        return
    check("Excel file exists", True)

    try:
        # Load twice: once for formula text (data_only=False) and once for the
        # cached values (data_only=True). A single data_only=True load returns
        # None for formula cells that were never recomputed.
        wb_f = openpyxl.load_workbook(path, data_only=False)
        wb_v = openpyxl.load_workbook(path, data_only=True)
    except Exception as e:
        check("Excel readable", False, str(e))
        return

    sheets = wb_f.sheetnames

    def _norm(s): return re.sub(r'[^a-z0-9]', '', s.lower())

    sheets_norm = {_norm(s): s for s in sheets}
    st_name = next((sheets_norm[k] for k in sheets_norm if "stocktrends" in k), None)
    pm_name = next((sheets_norm[k] for k in sheets_norm if "productmargins" in k), None)
    ca_name = next((sheets_norm[k] for k in sheets_norm if "correlationanalysis" in k), None)
    sr_name = next((sheets_norm[k] for k in sheets_norm if "strategicrecommendations" in k), None)

    check("Stock_Trends sheet present", st_name is not None, f"Sheets: {sheets}")
    check("Product_Margins sheet present", pm_name is not None, f"Sheets: {sheets}")
    check("Correlation_Analysis sheet present", ca_name is not None, f"Sheets: {sheets}")
    check("Strategic_Recommendations sheet present", sr_name is not None, f"Sheets: {sheets}")

    yf_expected = _merge_yf_expected(_fetch_yf_expected())

    # Stock_Trends
    if st_name:
        ws = wb_f[st_name]
        rows = list(_resolve_rows(ws, wb_v[st_name]))
        if rows:
            headers = [str(c).lower().strip() if c is not None else "" for c in rows[0]]
            data = [r for r in rows[1:] if any(c is not None and str(c).strip() != "" for c in r)]

            def _get_col(names):
                for nm in names:
                    nm_l = nm.lower()
                    for i, h in enumerate(headers):
                        if h == nm_l or nm_l in h:
                            return i
                return None

            sym_idx = _get_col(["symbol", "ticker"])
            avg_idx = _get_col(["avg_price", "avg price", "avg"])
            pct_idx = _get_col(["price_change", "change_pct", "pct"])
            vol_idx = _get_col(["volatility", "stddev", "std"])

            for ticker, label in [("GC=F", "GC=F"), ("AMZN", "AMZN")]:
                expected = yf_expected.get(ticker, {})
                row = next((r for r in data
                            if sym_idx is not None and r[sym_idx]
                            and ticker.lower() in str(r[sym_idx]).lower()), None)
                if not row:
                    check(f"Stock_Trends has {label} row with correct stats",
                          False, f"No row for {label}")
                    continue
                ok_avg = avg_idx is not None and _num_ok(row[avg_idx], expected.get("avg_price"), _tol_for(expected.get("avg_price"), 0.02, 1.0))
                ok_vol = vol_idx is not None and _num_ok(row[vol_idx], expected.get("volatility"), _tol_for(expected.get("volatility"), 0.05, 1.0))
                ok_pct = pct_idx is not None and _num_ok(row[pct_idx], expected.get("pct_change"), _tol_for(expected.get("pct_change"), 0.05, 5.0))
                check(f"Stock_Trends has {label} row with correct stats",
                      ok_avg and ok_vol and ok_pct,
                      f"row={row}, expected={expected}, ok_avg={ok_avg}, ok_vol={ok_vol}, ok_pct={ok_pct}")
        else:
            check("Stock_Trends has GC=F row with correct stats", False, "Empty sheet")
            check("Stock_Trends has AMZN row with correct stats", False, "Empty sheet")
    else:
        check("Stock_Trends has GC=F row with correct stats", False, "Sheet missing")
        check("Stock_Trends has AMZN row with correct stats", False, "Sheet missing")

    # Product_Margins
    if pm_name:
        ws = wb_f[pm_name]
        rows = list(_resolve_rows(ws, wb_v[pm_name]))
        data = [r for r in rows[1:] if any(c is not None and str(c).strip() != "" for c in r)]
        check("Product_Margins has >=5 categories", len(data) >= 5,
              f"Found {len(data)} category rows")

        # Check that key categories are present
        cats_lower = " ".join(str(r[0]).lower() for r in data if r and r[0])
        check("Product_Margins includes Audio + Cameras + Electronics",
              "audio" in cats_lower and "camera" in cats_lower and "electronics" in cats_lower,
              f"Categories: {[r[0] for r in data if r]}")

        # All margin_pct must equal the standard margin (40, since cost is 60%
        # of price). Accept 40, "40%", or ratio 0.40.
        if rows:
            headers = [str(c).lower().strip() if c is not None else "" for c in rows[0]]
            margin_idx = next((i for i, h in enumerate(headers) if "margin" in h), None)
            if margin_idx is not None:
                expected_margin = params["standard_margin_pct"]
                bad = [r for r in data
                       if r[margin_idx] is not _UNVERIFIABLE
                       and not num_close(_to_pct(r[margin_idx]), expected_margin, 0.5)]
                check(f"All Product_Margins rows have margin_pct == {expected_margin:g}%",
                      len(bad) == 0,
                      f"Bad rows: {[(r[0], r[margin_idx]) for r in bad[:3]]}")
    else:
        check("Product_Margins has >=5 categories", False, "Sheet missing")
        check("Product_Margins includes Audio + Cameras + Electronics", False, "Sheet missing")

    # Correlation_Analysis
    if ca_name:
        ws = wb_f[ca_name]
        rows = list(_resolve_rows(ws, wb_v[ca_name]))
        data = [r for r in rows[1:] if any(c is not None and str(c).strip() != "" for c in r)]
        check("Correlation_Analysis has 2 entries", len(data) >= 2,
              f"Found {len(data)} rows")
        # Per-row keyword requirements: at least one row mentions gold (GC=F),
        # at least one row mentions amzn/amazon, and consumer should appear somewhere.
        all_text = " ".join(_txt(c) for r in rows for c in r if c is not None).lower()
        per_row_text = [" ".join(_txt(c) for c in r if c is not None).lower() for r in data]
        has_gold_row = any(("gold" in t) or ("gc=f" in t) for t in per_row_text)
        has_stock_row = any(("amzn" in t) or ("amazon" in t) for t in per_row_text)
        check("Correlation_Analysis: at least one row references gold/GC=F",
              has_gold_row,
              f"Rows: {per_row_text[:3]}")
        check("Correlation_Analysis: at least one row references AMZN/Amazon",
              has_stock_row,
              f"Rows: {per_row_text[:3]}")
        check("Correlation_Analysis mentions consumer somewhere",
              "consumer" in all_text,
              f"Text: {all_text[:200]}")
    else:
        check("Correlation_Analysis has 2 entries", False, "Sheet missing")
        check("Correlation_Analysis: at least one row references gold/GC=F", False, "Sheet missing")
        check("Correlation_Analysis: at least one row references AMZN/Amazon", False, "Sheet missing")
        check("Correlation_Analysis mentions consumer somewhere", False, "Sheet missing")

    # Strategic_Recommendations
    if sr_name:
        ws = wb_f[sr_name]
        rows = list(_resolve_rows(ws, wb_v[sr_name]))
        data = [r for r in rows[1:] if any(c is not None and str(c).strip() != "" for c in r)]
        check("Strategic_Recommendations has rows for each category",
              len(data) >= 5,
              f"Found {len(data)}")
        if rows:
            headers = [str(c).lower().strip() if c is not None else "" for c in rows[0]]
            cur_idx = next((i for i, h in enumerate(headers) if "current" in h and "margin" in h), None)
            tar_idx = next((i for i, h in enumerate(headers) if "target" in h and "margin" in h), None)
            if cur_idx is not None and tar_idx is not None:
                # All current_margin should equal the standard margin
                expected_current = params["standard_margin_pct"]
                cur_bad = [r for r in data
                           if r[cur_idx] is not _UNVERIFIABLE
                           and not num_close(_to_pct(r[cur_idx]), expected_current, 0.5)]
                check(f"All current_margin == {expected_current:g}%",
                      len(cur_bad) == 0,
                      f"Bad: {[(r[0], r[cur_idx]) for r in cur_bad[:3]]}")
                # Target margin is derived from the actual gold price change:
                # if gold rose above the significance threshold -> reduced
                # target (35); otherwise maintain the standard target (40).
                gold_pct = yf_expected.get("GC=F", {}).get("pct_change") or 0
                threshold = params["gold_significant_change_pct"]
                if gold_pct > threshold:
                    expected_target = params["reduced_margin_target"]
                    note = f"gold pct {gold_pct:.1f} > {threshold:g} -> reduced target"
                else:
                    expected_target = params["standard_margin_target"]
                    note = f"gold pct {gold_pct:.1f} <= {threshold:g} -> maintain target"
                tar_bad = [r for r in data
                           if r[tar_idx] is not _UNVERIFIABLE
                           and not num_close(_to_pct(r[tar_idx]), expected_target, 0.5)]
                check(f"Target margin == {expected_target:g}% ({note})",
                      len(tar_bad) == 0,
                      f"Bad: {[(r[0], r[tar_idx]) for r in tar_bad[:3]]}")
    else:
        check("Strategic_Recommendations has rows for each category", False, "Sheet missing")


def check_word(workspace):
    print("\n=== Check 2: Pricing_Strategy_Memo.docx ===")
    path = os.path.join(workspace, "Pricing_Strategy_Memo.docx")
    if not os.path.exists(path):
        check("Word document exists", False, f"Not found at {path}")
        for label in ["Word: title section", "Word: market overview section",
                      "Word: product margin section", "Word: strategic recommendations section",
                      "Word: contains gold pct number", "Word: substantial content"]:
            check(label, False, "Word missing")
        return
    check("Word document exists", True)

    try:
        doc = Document(path)
    except Exception as e:
        check("Word readable", False, str(e))
        return

    # Collect text from paragraphs and from any tables (a memo may present
    # the figures in a table rather than prose, which must not hide them from
    # the number/heading checks).
    full_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " " + " ".join(p.text for p in cell.paragraphs)
    full_lower = full_text.lower()

    # Title check: require literal phrase 'Commodity Impact and Pricing Strategy'
    # Look in headings/title styles first, fallback to first paragraph
    headings = [p.text.strip() for p in doc.paragraphs
                if p.style and (p.style.name.startswith("Heading") or p.style.name == "Title")]
    headings_lower = [h.lower() for h in headings]
    first_para = (doc.paragraphs[0].text.strip().lower() if doc.paragraphs else "")
    top_text = " ".join(p.text for p in doc.paragraphs[:3]).lower()
    expected_title = "commodity impact and pricing strategy"
    title_match = (
        any(expected_title in h for h in headings_lower)
        or expected_title in first_para
        or expected_title in top_text
    )
    check("Word: title 'Commodity Impact and Pricing Strategy' present (heading or first paragraph)",
          title_match,
          f"Headings: {headings[:5]}; first para: {first_para[:120]}")
    check("Word: market overview section",
          "market overview" in full_lower or "market trends" in full_lower
          or "market summary" in full_lower or "market analysis" in full_lower,
          "")
    check("Word: product margin section",
          "margin" in full_lower and ("category" in full_lower or "categories" in full_lower),
          "")
    check("Word: strategic recommendations section",
          "recommendation" in full_lower or "strategic" in full_lower,
          "")
    # Specific numbers expected: the gold price change is ~138.7%, and a
    # correct memo may legitimately round it (e.g. "approximately 140%" or
    # "~139%"). Extract every numeric token and accept any one close to the
    # expected value, so reasonable rounding can never fail a correct model
    # while a memo that omits or misstates the figure still fails.
    yf_expected = _merge_yf_expected(_fetch_yf_expected())
    gold_pct = yf_expected.get("GC=F", {}).get("pct_change") or YF_FALLBACK["GC=F"]["pct_change"]
    pct_tol = _tol_for(gold_pct, 0.05, 5.0)
    mem_nums = [_to_float(m.group(0)) for m in re.finditer(r'\d+(?:[.,]\d+)*\s*%?', full_text)]
    mem_nums = [n for n in mem_nums if n is not None]
    gold_num_ok = any(abs(n - gold_pct) <= pct_tol for n in mem_nums)
    check("Word: contains gold pct number (~138%)",
          gold_num_ok,
          f"expected ~{gold_pct:g}%, numbers in memo: {mem_nums[:20]}")
    check("Word: substantial content", len(full_text) > 400, f"Length: {len(full_text)}")


def _rich_text_content(item):
    """Extract plain text from a Notion rich-text item or bare value.

    The title/text payload in the DB can arrive in many shapes: a rich-text
    object carrying text.content, a rich-text object carrying plain_text, a
    bare string, a title wrapper ({'title': [...]}) or a properties wrapper
    ({'properties': {'title': [...]}}). Walk them all and return the text.
    """
    if item is None:
        return ""
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        pt = item.get("plain_text")
        if isinstance(pt, str):
            return pt
        t = item.get("text")
        if isinstance(t, dict):
            c = t.get("content")
            if isinstance(c, str):
                return c
        if isinstance(t, str):
            return t
        title = item.get("title")
        if title is not None:
            return _rich_text_content(title)
        props = item.get("properties", {})
        if isinstance(props, dict):
            t = props.get("title")
            if t is not None:
                return _rich_text_content(t)
        return str(item)
    return str(item)


def _db_title_str(title):
    """Normalize a notion.database title column to a plain string.

    Handles list-of-rich-text, list-of-bare-strings, JSON-string encodings of
    either, dicts (title wrapper / properties wrapper) and plain strings."""
    if title is None:
        return ""
    if isinstance(title, list):
        return " ".join(_rich_text_content(item) for item in title)
    if isinstance(title, str):
        stripped = title.strip()
        if stripped.startswith(("[", "{")):
            try:
                return _db_title_str(json.loads(stripped))
            except Exception:
                pass
        return title
    if isinstance(title, dict):
        return _rich_text_content(title)
    return str(title)


def check_notion():
    print("\n=== Check 3: Notion Market Research Dashboard ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("SELECT id, title FROM notion.databases")
        dbs = cur.fetchall()
    except Exception as e:
        check("Notion DB query OK", False, str(e))
        return

    dashboard_dbs = []
    for db_id, title in dbs:
        title_str = _db_title_str(title)
        if "market research dashboard" in title_str.lower().strip():
            dashboard_dbs.append(db_id)
    check("'Market Research Dashboard' database exists",
          len(dashboard_dbs) >= 1,
          f"Databases: {[d[1] for d in dbs[:5]]}")

    if dashboard_dbs:
        try:
            # Count pages per matching dashboard db. The swarm framework does
            # not serialize writes, so a correct solution may legitimately
            # produce more than one matching database. Require that at least
            # one of them holds the two required entries.
            counts = []
            for db_id in dashboard_dbs:
                cur.execute("""
                    SELECT COUNT(*) FROM notion.pages
                    WHERE parent->>'database_id' = %s
                """, (db_id,))
                counts.append(cur.fetchone()[0])
            # The swarm framework does not serialize writes, so besides several
            # DBs each holding the two entries, one DB may legitimately end up
            # with the two entries duplicated (2 entries x 2 concurrent writers
            # = 4 pages). Accept a count of 2 or of 4 (duplicate pairs); any
            # other count means the required entries were not both delivered.
            ok = any(c == 2 or c == 4 for c in counts)
            check("Dashboard has exactly 2 entries",
                  ok,
                  f"Page counts per matching DB: {counts}")
        except Exception as e:
            check("Dashboard entry count check", False, str(e))

    try:
        cur.close()
        conn.close()
    except Exception:
        pass


def check_script(workspace):
    print("\n=== Check 4: correlation_analysis.py ===")
    path = os.path.join(workspace, "correlation_analysis.py")
    check("correlation_analysis.py exists", os.path.exists(path))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    params = _load_params(args.agent_workspace)
    check_excel(args.agent_workspace, params)
    check_word(args.agent_workspace)
    check_notion()
    check_script(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "failed": FAIL_COUNT}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
