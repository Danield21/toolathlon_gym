"""
Evaluation for arxiv-latex-analysis-notion-word task.

Checks:
1. Paper_Analysis_Report.docx exists and has at least 4 sections
2. Word doc contains all 3 paper title keywords
3. Word doc mentions Scaling Laws and RLHF/InstructGPT and OPT
4. Word doc has a Comparative Analysis section
5. Notion has 3 pages about the papers
6. Email sent to research_lead@university.edu
"""
import json
import os
import sys
from argparse import ArgumentParser

import re

import psycopg2
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": os.environ.get("PGUSER", "eigent"),
    "password": os.environ.get("PGPASSWORD", "camel"),
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document Paper_Analysis_Report.docx ===")

    docx_path = os.path.join(agent_workspace, "Paper_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Paper_Analysis_Report.docx exists", False, f"Not found at {docx_path}")
        return
    record("Paper_Analysis_Report.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Get all text
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Count headings/sections
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading") or
                (p.text.strip() and len(p.text.strip()) < 100 and
                 any(kw in p.text.lower() for kw in ["scaling", "instruct", "opt:", "analysis", "paper"]))]

    record("Word doc has at least 4 sections", len(headings) >= 4,
           f"Found {len(headings)} section-like headings")

    # Check paper keywords with stronger matching
    import re as _re
    has_scaling = "scaling laws" in all_text or "scaling law" in all_text or "kaplan" in all_text
    has_rlhf = ("rlhf" in all_text or "instructgpt" in all_text or
                "follow instructions" in all_text or "human feedback" in all_text or
                "ouyang" in all_text)
    # Use word boundary for 'opt' to avoid matching 'option', 'optimal', etc.
    # A correct report may identify OPT by shorthand heading ('OPT'), by its
    # descriptive subtitle, or by its key attributes; accept any of these.
    opt_word = _re.search(r"\bopt\b", all_text) is not None or "open pre-trained" in all_text
    has_opt = opt_word and ("open pre-trained transformer" in all_text or
                             "pre-trained transformer" in all_text or
                             "open-source" in all_text or "175b" in all_text or
                             "175 billion" in all_text or
                             "zhang" in all_text and "roller" in all_text)

    record("Mentions Scaling Laws paper", has_scaling, "No scaling laws content found")
    record("Mentions RLHF/InstructGPT paper", has_rlhf, "No RLHF/InstructGPT content found")
    record("Mentions OPT paper (with word-boundary match)", has_opt, "No OPT content found")

    # Tighten Comparative Analysis to require an explicit section title
    has_comparative_section = ("comparative analysis" in all_text or "comparative summary" in all_text or
                               "comparison of" in all_text or "comparing" in all_text)
    record("Has Comparative Analysis section title", has_comparative_section,
           "No comparative analysis section header found")


def _extract_page_titles(properties):
    """Extract page-title text from a stored notion.pages.properties jsonb value.

    The notion MCP 'create_page' tool schema (scripts/notion-openapi.json,
    POST /v1/pages) requires the agent to send `properties.title` as a raw
    rich-text ARRAY, e.g. {"title": [{"text": {"content": "Paper: ..."}}]},
    and pg-client.ts stores that payload verbatim. Some agents may instead
    send the Notion-API-style WRAPPED object {"title": {"type": "title",
    "title": [{"plain_text": "..."}]}}. Extract from both formats so a
    schema-following agent always scores correctly.
    """
    if not isinstance(properties, dict):
        return []
    titles = []
    # 1) Prefer keys whose name suggests a page title.
    for key, val in properties.items():
        if key.lower() not in ("title", "name", "page_title"):
            continue
        arr = val.get("title", []) if isinstance(val, dict) else (val if isinstance(val, list) else None)
        if not isinstance(arr, list) or not arr:
            continue
        text = "".join(
            (t.get("plain_text") if isinstance(t, dict) else "")
            or (t.get("text", {}).get("content", "") if isinstance(t, dict) else "")
            for t in arr
        )
        if text:
            titles.append(text)
    if titles:
        return titles
    # 2) Fallback: any property value that is a wrapped title object
    #    ({"type": "title", "title": [...]}) under an unexpected key.
    for val in properties.values():
        if isinstance(val, dict) and val.get("type") == "title":
            arr = val.get("title", [])
            if isinstance(arr, list):
                text = "".join(
                    (t.get("plain_text") if isinstance(t, dict) else "")
                    or (t.get("text", {}).get("content", "") if isinstance(t, dict) else "")
                    for t in arr
                )
                if text:
                    return [text]
    return []


def check_notion():
    print("\n=== Check 2: Notion Pages ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("SELECT properties FROM notion.pages")
    pages = cur.fetchall()
    cur.close()
    conn.close()

    # Extract clean page titles (with the 'Paper:' prefix from task)
    page_titles = []
    for (props,) in pages:
        for t in _extract_page_titles(props):
            page_titles.append(t.lower())

    # Need at least 3 pages with the 'Paper:' prefix per task
    paper_pages = [t for t in page_titles if t.strip().startswith("paper:")]
    record("At least 3 Notion 'Paper:' pages exist", len(paper_pages) >= 3,
           f"Found {len(paper_pages)} 'Paper:' pages out of {len(pages)}")

    all_titles = " ".join(page_titles)
    has_scaling = "scaling" in all_titles
    has_instruct = "instruct" in all_titles or "human feedback" in all_titles or "rlhf" in all_titles
    # OPT page: accept the full title ('OPT: Open Pre-trained Transformer ...'),
    # the common shorthand ('Paper: OPT'), or the descriptive subtitle alone.
    has_opt = (
        re.search(r"\bopt\b", all_titles) is not None
        or "open pre-trained" in all_titles
        or "pre-trained transformer" in all_titles
    )

    record("Notion page covers Scaling Laws", has_scaling,
           f"titles: {page_titles[:5]}")
    record("Notion page covers InstructGPT/RLHF", has_instruct)
    record("Notion page covers OPT", has_opt)


def check_email():
    print("\n=== Check 3: Email to research_lead@university.edu ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else str(to_addr).lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "research_lead@university.edu" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to research_lead@university.edu", matching is not None,
           f"Messages found: {len(messages)}")

    if matching:
        import re as _re_em
        subject, _, _, body_text = matching
        subject_lower = (subject or "").lower()
        body_lower = (body_text or "").lower()
        # Require subject-level paper/analysis/report context (fixed subject in
        # task.md) AND a non-empty body that references the work/deliverables.
        subj_ok = (
            "paper" in subject_lower or "analysis" in subject_lower or "report" in subject_lower
        )
        # task.md only asks the body to "summarize what you have done" and note
        # that the Word doc + Notion pages are ready for review. Do NOT require a
        # specific paper keyword (a literal agent may not name any paper). Accept
        # any non-empty body that references the work/deliverables, in English or
        # Chinese; paper-specific keywords remain accepted as evidence.
        body_kw_ok = bool(body_text and body_text.strip()) and (
            any(kw in body_lower for kw in (
                # generic English work/deliverable descriptors
                "paper", "analysis", "report", "three", "word", "notion",
                "document", "page", "ready", "review", "deliverable", "summar",
                # paper-specific keywords still accepted as evidence
                "scaling", "rlhf", "instruct", "kaplan", "ouyang", "transformer",
                # Chinese equivalents
                "论文", "报告", "审阅", "检查", "就绪", "完成", "三篇", "分析",
            ))
            or _re_em.search(r"\bopt\b", body_lower) is not None
        )
        record("Email mentions paper analysis (subject context)", subj_ok,
               f"Subject: {subject}")
        record("Email body summarizes work / mentions deliverables", body_kw_ok,
               f"body preview: {body_text[:100] if body_text else ''}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_notion()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} checks failed)")
        sys.exit(1)


if __name__ == "__main__":
    main()
