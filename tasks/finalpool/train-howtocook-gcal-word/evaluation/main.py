"""
Evaluation for train-howtocook-gcal-word task.

Checks (per task.md):
1. Qufu_Trip_Itinerary.docx exists with required sections + train/recipe content
2. GCal has 3 events: outbound train (03-12 17:30), return train (03-15 15:00),
   cultural visit (03-13).
"""
import json
import os
import sys
import re
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def _query_train_codes():
    """Return (outbound_codes, return_codes) — eligible high-speed trains.
    Outbound: Beijing -> Qufu on 2026-03-12 evening (>= 17:00).
    Return:   Qufu -> Beijing on 2026-03-15 afternoon (12:00-18:00).
    """
    fallback = (["G235"], ["G236"])
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT station_train_code, start_time
            FROM train.trains
            WHERE from_station_telecode IN ('BJP','VNP','VAP','BOP')
              AND to_station_telecode = 'QFB'
              AND depart_date = '2026-03-12'
              AND station_train_code LIKE 'G%'
        """)
        out = []
        for code, st in cur.fetchall():
            try:
                hour = int(str(st).split(":")[0])
                if hour >= 17:
                    out.append(code)
            except (TypeError, ValueError):
                continue

        cur.execute("""
            SELECT station_train_code, start_time
            FROM train.trains
            WHERE from_station_telecode = 'QFB'
              AND to_station_telecode IN ('BJP','VNP','VAP','BOP')
              AND depart_date = '2026-03-15'
              AND station_train_code LIKE 'G%'
        """)
        ret = []
        for code, st in cur.fetchall():
            try:
                hour = int(str(st).split(":")[0])
                if 12 <= hour < 18:
                    ret.append(code)
            except (TypeError, ValueError):
                continue
        cur.close()
        conn.close()
        return (out or fallback[0], ret or fallback[1])
    except Exception as e:
        print(f"[fallback] train DB query: {e}")
        return fallback


OUT_CODES, RET_CODES = _query_train_codes()


def check_word(agent_workspace):
    print("\n=== Check 1: Qufu_Trip_Itinerary.docx ===")
    docx_path = os.path.join(agent_workspace, "Qufu_Trip_Itinerary.docx")
    if not os.path.exists(docx_path):
        record("Qufu_Trip_Itinerary.docx exists", False, f"Not found at {docx_path}")
        # Account for missing-file impact on remaining checks
        for name in ["Word readable", "Trip Overview section", "Outbound Journey section",
                     "Activities in Qufu section", "Return Journey section",
                     "Meal Plan section", "Outbound train code mentioned",
                     "Return train code mentioned", "Mentions Qufu/Confucius",
                     "Mentions congee/porridge recipe", "Mentions stir-fry recipe"]:
            record(name, False, "missing docx")
        return
    record("Qufu_Trip_Itinerary.docx exists", True)

    try:
        import docx
        doc = docx.Document(docx_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        record("Word readable", False, str(e))
        return
    record("Word readable", True)

    text_lower = full_text.lower()

    # Required sections (per task.md). Prefer matching as a Heading-style paragraph;
    # fallback to any paragraph that starts with the section title (line-start match).
    heading_paragraphs = [
        (p.text or "").strip().lower()
        for p in doc.paragraphs
        if p.style and p.style.name and p.style.name.startswith("Heading")
    ]
    line_starts = [(p.text or "").strip().lower() for p in doc.paragraphs]

    def _section_present(alts):
        for a in alts:
            a_l = a.lower()
            # Heading-style match (substring within heading)
            if any(a_l in h for h in heading_paragraphs):
                return True
            # Paragraph that starts with the section title
            if any(line == a_l or line.startswith(a_l) for line in line_starts):
                return True
        return False

    section_titles = [
        ("Trip Overview", ["trip overview"]),
        ("Outbound Journey", ["outbound journey"]),
        ("Activities in Qufu", ["activities in qufu"]),
        ("Return Journey", ["return journey"]),
        ("Meal Plan", ["meal plan"]),
    ]
    for label, alts in section_titles:
        ok = _section_present(alts)
        record(f"{label} section present (heading or line-start)", ok, f"checked alts: {alts}")

    # Outbound train code
    out_ok = any(re.search(rf"\b{re.escape(code.lower())}\b", text_lower) for code in OUT_CODES)
    record(f"Outbound train code mentioned (any of {OUT_CODES})", out_ok)

    # Return train code
    ret_ok = any(re.search(rf"\b{re.escape(code.lower())}\b", text_lower) for code in RET_CODES)
    record(f"Return train code mentioned (any of {RET_CODES})", ret_ok)

    # Qufu / Confucius
    record("Mentions Qufu/Confucius", "qufu" in text_lower or "confucius" in text_lower)

    # Recipe content: must mention congee OR rice porridge in Meal Plan
    has_congee = ("congee" in text_lower) or ("rice porridge" in text_lower) or ("porridge" in text_lower)
    record("Mentions congee/porridge recipe", has_congee)

    # Stir-fry recipe
    has_stirfry = re.search(r"stir[\s\-]?fry", text_lower) is not None
    record("Mentions stir-fry recipe", has_stirfry)


def check_gcal():
    print("\n=== Check 2: Google Calendar Events ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # ---- Outbound train (03-12) ----
        cur.execute("""
            SELECT summary, start_datetime, end_datetime
            FROM gcal.events
            WHERE start_datetime >= '2026-03-12' AND start_datetime < '2026-03-13'
        """)
        events_mar12 = cur.fetchall()
        # Match the title pattern "Beijing to Qufu - <code>" — accept any code in OUT_CODES
        outbound = None
        for ev in events_mar12:
            sm = (ev[0] or "").lower()
            if ("beijing" in sm and "qufu" in sm) and any(c.lower() in sm for c in OUT_CODES):
                outbound = ev
                break
        if outbound is None:
            # fallback: find by code only
            for ev in events_mar12:
                sm = (ev[0] or "").lower()
                if any(c.lower() in sm for c in OUT_CODES):
                    outbound = ev
                    break
        record("Outbound train event on 2026-03-12 with correct title",
               outbound is not None,
               f"events on Mar12: {[e[0] for e in events_mar12]}")

        if outbound is not None:
            start = outbound[1]
            # Task says 'evening' — accept 17:00-19:59 (covers G235 17:30 and other evening trains)
            ok_time = start is not None and 17 <= start.hour <= 19
            record("Outbound event starts in evening (17:00-19:59)",
                   ok_time, f"start: {start}")

        # ---- Cultural visit (03-13) ----
        cur.execute("""
            SELECT summary, start_datetime, end_datetime
            FROM gcal.events
            WHERE start_datetime >= '2026-03-13' AND start_datetime < '2026-03-14'
        """)
        events_mar13 = cur.fetchall()
        cultural = None
        for ev in events_mar13:
            sm = (ev[0] or "").lower()
            if ("confucius" in sm and ("temple" in sm or "mansion" in sm)) or "cultural visit" in sm:
                cultural = ev
                break
        record("Cultural visit event on 2026-03-13 (Confucius Temple/Mansion)",
               cultural is not None,
               f"events on Mar13: {[e[0] for e in events_mar13]}")

        if cultural is not None:
            start = cultural[1]
            end = cultural[2]
            # Task: starts 09:00, ends 17:00
            ok_start = start is not None and start.hour == 9
            ok_end = end is not None and end.hour == 17
            record("Cultural visit starts at 09:00", ok_start, f"start: {start}")
            record("Cultural visit ends at 17:00", ok_end, f"end: {end}")

        # ---- Return train (03-15) ----
        cur.execute("""
            SELECT summary, start_datetime, end_datetime
            FROM gcal.events
            WHERE start_datetime >= '2026-03-15' AND start_datetime < '2026-03-16'
        """)
        events_mar15 = cur.fetchall()
        return_ev = None
        for ev in events_mar15:
            sm = (ev[0] or "").lower()
            if ("qufu" in sm and "beijing" in sm) and any(c.lower() in sm for c in RET_CODES):
                return_ev = ev
                break
        if return_ev is None:
            for ev in events_mar15:
                sm = (ev[0] or "").lower()
                if any(c.lower() in sm for c in RET_CODES):
                    return_ev = ev
                    break
        record("Return train event on 2026-03-15 with correct title",
               return_ev is not None,
               f"events on Mar15: {[e[0] for e in events_mar15]}")

        if return_ev is not None:
            start = return_ev[1]
            # Task says "afternoon (12:00-18:00)"; accept any hour in that window.
            ok_time = start is not None and 12 <= start.hour < 18
            record("Return event starts in afternoon window 12:00-17:59",
                   ok_time, f"start: {start}")

        cur.close()
        conn.close()
    except Exception as e:
        record("GCal check completed", False, f"DB error: {e}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
