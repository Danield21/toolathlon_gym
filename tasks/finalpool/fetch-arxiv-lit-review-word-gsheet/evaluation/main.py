"""Evaluation script for fetch-arxiv-lit-review-word-gsheet."""
import os
import argparse, json, os, sys
import openpyxl

# --- verify_v2 smart primitives ---
_EVAL_DIR = os.path.dirname(os.path.abspath(__file__))
_PACK_ROOT = os.path.abspath(os.path.join(_EVAL_DIR, "..", "..", "..", ".."))
if _PACK_ROOT not in sys.path:
    sys.path.insert(0, _PACK_ROOT)
try:
    from utils.verify_v2 import smart_column_exists
    from utils.verify_v2.eval_helpers import get_sheet_rows_as_dicts, get_gt_column_values
    _HAS_VERIFY_V2 = True
except Exception:
    _HAS_VERIFY_V2 = False

TASK_NAME = "fetch-arxiv-lit-review-word-gsheet"


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0
WARN_COUNT = 0
IS_GT_SELF_TEST = False

def check(name, condition, detail="", db_side=False):
    global PASS_COUNT, FAIL_COUNT, WARN_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        if IS_GT_SELF_TEST and db_side:
            WARN_COUNT += 1
            detail_str = str(detail)[:200] if detail else ""
            print(f"  [WARN] {name} (GT self-test, DB-side): {detail_str}")
        else:
            FAIL_COUNT += 1
            detail_str = str(detail)[:200] if detail else ""
            print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT, IS_GT_SELF_TEST
    PASS_COUNT = 0
    FAIL_COUNT = 0
    try:
        if (agent_workspace and groundtruth_workspace
                and os.path.realpath(agent_workspace) ==
                os.path.realpath(groundtruth_workspace)):
            IS_GT_SELF_TEST = True
        else:
            IS_GT_SELF_TEST = False
    except Exception:
        IS_GT_SELF_TEST = False


    excel_path = os.path.join(agent_workspace, "Lit_Review_Report.xlsx")
    check("Lit_Review_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Lit_Review_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        def check_columns(sheet_name, expected_cols, min_rows):
            """Verify sheet exists, >= min_rows, contains required columns.
            Uses LLM semantic mapping via verify_v2, falls back to strict
            header match."""
            check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
            if sheet_name not in wb.sheetnames:
                return
            _ws = wb[sheet_name]
            _data_rows = list(_ws.iter_rows(min_row=2, values_only=True))
            check(f"{sheet_name} has >= {min_rows} rows",
                  len(_data_rows) >= min_rows, f"got {len(_data_rows)}")
            if _HAS_VERIFY_V2 and gt_wb is not None:
                _raw_headers, _agent_rows = get_sheet_rows_as_dicts(wb, sheet_name)
                for _exp in expected_cols:
                    _gt_vals = get_gt_column_values(gt_wb, sheet_name, _exp)
                    _ok, _matched, _reason = smart_column_exists(
                        expected_col=_exp, agent_headers=_raw_headers,
                        gt_samples=_gt_vals[:3], agent_rows=_agent_rows,
                        task_name=TASK_NAME,
                    )
                    _detail = _reason
                    if _ok and _matched and _matched.lower() != _exp.lower():
                        _detail = f"LLM-mapped to {_matched!r}"
                    check(f"{sheet_name} has {_exp} column", _ok, _detail)
            else:
                _headers = [str(c.value).strip().lower() if c.value else "" for c in _ws[1]]
                for _exp in expected_cols:
                    check(f"{sheet_name} has {_exp} column",
                          _exp.lower() in _headers, f"headers: {_headers[:8]}")

        check("Data_Analysis sheet exists", "Data_Analysis" in wb.sheetnames)
        if "Data_Analysis" in wb.sheetnames:
            ws = wb["Data_Analysis"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Data_Analysis has >= 4 rows", len(data_rows) >= 4, f"got {len(data_rows)}")

            check_columns('Data_Analysis', ['Paper_ID', 'Title', 'Area', 'Citations', 'Relevance_Score'], 4)

            # Content validation: must include expected papers, must NOT include noise paper
            EXPECTED_PAPER_IDS = {"2301.01234", "2302.05678", "2303.09012", "2304.03456"}
            NOISE_ID = "9999.99999"
            all_text = " ".join(str(c) for r in data_rows for c in r if c is not None).lower()
            for pid in EXPECTED_PAPER_IDS:
                check(f"Data_Analysis contains expected paper {pid}",
                      pid in all_text, "missing")
            check(f"Data_Analysis does NOT contain noise paper {NOISE_ID}",
                  NOISE_ID not in all_text, "noise paper leaked into analysis")
        check("Metrics sheet exists", "Metrics" in wb.sheetnames)
        if "Metrics" in wb.sheetnames:
            ws = wb["Metrics"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Metrics has >= 4 rows", len(data_rows) >= 4, f"got {len(data_rows)}")

            check_columns('Metrics', ['Metric', 'Value'], 4)
        check("Recommendations sheet exists", "Recommendations" in wb.sheetnames)
        if "Recommendations" in wb.sheetnames:
            ws = wb["Recommendations"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Recommendations has >= 2 rows", len(data_rows) >= 2, f"got {len(data_rows)}")

            check_columns('Recommendations', ['Priority', 'Action', 'Area'], 2)
        try:
            conn = get_conn()
            cur = conn.cursor()
            # Tightened: title must start with 'Arxiv Litreview' (case-insensitive)
            cur.execute("SELECT title FROM gsheet.spreadsheets WHERE title ILIKE %s", ('Arxiv Litreview%',))
            sheets = cur.fetchall()
            check("Google Sheet 'Arxiv Litreview Tracker' exists",
                  len(sheets) >= 1, f"titles found: {[s[0] for s in sheets]}", db_side=True)
            conn.close()
        except Exception as e:
            check("GSheet check", False, str(e), db_side=True)

        # Check Word document - require Lit_Review_Analysis.docx with the 3 sections
        docx_path = os.path.join(agent_workspace, "Lit_Review_Analysis.docx")
        check("Lit_Review_Analysis.docx exists", os.path.exists(docx_path),
              f"missing: {docx_path}")
        if os.path.exists(docx_path):
            from docx import Document
            doc = Document(docx_path)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            heading_text = " ".join(p.text.lower() for p in doc.paragraphs
                                    if p.style.name.startswith('Heading'))
            check("Word has substantial content (>= 100 chars)",
                  len(text) >= 100, f"text length: {len(text)}")
            for sec in ["executive summary", "key findings", "recommendations"]:
                check(f"Word has '{sec}' section",
                      sec in heading_text or sec in text,
                      f"text sample: {text[:200]}")

        check("arxiv_litreview_processor.py exists", os.path.exists(os.path.join(agent_workspace, "arxiv_litreview_processor.py")))


    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
