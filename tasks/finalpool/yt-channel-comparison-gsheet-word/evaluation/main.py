"""
Evaluation for yt-channel-comparison-gsheet-word task.

Checks:
1. Channel_Analysis_Report.docx exists with 4 required headings and channel mentions.
2. Word doc surfaces specific computed numerical metrics for both channels.
3. GSheet "Channel Comparison Analysis" exists with two sheets (Comparison, Monthly_Fireship).
4. Comparison sheet has 2 data rows with correct metric values for Fireship and Veritasium.
5. Monthly_Fireship sheet has correct row count and at least anchor month value.
6. Email sent to media@company.com with proper subject.
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

# Computed from authoritative DB query (see preprocess docstring)
EXPECTED = {
    "Fireship": {
        "Video_Count": 104,
        "Total_Views": 121886817,
        "Total_Likes": 4837913,
        "Avg_Views": 1171989,
        "Avg_Duration_Sec": 291,
        "Avg_Like_Rate": 3.97,
        "Top_Video_Title": "Big Tech in panic mode... Did DeepSeek R1 just pop the AI bubble?",
        "Top_Video_Views": 3878491,
    },
    "Veritasium": {
        "Video_Count": 30,
        "Total_Views": 365002206,
        "Total_Likes": 8171292,
        "Avg_Views": 12166740,
        "Avg_Duration_Sec": 1664,
        "Avg_Like_Rate": 2.24,
        "Top_Video_Title": "How One Company Secretly Poisoned The Planet",
        "Top_Video_Views": 32716281,
    },
}

# Monthly Fireship counts
MONTHLY_FIRESHIP = {
    "2024-06": 9, "2024-07": 10, "2024-08": 10, "2024-09": 6, "2024-10": 9,
    "2024-11": 10, "2024-12": 4, "2025-01": 9, "2025-02": 9, "2025-03": 9,
    "2025-04": 6, "2025-05": 6, "2025-06": 7,
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, rel_tol=0.005, abs_tol=1):
    """Loose tolerance for live YT data drift, but tight for these snapshots."""
    try:
        a = float(a); b = float(b)
        return abs(a - b) <= max(abs_tol, abs(b) * rel_tol)
    except (TypeError, ValueError):
        return False


def get_int(v):
    try:
        if v is None:
            return None
        s = str(v).strip().replace(",", "")
        return int(float(s))
    except (TypeError, ValueError):
        return None


def get_float(v):
    try:
        if v is None:
            return None
        s = str(v).strip().replace(",", "").replace("%", "")
        return float(s)
    except (TypeError, ValueError):
        return None


def check_word(agent_workspace):
    print("\n=== Check 1: Channel_Analysis_Report.docx ===")
    docx_path = os.path.join(agent_workspace, "Channel_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Channel_Analysis_Report.docx exists", False, f"Not found at {docx_path}")
        return
    record("Channel_Analysis_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith('Heading')]
    heading_text = " ".join(headings).lower()
    record("Has Executive Summary heading",
           "executive summary" in heading_text,
           f"Headings: {headings}")
    record("Has Channel Metrics Comparison heading",
           "channel metrics comparison" in heading_text or "metrics comparison" in heading_text,
           f"Headings: {headings}")
    record("Has Key Findings heading",
           "key findings" in heading_text,
           f"Headings: {headings}")
    record("Has Recommendations heading",
           "recommendation" in heading_text,
           f"Headings: {headings}")

    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    record("Doc mentions Fireship", "fireship" in full_text, "Fireship not found in document")
    record("Doc mentions Veritasium", "veritasium" in full_text, "Veritasium not found")

    # Doc must surface at least one specific computed numerical metric per channel.
    # Use Total_Views as anchor - large distinctive numbers.
    fireship_views_str = "121886817"
    fireship_views_alt = "121,886,817"
    record("Doc cites Fireship Total_Views",
           fireship_views_str in full_text or fireship_views_alt in full_text,
           "Expected Fireship Total_Views numeric reference")
    veritasium_views_str = "365002206"
    veritasium_views_alt = "365,002,206"
    record("Doc cites Veritasium Total_Views",
           veritasium_views_str in full_text or veritasium_views_alt in full_text,
           "Expected Veritasium Total_Views numeric reference")


def check_gsheet():
    print("\n=== Check 2: GSheet 'Channel Comparison Analysis' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE title ILIKE %s
    """, ("%Channel Comparison%",))
    sheets = cur.fetchall()
    record("GSheet 'Channel Comparison Analysis' exists", len(sheets) >= 1,
           f"Found: {[s[1] for s in sheets]}")

    if not sheets:
        cur.close(); conn.close()
        return

    ss_id = sheets[0][0]
    cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
    sheet_rows = cur.fetchall()
    sheet_lookup = {r[1].lower(): r[0] for r in sheet_rows}

    comp_id = next((sheet_lookup[t] for t in sheet_lookup if "comparison" in t), None)
    monthly_id = next((sheet_lookup[t] for t in sheet_lookup if "monthly" in t and "fireship" in t), None)
    if not monthly_id:
        monthly_id = next((sheet_lookup[t] for t in sheet_lookup if "monthly" in t or "fireship" in t), None)

    record("Comparison sheet exists in GSheet", comp_id is not None,
           f"Sheets: {list(sheet_lookup.keys())}")
    record("Monthly_Fireship sheet exists in GSheet", monthly_id is not None,
           f"Sheets: {list(sheet_lookup.keys())}")

    if comp_id is not None:
        cur.execute("""
            SELECT row_index, col_index, value FROM gsheet.cells
            WHERE sheet_id = %s ORDER BY row_index, col_index
        """, (comp_id,))
        cells = cur.fetchall()
        # Build matrix
        grid = {}
        for r, c, v in cells:
            grid.setdefault(r, {})[c] = v
        # Identify header row (first row with content) - allow row 0 or 1
        if grid:
            min_r = min(grid.keys())
            header_cells = [(c, grid[min_r][c]) for c in sorted(grid[min_r].keys())]
            headers = [str(v).strip().lower() if v is not None else "" for _, v in header_cells]
            # Build header->col index
            hcol = {}
            for c, val in header_cells:
                if val is not None:
                    hcol[str(val).strip().lower()] = c
            required_headers = ["channel", "video_count", "total_views", "total_likes",
                                "avg_views", "avg_duration_sec", "avg_like_rate",
                                "top_video_title", "top_video_views"]
            for h in required_headers:
                record(f"Comparison sheet has '{h}' column", h in hcol,
                       f"Got headers: {headers}")
            # Check 2 data rows
            data_rows = sorted([r for r in grid if r > min_r])
            record("Comparison sheet has >= 2 data rows",
                   len(data_rows) >= 2,
                   f"Got {len(data_rows)}")
            # Match each channel
            ch_col = hcol.get("channel")
            row_by_channel = {}
            for r in data_rows:
                if ch_col is not None and ch_col in grid[r]:
                    name = str(grid[r][ch_col]).strip()
                    row_by_channel[name.lower()] = grid[r]
            for ch, exp in EXPECTED.items():
                if ch.lower() not in row_by_channel:
                    record(f"Comparison: {ch} row present", False,
                           f"Channels found: {list(row_by_channel.keys())}")
                    continue
                record(f"Comparison: {ch} row present", True)
                row = row_by_channel[ch.lower()]
                for metric, exp_val in exp.items():
                    col = hcol.get(metric.lower())
                    if col is None:
                        continue
                    actual = row.get(col)
                    if metric == "Top_Video_Title":
                        match = (str(actual).strip().lower() == str(exp_val).strip().lower())
                        record(f"Comparison: {ch}.{metric}",
                               match, f"got {actual!r}, expected {exp_val!r}")
                    elif metric == "Avg_Like_Rate":
                        af = get_float(actual)
                        record(f"Comparison: {ch}.{metric} ~= {exp_val}",
                               af is not None and abs(af - exp_val) <= 0.05,
                               f"got {actual}")
                    else:
                        ai = get_int(actual)
                        record(f"Comparison: {ch}.{metric} == {exp_val}",
                               ai is not None and num_close(ai, exp_val),
                               f"got {actual}")

    if monthly_id is not None:
        cur.execute("""
            SELECT row_index, col_index, value FROM gsheet.cells
            WHERE sheet_id = %s ORDER BY row_index, col_index
        """, (monthly_id,))
        cells = cur.fetchall()
        grid = {}
        for r, c, v in cells:
            grid.setdefault(r, {})[c] = v
        if grid:
            min_r = min(grid.keys())
            data_rows = sorted([r for r in grid if r > min_r])
            record("Monthly_Fireship has >= 12 data rows",
                   len(data_rows) >= 12,
                   f"Got {len(data_rows)}")
            # Build month->count map (any 2 columns)
            mcols = sorted(grid[min_r].keys())
            mcol = mcols[0] if mcols else None
            ccol = mcols[1] if len(mcols) > 1 else None
            month_map = {}
            if mcol is not None and ccol is not None:
                for r in data_rows:
                    m = grid[r].get(mcol)
                    c = grid[r].get(ccol)
                    if m is not None:
                        month_map[str(m).strip()] = get_int(c)
            # Verify a couple of distinctive months
            for anchor_m, anchor_c in [("2024-06", 9), ("2025-06", 7), ("2024-12", 4)]:
                actual = month_map.get(anchor_m)
                record(f"Monthly_Fireship: {anchor_m} = {anchor_c}",
                       actual == anchor_c, f"got {actual}; map: {dict(list(month_map.items())[:5])}")

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 3: Email sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT m.to_addr, m.subject FROM email.messages m
        WHERE m.to_addr::text ILIKE %s
        ORDER BY m.date DESC LIMIT 5
    """, ("%media@company.com%",))
    emails = cur.fetchall()

    cur.close()
    conn.close()

    record("Email sent to media@company.com", len(emails) >= 1,
           f"Found: {emails}")
    if emails:
        subject = str(emails[0][1]).lower() if emails[0][1] else ""
        # Tightened: must contain both 'youtube' and 'channel' OR 'channel comparison'
        ok = ("youtube channel comparison" in subject) or \
             ("channel comparison" in subject and "report" in subject) or \
             (("comparison" in subject) and ("channel" in subject) and ("youtube" in subject))
        record("Email subject matches 'YouTube Channel Comparison Report'",
               ok, f"Subject: {emails[0][1]}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    print(f"\nOverall: {PASS_COUNT}/{total} checks passed")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "failed": FAIL_COUNT,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print(f"FAIL ({FAIL_COUNT} failures)")
        sys.exit(1)


if __name__ == "__main__":
    main()
