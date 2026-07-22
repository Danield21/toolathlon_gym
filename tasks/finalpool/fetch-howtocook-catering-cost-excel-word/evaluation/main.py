"""Evaluation script for fetch-howtocook-catering-cost-excel-word."""
import os
import argparse, json, os, sys
import openpyxl

# --- verify_v2 smart primitives ---
_EVAL_DIR = os.path.dirname(os.path.abspath(__file__))
_PACK_ROOT = os.path.abspath(os.path.join(_EVAL_DIR, "..", "..", "..", ".."))
if _PACK_ROOT not in sys.path:
    sys.path.insert(0, _PACK_ROOT)
try:
    from utils.verify_v2 import smart_column_exists
    from utils.verify_v2.eval_helpers import get_sheet_rows_as_dicts, get_gt_column_values
    _HAS_VERIFY_V2 = True
except Exception:
    _HAS_VERIFY_V2 = False

TASK_NAME = "fetch-howtocook-catering-cost-excel-word"


DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    
    excel_path = os.path.join(agent_workspace, "Catering_Cost_Report.xlsx")
    check("Catering_Cost_Report.xlsx exists", os.path.exists(excel_path))
    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)
        gt_path = os.path.join(groundtruth_workspace, "Catering_Cost_Report.xlsx")
        gt_wb = openpyxl.load_workbook(gt_path) if os.path.exists(gt_path) else None

        def check_columns(sheet_name, expected_cols, min_rows):
            """Verify sheet exists, >= min_rows, contains required columns.
            Uses LLM semantic mapping via verify_v2, falls back to strict
            header match."""
            check(f"{sheet_name} sheet exists", sheet_name in wb.sheetnames)
            if sheet_name not in wb.sheetnames:
                return
            _ws = wb[sheet_name]
            _data_rows = list(_ws.iter_rows(min_row=2, values_only=True))
            check(f"{sheet_name} has >= {min_rows} rows",
                  len(_data_rows) >= min_rows, f"got {len(_data_rows)}")
            if _HAS_VERIFY_V2 and gt_wb is not None:
                _raw_headers, _agent_rows = get_sheet_rows_as_dicts(wb, sheet_name)
                for _exp in expected_cols:
                    _gt_vals = get_gt_column_values(gt_wb, sheet_name, _exp)
                    _ok, _matched, _reason = smart_column_exists(
                        expected_col=_exp, agent_headers=_raw_headers,
                        gt_samples=_gt_vals[:3], agent_rows=_agent_rows,
                        task_name=TASK_NAME,
                    )
                    _detail = _reason
                    if _ok and _matched and _matched.lower() != _exp.lower():
                        _detail = f"LLM-mapped to {_matched!r}"
                    check(f"{sheet_name} has {_exp} column", _ok, _detail)
            else:
                _headers = [str(c.value).strip().lower() if c.value else "" for c in _ws[1]]
                for _exp in expected_cols:
                    check(f"{sheet_name} has {_exp} column",
                          _exp.lower() in _headers, f"headers: {_headers[:8]}")

        check("Data_Analysis sheet exists", "Data_Analysis" in wb.sheetnames)
        if "Data_Analysis" in wb.sheetnames and gt_wb is not None:
            ws = wb["Data_Analysis"]
            gt_ws = gt_wb["Data_Analysis"] if "Data_Analysis" in gt_wb.sheetnames else None
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            if gt_ws is not None:
                gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                check(f"Data_Analysis has == {len(gt_rows)} data rows",
                      len(data_rows) == len(gt_rows), f"got {len(data_rows)}, gt={len(gt_rows)}")

            check_columns('Data_Analysis', ['Recipe', 'Calories', 'Protein_g', 'Meets_Guidelines'], len(gt_rows) if gt_ws is not None else 5)

            # Validate Data_Analysis structure (recipes are agent's choice; values are agent-derived)
            if gt_ws is not None:
                gt_headers = [str(c.value).strip().lower() if c.value else "" for c in gt_ws[1]]
                headers = [str(c.value).strip().lower() if c.value else "" for c in ws[1]]
                header_map = {h: i for i, h in enumerate(headers)}
                # Sanity: every row should have a non-empty recipe and reasonable calorie/protein values
                agent_recipe_idx = header_map.get("recipe", 0)
                agent_calories_idx = header_map.get("calories")
                agent_protein_idx = header_map.get("protein_g")
                agent_meets_idx = header_map.get("meets_guidelines")
                rows_with_recipe = 0
                rows_with_calories = 0
                rows_with_protein = 0
                rows_with_meets = 0
                for r in data_rows:
                    if agent_recipe_idx is not None and agent_recipe_idx < len(r) and r[agent_recipe_idx] and str(r[agent_recipe_idx]).strip():
                        rows_with_recipe += 1
                    if agent_calories_idx is not None and agent_calories_idx < len(r):
                        cf = safe_float(r[agent_calories_idx])
                        if cf is not None and cf > 0:
                            rows_with_calories += 1
                    if agent_protein_idx is not None and agent_protein_idx < len(r):
                        pf = safe_float(r[agent_protein_idx])
                        if pf is not None and pf >= 0:
                            rows_with_protein += 1
                    if agent_meets_idx is not None and agent_meets_idx < len(r) and r[agent_meets_idx]:
                        mv = str(r[agent_meets_idx]).strip().lower()
                        if mv in ("yes", "no", "true", "false", "1", "0"):
                            rows_with_meets += 1
                check("Data_Analysis has non-empty Recipe values",
                      rows_with_recipe >= len(gt_rows), f"{rows_with_recipe}/{len(gt_rows)}")
                check("Data_Analysis has positive Calories values",
                      rows_with_calories >= len(gt_rows), f"{rows_with_calories}/{len(gt_rows)}")
                check("Data_Analysis has Protein_g values",
                      rows_with_protein >= len(gt_rows), f"{rows_with_protein}/{len(gt_rows)}")
                check("Data_Analysis has Meets_Guidelines yes/no values",
                      rows_with_meets >= len(gt_rows), f"{rows_with_meets}/{len(gt_rows)}")

        check("Metrics sheet exists", "Metrics" in wb.sheetnames)
        if "Metrics" in wb.sheetnames and gt_wb is not None:
            ws = wb["Metrics"]
            gt_ws = gt_wb["Metrics"] if "Metrics" in gt_wb.sheetnames else None
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            if gt_ws is not None:
                gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                check(f"Metrics has == {len(gt_rows)} rows",
                      len(data_rows) == len(gt_rows), f"got {len(data_rows)}")

            check_columns('Metrics', ['Metric', 'Value'], len(gt_rows) if gt_ws is not None else 4)

            # Validate metric NAMES present (values are agent-derived from agent's chosen recipes,
            # so don't compare numeric values strictly).
            if gt_ws is not None:
                gt_metric_names = [str(r[0]).strip().lower() for r in gt_rows if r and r[0]]
                agent_metric_map = {str(r[0]).strip().lower(): r[1] for r in data_rows if r and r[0]}
                for gt_m in gt_metric_names:
                    # Allow loose substring match (e.g. "average calories" satisfies "avg_calories")
                    norm_m = gt_m.replace("_", "").replace(" ", "")
                    found = any(norm_m in k.replace("_", "").replace(" ", "") or
                                k.replace("_", "").replace(" ", "") in norm_m
                                for k in agent_metric_map.keys())
                    check(f"Metrics has '{gt_m}' (or equivalent)", found,
                          f"agent metrics: {list(agent_metric_map.keys())[:6]}")

                # Cross-validate metric VALUES against agent's own Data_Analysis rows.
                # The metrics are deterministic functions of the chosen recipes; if the
                # agent's Metrics sheet contradicts its own Data_Analysis sheet, that is wrong.
                if "Data_Analysis" in wb.sheetnames:
                    da_ws = wb["Data_Analysis"]
                    da_headers = [str(c.value).strip().lower() if c.value else "" for c in da_ws[1]]
                    da_map = {h: i for i, h in enumerate(da_headers)}
                    da_rows = list(da_ws.iter_rows(min_row=2, values_only=True))
                    da_rows = [r for r in da_rows if r and any(c is not None for c in r)]
                    # Helper: pick agent metric value by fuzzy name
                    def _find_metric(target_norm):
                        for k, v in agent_metric_map.items():
                            kn = k.replace("_", "").replace(" ", "")
                            if target_norm in kn or kn in target_norm:
                                return v
                        return None
                    # Total_Recipes: should equal number of Data_Analysis rows
                    total_v = safe_float(_find_metric("totalrecipes"))
                    if total_v is not None:
                        check(f"Metrics Total_Recipes matches Data_Analysis row count ({len(da_rows)})",
                              abs(total_v - len(da_rows)) <= 0,
                              f"got {total_v}, da_rows={len(da_rows)}")
                    # Recipes_Meeting_Guidelines: count of rows with Meets_Guidelines == yes/true
                    if "meets_guidelines" in da_map:
                        mi = da_map["meets_guidelines"]
                        meets_count = sum(
                            1 for r in da_rows
                            if mi < len(r) and r[mi] is not None
                            and str(r[mi]).strip().lower() in ("yes", "true", "1")
                        )
                        meet_v = safe_float(_find_metric("recipesmeetingguidelines"))
                        if meet_v is not None:
                            check(f"Metrics Recipes_Meeting_Guidelines matches count of Yes ({meets_count})",
                                  abs(meet_v - meets_count) <= 0,
                                  f"got {meet_v}, da count={meets_count}")
                    # Avg_Calories: mean of Calories column
                    if "calories" in da_map:
                        ci = da_map["calories"]
                        cvals = [safe_float(r[ci]) for r in da_rows if ci < len(r)]
                        cvals = [v for v in cvals if v is not None]
                        if cvals:
                            avg_c = sum(cvals) / len(cvals)
                            ac_v = safe_float(_find_metric("avgcalories"))
                            if ac_v is not None:
                                tol = max(2.0, abs(avg_c) * 0.05)
                                check(f"Metrics Avg_Calories matches Data_Analysis mean (~{avg_c:.1f})",
                                      abs(ac_v - avg_c) <= tol,
                                      f"got {ac_v}, computed={avg_c}")
                    # Avg_Protein: mean of Protein_g column
                    if "protein_g" in da_map:
                        pi = da_map["protein_g"]
                        pvals = [safe_float(r[pi]) for r in da_rows if pi < len(r)]
                        pvals = [v for v in pvals if v is not None]
                        if pvals:
                            avg_p = sum(pvals) / len(pvals)
                            ap_v = safe_float(_find_metric("avgprotein"))
                            if ap_v is not None:
                                tol = max(0.5, abs(avg_p) * 0.05)
                                check(f"Metrics Avg_Protein matches Data_Analysis mean (~{avg_p:.2f})",
                                      abs(ap_v - avg_p) <= tol,
                                      f"got {ap_v}, computed={avg_p}")

        check("Recommendations sheet exists", "Recommendations" in wb.sheetnames)
        if "Recommendations" in wb.sheetnames:
            ws = wb["Recommendations"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            check_columns('Recommendations', ['Priority', 'Action'], 2)
            gt_ws = gt_wb["Recommendations"] if (gt_wb is not None and "Recommendations" in gt_wb.sheetnames) else None
            if gt_ws is not None:
                gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
                check(f"Recommendations has >= {len(gt_rows)} rows",
                      len(data_rows) >= len(gt_rows), f"got {len(data_rows)}")

        # Check Word document - require executive summary, key findings, recommendations
        word_path = os.path.join(agent_workspace, "Catering_Cost_Analysis.docx")
        check("Catering_Cost_Analysis.docx exists", os.path.exists(word_path),
              f"expected at {word_path}")
        if os.path.exists(word_path):
            from docx import Document
            doc = Document(word_path)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            # Task asks for executive summary, key findings, recommendations.
            # GT uses 'Overview' (acceptable as exec-summary equivalent) + 'Key Findings' + 'Recommendations'.
            # Prefer to match against actual heading paragraphs to ensure structural sections, not just keywords in body.
            heading_texts = [p.text.strip().lower() for p in doc.paragraphs if p.style.name.startswith("Heading")]
            heading_blob = " ||| ".join(heading_texts)
            # Section 1: executive summary / overview
            has_exec = (
                "executive summary" in heading_blob or "summary" in heading_blob or "overview" in heading_blob
            ) if heading_texts else (
                "executive summary" in text or "summary" in text or "overview" in text
            )
            check("Word doc has executive summary or overview heading",
                  has_exec, f"headings: {heading_texts[:5]}")
            # Section 2: key findings
            has_findings = (
                "key findings" in heading_blob or "findings" in heading_blob
            ) if heading_texts else (
                "key findings" in text or "findings" in text
            )
            check("Word doc has key findings heading",
                  has_findings, f"headings: {heading_texts[:5]}")
            # Section 3: recommendations
            has_recs = (
                "recommendation" in heading_blob
            ) if heading_texts else ("recommendation" in text)
            check("Word doc has recommendations heading",
                  has_recs, f"headings: {heading_texts[:5]}")
            # Substantive body
            check("Word doc has substantive content",
                  len(text) > 200, f"text length: {len(text)}")
            # At least 3 distinct heading paragraphs (forces structured doc, not a single paragraph)
            check("Word doc has >= 3 heading-styled sections",
                  len(heading_texts) >= 3, f"found {len(heading_texts)} headings")

        check("cook_catering_processor.py exists", os.path.exists(os.path.join(agent_workspace, "cook_catering_processor.py")))


    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
