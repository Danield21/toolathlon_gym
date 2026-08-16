"""Evaluation for sf-support-sla-coaching-word-gcal."""
import argparse
import json
import os
import re
import sys

import psycopg2

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# gcal.events.start_datetime is TIMESTAMPTZ; bare sdt.hour silently compares
# wrong in non-UTC PG sessions. Use gcal_helpers.
# ──────────────────────────────────────────────────────────────────────────
# task.md line 5: "10:00 AM to 11:00 AM" (enterprise support, ET default)
EXPECTED_TIMEZONE = "America/New_York"

from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol):
    try:
        return abs(float(a) - float(b)) <= tol
    except Exception:
        return False


def _fetch_priority_stats():
    """Fetch ticket counts, avg response, sla hours, compliance per priority."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT "PRIORITY", COUNT(*) as ticket_count,
                   AVG("RESPONSE_TIME_HOURS")::float as avg_hr,
                   AVG("SLA_HOURS")::float as sla_hr,
                   SUM(CASE WHEN "RESPONSE_TIME_HOURS" <= "SLA_HOURS" THEN 1 ELSE 0 END)::float / COUNT(*) as compliance
            FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
            GROUP BY "PRIORITY"
            ORDER BY "PRIORITY"
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        out = {}
        for prio, cnt, avg_hr, sla_hr, comp in rows:
            out[prio] = {
                "count": int(cnt),
                "avg_hr": float(avg_hr),
                "sla_hr": float(sla_hr),
                "compliance": float(comp),
            }
        return out
    except Exception as e:
        # Fail loudly rather than silently falling back to stale constants —
        # masking DB drift would let invalid eval pass.
        record(f"DB fetch for SLA priority stats", False, str(e))
        return {}


def check_word(agent_ws):
    print("Checking Word document...")
    docx_path = os.path.join(agent_ws, "SLA_Coaching_Plan.docx")
    if not os.path.exists(docx_path):
        record("Word file SLA_Coaching_Plan.docx exists", False, docx_path)
        return
    record("Word file SLA_Coaching_Plan.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word file readable", False, str(e))
        return
    record("Word file readable", True)

    # Body text only (paragraphs); table contents separately
    para_text = "\n".join(p.text for p in doc.paragraphs)
    para_text_lower = para_text.lower()

    # Title
    record("Title 'SLA Performance Coaching Plan' present",
           "sla performance coaching plan" in para_text_lower,
           f"para text head: {para_text[:300]}")

    # Required section headers
    for sec in ["SLA Performance Summary", "Priority Analysis", "Coaching Recommendations"]:
        record(f"Section '{sec}' present",
               sec.lower() in para_text_lower,
               f"para text head: {para_text[:300]}")

    # Table content
    if not doc.tables:
        record("Document has at least 1 table", False)
        return
    record("Document has at least 1 table", True)

    # Find a table with these columns
    target_table = None
    expected_cols = ["priority", "total_tickets", "avg_response_hours", "sla_hours", "compliance_rate"]
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip().lower().replace(" ", "_") for c in table.rows[0].cells]
        if all(any(ec.replace(" ", "_") in h.replace(" ", "_") for h in header) for ec in expected_cols):
            target_table = table
            break
        # Looser match: each expected col appears as substring in any header
        normalized_header = " | ".join(header)
        if all(ec in normalized_header for ec in expected_cols):
            target_table = table
            break
    record("Table with Priority/Total_Tickets/Avg_Response_Hours/SLA_Hours/Compliance_Rate columns",
           target_table is not None)

    # Verify per-priority values in any table cell text
    full_text = para_text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    full_lower = full_text.lower()

    stats = _fetch_priority_stats()
    for priority, s in stats.items():
        # priority name
        record(f"Doc mentions priority '{priority}'",
               priority.lower() in full_lower,
               "")
        # Total_Tickets count
        record(f"Doc mentions Total_Tickets {s['count']} for {priority}",
               re.search(rf"\b{s['count']:,}\b|\b{s['count']}\b", full_text) is not None,
               f"Expected count {s['count']}")
        # SLA hours integer
        record(f"Doc mentions SLA hours {int(s['sla_hr'])} for {priority}",
               re.search(rf"\b{int(s['sla_hr'])}\b", full_text) is not None,
               "")

    # At least 3 recommendation paragraphs/bullets, each starting with a distinct
    # action verb (heuristic: first token after any leading bullet/number).
    rec_idx = para_text_lower.find("coaching recommendations")
    rec_paragraphs = []
    if rec_idx >= 0:
        idx_in_paras = 0
        for i, p in enumerate(doc.paragraphs):
            if "coaching recommendations" in p.text.lower():
                idx_in_paras = i
                break
        rec_paragraphs = [p.text.strip() for p in doc.paragraphs[idx_in_paras + 1:]
                          if p.text.strip()]
    # Extract first non-bullet token from each line/bullet
    distinct_verbs = set()
    bullet_re = re.compile(r"^\s*(?:[-*•]|\d+[\.\)])\s*")
    for para in rec_paragraphs:
        for line in para.splitlines():
            line = bullet_re.sub("", line.strip())
            if not line:
                continue
            first = line.split()[0].lower().rstrip(",.;:")
            distinct_verbs.add(first)
    # Also check bullet lines in the joined recommendations text in case paras were split
    rec_text = para_text[rec_idx:] if rec_idx >= 0 else ""
    for m in re.finditer(r"^\s*(?:[-*•]|\d+[\.\)])\s+(\w+)",
                          rec_text, re.MULTILINE):
        distinct_verbs.add(m.group(1).lower().rstrip(",.;:"))
    record("At least 3 distinct-verb recommendations",
           len(distinct_verbs) >= 3,
           f"Distinct first words: {sorted(distinct_verbs)[:10]}")


def check_gcal():
    print("\nChecking Google Calendar events...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        quarters = [
            ("Q1 SLA Coaching Review", "2026-03-15"),
            ("Q2 SLA Coaching Review", "2026-06-15"),
            ("Q3 SLA Coaching Review", "2026-09-15"),
            ("Q4 SLA Coaching Review", "2026-12-15"),
        ]
        for q_title, exp_date in quarters:
            cur.execute("""
                SELECT summary, start_datetime, end_datetime
                FROM gcal.events
                WHERE start_datetime::date = %s
            """, (exp_date,))
            rows = cur.fetchall()
            match = None
            for summary, sdt, edt in rows:
                if (summary or "").strip().lower() == q_title.lower():
                    match = (summary, sdt, edt)
                    break
            record(f"Event '{q_title}' on {exp_date} found", match is not None,
                   f"Events on date: {[r[0] for r in rows]}")
            if match:
                _, sdt, edt = match
                # gcal.events.start_datetime is TIMESTAMPTZ; use session-tz-
                # independent helper to extract ET hour.
                _sd, s_h, _sm = get_zone_components(sdt, EXPECTED_TIMEZONE)
                _ed2, e_h, _em = get_zone_components(edt, EXPECTED_TIMEZONE)
                hours_ok = (s_h is not None and e_h is not None
                            and s_h == 10 and e_h == 11)
                record(f"'{q_title}' time 10:00-11:00",
                       hours_ok, f"start={sdt}, end={edt}")

        cur.close()
        conn.close()
    except Exception as e:
        record("GCal check error", False, str(e))


def check_email():
    print("\nChecking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
        emails = cur.fetchall()
        cur.close()
        conn.close()

        # Filter to those addressed to the manager
        candidates = []
        for subj, from_addr, to_addr, body in emails:
            to_str = json.dumps(to_addr).lower() if isinstance(to_addr, (list, dict)) else str(to_addr or "").lower()
            if "support_manager@company.com" in to_str:
                candidates.append((subj, from_addr, to_addr, body))
        record("Email to support_manager@company.com sent", len(candidates) >= 1,
               f"Found {len(candidates)} candidates")
        if candidates:
            subj, _, _, body = candidates[0]
            sl = (subj or "").lower()
            record("Email subject references coaching plan",
                   "coaching" in sl and ("plan" in sl or "sla" in sl),
                   f"Subject: {subj}")
            bl = (body or "").lower()
            record("Body summarizes SLA findings (mentions priorities)",
                   "high" in bl and "medium" in bl and "low" in bl,
                   f"Body head: {bl[:300]}")
            record("Body mentions quarterly review calendar events",
                   ("quarterly" in bl or "quarter" in bl or "q1" in bl) and ("calendar" in bl or "schedul" in bl or "review" in bl),
                   f"Body head: {bl[:300]}")
    except Exception as e:
        record("Email check error", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    agent_ws = args.agent_workspace or task_root

    check_word(agent_ws)
    check_gcal()
    check_email()

    all_passed = (FAIL_COUNT == 0)
    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_passed else 'FAIL'}")
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": all_passed}, f)
    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
