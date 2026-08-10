#!/usr/bin/env python3
"""Evaluation for supplier-contract-compliance-audit.

Verifies the agent produced concrete supplier compliance audit deliverables:
  - XLSX compliance assessment (per-supplier scores) and corrective actions
  - DOCX audit findings and summary
  - PDF compliance report (or referenced compliance_report file)
  - Email audit reports to supplier contacts
  - Calendar meetings with under-performing suppliers
"""
import argparse
import datetime
import json
import os
import sys

import psycopg2

DB = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": os.environ.get("PGUSER", "eigent"),
    "password": os.environ.get("PGPASSWORD", "camel"),
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def _to_float(v):
    """Robustly coerce a value to float, or None if not parseable.

    Handles int/float, None, and strings with thousands separators, currency
    symbols ($, ¥, €), percent signs, and surrounding whitespace.
    """
    if v is None:
        return None
    if isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    for ch in (",", "$", "¥", "€", "%", " "):
        s = s.replace(ch, "")
    try:
        return float(s)
    except ValueError:
        return None


def num_close(a, b, tol=1.0):
    fa, fb = _to_float(a), _to_float(b)
    if fa is not None and fb is not None:
        return abs(fa - fb) <= tol
    # One side not numeric: fall back to case-insensitive string comparison
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


# Distinctive case-insensitive tokens that identify a supplier. A correct
# deliverable may refer to "TechServices" or "Consulting Group" instead of
# the exact full name in supplier_list.csv, so matching must accept these.
SUPPLIER_TOKENS = {
    "TechServices Inc": ["techservices", "tech services"],
    "Global Logistics": ["global logistics"],
    "Quality Manufacturing": ["quality manufacturing"],
    "Consulting Group Ltd": ["consulting group", "consulting group ltd"],
    "Support Solutions": ["support solutions"],
}


def _supplier_tokens(supplier):
    return SUPPLIER_TOKENS.get(supplier, [supplier.lower()])


def _supplier_mentioned(text_lower, supplier):
    return any(tok in text_lower for tok in _supplier_tokens(supplier))


def _mentioned_suppliers(text_lower):
    return {s for s in SUPPLIERS if _supplier_mentioned(text_lower, s)}


# The GT column labels are not fully derivable from task.md (e.g. the task's
# phase-2/4 dimensions differ from "Scope Clarity"/"IP Rights"/...), so the
# xlsx column check must accept reasonable synonyms rather than byte-equal
# GT headers.
#   - compliance_assessment.xlsx: only the supplier identity and an overall
#     score column are mandated; dimension names are free-form.
#   - corrective_actions.xlsx: all seven columns are explicitly listed in the
#     task wording, so each must appear under some reasonable label.
COLUMN_ALIASES = {
    "assessment": {
        "supplier": ["supplier", "vendor", "company", "provider"],
        # Identifies the overall-score column. Dimension columns named
        # "SLA Score" etc. are excluded here on purpose (see _is_score_col).
        "score": ["overall", "compliance", "rating", "grade", "total"],
    },
    "actions": {
        "supplier": ["supplier", "vendor", "company", "provider"],
        "issue": ["issue", "finding", "problem", "non-compliance",
                  "noncompliance", "non compliance", "gap"],
        "required action": ["action", "remediation", "remedial"],
        "owner": ["owner", "responsible", "assignee", "accountable", "point of contact"],
        "target date": ["target date", "date", "deadline", "due"],
        "est. cost": ["cost", "expense", "amount", "estimate", "budget"],
        "status": ["status", "state", "progress"],
    },
}


def _header_has(hdr_low, aliases):
    return any(a in hdr_low for a in aliases)


def _is_supplier_col(h):
    return _header_has(h, COLUMN_ALIASES["assessment"]["supplier"])


def _is_score_col(h):
    hl = h.strip().lower()
    if _header_has(h, COLUMN_ALIASES["assessment"]["score"]):
        return True
    # A bare "Score"/"Overall Score" column is the overall-score column; a
    # dimension column is usually qualified ("SLA Score"), which is handled
    # by the aliases above (it won't contain overall/compliance/rating/...).
    return hl in ("score", "overall score", "total score", "final score")


def _docx_text(doc):
    """Full document text, including table cell content (python-docx
    paragraphs alone miss text stored inside tables)."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _launch_window(launch_time):
    """Return a datetime marking the earliest acceptable entry for this run,
    or None when launch_time is unavailable/unparseable (fall back to a full
    table scan, preserving the old behaviour)."""
    if not launch_time:
        return None
    s = str(launch_time).strip().strip('"')
    if not s:
        return None
    for fmt in (
        "%Y-%m-%d %H:%M:%S %A",       # harness default, e.g. '2026-08-07 10:00:00 Thursday'
        "%Y-%m-%dT%H:%M:%S.%f%z",
        "%Y-%m-%dT%H:%M:%S%z",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%d",
    ):
        try:
            dt = datetime.datetime.strptime(s, fmt)
            break
        except ValueError:
            continue
    else:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=datetime.timezone.utc)
    # Generous margin absorbs timezone ambiguity and clock skew so all
    # entries created during the run are kept while clearly-stale leftovers
    # from other tasks are excluded.
    return dt - datetime.timedelta(days=1)


# Suppliers from initial_workspace/supplier_list.csv (kept stable for audit topic)
SUPPLIERS = [
    "TechServices Inc",
    "Global Logistics",
    "Quality Manufacturing",
    "Consulting Group Ltd",
    "Support Solutions",
]


def _find_xlsx_by_keywords(workspace, keywords, exclude=None,
                            anti_keywords=None):
    """Score xlsx files by keyword matches; filename match weighs 10x more than content.

    anti_keywords: filename tokens that, when matched, subtract 5 from score
    so a file whose name overwhelmingly belongs to a *different* target
    can't accidentally win. Used to disambiguate compliance_assessment vs
    corrective_actions when both contain 'compliance' content keywords.
    """
    import glob
    import openpyxl
    exclude = exclude or set()
    anti_keywords = anti_keywords or []
    scored = []
    for path in sorted(glob.glob(os.path.join(workspace, "*.xlsx"))):
        if path in exclude:
            continue
        fname_low = os.path.basename(path).lower()
        if fname_low.startswith("~$") or fname_low.startswith("."):
            continue
        score = 0
        for kw in keywords:
            if kw in fname_low:
                score += 10
        for ak in anti_keywords:
            if ak in fname_low:
                score -= 5
        try:
            wb = openpyxl.load_workbook(path, data_only=False, read_only=True)
            content_text = " ".join(s.lower() for s in wb.sheetnames)
            for ws in wb.worksheets:
                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    if row_count > 5:
                        break
                    row_count += 1
                    for cell in row:
                        if cell is not None:
                            content_text += " " + str(cell).lower()
            wb.close()
            for kw in keywords:
                if kw in content_text:
                    score += 1
        except Exception:
            pass
        if score > 0:
            scored.append((score, path))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [p for _, p in scored]


def _resolve_deliverable(workspace, gt_fname, find_fn, keywords, anti_keywords,
                         used_paths, glob_pattern):
    """Resolve the file for a deliverable target with decreasing strictness:
    1) exact GT filename, 2) best keyword match not already used,
    3) best keyword match with reuse allowed (a single merged file can cover
    several targets), 4) any file of the type.
    Content checks later still gate quality, so these permissive fallbacks
    only broaden *existence* detection and never let an unrelated file pass
    on content.
    """
    exact = os.path.join(workspace, gt_fname)
    if os.path.isfile(exact):
        return exact
    cands = find_fn(workspace, keywords, anti_keywords=anti_keywords)
    for c in cands:
        if c not in used_paths:
            return c
    if cands:
        return cands[0]
    import glob
    any_files = sorted(glob.glob(os.path.join(workspace, glob_pattern)))
    for p in any_files:
        base = os.path.basename(p)
        if base.startswith("~$") or base.startswith("."):
            continue
        # Last resort allows reuse of an already-picked file: content checks
        # later still gate quality, so this only broadens existence detection.
        return p
    return None


def check_xlsx_content(workspace, gt_workspace="."):
    """Locate compliance_assessment + corrective_actions xlsx by filename or content."""
    print("\n=== Check: XLSX files ===")
    import openpyxl

    targets = [
        ("compliance_assessment.xlsx", "compliance assessment xlsx",
         ["compliance_assessment", "compliance assessment", "compliance",
          "assessment", "scorecard"],
         ["corrective", "remediation", "action_plan"]),
        ("corrective_actions.xlsx", "corrective actions xlsx",
         ["corrective_actions", "corrective actions", "corrective",
          "action", "remediation"],
         ["assessment", "scorecard"]),
    ]

    used_paths = set()
    for gt_fname, label, keywords, anti in targets:
        chosen = _resolve_deliverable(workspace, gt_fname,
                                      _find_xlsx_by_keywords, keywords, anti,
                                      used_paths, "*.xlsx")
        if chosen is None:
            record(f"xlsx for {label} exists", False,
                   f"No xlsx with keywords {keywords[:3]} found")
            continue
        used_paths.add(chosen)
        record(f"xlsx for {label} exists ({os.path.basename(chosen)})", True)
        fname = os.path.basename(chosen)

        try:
            wb = openpyxl.load_workbook(chosen, data_only=False)
        except Exception as e:
            record(f"xlsx {fname} readable", False, str(e))
            continue
        # All cells flattened to lower-case for content scan
        all_text = ""
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        all_text += str(cell) + "\n"
        all_text_lower = all_text.lower()
        mentioned = _mentioned_suppliers(all_text_lower)
        if "compliance" in label and "assessment" in label:
            record(f"{fname} mentions every supplier",
                   len(mentioned) == len(SUPPLIERS),
                   f"Missing: {[s for s in SUPPLIERS if s not in mentioned]}")
        else:
            record(f"{fname} mentions at least 4 of 5 suppliers",
                   len(mentioned) >= 4,
                   f"Mentioned {len(mentioned)} suppliers")

        gt_path = os.path.join(gt_workspace, gt_fname)
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=False)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None and len(wb.sheetnames) == 1 and len(gt_wb.sheetnames) == 1:
                    a_ws = wb[wb.sheetnames[0]]
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' present in {fname}", False, f"Available: {wb.sheetnames}")
                    continue
                record(f"GT sheet '{gt_sname}' present in {fname}", True)

                a_rows = list(a_ws.iter_rows(values_only=True))
                gt_rows = list(gt_ws.iter_rows(values_only=True))

                def find_header_row(rows):
                    for i, r in enumerate(rows):
                        non_none = [v for v in r if v is not None]
                        if len(non_none) >= 4 and all(isinstance(v, str) for v in non_none):
                            return i
                    return 0

                a_hdr_idx = find_header_row(a_rows)
                gt_hdr_idx = find_header_row(gt_rows)
                a_hdrs = [str(v).strip().lower() if v else "" for v in a_rows[a_hdr_idx]]
                gt_hdrs = [str(v).strip().lower() if v else "" for v in gt_rows[gt_hdr_idx]]
                hdrs = [h for h in a_hdrs if h]
                if "assessment" in gt_fname:
                    # compliance_assessment.xlsx: dimension names are not
                    # mandated, so require only supplier + an overall score
                    # column + a couple of dimension columns.
                    has_supplier = any(_is_supplier_col(h) for h in hdrs)
                    has_score = any(_is_score_col(h) for h in hdrs)
                    dims = [h for h in hdrs if not _is_supplier_col(h) and not _is_score_col(h)]
                    ok = has_supplier and has_score and len(dims) >= 2
                    detail = (f"supplier_col={has_supplier}, score_col={has_score}, "
                              f"dimension_cols={len(dims)} (need >= 2)")
                else:
                    # corrective_actions.xlsx: all seven columns are explicitly
                    # required by the task wording; accept any reasonable label.
                    missing_hdrs = [
                        gt_h for gt_h, aliases in COLUMN_ALIASES["actions"].items()
                        if not any(_header_has(h, aliases) for h in hdrs)
                    ]
                    ok = len(missing_hdrs) == 0
                    detail = (f"missing {missing_hdrs}" if missing_hdrs
                              else "all 7 required columns present")
                record(f"{fname}/{gt_sname}: required columns present", ok, detail)

                a_data_count = sum(1 for r in a_rows[a_hdr_idx + 1:] if any(v is not None for v in r))
                record(f"{fname}/{gt_sname}: at least 5 data rows", a_data_count >= 5,
                       f"got {a_data_count}")
            gt_wb.close()
        wb.close()


def _find_docx_by_keywords(workspace, keywords, exclude=None,
                            anti_keywords=None):
    import glob
    from docx import Document
    exclude = exclude or set()
    anti_keywords = anti_keywords or []
    scored = []
    for path in sorted(glob.glob(os.path.join(workspace, "*.docx"))):
        if path in exclude:
            continue
        fname_low = os.path.basename(path).lower()
        if fname_low.startswith("~$") or fname_low.startswith("."):
            continue
        score = 0
        for kw in keywords:
            if kw in fname_low:
                score += 10
        for ak in anti_keywords:
            if ak in fname_low:
                score -= 5
        try:
            doc = Document(path)
            text_low = _docx_text(doc).lower()
            for kw in keywords:
                if kw in text_low:
                    score += 1
        except Exception:
            pass
        if score > 0:
            scored.append((score, path))
    scored.sort(key=lambda x: (-x[0], x[1]))
    return [p for _, p in scored]


DOCX_RULES = [
    {
        "gt_fname": "audit_findings.docx",
        "label": "audit findings docx",
        "keywords": ["audit_findings", "audit findings", "findings", "detailed", "narrative"],
        "anti": ["summary", "overview"],
        "min_len": 300,
        "require_all_suppliers": True,
        "topic": ["compliance", "audit"],
        "signal": ["find", "recommend", "issue", "gap", "risk", "impact", "evidence"],
    },
    {
        "gt_fname": "audit_summary.docx",
        "label": "audit summary docx",
        "keywords": ["audit_summary", "audit summary", "summary", "executive", "overview"],
        "anti": ["findings_detail", "narrative"],
        "min_len": 300,
        "require_all_suppliers": False,
        "topic": ["compliance", "audit"],
        "signal": ["find", "recommend", "issue", "gap", "risk", "overview", "summary"],
    },
]


def check_docx_content(workspace):
    """Locate audit findings + audit summary docx by filename or content keyword.

    The findings document is the detailed per-supplier report, so every
    supplier must be named (matching tolerates partial names) and the text
    must be substantive. The summary is an executive overview; task.md only
    asks for "overall compliance status, key findings, and strategic
    recommendations", so it is NOT required to name every supplier nor to be
    800+ chars -- it is held to topical keywords and a modest length floor.
    """
    print("\n=== Check: DOCX files ===")
    from docx import Document

    used_paths = set()
    for rule in DOCX_RULES:
        gt_fname, label = rule["gt_fname"], rule["label"]
        chosen = _resolve_deliverable(workspace, gt_fname,
                                      _find_docx_by_keywords, rule["keywords"],
                                      rule["anti"], used_paths, "*.docx")
        if chosen is None:
            record(f"docx for {label} exists", False,
                   f"No docx with keywords {rule['keywords'][:3]} found")
            continue
        used_paths.add(chosen)
        fname = os.path.basename(chosen)
        record(f"docx for {label} exists ({fname})", True)
        try:
            doc = Document(chosen)
        except Exception as e:
            record(f"docx {fname} readable", False, str(e))
            continue
        text = _docx_text(doc)
        text_lower = text.lower()
        if rule["require_all_suppliers"]:
            mentioned = _mentioned_suppliers(text_lower)
            for s in SUPPLIERS:
                record(f"{fname} mentions supplier '{s}'", s in mentioned, "missing")
        record(f"{fname} has >= {rule['min_len']} chars",
               len(text) >= rule["min_len"], f"got {len(text)}")
        topic_ok = any(k in text_lower for k in rule["topic"])
        record(f"{fname} mentions {'/'.join(rule['topic'])}", topic_ok, "missing")
        signal_ok = any(k in text_lower for k in rule["signal"])
        record(f"{fname} discusses findings/recommendations", signal_ok, "missing")


def check_pdf(workspace):
    print("\n=== Check: PDF compliance report ===")
    import glob
    pdf_candidates = []
    for f in os.listdir(workspace):
        if f.lower().endswith(".pdf") and "compliance" in f.lower():
            pdf_candidates.append(f)
    # "Professional audit report format" can legitimately be delivered as a
    # docx (audit_findings / audit_summary) without a derived PDF. A
    # compliance PDF is therefore only *required* when no report docx exists.
    report_docx = [p for p in glob.glob(os.path.join(workspace, "*.docx"))
                   if not os.path.basename(p).startswith("~$")
                   and not os.path.basename(p).startswith(".")]
    satisfied = len(pdf_candidates) >= 1 or len(report_docx) >= 1
    record("Compliance PDF report present (or docx audit report)",
           satisfied,
           f"pdfs: {pdf_candidates}, report docx: {len(report_docx)}")


def _read_supplier_contacts(workspace):
    """Parse supplier_list.csv to derive expected supplier email domains and
    full email addresses dynamically. Returns ([domains], [emails])."""
    import csv
    sup_csv = os.path.join(workspace, "supplier_list.csv")
    domains, emails = [], []
    if not os.path.isfile(sup_csv):
        return domains, emails
    try:
        with open(sup_csv, newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                addr = (row.get("Contact_Email") or "").strip().lower()
                if not addr or "@" not in addr:
                    continue
                emails.append(addr)
                domains.append(addr.split("@", 1)[1])
    except Exception:
        pass
    return domains, emails


def check_email(agent_workspace, launch_time):
    print("\n=== Check: Audit emails to suppliers ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
    except psycopg2.Error as e:
        record("DB connect", False, str(e))
        return
    cur.execute("SELECT id FROM email.folders WHERE LOWER(name) IN ('sent','drafts')")
    folder_ids = [r[0] for r in cur.fetchall()]
    if not folder_ids:
        record("Email folders for sent/drafts exist", False, "")
        cur.close()
        conn.close()
        return
    placeholders = ",".join(["%s"] * len(folder_ids))
    sql = (f"SELECT subject, to_addr, body_text FROM email.messages "
           f"WHERE folder_id IN ({placeholders})")
    params = list(folder_ids)
    win = _launch_window(launch_time)
    if win is not None:
        # Scope to this run's messages so stale rows from other tasks /
        # earlier runs don't produce false PASS.
        sql += " AND created_at >= %s"
        params.append(win)
    cur.execute(sql, params)
    sent = cur.fetchall()
    cur.close()
    conn.close()

    # An email that carries the audit report should be recognizable either by
    # its subject or its body (the report text itself contains audit/compliance
    # language), so a generic subject on an otherwise-correct report email does
    # not produce a false negative.
    audit_subject_or_body = any(
        ("audit" in ((s or "") + (b or "")).lower())
        or ("compliance" in ((s or "") + (b or "")).lower())
        for s, _, b in sent
    )
    record("At least 1 audit-related email sent", audit_subject_or_body,
           f"found {len(sent)} sent emails")
    # Derive expected supplier domains/emails from supplier_list.csv.
    # Falls back to the supplierN.com convention (which the bundled CSV uses).
    derived_domains, derived_emails = _read_supplier_contacts(agent_workspace)
    if not derived_domains:
        derived_domains = ["supplier1.com", "supplier2.com", "supplier3.com",
                           "supplier4.com", "supplier5.com"]
    matched_domains = set()
    for _, to, _ in sent:
        to_str = json.dumps(to) if not isinstance(to, str) else to
        to_low = to_str.lower()
        # Match either the full email address (preferred) or domain
        for em in derived_emails:
            if em in to_low:
                matched_domains.add(em.split("@", 1)[1])
        for dom in derived_domains:
            if dom in to_low:
                matched_domains.add(dom)
    record("Emails reach at least 2 supplier contacts",
           len(matched_domains) >= 2,
           f"found supplier domains: {matched_domains}")


def check_gcal(launch_time):
    print("\n=== Check: Calendar meetings with suppliers ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
    except psycopg2.Error as e:
        record("DB connect", False, str(e))
        return
    sql = "SELECT id, summary, description FROM gcal.events"
    params = []
    win = _launch_window(launch_time)
    if win is not None:
        # Scope to this run's events so stale rows from other tasks / earlier
        # runs don't produce false PASS.
        sql += " WHERE created >= %s"
        params.append(win)
    cur.execute(sql, params)
    events = cur.fetchall()
    cur.close()
    conn.close()
    record("At least 1 calendar event created", len(events) >= 1, f"found {len(events)}")
    matched = 0
    for eid, summary, desc in events:
        text = f"{summary or ''} {desc or ''}".lower()
        if any(s.lower() in text for s in SUPPLIERS) or "supplier" in text or "vendor" in text:
            matched += 1
    record("At least 1 calendar event references a supplier/vendor",
           matched >= 1, f"matched {matched}/{len(events)}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    ws = args.agent_workspace
    if not os.path.isdir(ws):
        print(f"Agent workspace not found: {ws}")
        sys.exit(1)

    # GT self-test toleration: when agent==groundtruth (V1 GT-only smoke
    # test), the email/gcal checks would fail because no email/event was
    # actually sent.
    is_gt_self_test = (
        args.agent_workspace and args.groundtruth_workspace
        and os.path.abspath(args.agent_workspace) == os.path.abspath(args.groundtruth_workspace)
    )

    check_xlsx_content(ws, args.groundtruth_workspace)
    check_docx_content(ws)
    check_pdf(ws)
    if not is_gt_self_test:
        check_email(ws, args.launch_time)
        check_gcal(args.launch_time)
    else:
        print("  [SKIP] email and gcal checks (GT self-test mode)")

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    all_ok = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if all_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": all_ok}, f)

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
