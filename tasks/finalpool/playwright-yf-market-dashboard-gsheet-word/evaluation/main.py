"""
Evaluation script for playwright-yf-market-dashboard-gsheet-word task.

Checks:
1. Google Sheet "Weekly_Market_Analysis" with sector and stock data (numeric values validated)
2. Weekly_Market_Report.docx with executive summary (specific sectors / numbers required)
3. Email sent to committee@investmentfirm.com with analysis (recipient + top sector validated)
"""

import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=0.5):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def to_float(v):
    if v is None:
        return None
    s = str(v).strip().replace("%", "").replace(",", "")
    try:
        return float(s)
    except ValueError:
        return None


# Expected sector data from dashboard (single source of truth)
SECTOR_DATA = {
    "communication services": {"weekly": 2.8, "index": 118.5, "ytd": 12.4},
    "consumer cyclical":      {"weekly": -1.2, "index": 105.3, "ytd": 5.7},
    "financial services":     {"weekly": 1.5, "index": 112.8, "ytd": 8.9},
    "healthcare":             {"weekly": 0.3, "index": 98.7, "ytd": -2.1},
    "energy":                 {"weekly": 3.1, "index": 122.4, "ytd": 15.2},
}

# Top performer (highest weekly_return) for cross-validation
TOP_SECTOR_NAME = "Energy"
TOP_SECTOR_WEEKLY = 3.1
NEGATIVE_SECTORS = ["Consumer Cyclical"]  # weekly < 0

STOCK_TO_SECTOR = {
    "GOOGL": "Communication Services",
    "AMZN": "Consumer Cyclical",
    "JPM": "Financial Services",
    "JNJ": "Healthcare",
    "XOM": "Energy",
}


def _build_grid(cells):
    grid = {}
    for row, col, val in cells:
        grid.setdefault(row, {})[col] = val
    return grid


def _row_starts_with(row_dict, sector_lower):
    """Return True if column 0 of the row matches the given sector name (case-insensitive)."""
    val = row_dict.get(0) or row_dict.get(1)  # try col 0 first
    if val is None:
        return False
    return sector_lower in str(val).strip().lower()


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()

    found_ss = None
    found_title = None
    for ss_id, title in spreadsheets:
        t = (title or "").strip().lower()
        # Strict: title must equal 'Weekly_Market_Analysis' or 'Weekly Market Analysis'
        if t in ("weekly_market_analysis", "weekly market analysis"):
            found_ss = ss_id
            found_title = title
            break

    if not found_ss:
        record("Google Sheet 'Weekly_Market_Analysis' exists", False,
               f"Found: {[t for _, t in spreadsheets]}")
        cur.close()
        conn.close()
        return

    record(f"Google Sheet exists (title: {found_title})", True)

    cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (found_ss,))
    sheets = cur.fetchall()
    sheet_titles = [(sid, t) for sid, t in sheets]

    sector_sheet_id = None
    stock_sheet_id = None
    for s_id, s_title in sheet_titles:
        t_norm = (s_title or "").strip().lower()
        if t_norm == "sector performance":
            sector_sheet_id = s_id
        if t_norm == "stock vs sector":
            stock_sheet_id = s_id

    # Allow loose match if exact not found
    if sector_sheet_id is None:
        for s_id, s_title in sheet_titles:
            if "sector" in (s_title or "").lower() and "perform" in (s_title or "").lower():
                sector_sheet_id = s_id
                break
    if stock_sheet_id is None:
        for s_id, s_title in sheet_titles:
            t = (s_title or "").lower()
            if "stock" in t and "sector" in t:
                stock_sheet_id = s_id
                break

    record("Has 'Sector Performance' sheet", sector_sheet_id is not None,
           f"Sheets: {[t for _, t in sheet_titles]}")
    record("Has 'Stock vs Sector' sheet", stock_sheet_id is not None,
           f"Sheets: {[t for _, t in sheet_titles]}")

    # === Sector sheet validation ===
    if sector_sheet_id:
        cur.execute("""
            SELECT row_index, col_index, COALESCE(formatted_value, value) FROM gsheet.cells
            WHERE spreadsheet_id = %s AND sheet_id = %s
            ORDER BY row_index, col_index
        """, (found_ss, sector_sheet_id))
        cells = cur.fetchall()
        grid = _build_grid(cells)

        # Need at least 5 data rows
        data_rows = sorted(r for r in grid if r > 0)
        record(f"Sector sheet has >= 5 data rows", len(data_rows) >= 5,
               f"Found {len(data_rows)}")

        # Header row at index 0
        header_row = grid.get(0, {})
        header_cols = [str(header_row.get(i, "") or "").strip().lower() for i in sorted(header_row.keys())]
        # Map column names to indices
        wanted = {"sector": None, "weekly_return": None, "sector_index": None,
                  "ytd_return": None, "volatility": None, "consensus": None, "target_upside": None}
        for i, h in enumerate(header_cols):
            hn = h.replace(" ", "_")
            for k in list(wanted.keys()):
                if k == hn or k.replace("_", "") == hn.replace("_", ""):
                    wanted[k] = i
                    break

        record("Sector sheet has 'Sector' header column", wanted["sector"] is not None,
               f"Headers: {header_cols}")
        for col_key in ["weekly_return", "sector_index", "ytd_return"]:
            record(f"Sector sheet has '{col_key}' column", wanted[col_key] is not None,
                   f"Headers: {header_cols}")

        # Validate per-sector numeric values
        for sector_lower, expected in SECTOR_DATA.items():
            row = None
            for r in data_rows:
                rd = grid[r]
                v = rd.get(wanted["sector"]) if wanted["sector"] is not None else rd.get(0)
                if v is not None and sector_lower in str(v).strip().lower():
                    row = rd
                    break
            if row is None:
                record(f"Sector '{sector_lower}' present in sheet", False)
                continue
            # weekly
            if wanted["weekly_return"] is not None:
                v = to_float(row.get(wanted["weekly_return"]))
                record(f"{sector_lower} weekly_return ~ {expected['weekly']}",
                       num_close(v, expected['weekly'], 0.2),
                       f"Got: {row.get(wanted['weekly_return'])}")
            if wanted["sector_index"] is not None:
                v = to_float(row.get(wanted["sector_index"]))
                record(f"{sector_lower} sector_index ~ {expected['index']}",
                       num_close(v, expected['index'], 0.5),
                       f"Got: {row.get(wanted['sector_index'])}")
            if wanted["ytd_return"] is not None:
                v = to_float(row.get(wanted["ytd_return"]))
                record(f"{sector_lower} ytd_return ~ {expected['ytd']}",
                       num_close(v, expected['ytd'], 0.5),
                       f"Got: {row.get(wanted['ytd_return'])}")

    # === Stock sheet validation ===
    if stock_sheet_id:
        cur.execute("""
            SELECT row_index, col_index, COALESCE(formatted_value, value) FROM gsheet.cells
            WHERE spreadsheet_id = %s AND sheet_id = %s
            ORDER BY row_index, col_index
        """, (found_ss, stock_sheet_id))
        cells = cur.fetchall()
        grid = _build_grid(cells)

        data_rows = sorted(r for r in grid if r > 0)
        record("Stock sheet has >= 5 data rows", len(data_rows) >= 5,
               f"Found {len(data_rows)}")

        header_row = grid.get(0, {})
        header_cols = [str(header_row.get(i, "") or "").strip().lower() for i in sorted(header_row.keys())]
        wanted = {"ticker": None, "stock_price": None, "sector": None,
                  "market_cap_b": None, "pe_ratio": None,
                  "dividend_yield_pct": None,
                  "sector_weekly_return": None, "sector_ytd_return": None}
        for i, h in enumerate(header_cols):
            hn = h.replace(" ", "_")
            for k in list(wanted.keys()):
                if k == hn or k.replace("_", "") == hn.replace("_", ""):
                    wanted[k] = i
                    break

        # Find each ticker row
        for ticker, expected_sector in STOCK_TO_SECTOR.items():
            row = None
            for r in data_rows:
                rd = grid[r]
                v = rd.get(wanted["ticker"]) if wanted["ticker"] is not None else rd.get(0)
                if v is not None and ticker.upper() == str(v).strip().upper():
                    row = rd
                    break
            if row is None:
                record(f"Ticker {ticker} present in stock sheet", False)
                continue
            record(f"Ticker {ticker} present in stock sheet", True)
            # Sector mapping
            if wanted["sector"] is not None:
                v = row.get(wanted["sector"])
                record(f"{ticker} sector = '{expected_sector}'",
                       v is not None and expected_sector.lower() in str(v).lower(),
                       f"Got: {v}")
            # Sector weekly return — should match expected_sector's weekly return
            if wanted["sector_weekly_return"] is not None:
                v = to_float(row.get(wanted["sector_weekly_return"]))
                expected_weekly = SECTOR_DATA[expected_sector.lower()]["weekly"]
                record(f"{ticker} sector_weekly_return ~ {expected_weekly}",
                       num_close(v, expected_weekly, 0.2),
                       f"Got: {row.get(wanted['sector_weekly_return'])}")

    cur.close()
    conn.close()


def check_word(agent_workspace):
    print("\n=== Checking Word Output ===")
    fpath = os.path.join(agent_workspace, "Weekly_Market_Report.docx")
    if not os.path.isfile(fpath):
        record("Word file exists", False, f"Not found: {fpath}")
        return
    record("Word file exists", True)

    try:
        from docx import Document
        doc = Document(fpath)
    except Exception as e:
        record("Word file readable", False, str(e))
        return

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    text_lower = full_text.lower()

    record("Word doc length >= 500 chars",
           len(full_text.strip()) >= 500,
           f"Length: {len(full_text)}")

    # All 5 sectors mentioned
    for sec in ["communication services", "consumer cyclical", "financial services", "healthcare", "energy"]:
        record(f"Word mentions '{sec}'",
               sec in text_lower,
               f"Not found")

    # All 5 stock tickers mentioned
    for ticker in ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]:
        record(f"Word mentions ticker {ticker}",
               ticker in full_text,
               f"Not found")

    # Best/worst conclusion
    record("Word identifies Energy as best/top sector",
           ("energy" in text_lower) and any(k in text_lower for k in ["best", "top", "highest", "leading"]),
           "Expected energy + best/top descriptor")
    record("Word identifies Consumer Cyclical as worst/bottom sector",
           ("consumer cyclical" in text_lower) and any(k in text_lower for k in ["worst", "bottom", "lowest", "underperform", "weakest"]),
           "Expected consumer cyclical + worst/bottom descriptor")


def check_email():
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("SELECT subject, from_addr, to_addr, COALESCE(body_text, body_html, '') FROM email.messages")
        emails = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Email DB accessible", False, str(e))
        return

    target_subj = "weekly market analysis report"
    target_from = "analyst@investmentfirm.com"
    target_to = "committee@investmentfirm.com"

    match = None
    for subject, from_addr, to_addr, body in emails:
        subj_lower = (subject or "").strip().lower()
        from_lower = (from_addr or "").strip().lower()
        to_list = []
        if isinstance(to_addr, list):
            to_list = [str(x).lower() for x in to_addr]
        elif isinstance(to_addr, str):
            try:
                p = json.loads(to_addr)
                to_list = [str(x).lower() for x in p] if isinstance(p, list) else [to_addr.lower()]
            except Exception:
                to_list = [to_addr.lower()]
        if (target_subj in subj_lower
                and target_from == from_lower
                and any(target_to in t for t in to_list)):
            match = (subject, from_addr, to_addr, body)
            break

    if not match:
        record(f"Email subject '{target_subj}' from {target_from} to {target_to}",
               False,
               f"Found {len(emails)} emails; none match all 3 criteria")
        return
    record(f"Email subject '{target_subj}' from {target_from} to {target_to}", True)

    body_lower = (match[3] or "").lower()
    record("Email body mentions Energy (top sector)",
           "energy" in body_lower,
           f"Body preview: {(match[3] or '')[:200]}")
    record("Email body mentions Consumer Cyclical (negative sector)",
           "consumer cyclical" in body_lower or "consumer" in body_lower,
           f"Body preview: {(match[3] or '')[:200]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_gsheet()
    check_word(args.agent_workspace)
    check_email()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
