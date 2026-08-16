"""
Evaluation for 12306-beijing-shanghai-trip-notion-gcal-word.
"""
import json
import os
import re
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def has_word_boundary(token, text):
    """Check if `token` appears as a word in `text` (case-insensitive).
    Word boundaries are non-alphanumeric chars or start/end.
    """
    pattern = r'(?<![A-Za-z0-9])' + re.escape(token) + r'(?![A-Za-z0-9])'
    return re.search(pattern, text, re.IGNORECASE) is not None


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Travel_Plan.docx ===")

    # Require exact filename
    doc_path = os.path.join(agent_workspace, "Travel_Plan.docx")
    if not os.path.exists(doc_path):
        record("Travel_Plan.docx exists", False,
               f"Not found at {doc_path}")
        for _ in range(10):
            record("Word sub-check (skipped)", False)
        return
    record("Travel_Plan.docx exists", True)

    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    record("Word doc has at least 3 headings", len(headings) >= 3,
           f"Found {len(headings)} headings")

    # Title check
    title_text = " ".join(p.text for p in doc.paragraphs
                          if p.style.name == "Title" or p.style.name.startswith("Heading"))
    record("Title 'Beijing-Shanghai Business Trip Travel Plan'",
           "beijing-shanghai" in title_text.lower() and
           "travel plan" in title_text.lower())

    # Section heading checks
    heading_text = " ".join(p.text.lower() for p in headings)
    record("Section: Outbound Journey", "outbound" in heading_text)
    record("Section: Return Journey", "return" in heading_text)
    record("Section: Booking Summary", "booking" in heading_text)

    # Combine paragraph + table text
    full_text = "\n".join(p.text for p in doc.paragraphs)
    table_text = ""
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                table_text += c.text + " "
    combined = full_text + "\n" + table_text

    # Word-boundary G11 and G2
    has_g11 = has_word_boundary("G11", combined)
    has_g2 = has_word_boundary("G2", combined)
    record("Contains train number G11 (word boundary)", has_g11)
    record("Contains train number G2 (word boundary)", has_g2)

    # Beijing and Shanghai
    record("Mentions Beijing", "beijing" in combined.lower())
    record("Mentions Shanghai", "shanghai" in combined.lower())

    # Total cost: 1106 (sum of two 553 tickets) - word boundary
    has_total = has_word_boundary("1106", combined)
    record("Total cost 1106 mentioned (word boundary)", has_total)
    # Per-ticket cost: 553 - word boundary
    has_553 = has_word_boundary("553", combined)
    record("Per-ticket cost 553 mentioned (word boundary)", has_553)

    # At least 2 tables (one per train)
    record("Has at least 2 tables (one per train)", len(doc.tables) >= 2,
           f"got {len(doc.tables)} tables")


def check_notion():
    print("\n=== Check 2: Notion page ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
    except Exception as e:
        record("Notion DB connection", False, str(e))
        return
    cur = conn.cursor()
    cur.execute("SELECT id, properties FROM notion.pages")
    rows = cur.fetchall()
    cur.close()
    conn.close()

    def extract_title(props):
        """Pull the actual title text from notion properties dict."""
        if not isinstance(props, dict):
            return ""
        # Properties typically have a 'title' field with rich-text array
        for key, val in props.items():
            if isinstance(val, dict) and val.get("type") == "title":
                arr = val.get("title", [])
                if isinstance(arr, list):
                    parts = []
                    for it in arr:
                        if isinstance(it, dict):
                            t = it.get("text") or {}
                            if isinstance(t, dict) and "content" in t:
                                parts.append(str(t.get("content", "")))
                            elif "plain_text" in it:
                                parts.append(str(it.get("plain_text", "")))
                    if parts:
                        return " ".join(parts)
        return ""

    target_title_keys = ["beijing-shanghai", "beijing - shanghai", "beijing shanghai"]
    matching = 0
    for pid, props in rows:
        title_text = extract_title(props).lower()
        # Title field must contain Beijing-Shanghai + Business Trip + March 10
        title_ok = any(k in title_text for k in target_title_keys)
        biz_ok = "business trip" in title_text
        march_ok = "march 10" in title_text or "march-10" in title_text
        if title_ok and biz_ok and march_ok:
            matching += 1

    record("Notion page 'Beijing-Shanghai Business Trip - March 10' exists (title field)",
           matching >= 1, f"Found {matching} matching pages out of {len(rows)} total")


def check_gcal():
    print("\n=== Check 3: Google Calendar events ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
    except Exception as e:
        record("GCal DB connection", False, str(e))
        return
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime, end_datetime FROM gcal.events
        WHERE start_datetime >= '2026-03-10'
          AND start_datetime < '2026-03-11'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Tighten: must have BOTH outbound and return events with specific summaries
    outbound_events = [e for e in events
                       if "business trip" in (e[0] or "").lower()
                       and ("beijing to shanghai" in (e[0] or "").lower()
                            or "to shanghai" in (e[0] or "").lower())]
    return_events = [e for e in events
                     if "business trip" in (e[0] or "").lower()
                     and ("return" in (e[0] or "").lower()
                          or "to beijing" in (e[0] or "").lower())]
    record("Outbound calendar event 'Business Trip - Beijing to Shanghai'",
           len(outbound_events) >= 1,
           f"got: {[e[0] for e in events]}")
    record("Return calendar event 'Business Trip - Return to Beijing'",
           len(return_events) >= 1,
           f"got: {[e[0] for e in events]}")

    # Ordering check: outbound start_datetime must be earlier than return start_datetime
    # (prevents agents that swap directions from passing on summary keywords alone)
    if outbound_events and return_events:
        try:
            ob_start = outbound_events[0][1]
            rt_start = return_events[0][1]
            record("Outbound starts before return",
                   ob_start < rt_start,
                   f"outbound={ob_start}, return={rt_start}")
        except Exception as e:
            record("Outbound starts before return", False, str(e))


def check_email():
    print("\n=== Check 4: Email to travel@consulting.com ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, to_addr, body_text FROM email.messages
            WHERE to_addr::text ILIKE '%travel@consulting.com%'
              AND from_addr NOT ILIKE '%travel@consulting.com%'
        """)
        emails = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Email DB query", False, str(e))
        return

    record("Email to travel@consulting.com sent",
           len(emails) >= 1, f"got {len(emails)} email(s)")
    if emails:
        subj, to_addr, body = emails[0]
        subj_lower = (subj or "").lower()
        # Subject must indicate business trip travel confirmation
        ok_subj = ("business trip" in subj_lower
                   or "travel confirmation" in subj_lower
                   or ("travel" in subj_lower and "confirm" in subj_lower)
                   or ("trip" in subj_lower and "confirm" in subj_lower))
        record("Subject indicates business trip travel confirmation",
               ok_subj, f"got: '{subj}'")
        # Body should mention G11 and G2
        body = body or ""
        ok_g11 = has_word_boundary("G11", body)
        ok_g2 = has_word_boundary("G2", body)
        record("Email body mentions G11", ok_g11, f"body sample: {body[:200]}")
        record("Email body mentions G2", ok_g2, f"body sample: {body[:200]}")
        # Body should mention 2026-03-10 or March 10
        ok_date = ("2026-03-10" in body or "march 10" in body.lower()
                   or "march-10" in body.lower())
        record("Email body mentions March 10 / 2026-03-10",
               ok_date, f"body sample: {body[:200]}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_notion()
    check_gcal()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed")

    result = {"total_passed": PASS_COUNT, "total_checks": total}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
