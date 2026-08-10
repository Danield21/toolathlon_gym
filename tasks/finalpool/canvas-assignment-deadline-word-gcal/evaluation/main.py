"""
Evaluation script for canvas-assignment-deadline-word-gcal task.

Checks:
1. Excel file Assignment_Deadlines_FFF2013J.xlsx - 2 sheets with correct data
2. Word document Assignment_Schedule_FFF2013J.docx exists with table
3. Google Calendar has assignment reminder events
4. Email sent to fff2013j.students@university.edu
"""

import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def _date_match(a, b):
    """Compare two date-ish values by their YYYY-MM-DD portion.

    A datetime cell ("2013-10-20 00:00:00") written by openpyxl must still match
    the groundtruth's plain "2013-10-20" string. Non-date strings fall back to a
    case-insensitive exact comparison.
    """
    def norm(v):
        if v is None:
            return None
        s = str(v).strip()
        if len(s) >= 10 and s[4] == "-" and s[7] == "-":
            return s[:10]
        return s.lower()
    return norm(a) == norm(b)


def load_sheet_by_name(wb, name):
    for sname in wb.sheetnames:
        if sname.strip().lower() == name.strip().lower():
            return [[cell.value for cell in row] for row in wb[sname].iter_rows()]
    return None


# ============================================================================
# Check 1: Excel file
# ============================================================================

def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Checking Assignment_Deadlines_FFF2013J.xlsx ===")

    try:
        import openpyxl
    except ImportError:
        record("openpyxl available", False, "pip install openpyxl")
        return False

    agent_file = os.path.join(agent_workspace, "Assignment_Deadlines_FFF2013J.xlsx")
    gt_file = os.path.join(groundtruth_workspace, "Assignment_Deadlines_FFF2013J.xlsx")

    if not os.path.isfile(agent_file):
        record("Excel file exists", False, f"Not found: {agent_file}")
        return False
    record("Excel file exists", True)

    if not os.path.isfile(gt_file):
        record("Groundtruth file exists", False, f"Not found: {gt_file}")
        return False

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_ok = True

    # Check All Assignments sheet
    a_all = load_sheet_by_name(agent_wb, "All Assignments")
    g_all = load_sheet_by_name(gt_wb, "All Assignments")
    record("Sheet 'All Assignments' exists", a_all is not None)

    if a_all is not None and g_all is not None:
        a_data = [r for r in a_all[1:] if any(v is not None for v in r)]
        g_data = [r for r in g_all[1:] if any(v is not None for v in r)]
        record("All Assignments row count matches",
               len(a_data) == len(g_data),
               f"Expected {len(g_data)}, got {len(a_data)}")

        # Build lookup by assignment name
        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row

        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            name = str(g_row[0]).strip()
            key = name.lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                record(f"Assignment row: {name}", False, "Not found")
                all_ok = False
                continue
            record(f"Assignment row: {name}", True)

            # Points_Possible (col 1)
            if len(g_row) > 1 and len(a_row) > 1:
                record(f"{name}: Points_Possible correct",
                       num_close(a_row[1], g_row[1], 0.01),
                       f"got {a_row[1]}, expected {g_row[1]}")

            # Due_Date (col 2), normalized so datetime cells still match.
            if len(g_row) > 2 and len(a_row) > 2:
                gt_due = g_row[2]
                if gt_due is not None and str(gt_due).strip() not in {"", "TBD"}:
                    record(f"{name}: Due_Date correct",
                           _date_match(a_row[2], gt_due),
                           f"got {a_row[2]}, expected {gt_due}")

    # Check Summary sheet
    a_summ = load_sheet_by_name(agent_wb, "Summary")
    g_summ = load_sheet_by_name(gt_wb, "Summary")
    record("Sheet 'Summary' exists", a_summ is not None)

    if a_summ is not None and g_summ is not None:
        a_data = [r for r in a_summ[1:] if any(v is not None for v in r)]
        g_data = [r for r in g_summ[1:] if any(v is not None for v in r)]

        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row

        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                record(f"Summary row: {g_row[0]}", False, "Not found")
                all_ok = False
                continue
            record(f"Summary row: {g_row[0]}", True)

            if key == "total_assignments":
                record("Total_Assignments = 13",
                       num_close(a_row[1], g_row[1], 0),
                       f"got {a_row[1]}, expected {g_row[1]}")
            elif key == "total_points_possible":
                record("Total_Points_Possible = 900",
                       num_close(a_row[1], g_row[1], 1.0),
                       f"got {a_row[1]}, expected {g_row[1]}")
            elif key == "avg_points_per_assignment":
                record("Avg_Points_Per_Assignment correct",
                       num_close(a_row[1], g_row[1], 1.0),
                       f"got {a_row[1]}, expected {g_row[1]}")

    return all_ok


# ============================================================================
# Check 2: Word document
# ============================================================================

def check_word(agent_workspace, groundtruth_workspace=None):
    print("\n=== Checking Assignment_Schedule_FFF2013J.docx ===")

    docx_path = os.path.join(agent_workspace, "Assignment_Schedule_FFF2013J.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, f"Not found: {docx_path}")
        return False
    record("Word file exists", True)

    # Load groundtruth assignments (name -> (points, due_date)) for the
    # per-row Word-table verification. Empty if GT unavailable.
    gt_assignments = {}
    if groundtruth_workspace:
        gt_path = os.path.join(groundtruth_workspace, "Assignment_Deadlines_FFF2013J.xlsx")
        if os.path.isfile(gt_path):
            try:
                import openpyxl
                gwb = openpyxl.load_workbook(gt_path, data_only=True)
                for sn in gwb.sheetnames:
                    if sn.strip().lower() == "all assignments":
                        ws = gwb[sn]
                        for r in ws.iter_rows(min_row=2, values_only=True):
                            if r and r[0] is not None:
                                due = str(r[2]).strip() if r[2] is not None else None
                                if due in {"", "TBD"}:
                                    due = None
                                gt_assignments[str(r[0]).strip()] = (r[1], due)
                        break
            except Exception:
                gt_assignments = {}
    gt_lookup = {k.lower(): v for k, v in gt_assignments.items()}

    try:
        from docx import Document
        doc = Document(docx_path)
        para_text = " ".join(p.text for p in doc.paragraphs)
        # Content includes table cell text too: the task's required output is
        # a heading plus an assignment table, so a correct document may have
        # almost no paragraph text beyond the heading.
        table_text = " ".join(
            cell.text
            for tbl in doc.tables
            for row in tbl.rows
            for cell in row.cells
        )
        all_text = (para_text + " " + table_text).lower()
        headings_text = " ".join(p.text for p in doc.paragraphs
                                 if p.style.name.startswith("Heading")).lower()

        text_len = len(all_text.strip())
        record("Word doc has content", text_len >= 100,
               f"Content length: {text_len}")
        # Diagnostics to locate why the document is (nearly) empty:
        # generation truncation vs. Word MCP write failure.
        if text_len < 100:
            file_size = os.path.getsize(docx_path)
            n_paras = len(doc.paragraphs)
            n_tables = len(doc.tables)
            print(f"  [DIAG] Word doc content below threshold: "
                  f"file_size={file_size} bytes, paragraphs={n_paras}, "
                  f"tables={n_tables}, paragraph_text_len={len(para_text.strip())}, "
                  f"table_text_len={len(table_text.strip())}")
            print(f"  [DIAG] First 200 chars of content: "
                  f"{all_text.strip()[:200]!r}")
            print("  [DIAG] Hint: empty content may be caused by generation "
                  "truncation or a Word API issue.")
        # Heading must reference 'Foundations of Finance' AND 'Fall 2013' to
        # avoid passing on a generic 'Assignment' heading.
        head_combined = headings_text + " " + all_text
        record("Word doc heading 'Assignment Schedule for Foundations of Finance Fall 2013'",
               (("foundations of finance" in head_combined or "fff" in head_combined)
                and ("fall 2013" in head_combined or "2013j" in head_combined)
                and "assignment schedule" in head_combined),
               f"Heading text: {headings_text[:200]}")

        tables = doc.tables
        record("Word doc has at least 1 table", len(tables) >= 1,
               f"Found {len(tables)} tables")

        table_rows_ok = True
        # Check table has assignment data and at least 13 rows + header.
        if tables:
            tbl = tables[0]
            n_rows = len(tbl.rows)
            n_cols = len(tbl.columns) if n_rows else 0
            table_text = " ".join(
                cell.text.lower()
                for row in tbl.rows
                for cell in row.cells
            )
            record("Table has TMA assignments",
                   "tma" in table_text,
                   f"Table text: {table_text[:200]}")
            # Task: each assignment listed (13 in GT) + header row -> >=14 rows.
            record("Table has >= 14 rows (header + 13 assignments)",
                   n_rows >= 14,
                   f"rows={n_rows}")
            # Task says columns: name, due date, points (3 cols, but allow >=3).
            record("Table has >= 3 columns",
                   n_cols >= 3,
                   f"cols={n_cols}")

            # Per-row verification: each GT assignment should appear in the
            # table with a matching due date (YYYY-MM-DD) and points value.
            # Tolerate at most one GT row missing/mismatched (13 -> >=12).
            if gt_lookup:
                import re
                matched_rows = 0
                for row in tbl.rows:
                    cell_texts = [c.text.strip() for c in row.cells]
                    name = next((t for t in cell_texts if t.lower() in gt_lookup), None)
                    if name is None:
                        continue
                    gt_pts, gt_due = gt_lookup[name.lower()]
                    date_cell = next((t for t in cell_texts
                                      if re.fullmatch(r"\d{4}-\d{2}-\d{2}", t)), None)
                    pts_cell = next((t for t in cell_texts
                                     if re.fullmatch(r"\d+(?:\.\d+)?", t)), None)
                    date_ok = gt_due is not None and date_cell is not None and _date_match(date_cell, gt_due)
                    pts_ok = pts_cell is not None and num_close(float(pts_cell), gt_pts, 0.01)
                    if date_ok and pts_ok:
                        matched_rows += 1
                expected_rows = len(gt_assignments)
                table_rows_ok = matched_rows >= expected_rows - 1
                record(f"Word table rows verified ({matched_rows}/{expected_rows} assignments with correct date + points)",
                       table_rows_ok,
                       "At most 1 GT assignment may be missing or mismatched")

        return True and table_rows_ok

    except ImportError:
        size = os.path.getsize(docx_path)
        record("Word file has content (>3KB)", size > 3000, f"Size: {size} bytes")
        return size > 3000
    except Exception as e:
        record("Word file readable", False, str(e))
        return False


# ============================================================================
# Check 3: Google Calendar
# ============================================================================

def check_gcal():
    print("\n=== Checking Google Calendar ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT summary, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_gcal] Found {len(events)} calendar events.")

    # Expect events with "Assignment Due:" in title.
    # Determine GT count from groundtruth excel: count rows with non-null due date.
    gt_due_count = 0
    try:
        import openpyxl
        gt_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "groundtruth_workspace", "Assignment_Deadlines_FFF2013J.xlsx",
        )
        if os.path.isfile(gt_path):
            gwb = openpyxl.load_workbook(gt_path, data_only=True)
            for sn in gwb.sheetnames:
                if sn.strip().lower() == "all assignments":
                    ws = gwb[sn]
                    rows = list(ws.iter_rows(min_row=2, values_only=True))
                    for r in rows:
                        # Due_Date is col index 2 (0-based: name, points, due_date).
                        if len(r) > 2 and r[2] is not None and str(r[2]).strip() not in {"", "TBD"}:
                            gt_due_count += 1
                    break
    except Exception:
        gt_due_count = 0
    if gt_due_count <= 0:
        # Fallback to the documented expectation.
        gt_due_count = 13

    assignment_events = [e for e in events
                         if e[0] and ("assignment due" in e[0].lower()
                                      or "tma" in e[0].lower()
                                      or "cma" in e[0].lower())]
    record(
        f"Assignment reminder events created (>= {gt_due_count})",
        len(assignment_events) >= gt_due_count,
        f"Found {len(assignment_events)} assignment events; expected >= {gt_due_count}",
    )

    # Validate at least one event respects the '7 days before due date at 8am ET'
    # constraint. 8am ET = 12:00 UTC (EST UTC-5: 13:00 / EDT UTC-4: 12:00).
    # Build set of expected reminder dates from GT due dates minus 7 days.
    from datetime import datetime, timedelta

    gt_due_dates = []
    try:
        import openpyxl
        gt_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "groundtruth_workspace", "Assignment_Deadlines_FFF2013J.xlsx",
        )
        if os.path.isfile(gt_path):
            gwb = openpyxl.load_workbook(gt_path, data_only=True)
            for sn in gwb.sheetnames:
                if sn.strip().lower() == "all assignments":
                    ws = gwb[sn]
                    for r in ws.iter_rows(min_row=2, values_only=True):
                        if len(r) > 2 and r[2]:
                            try:
                                if isinstance(r[2], str):
                                    d = datetime.strptime(r[2], "%Y-%m-%d").date()
                                else:
                                    d = r[2].date() if hasattr(r[2], "date") else r[2]
                                gt_due_dates.append(d)
                            except Exception:
                                pass
                    break
    except Exception:
        pass

    expected_reminder_dates = {d - timedelta(days=7) for d in gt_due_dates}
    valid_reminder_events = 0
    for summary, start_dt, end_dt in assignment_events:
        if start_dt is None:
            continue
        # start_dt is a datetime (or string)
        try:
            if isinstance(start_dt, str):
                sd = datetime.fromisoformat(start_dt.replace("Z", "+00:00"))
            else:
                sd = start_dt
            event_date = sd.date() if hasattr(sd, "date") else None
            event_hour = sd.hour if hasattr(sd, "hour") else None
        except Exception:
            continue
        if event_date in expected_reminder_dates and event_hour in {12, 13}:
            valid_reminder_events += 1
    record(
        "At least 1 reminder event scheduled 7 days before a due date at 8am ET (12:00/13:00 UTC)",
        valid_reminder_events >= 1,
        f"Found {valid_reminder_events} matching reminder events",
    )

    return len(assignment_events) >= gt_due_count and valid_reminder_events >= 1


# ============================================================================
# Check 4: Email
# ============================================================================

def check_emails():
    print("\n=== Checking Emails ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_emails] Found {len(all_emails)} total emails.")
    record("At least 1 email sent", len(all_emails) >= 1, f"Found {len(all_emails)}")

    found_email = False
    for subject, from_addr, to_addr, body_text in all_emails:
        to_str = str(to_addr or "").lower()
        subject_lower = (subject or "").lower()
        # Recipient must strictly match; both subject and recipient required (AND).
        if ("fff2013j.students@university.edu" in to_str
                and ("foundations of finance" in subject_lower
                     or ("assignment" in subject_lower and "deadline" in subject_lower)
                     or ("assignment" in subject_lower and "reminder" in subject_lower))):
            found_email = True
            record("Email to fff2013j.students@university.edu with correct subject", True)

            body_lower = (body_text or "").lower()
            record("Email body lists assignments",
                   any(term in body_lower for term in ["tma", "cma", "assignment", "due date"]),
                   "Body missing assignment list")
            break

    if not found_email:
        record("Assignment deadline email to fff2013j.students@university.edu", False,
               f"Emails: {[(e[0], str(e[2])[:60]) for e in all_emails[:3]]}")

    return found_email


# ============================================================================
# Main
# ============================================================================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    excel_ok = check_excel(args.agent_workspace, gt_dir)
    word_ok = check_word(args.agent_workspace, gt_dir)
    gcal_ok = check_gcal()
    email_ok = check_emails()

    # Aggregate gate must include FAIL_COUNT==0 so per-row record() failures
    # (e.g. wrong Points_Possible) propagate to overall PASS/FAIL.
    all_passed = (
        excel_ok and word_ok and gcal_ok and email_ok and FAIL_COUNT == 0
    )

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_passed else 'FAIL'}")

    if args.res_log_file:
        result = {
            "passed": PASS_COUNT,
            "failed": FAIL_COUNT,
            "success": all_passed,
        }
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
