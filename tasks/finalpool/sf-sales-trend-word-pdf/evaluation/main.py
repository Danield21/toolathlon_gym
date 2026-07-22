"""
Evaluation for sf-sales-trend-word-pdf task.

Checks:
1. Word document Sales_Trend_Analysis.docx exists with correct data
2. PDF file Sales_Trend_Analysis.pdf exists with correct data
"""

import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=500.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def get_expected_data():
    """Fetch expected monthly revenue data from the database."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT "MONTH_KEY", "TOTAL_REVENUE", "ORDER_COUNT", "UNIQUE_CUSTOMERS", "AVG_ORDER_VALUE"
        FROM sf_data."SALES_DW__ANALYTICS__MONTHLY_REVENUE"
        WHERE "MONTH_KEY" >= '2025-01-01' AND "MONTH_KEY" <= '2025-12-31'
        ORDER BY "MONTH_KEY"
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return rows


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Sales_Trend_Analysis.docx ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Sales_Trend_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Not found: {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)
    expected = get_expected_data()

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "sales trend" in p.text.lower() and "2025" in p.text:
            has_heading = True
            break
    check("Document has Sales Trend 2025 heading", has_heading)

    # Check table
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    data_rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(cells)

    check("Table has 12 month rows", len(data_rows) == 12, f"Got {len(data_rows)} rows")

    # Verify revenue and related metrics for each month
    total_revenue = sum(float(r[1]) for r in expected)
    matched_months = 0
    matched_orders = 0
    matched_customers = 0
    for exp_row in expected:
        month_str = exp_row[0].strftime("%Y-%m")
        exp_revenue = float(exp_row[1])
        exp_orders = int(exp_row[2]) if exp_row[2] is not None else None
        exp_customers = int(exp_row[3]) if exp_row[3] is not None else None
        found = False
        orders_ok = False
        customers_ok = False
        for dr in data_rows:
            if month_str in dr[0]:
                def parse_num(txt):
                    try:
                        return float(str(txt).replace("$","").replace(",","").strip())
                    except Exception:
                        return None
                rev_val = parse_num(dr[1]) if len(dr) > 1 else None
                if rev_val is not None and num_close(rev_val, exp_revenue, max(100.0, exp_revenue*0.01)):
                    found = True
                # Check order_count (col 2) & unique_customers (col 3)
                if len(dr) > 2 and exp_orders is not None:
                    v = parse_num(dr[2])
                    if v is not None and abs(v - exp_orders) <= max(1, exp_orders*0.02):
                        orders_ok = True
                if len(dr) > 3 and exp_customers is not None:
                    v = parse_num(dr[3])
                    if v is not None and abs(v - exp_customers) <= max(1, exp_customers*0.02):
                        customers_ok = True
                break
        if found:
            matched_months += 1
        if orders_ok:
            matched_orders += 1
        if customers_ok:
            matched_customers += 1

    # Stricter: all 12 months must match
    check("All 12 months have correct revenue", matched_months == 12,
          f"Matched {matched_months}/12")
    check("All 12 months have correct order_count", matched_orders == 12,
          f"Matched {matched_orders}/12")
    check("All 12 months have correct unique_customers", matched_customers == 12,
          f"Matched {matched_customers}/12")

    # Avg_Order_Value per month (column 4 in table)
    matched_aov = 0
    for exp_row in expected:
        month_str = exp_row[0].strftime("%Y-%m")
        exp_aov = float(exp_row[4]) if exp_row[4] is not None else None
        for dr in data_rows:
            if month_str in dr[0]:
                def parse_num(txt):
                    try:
                        return float(str(txt).replace("$","").replace(",","").strip())
                    except Exception:
                        return None
                if len(dr) > 4 and exp_aov is not None:
                    v = parse_num(dr[4])
                    if v is not None and abs(v - exp_aov) <= 1.0:
                        matched_aov += 1
                break
    check("All 12 months have correct avg_order_value", matched_aov == 12,
          f"Matched {matched_aov}/12")

    # Check summary section
    full_text = " ".join(p.text for p in doc.paragraphs)
    check("Document mentions total revenue",
          "total" in full_text.lower() and ("revenue" in full_text.lower() or "annual" in full_text.lower()),
          "Expected total revenue in summary")

    # Check H1/H2 analysis
    full_lower = full_text.lower()
    has_h1h2 = "h1" in full_lower or "first half" in full_lower or "second half" in full_lower
    check("Document has H1/H2 trend analysis", has_h1h2)

    # Check Best/Worst month (strict: must mention YYYY-MM or full month name)
    if expected:
        best = max(expected, key=lambda r: float(r[1]))
        worst = min(expected, key=lambda r: float(r[1]))
        best_month = best[0].strftime("%Y-%m")
        worst_month = worst[0].strftime("%Y-%m")
        best_name = best[0].strftime("%B").lower()
        worst_name = worst[0].strftime("%B").lower()
        best_hit = (best_month in full_text) or (best_name in full_lower)
        worst_hit = (worst_month in full_text) or (worst_name in full_lower)
        check(f"Document mentions best month ({best_month}/{best_name})", best_hit,
              f"Expected {best_month} or {best_name}")
        check(f"Document mentions worst month ({worst_month}/{worst_name})", worst_hit,
              f"Expected {worst_month} or {worst_name}")

        # Check Best/Worst revenue values present
        best_rev = float(best[1])
        worst_rev = float(worst[1])
        import re
        nums_in_text = []
        for n in re.findall(r"\$?([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)", full_text):
            try:
                nums_in_text.append(float(n.replace(",", "")))
            except Exception:
                pass
        check(f"Document mentions best revenue (~${best_rev:,.2f})",
              any(abs(v - best_rev) < max(10.0, best_rev * 0.005) for v in nums_in_text),
              f"Expected best revenue {best_rev}")
        check(f"Document mentions worst revenue (~${worst_rev:,.2f})",
              any(abs(v - worst_rev) < max(10.0, worst_rev * 0.005) for v in nums_in_text),
              f"Expected worst revenue {worst_rev}")

        # Check H1/H2 revenue totals + difference + percentage
        h1_revenue = sum(float(r[1]) for r in expected[:6])
        h2_revenue = sum(float(r[1]) for r in expected[6:])
        diff = h2_revenue - h1_revenue
        pct = round(abs(diff) / h1_revenue * 100, 1) if h1_revenue else 0.0
        check(f"Document mentions H1 revenue (~${h1_revenue:,.2f})",
              any(abs(v - h1_revenue) < max(50.0, h1_revenue * 0.01) for v in nums_in_text),
              f"Expected H1 revenue {h1_revenue}")
        check(f"Document mentions H2 revenue (~${h2_revenue:,.2f})",
              any(abs(v - h2_revenue) < max(50.0, h2_revenue * 0.01) for v in nums_in_text),
              f"Expected H2 revenue {h2_revenue}")
        # diff
        check(f"Document mentions H1/H2 revenue difference (~${abs(diff):,.2f})",
              any(abs(v - abs(diff)) < max(50.0, abs(diff) * 0.02) for v in nums_in_text),
              f"Expected diff ~{abs(diff)}")
        # pct
        check(f"Document mentions H1/H2 pct change (~{pct}%)",
              any(abs(v - pct) < 0.3 for v in nums_in_text) or f"{pct}" in full_text,
              f"Expected pct {pct}")

    return True


def check_pdf(agent_workspace):
    """Check the PDF file exists and has content."""
    print("\n=== Checking Sales_Trend_Analysis.pdf ===")

    pdf_path = os.path.join(agent_workspace, "Sales_Trend_Analysis.pdf")
    check("PDF file exists", os.path.isfile(pdf_path), f"Not found: {pdf_path}")
    if not os.path.isfile(pdf_path):
        return False

    size = os.path.getsize(pdf_path)
    check("PDF file size reasonable", size > 1024, f"Size: {size} bytes")

    # Try to read PDF text
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as p:
                text = ""
                for page in p.pages:
                    text += page.extract_text() or ""
        except ImportError:
            print("  [WARN] No PDF reader available. Checking file existence only.")
            return True

    text_lower = text.lower()
    check("PDF contains Sales Trend title", "sales trend" in text_lower or "sales" in text_lower)

    # Check all months appear (YYYY-MM, allow YYYY-MM or full name)
    expected = get_expected_data()
    months_expected = [r[0].strftime("%Y-%m") for r in expected]
    month_names = [r[0].strftime("%B").lower() for r in expected]
    months_found = sum(1 for m, n in zip(months_expected, month_names)
                       if m in text or n in text_lower)
    check("PDF contains at least 10/12 months", months_found >= 10,
          f"Found {months_found}/12 months")

    # Check revenue numbers for a sample (best and worst)
    best = max(expected, key=lambda r: float(r[1]))
    worst = min(expected, key=lambda r: float(r[1]))
    import re
    nums_in_text = []
    for n in re.findall(r"\$?([0-9]{1,3}(?:,[0-9]{3})*(?:\.[0-9]+)?)", text):
        try:
            nums_in_text.append(float(n.replace(",", "")))
        except Exception:
            pass
    best_rev = float(best[1])
    worst_rev = float(worst[1])
    check(f"PDF mentions best revenue (~${best_rev:,.2f})",
          any(abs(v - best_rev) < max(10.0, best_rev * 0.01) for v in nums_in_text),
          f"Expected best revenue {best_rev}")
    check(f"PDF mentions worst revenue (~${worst_rev:,.2f})",
          any(abs(v - worst_rev) < max(10.0, worst_rev * 0.01) for v in nums_in_text),
          f"Expected worst revenue {worst_rev}")

    # H1/H2 totals
    h1_revenue = sum(float(r[1]) for r in expected[:6])
    h2_revenue = sum(float(r[1]) for r in expected[6:])
    check(f"PDF mentions H1 revenue (~${h1_revenue:,.2f})",
          any(abs(v - h1_revenue) < max(50.0, h1_revenue * 0.01) for v in nums_in_text),
          f"Expected H1 revenue {h1_revenue}")
    check(f"PDF mentions H2 revenue (~${h2_revenue:,.2f})",
          any(abs(v - h2_revenue) < max(50.0, h2_revenue * 0.01) for v in nums_in_text),
          f"Expected H2 revenue {h2_revenue}")

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SALES TREND WORD PDF - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_pdf(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": overall}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
