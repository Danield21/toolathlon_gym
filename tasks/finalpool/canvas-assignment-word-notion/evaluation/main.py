"""
Evaluation for canvas-assignment-word-notion task.

Checks:
1. Word document Assignment_Guide.docx with correct assignment data
2. Notion page created (non-blocking)
"""

import argparse
import json
import os
import re
import sys

import psycopg2

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


def _levenshtein(a, b):
    """Pure-Python Levenshtein distance for short strings."""
    if a == b:
        return 0
    if not a:
        return len(b)
    if not b:
        return len(a)
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a, 1):
        cur = [i]
        for j, cb in enumerate(b, 1):
            cur.append(min(prev[j] + 1, cur[j - 1] + 1, prev[j - 1] + (ca != cb)))
        prev = cur
    return prev[-1]


def extract_page_title(props):
    """Extract a Notion page title from its properties JSON blob."""
    if isinstance(props, str):
        try:
            props = json.loads(props)
        except Exception:
            return ""
    if not isinstance(props, dict):
        return ""
    for prop in props.values():
        if isinstance(prop, dict) and isinstance(prop.get("title"), list) and prop["title"]:
            return "".join(
                (p.get("plain_text", "") or (p.get("text") or {}).get("content", ""))
                for p in prop["title"] if isinstance(p, dict)
            ).strip()
    return ""


def loose_title_match(actual, expected, keywords=None):
    """Lenient title match: after normalization, pass if either string contains
    the other, Levenshtein distance <= 3, or all keywords appear."""
    a = " ".join(str(actual or "").lower().split())
    e = " ".join(str(expected or "").lower().split())
    if not a or not e:
        return False
    if e in a or a in e:
        return True
    if _levenshtein(a, e) <= 3:
        return True
    if keywords and all(k.lower() in a for k in keywords):
        return True
    return False


def find_labeled_number(paragraphs, label):
    """Find a paragraph containing `label` (case-insensitive) and extract its
    numeric value. Prefers the first number that appears at or after the label,
    so a combined line like "Total assignments: 10, Total points: 300" yields 10
    for the assignments label and 300 for the points label. Returns None if no
    such value exists."""
    label = label.lower()
    for p in paragraphs:
        text = p.text or ""
        low = text.lower()
        if label in low:
            m = re.search(r"-?\d+(?:\.\d+)?", text[low.find(label):])
            if not m:
                m = re.search(r"-?\d+(?:\.\d+)?", text)
            if m:
                return float(m.group(0))
    return None


def find_number_loose(paragraphs, nouns, qualifiers):
    """Fallback for the total-assignments / total-points lines.

    A correct agent is instructed to use the exact phrasing "Total Assignments:
    N" / "Total Points: X"; the exact-string finder above handles that. This
    helper catches reasonable paraphrases — any paragraph that mentions one of
    `nouns` (e.g. "assignment"/"point") together with one of `qualifiers`
    (e.g. "total"/"sum"/"count") — and returns the number appearing at or after
    the noun. Several paragraphs may match (a per-group breakdown line like
    "Assignment count by group: 4" qualifies too), so among all matches we take
    the LARGEST value: a total is, by definition, at least as large as any
    constituent count, and a point total (300) exceeds any single assignment's
    points (<=100). A paragraph like "Assignments without due dates: 2" is
    excluded because it lacks a qualifier, so it never competes.
    """
    nouns = [n.lower() for n in nouns]
    qualifiers = [q.lower() for q in qualifiers]
    best = None
    for p in paragraphs:
        text = p.text or ""
        low = text.lower()
        if any(n in low for n in nouns) and any(q in low for q in qualifiers):
            # Prefer a number that appears right after the matched noun.
            best_pos = len(text)
            for n in nouns:
                pos = low.find(n)
                if pos != -1 and pos < best_pos:
                    best_pos = pos
            m = re.search(r"-?\d+(?:\.\d+)?", text[best_pos:])
            if not m:
                m = re.search(r"-?\d+(?:\.\d+)?", text)
            if m:
                val = float(m.group(0))
                if best is None or val > best:
                    best = val
    return best


DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

# Expected assignments for course_id=7 (Creative Computing & Culture Fall 2014)
EXPECTED_ASSIGNMENT_COUNT = 10
EXPECTED_TOTAL_POINTS = 300.0
EXPECTED_ASSIGNMENT_NAMES = [
    "CMA 24295", "CMA 24296", "CMA 24297", "CMA 24298",
    "TMA 24291", "TMA 24292", "TMA 24293", "TMA 24294",
    "Final Exam 24299", "Final Exam 40088"
]


def check(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Assignment_Guide.docx ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Assignment_Guide.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Not found: {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Aggregate paragraph text once; heading tokens may legitimately be split
    # across paragraphs, so the heading check uses the whole document text.
    full_text = " ".join(p.text for p in doc.paragraphs)
    check("Document has correct heading",
          "creative computing" in full_text.lower() and "assignment guide" in full_text.lower())

    # Check course code
    check("Document mentions CCC-2014J", "CCC-2014J" in full_text)

    # Check table
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    data_rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(cells)

    check("Table has 10 assignment rows", len(data_rows) == EXPECTED_ASSIGNMENT_COUNT,
          f"Got {len(data_rows)} rows")

    # Validate 4-column structure (Assignment Name, Due Date, Points, Status)
    if len(table.rows) > 0:
        ncols = len(table.rows[0].cells)
        check("Table has 4 columns", ncols == 4, f"Got {ncols} columns")

    # Validate sort order: due_date ascending, no-due-date entries last,
    # ties broken by assignment id ascending.
    def _sort_key(r):
        # r is list of cell strings; due-date is index 1
        date_str = r[1] if len(r) > 1 else ""
        no_date = "no due date" in date_str.lower()
        # extract id: assignment names are "TYPE NNNNN" - use trailing number
        name = r[0] if r else ""
        try:
            asg_id = int("".join(c for c in name.split()[-1] if c.isdigit()) or "0")
        except Exception:
            asg_id = 0
        return (1 if no_date else 0, date_str if not no_date else "", asg_id)

    expected_order = sorted(data_rows, key=_sort_key)
    sort_ok = data_rows == expected_order
    check(
        "Table rows sorted by due_date asc (no-due-date last, ties by id asc)",
        sort_ok,
        f"Actual: {[r[0] for r in data_rows]} | Expected: {[r[0] for r in expected_order]}",
    )

    # Validate Status column contains expected values
    if len(table.rows) > 0 and len(table.rows[0].cells) >= 4:
        status_vals = [r[3].strip().lower() for r in data_rows if len(r) > 3]
        valid_statuses = {"past due", "no due date"}
        all_valid = all(s in valid_statuses for s in status_vals)
        check(
            "Status column contains only 'Past Due' / 'No due date'",
            all_valid,
            f"Statuses: {status_vals}",
        )

    # Check assignment names appear
    found_names = 0
    row_texts = " ".join(str(cell) for row in data_rows for cell in row)
    for name in EXPECTED_ASSIGNMENT_NAMES:
        if name in row_texts:
            found_names += 1
    check("At least 8/10 assignment names found", found_names >= 8,
          f"Found {found_names}/10")

    # Check some due dates
    has_dates = sum(1 for r in data_rows if "2014" in r[1] or "2015" in r[1])
    has_no_date = sum(1 for r in data_rows if "no due date" in r[1].lower())
    check("Table has assignments with dates", has_dates >= 8, f"Found {has_dates}")
    check("Table has assignments without dates", has_no_date >= 1, f"Found {has_no_date}")

    # Check total assignments and points: find the labeled line, extract its
    # numeric value, compare numerically (formatting variants like "300" vs
    # "300.0" are accepted; value must still match). Fall back to a looser scan
    # for paraphrased labels (e.g. "Total number of assignments").
    total_assignments = find_labeled_number(doc.paragraphs, "total assignments")
    if total_assignments is None:
        total_assignments = find_number_loose(
            doc.paragraphs, ["assignment"], ["total", "sum", "count", "overall", "number of"])
    check("Document mentions total assignments count",
          total_assignments is not None
          and num_close(total_assignments, EXPECTED_ASSIGNMENT_COUNT, rel_tol=0.0, abs_tol=0.5),
          f"Expected a line like 'Total Assignments: {EXPECTED_ASSIGNMENT_COUNT}', got {total_assignments}")

    total_points = find_labeled_number(doc.paragraphs, "total points")
    if total_points is None:
        total_points = find_number_loose(
            doc.paragraphs, ["point"], ["total", "sum", "overall", "possible", "aggregate"])
    check("Document mentions total points (300.0)",
          total_points is not None
          and num_close(total_points, EXPECTED_TOTAL_POINTS, rel_tol=0.0, abs_tol=0.5),
          f"Expected a line like 'Total Points: {EXPECTED_TOTAL_POINTS}', got {total_points}")

    return True


def check_notion():
    """Check Notion page - BLOCKING. Page titled 'CCC-2014J Assignment Overview' required."""
    print("\n=== Checking Notion ===")

    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()

        cur.execute(
            "SELECT id, properties FROM notion.pages "
            "WHERE (archived IS NULL OR archived = false) "
            "  AND (in_trash IS NULL OR in_trash = false)"
        )
        pages = cur.fetchall()

        # Find page whose title loosely matches 'CCC-2014J Assignment Overview':
        # normalized containment, Levenshtein <= 3, or keyword coverage.
        target = None
        titles_seen = []
        for page_id, props in pages:
            props_str = json.dumps(props) if isinstance(props, dict) else str(props)
            title = extract_page_title(props)
            if title:
                titles_seen.append(title)
            haystack = title if title else props_str
            if loose_title_match(haystack, "CCC-2014J Assignment Overview",
                                 keywords=["ccc-2014j", "assignment", "overview"]):
                target = (page_id, props_str)
                break
        check(
            "Notion page 'CCC-2014J Assignment Overview' exists (fuzzy match)",
            target is not None,
            f"{len(pages)} pages scanned, titles: {titles_seen[:10]}",
        )

        if target is None:
            cur.close()
            conn.close()
            return

        page_id, _props_str = target
        # Verify page body text mentions course name, assignment count, total points.
        # Notion blocks store text inside JSONB block_data, addressable via parent_id.
        cur.execute(
            "SELECT block_data FROM notion.blocks "
            "WHERE parent_id = %s",
            (page_id,),
        )
        rows = cur.fetchall()
        body_parts = []
        for (bd,) in rows:
            if bd is None:
                continue
            body_parts.append(json.dumps(bd) if isinstance(bd, dict) else str(bd))
        body = (" ".join(body_parts) + " " + _props_str).lower()

        check(
            "Notion page mentions course name 'Creative Computing'",
            "creative computing" in body,
            f"Body excerpt: {body[:200]}",
        )
        check(
            "Notion page mentions number of assignments (10)",
            "10" in body,
            f"Body excerpt: {body[:200]}",
        )
        check(
            "Notion page mentions total points (300)",
            "300" in body,
            f"Body excerpt: {body[:200]}",
        )

        cur.close()
        conn.close()

    except Exception as e:
        check("Notion check completed", False, f"Error: {e}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("CANVAS ASSIGNMENT WORD NOTION - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_notion()  # blocking — counts toward overall PASS/FAIL

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": overall}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
