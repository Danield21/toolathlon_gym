"""Evaluation for terminal-wc-yf-commodity-gsheet-word."""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432,
          dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
          user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

def get_expected_categories():
    """Return the 6 categories that task.md explicitly lists. The DB
    (wc.product_categories) holds 8 categories; the task asks for the
    Audio/Cameras/Electronics/Home Appliances/TV & Home Theater/Watches set.
    Returning those 6 keeps the substring-match check stable regardless of
    whether the agent reports 6 or 8 rows."""
    return ["Audio", "Cameras", "Electronics", "Home Appliances",
            "TV & Home Theater", "Watches"]


EXPECTED_CATEGORIES = get_expected_categories()
SENSITIVITY_MAP = {"Watches": 40, "Electronics": 15, "Audio": 10,
                   "Cameras": 10, "TV & Home Theater": 5, "Home Appliances": 8}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{d}")


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(",", "").replace("$", "").replace("%", "").strip())
    except (ValueError, TypeError):
        return default


def check_gsheet():
    """Check Google Sheets spreadsheet."""
    print("\n=== Checking Google Sheets ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()

    target_ss = None
    for ss_id, title in spreadsheets:
        tl = (title or "").lower()
        # Require all three keywords to be present to avoid false matches
        if "commodity" in tl and "pricing" in tl and "dashboard" in tl:
            target_ss = ss_id
            break

    check("Commodity Pricing Dashboard spreadsheet exists", target_ss is not None,
          f"Spreadsheets: {[s[1] for s in spreadsheets]}")

    if target_ss:
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (target_ss,))
        sheets = cur.fetchall()
        sheet_names = {s[1].lower().replace(" ", "_"): s for s in sheets}

        # Check Product_Categories sheet (must contain "product" AND "categor")
        pc_sheet = None
        for name, (sid, stitle) in sheet_names.items():
            if "product" in name and "categor" in name:
                pc_sheet = sid
                break
        check("Product_Categories sheet exists", pc_sheet is not None,
              f"Sheets: {[s[1] for s in sheets]}")

        if pc_sheet:
            cur.execute("""
                SELECT value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2
                AND col_index = 1
            """, (target_ss, pc_sheet))
            cat_values = [r[0] for r in cur.fetchall() if r[0]]
            cats_found = sum(1 for c in cat_values
                           if any(ec.lower() in c.lower() for ec in EXPECTED_CATEGORIES))
            # >=6: accept agents who include the 6 task-listed categories
            # (and optionally also Headphones/Speakers, which don't substring-match).
            check("Product_Categories has all 6 categories", cats_found >= 6,
                  f"Found {cats_found} matching categories from {cat_values}")

        # Check Gold_Price_Trend sheet (must contain "gold")
        gp_sheet = None
        for name, (sid, stitle) in sheet_names.items():
            if "gold" in name:
                gp_sheet = sid
                break
        check("Gold_Price_Trend sheet exists", gp_sheet is not None,
              f"Sheets: {[s[1] for s in sheets]}")

        if gp_sheet:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2
            """, (target_ss, gp_sheet))
            row_count = cur.fetchone()[0]
            check("Gold_Price_Trend has exactly 20 data rows", row_count == 20,
                  f"Found {row_count} rows")

            # Validate dates are sorted (monotonic ascending)
            cur.execute("""
                SELECT row_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2 AND col_index = 1
                ORDER BY row_index
            """, (target_ss, gp_sheet))
            dates = [r[1] for r in cur.fetchall() if r[1]]
            sorted_ok = dates == sorted(dates)
            check("Gold_Price_Trend dates ascending order", sorted_ok,
                  f"Dates: {dates[:5]}...")

            # Validate Trend column = 'Up' if close > MA else 'Down' for each
            # row. Locate Gold_Close (col 2), Moving_Avg_20d (col 3), Trend (col 4).
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2
            """, (target_ss, gp_sheet))
            gp_rows = {}
            for ri, ci, v in cur.fetchall():
                gp_rows.setdefault(ri, {})[ci] = v
            trend_correct = 0
            trend_total = 0
            for ri, cells in gp_rows.items():
                close_v = safe_float(cells.get(2))
                ma_v = safe_float(cells.get(3))
                trend_str = (cells.get(4) or "").strip().lower()
                if close_v is None or ma_v is None or not trend_str:
                    continue
                trend_total += 1
                expected = "up" if close_v > ma_v else "down"
                if expected == trend_str:
                    trend_correct += 1
            check("Gold_Price_Trend Trend values match Close vs MA",
                  trend_total >= 15 and trend_correct == trend_total,
                  f"correct {trend_correct}/{trend_total}")

        # Check Pricing_Impact sheet (must contain "impact" or "pricing_impact")
        pi_sheet = None
        for name, (sid, stitle) in sheet_names.items():
            if "impact" in name:
                pi_sheet = sid
                break
        check("Pricing_Impact sheet exists", pi_sheet is not None,
              f"Sheets: {[s[1] for s in sheets]}")

        if pi_sheet:
            # Build category -> {action, avg, sens, est_cost} map from sheet.
            # Pricing_Impact columns: Category(1), Avg_Price(2), Sensitivity_Pct(3),
            # Estimated_Material_Cost(4), Current_Gold_Price(5), Recommended_Action(6)
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2
            """, (target_ss, pi_sheet))
            rows_data = {}
            for ri, ci, v in cur.fetchall():
                rows_data.setdefault(ri, {})[ci] = v
            cat_action = {}
            cat_data = {}
            for ri, cells in rows_data.items():
                cat = cells.get(1, "")
                action = cells.get(6, "")
                if cat:
                    cat_clean = cat.strip()
                    cat_action[cat_clean] = (action or "").strip().lower()
                    cat_data[cat_clean] = {
                        "avg": safe_float(cells.get(2)),
                        "sens": safe_float(cells.get(3)),
                        "est_cost": safe_float(cells.get(4)),
                    }
            # Validate Estimated_Material_Cost = Avg_Price * Sensitivity_Pct / 100
            # for each category. Tol: 5% of expected or 1.0 absolute.
            est_cost_correct = 0
            est_cost_total = 0
            for cat, d in cat_data.items():
                if d["avg"] is None or d["sens"] is None or d["est_cost"] is None:
                    continue
                expected = d["avg"] * d["sens"] / 100.0
                tol = max(1.0, abs(expected) * 0.05)
                est_cost_total += 1
                if abs(d["est_cost"] - expected) <= tol:
                    est_cost_correct += 1
            check("Pricing_Impact Estimated_Material_Cost = Avg_Price*Sens/100",
                  est_cost_total >= 5 and est_cost_correct == est_cost_total,
                  f"{est_cost_correct}/{est_cost_total} categories correct")

            # Watches must be "Review Pricing" (sensitivity 40 > 20)
            watches_action = next((v for k, v in cat_action.items()
                                   if "watch" in k.lower()), None)
            check("Watches recommended_action is 'Review Pricing'",
                  watches_action is not None and "review" in watches_action,
                  f"Got: {watches_action}")

            # All other categories should be "Monitor" (all <= 20% sensitivity)
            non_watches_actions = {k: v for k, v in cat_action.items()
                                   if "watch" not in k.lower()}
            non_watches_correct = all("monitor" in v for v in non_watches_actions.values()) and len(non_watches_actions) >= 5
            check("Non-Watches categories are 'Monitor'",
                  non_watches_correct,
                  f"Non-watches actions: {non_watches_actions}")

    conn.close()


def check_word(ws_path):
    """Check Pricing_Strategy_Report.docx."""
    print("\n=== Checking Word Document ===")
    path = os.path.join(ws_path, "Pricing_Strategy_Report.docx")
    if not os.path.isfile(path):
        check("Word document exists", False, f"Not found: {path}")
        return
    check("Word document exists", True)

    from docx import Document
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()

    check("Document mentions gold", "gold" in full_text)
    check("Document mentions watches", "watch" in full_text)
    check("Document mentions sensitivity or exposure",
          "sensitiv" in full_text or "exposure" in full_text)
    check("Document mentions trend or moving average",
          "trend" in full_text or "moving average" in full_text)
    check("Document mentions recommendations",
          "recommend" in full_text or "action" in full_text)
    check("Document length >= 800 chars", len(full_text) >= 800,
          f"Length: {len(full_text)}")
    # Per task.md: required sections include executive summary, category-level
    # exposure, gold price trend with 20-day moving average, recommendations.
    check("Document has executive summary section",
          "executive" in full_text and "summary" in full_text,
          "Missing executive summary section")
    check("Document mentions 20-day moving average",
          "20" in full_text and ("moving average" in full_text or "moving avg" in full_text),
          "Missing 20-day moving average reference")
    # Watches has 40% sensitivity and is highlighted as high-sensitivity in
    # task.md: the recommendations section must explicitly mention Watches
    # alongside review/pricing actions.
    has_watches_recommendation = (
        "watch" in full_text
        and ("review pricing" in full_text or "review price" in full_text
             or ("recommend" in full_text and "watch" in full_text))
    )
    check("Recommendations mention Watches review/pricing action",
          has_watches_recommendation,
          "Missing specific Watches recommendation")


def check_xlsx_content(workspace):
    """Check Commodity_Pricing_Dashboard.xlsx has valid content."""
    print("\n=== Checking XLSX Content ===")
    try:
        import openpyxl
    except ImportError:
        check("openpyxl available", False, "Cannot import openpyxl")
        return False

    xlsx_path = os.path.join(workspace, "Commodity_Pricing_Dashboard.xlsx")
    if not os.path.isfile(xlsx_path):
        check("Commodity_Pricing_Dashboard.xlsx exists", False, f"Not found: {xlsx_path}")
        return False
    check("Commodity_Pricing_Dashboard.xlsx exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        check("XLSX has at least one sheet", len(wb.worksheets) >= 1,
              f"Found {len(wb.worksheets)} sheets")
        all_ok = True
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            has_data = len(rows) >= 2
            check(f"XLSX sheet '{ws.title}' has data rows", has_data,
                  f"Only {len(rows)} rows")
            if not has_data:
                all_ok = False
        wb.close()
        return all_ok
    except Exception as e:
        check("XLSX readable", False, str(e))
        return False


def check_reverse_validation(workspace):
    """Verify things that should NOT exist in the output."""
    print("\n=== Reverse Validation ===")
    # Google Sheet: no duplicate spreadsheets with same commodity/pricing name
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("""
            SELECT title, COUNT(*) FROM gsheet.spreadsheets
            WHERE lower(title) LIKE '%%commodity%%' OR lower(title) LIKE '%%pricing%%'
            GROUP BY title HAVING COUNT(*) > 1
        """)
        dupes = cur.fetchall()
        check("No duplicate pricing spreadsheets", len(dupes) == 0,
              f"Duplicates: {dupes}")
        cur.close()
        conn.close()
    except Exception:
        pass

    # Word: document should not contain placeholder text
    path = os.path.join(workspace, "Pricing_Strategy_Report.docx")
    if os.path.isfile(path):
        try:
            from docx import Document
            doc = Document(path)
            full_text = " ".join(p.text for p in doc.paragraphs).lower()
            check("No placeholder text in Word doc",
                  "[insert" not in full_text and "todo" not in full_text and "xxx" not in full_text,
                  "Found placeholder text")
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("TERMINAL-WC-YF-COMMODITY-GSHEET-WORD - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    # NOTE: Removed check_xlsx_content — task.md only requires a cloud spreadsheet,
    # not a local Commodity_Pricing_Dashboard.xlsx file.
    check_reverse_validation(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
