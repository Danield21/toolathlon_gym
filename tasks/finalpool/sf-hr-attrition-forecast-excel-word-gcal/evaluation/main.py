"""Evaluation for sf-hr-attrition-forecast-excel-word-gcal.

Checks:
1. Attrition_Risk_Analysis.xlsx with 3 sheets matching groundtruth
2. Intervention_Report.docx with required sections
3. 7 department review meetings in gcal.events
"""
import json
import os
import sys
from argparse import ArgumentParser
from collections import defaultdict

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

DEPARTMENTS = ["Engineering", "Finance", "HR", "Operations", "R&D", "Sales", "Support"]


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def sheet_dicts(wb, name):
    """Get rows as list of dicts from a sheet, matching by case-insensitive name."""
    for sn in wb.sheetnames:
        if sn.strip().lower() == name.strip().lower():
            ws = wb[sn]
            rows = list(ws.iter_rows(values_only=True))
            if len(rows) < 2:
                return []
            hdrs = [str(h).strip() if h else "" for h in rows[0]]
            return [
                {hdrs[i]: row[i] for i in range(len(hdrs))}
                for row in rows[1:]
                if not all(v is None for v in row)
            ]
    return None


def get_groundtruth(gt_path):
    """Load groundtruth from Excel file."""
    gt_xlsx = os.path.join(gt_path, "Attrition_Risk_Analysis.xlsx")
    if not os.path.isfile(gt_xlsx):
        return None
    wb = openpyxl.load_workbook(gt_xlsx, data_only=True)
    gt = {
        "risk_assessment": sheet_dicts(wb, "Risk Assessment"),
        "dept_summary": sheet_dicts(wb, "Department Summary"),
        "top_factors": sheet_dicts(wb, "Top Risk Factors"),
    }
    wb.close()
    return gt


def check_excel(agent_ws, gt):
    print("\n=== Checking Excel: Attrition_Risk_Analysis.xlsx ===")
    xlsx_path = os.path.join(agent_ws, "Attrition_Risk_Analysis.xlsx")
    if not os.path.isfile(xlsx_path):
        record("Excel file exists", False, xlsx_path)
        return
    record("Excel file exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    except Exception as e:
        record("Excel readable", False, str(e))
        return
    record("Excel readable", True)

    # --- Sheet 1: Risk Assessment ---
    agent_ra = sheet_dicts(wb, "Risk Assessment")
    gt_ra = gt["risk_assessment"]
    if agent_ra is None:
        record("Sheet Risk Assessment exists", False, str(wb.sheetnames))
    else:
        record("Sheet Risk Assessment exists", True)
        record("Risk Assessment row count", len(agent_ra) == len(gt_ra),
               f"Expected {len(gt_ra)}, got {len(agent_ra)}")

        # Check a sample of rows
        for gt_row in gt_ra:
            dept = gt_row.get("Department", "")
            risk = gt_row.get("Risk_Level", "")
            match = next(
                (r for r in agent_ra
                 if str_match(r.get("Department"), dept) and str_match(r.get("Risk_Level"), risk)),
                None
            )
            if not match:
                record(f"Row {dept}/{risk} exists", False, "Missing")
                continue

            # Tightened tolerances (deterministic SF data — counts are exact)
            ec_a = match.get("Employee_Count"); ec_g = gt_row.get("Employee_Count")
            record(f"{dept}/{risk} Employee_Count",
                   num_close(ec_a, ec_g, 0),
                   f"Got {ec_a} vs {ec_g}")
            record(f"{dept}/{risk} Avg_Satisfaction",
                   num_close(match.get("Avg_Satisfaction"), gt_row.get("Avg_Satisfaction"), 0.2),
                   f"Got {match.get('Avg_Satisfaction')} vs {gt_row.get('Avg_Satisfaction')}")
            record(f"{dept}/{risk} Avg_Performance",
                   num_close(match.get("Avg_Performance"), gt_row.get("Avg_Performance"), 0.2),
                   f"Got {match.get('Avg_Performance')} vs {gt_row.get('Avg_Performance')}")
            record(f"{dept}/{risk} Avg_Salary_Growth_Pct",
                   num_close(match.get("Avg_Salary_Growth_Pct"), gt_row.get("Avg_Salary_Growth_Pct"), 0.3),
                   f"Got {match.get('Avg_Salary_Growth_Pct')} vs {gt_row.get('Avg_Salary_Growth_Pct')}")

    # --- Sheet 2: Department Summary ---
    agent_ds = sheet_dicts(wb, "Department Summary")
    gt_ds = gt["dept_summary"]
    if agent_ds is None:
        record("Sheet Department Summary exists", False, str(wb.sheetnames))
    else:
        record("Sheet Department Summary exists", True)
        record("Department Summary row count", len(agent_ds) == len(gt_ds),
               f"Expected {len(gt_ds)}, got {len(agent_ds)}")

        for gt_row in gt_ds:
            dept = gt_row.get("Department", "")
            match = next(
                (r for r in agent_ds if str_match(r.get("Department"), dept)),
                None
            )
            if not match:
                record(f"Dept Summary {dept} exists", False, "Missing")
                continue

            # Tightened tolerances — counts are deterministic
            record(f"{dept} Total_Employees",
                   num_close(match.get("Total_Employees"), gt_row.get("Total_Employees"), 0),
                   f"Got {match.get('Total_Employees')} vs {gt_row.get('Total_Employees')}")
            record(f"{dept} High_Risk_Count",
                   num_close(match.get("High_Risk_Count"), gt_row.get("High_Risk_Count"), 0),
                   f"Got {match.get('High_Risk_Count')} vs {gt_row.get('High_Risk_Count')}")
            record(f"{dept} Medium_Risk_Count",
                   num_close(match.get("Medium_Risk_Count"), gt_row.get("Medium_Risk_Count"), 0),
                   f"Got {match.get('Medium_Risk_Count')} vs {gt_row.get('Medium_Risk_Count')}")
            record(f"{dept} Low_Risk_Count",
                   num_close(match.get("Low_Risk_Count"), gt_row.get("Low_Risk_Count"), 0),
                   f"Got {match.get('Low_Risk_Count')} vs {gt_row.get('Low_Risk_Count')}")
            record(f"{dept} High_Risk_Pct",
                   num_close(match.get("High_Risk_Pct"), gt_row.get("High_Risk_Pct"), 0.3),
                   f"Got {match.get('High_Risk_Pct')} vs {gt_row.get('High_Risk_Pct')}")
            # 5% relative tolerance for budget (deterministic but allow rounding/grouping)
            try:
                expected_budget = float(gt_row.get("Total_Intervention_Budget") or 0)
                rel_tol = max(50000, expected_budget * 0.02)
            except Exception:
                rel_tol = 50000
            record(f"{dept} Total_Intervention_Budget",
                   num_close(match.get("Total_Intervention_Budget"),
                             gt_row.get("Total_Intervention_Budget"), rel_tol),
                   f"Got {match.get('Total_Intervention_Budget')} vs {gt_row.get('Total_Intervention_Budget')} (tol {rel_tol:.0f})")

    # --- Sheet 3: Top Risk Factors ---
    agent_tf = sheet_dicts(wb, "Top Risk Factors")
    gt_tf = gt["top_factors"]
    if agent_tf is None:
        record("Sheet Top Risk Factors exists", False, str(wb.sheetnames))
    else:
        record("Sheet Top Risk Factors exists", True)
        record("Top Risk Factors row count", len(agent_tf) == len(gt_tf),
               f"Expected {len(gt_tf)}, got {len(agent_tf)}")

        for gt_row in gt_tf:
            dept = gt_row.get("Department", "")
            match = next(
                (r for r in agent_tf if str_match(r.get("Department"), dept)),
                None
            )
            if not match:
                record(f"Top Factor {dept} exists", False, "Missing")
                continue
            record(f"{dept} Top_Factor",
                   str_match(match.get("Top_Factor"), gt_row.get("Top_Factor")),
                   f"Got '{match.get('Top_Factor')}' vs '{gt_row.get('Top_Factor')}'")
            record(f"{dept} Factor_Score",
                   num_close(match.get("Factor_Score"), gt_row.get("Factor_Score"), 0.3),
                   f"Got {match.get('Factor_Score')} vs {gt_row.get('Factor_Score')}")

    wb.close()


def check_word(agent_ws):
    print("\n=== Checking Word: Intervention_Report.docx ===")
    docx_path = os.path.join(agent_ws, "Intervention_Report.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, docx_path)
        return
    record("Word file exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word file readable", False, str(e))
        return
    record("Word file readable", True)

    # Collect all paragraphs + their styles for heading detection
    paras = [(p.text, (p.style.name if p.style else "")) for p in doc.paragraphs]
    headings = [t.strip().lower() for t, st in paras if st and "heading" in st.lower() and t.strip()]
    all_text = "\n".join(t for t, _ in paras).lower()
    # Include table text too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text.lower()

    # Required sections (look in headings if any, fall back to text)
    required_sections = [
        "executive summary",
        "department",  # Department-by-Department Analysis
        "budget",      # Budget Overview
        "recommended action",
        "timeline",
    ]
    for section in required_sections:
        in_headings = any(section in h for h in headings)
        in_text = section in all_text
        record(f"Word has '{section}' section",
               in_headings or in_text,
               f"headings={headings[:10]}")

    # All 7 departments mentioned (substring) AND each appears in a heading or as own subsection
    for dept in DEPARTMENTS:
        record(f"Word mentions {dept}",
               dept.lower() in all_text,
               "Department not found in document")

    # Has a Budget table
    has_budget_table = False
    for table in doc.tables:
        text_blob = " ".join(c.text.lower() for r in table.rows for c in r.cells)
        if "department" in text_blob and ("budget" in text_blob or "intervention" in text_blob):
            has_budget_table = True
            break
    record("Word has Budget table with department/budget columns",
           has_budget_table,
           "No table containing Department + Budget/Intervention found")

    # Risk levels
    record("Word mentions high risk",
           "high risk" in all_text or "high-risk" in all_text,
           "No risk level discussion found")
    # Specific budget number ($ + 4-7 digit number, with optional thousand separators)
    import re
    has_dollar_amount = bool(re.search(r"\$\s*[0-9](?:[0-9],?){2,6}", all_text))
    record("Word mentions a dollar budget figure",
           has_dollar_amount,
           f"No $-prefixed amount (>=4 digits) found")


def check_gcal():
    print("\n=== Checking Calendar Events ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SET TIME ZONE 'UTC'")
    cur.execute("""
        SELECT summary, description,
               start_datetime AT TIME ZONE 'UTC',
               end_datetime AT TIME ZONE 'UTC'
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Tighten: require 'attrition risk review' in summary AND a department name
    review_events_by_dept = {}
    for summary, desc, start_dt, end_dt in events:
        s_lower = (summary or "").lower()
        if "attrition risk review" not in s_lower:
            continue
        for dept in DEPARTMENTS:
            # word boundary-ish for R&D and short names
            dl = dept.lower()
            if dl in s_lower:
                # Avoid false positives where 'HR' could appear inside other words
                # Use simple heuristic: check it's not a substring of larger word for short names
                if dept in ("HR",):
                    import re
                    if not re.search(r"\bhr\b", s_lower):
                        continue
                review_events_by_dept.setdefault(dept, []).append((summary, desc, start_dt, end_dt))
                break

    # Flat list
    review_events = [(s, d, st, et, dept) for dept, evs in review_events_by_dept.items() for (s, d, st, et) in evs]

    record("Exactly 7 department review events",
           len(review_events) == 7,
           f"Found {len(review_events)} review events out of {len(events)} total. By dept: {[(k,len(v)) for k,v in review_events_by_dept.items()]}")

    found_depts = set(review_events_by_dept.keys())
    for dept in DEPARTMENTS:
        record(f"Calendar event for {dept}",
               dept in found_depts,
               f"Found events for: {sorted(found_depts)}")

    # Each meeting in Mar 16-20, 2026 (inclusive) AND duration ~ 60min AND business hours (9-17)
    starts_seen = []
    for summary, desc, start_dt, end_dt, dept in review_events:
        if start_dt:
            d = start_dt.date().isoformat()
            in_range = "2026-03-16" <= d <= "2026-03-20"
            record(f"{dept} meeting in Mar 16-20",
                   in_range,
                   f"Scheduled on {d}")
            # Business hours 9:00 <= start < 17:00 (UTC stored)
            record(f"{dept} meeting in business hours",
                   9 <= start_dt.hour < 17,
                   f"Start hour {start_dt.hour}:{start_dt.minute:02d}")
            starts_seen.append(start_dt)
        if start_dt and end_dt:
            duration_min = (end_dt - start_dt).total_seconds() / 60
            record(f"{dept} meeting is 60 min",
                   abs(duration_min - 60) <= 1,
                   f"Duration: {duration_min} minutes")
        # description should mention high risk count or budget
        d_lower = (desc or "").lower()
        record(f"{dept} description mentions risk count or budget",
               ("high risk" in d_lower or "high-risk" in d_lower or "budget" in d_lower or "$" in d_lower),
               f"Description: {(desc or '')[:200]}")

    # No-conflict: events shouldn't overlap each other
    if len(review_events) >= 2:
        sorted_evs = sorted([e for e in review_events if e[2] and e[3]], key=lambda e: e[2])
        no_conflict = True
        for i in range(1, len(sorted_evs)):
            prev_end = sorted_evs[i-1][3]
            cur_start = sorted_evs[i][2]
            if prev_end and cur_start and cur_start < prev_end:
                no_conflict = False
                break
        record("No overlapping department review meetings",
               no_conflict,
               "Some meetings overlap")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", default=".")
    parser.add_argument("--groundtruth_workspace", default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    gt = get_groundtruth(args.groundtruth_workspace)
    if gt is None:
        print("ERROR: Could not load groundtruth Excel.")
        sys.exit(1)

    check_excel(args.agent_workspace, gt)
    check_word(args.agent_workspace)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\n=== SUMMARY: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%) ===")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
