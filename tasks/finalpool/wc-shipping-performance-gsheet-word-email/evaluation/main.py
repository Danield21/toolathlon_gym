"""Evaluation for wc-shipping-performance-gsheet-word-email.

Checks:
1. Google Sheet "Shipping Performance Dashboard" with zone data (at least 3 rows)
2. Word doc Shipping_Performance_Report.docx with heading, table, and paragraphs
3. Email to logistics@store.example.com from operations@store.example.com
   Subject: "Monthly Shipping Performance Report"
"""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=int(os.environ.get("PGPORT", "5432")), dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

# Expected shipping data (fallback if DB query fails)
FALLBACK_ZONES = [
    {"name": "standard shipping", "orders": 64, "revenue": 25369.91},
    {"name": "free shipping", "orders": 34, "revenue": 12734.76},
    {"name": "express shipping", "orders": 30, "revenue": 12673.21},
]


def get_expected_zones():
    """Dynamically query WC for shipping method aggregates."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        # Method-level aggregation from shipping_lines
        cur.execute("""
            SELECT LOWER(sl->>'method_title') AS method,
                   COUNT(*) AS orders,
                   SUM(o.total::float) AS revenue
            FROM wc.orders o, jsonb_array_elements(o.shipping_lines) AS sl
            WHERE o.status IN ('completed','processing','on-hold','pending')
            GROUP BY LOWER(sl->>'method_title')
        """)
        res = [{"name": r[0], "orders": int(r[1]), "revenue": round(float(r[2]), 2)} for r in cur.fetchall()]
        # Free shipping = orders with no shipping lines
        cur.execute("""
            SELECT COUNT(*), COALESCE(SUM(total::float), 0)
            FROM wc.orders
            WHERE status IN ('completed','processing','on-hold','pending')
              AND (jsonb_array_length(COALESCE(shipping_lines, '[]'::jsonb)) = 0
                   OR shipping_lines IS NULL)
        """)
        free_row = cur.fetchone()
        if free_row and free_row[0] > 0:
            res.append({"name": "free shipping", "orders": int(free_row[0]), "revenue": round(float(free_row[1]), 2)})
        cur.close()
        conn.close()
        if res:
            return res
    except Exception as e:
        print(f"  WARN: could not query dynamic zones, using fallback: {e}")
    return FALLBACK_ZONES

EXPECTED_ZONES = get_expected_zones()


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=5.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE LOWER(title) LIKE '%shipping%'
    """)
    sheets = cur.fetchall()
    check("Shipping Performance Dashboard spreadsheet exists", len(sheets) >= 1,
          f"Found {len(sheets)} matching spreadsheets")

    if not sheets:
        cur.close()
        conn.close()
        return

    ss_id = sheets[0][0]
    # Load cells as a 2D structured grid
    cur.execute("""
        SELECT sheet_id, row_index, col_index, value, formatted_value
        FROM gsheet.cells
        WHERE spreadsheet_id = %s
        ORDER BY sheet_id, row_index, col_index
    """, (ss_id,))
    raw_cells = cur.fetchall()
    cur.close()
    conn.close()

    # Group cells by sheet, find the one that has the Zone_Name header
    sheets_grid = {}
    for sid, ri, ci, val, fval in raw_cells:
        text = val if val is not None else fval
        sheets_grid.setdefault(sid, {})[(ri, ci)] = text

    target_grid = None
    for sid, grid in sheets_grid.items():
        if not grid:
            continue
        # Find the first non-empty row treated as the header
        max_row = max(r for r, _ in grid.keys())
        max_col = max(c for _, c in grid.keys())
        header_row_idx = None
        for r in range(max_row + 1):
            row_vals = [str(grid.get((r, c), "") or "").lower() for c in range(max_col + 1)]
            if any("zone" in v for v in row_vals):
                header_row_idx = r
                break
        if header_row_idx is not None:
            target_grid = (grid, header_row_idx, max_row, max_col)
            break

    if target_grid is None:
        check("GSheet has Zone_Name header", False, "No header row with 'zone' found")
        return

    grid, hr, max_row, max_col = target_grid
    header = {str(grid.get((hr, c), "") or "").strip().lower(): c
              for c in range(max_col + 1)}
    # Column indices
    zn_col = next((v for k, v in header.items() if "zone" in k or "name" in k), None)
    oc_col = next((v for k, v in header.items() if "order" in k and "count" in k), None)
    rev_col = next((v for k, v in header.items() if "revenue" in k), None)
    aov_col = next((v for k, v in header.items() if "avg" in k and ("value" in k or "order" in k)), None)
    check("GSheet header has Zone_Name column", zn_col is not None)
    check("GSheet header has Order_Count column", oc_col is not None)
    check("GSheet header has Total_Revenue column", rev_col is not None)
    check("GSheet header has Avg_Order_Value column", aov_col is not None)

    if zn_col is None or oc_col is None or rev_col is None:
        return

    # Collect data rows
    data_rows = []
    for r in range(hr + 1, max_row + 1):
        zn = str(grid.get((r, zn_col), "") or "").strip()
        if not zn:
            continue
        try:
            oc = float(str(grid.get((r, oc_col), "") or "").replace(',', ''))
        except (ValueError, TypeError):
            oc = None
        try:
            rev = float(str(grid.get((r, rev_col), "") or "").replace(',', '').replace('$', ''))
        except (ValueError, TypeError):
            rev = None
        aov = None
        if aov_col is not None:
            try:
                aov = float(str(grid.get((r, aov_col), "") or "").replace(',', '').replace('$', ''))
            except (ValueError, TypeError):
                aov = None
        data_rows.append({"name": zn, "orders": oc, "revenue": rev, "aov": aov})
    check("GSheet has at least 3 data rows", len(data_rows) >= 3,
          f"Got {len(data_rows)}")

    # Verify per-zone data matches EXPECTED_ZONES
    rev_prefixes = [str(int(z["revenue"])) for z in EXPECTED_ZONES]
    for exp in EXPECTED_ZONES:
        expected_name = exp["name"].lower()
        row = next((r for r in data_rows if expected_name in r["name"].lower()), None)
        check(f"GSheet row for '{exp['name']}' present", row is not None)
        if row is None:
            continue
        if row["orders"] is not None:
            check(f"GSheet {exp['name']} Order_Count={exp['orders']}",
                  abs(row["orders"] - exp["orders"]) <= 1,
                  f"Got {row['orders']}")
        if row["revenue"] is not None:
            check(f"GSheet {exp['name']} Total_Revenue~{exp['revenue']}",
                  abs(row["revenue"] - exp["revenue"]) <= 5.0,
                  f"Got {row['revenue']}")


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "Shipping_Performance_Report.docx")
    check("Shipping_Performance_Report.docx exists", os.path.isfile(doc_path),
          f"Expected at {doc_path}")
    if not os.path.isfile(doc_path):
        return

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        return

    full_text = " ".join(p.text.lower() for p in doc.paragraphs)
    headings = [p.text.lower() for p in doc.paragraphs
                if p.style.name.startswith("Heading")]

    check("Word doc has 'Shipping Performance Report' heading",
          any("shipping performance report" in h for h in headings) or
          "shipping performance report" in full_text,
          "Heading not found")

    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    check("Word doc has at least 3 paragraphs", para_count >= 3,
          f"Found {para_count} non-empty paragraphs")

    check("Word doc has a table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    if doc.tables:
        # Find the table that actually has Zone_Name / Order_Count headers
        target_tbl = None
        for tbl in doc.tables:
            if not tbl.rows:
                continue
            hdr_cells = [c.text.strip().lower() for c in tbl.rows[0].cells]
            if any("zone" in h for h in hdr_cells) and any("order" in h for h in hdr_cells):
                target_tbl = tbl
                break
        check("Shipping table with Zone/Order columns present", target_tbl is not None)
        if target_tbl is not None:
            hdr_cells = [c.text.strip().lower() for c in target_tbl.rows[0].cells]
            try:
                zn_idx = next(i for i, h in enumerate(hdr_cells) if "zone" in h or "name" in h)
                oc_idx = next(i for i, h in enumerate(hdr_cells) if "order" in h and "count" in h)
                rev_idx = next(i for i, h in enumerate(hdr_cells) if "revenue" in h)
            except StopIteration:
                zn_idx = oc_idx = rev_idx = None
            check("Word table has zone, order_count, revenue columns",
                  zn_idx is not None and oc_idx is not None and rev_idx is not None)
            if zn_idx is not None and oc_idx is not None and rev_idx is not None:
                # Extract data rows
                tbl_data = []
                for row in target_tbl.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    if not cells or len(cells) <= rev_idx:
                        continue
                    try:
                        oc = float(cells[oc_idx].replace(',', ''))
                    except ValueError:
                        oc = None
                    try:
                        rev = float(cells[rev_idx].replace(',', '').replace('$', ''))
                    except ValueError:
                        rev = None
                    tbl_data.append({"name": cells[zn_idx], "orders": oc, "revenue": rev})
                for exp in EXPECTED_ZONES:
                    match = next((r for r in tbl_data
                                  if exp["name"].lower() in r["name"].lower()), None)
                    check(f"Word table row for '{exp['name']}' present",
                          match is not None)
                    if match is None:
                        continue
                    if match["orders"] is not None:
                        check(f"Word table {exp['name']} Order_Count={exp['orders']}",
                              abs(match["orders"] - exp["orders"]) <= 1,
                              f"Got {match['orders']}")
                    if match["revenue"] is not None:
                        check(f"Word table {exp['name']} Total_Revenue~{exp['revenue']}",
                              abs(match["revenue"] - exp["revenue"]) <= 5.0,
                              f"Got {match['revenue']}")

    check("Word doc mentions Standard Shipping",
          "standard" in full_text, "Standard shipping not mentioned")
    check("Word doc mentions Express Shipping",
          "express" in full_text, "Express shipping not mentioned")


def check_emails():
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    conn.close()

    def parse_recipients(to_addr):
        if to_addr is None:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        to_str = str(to_addr).strip()
        try:
            parsed = json.loads(to_str)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [to_str.lower()]
        except (json.JSONDecodeError, TypeError):
            return [to_str.lower()]

    target = "logistics@store.example.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if target in recipients:
            found = (subj, from_addr, to_addr, body)
            break

    check("Email sent to logistics@store.example.com", found is not None)
    if found:
        subj, from_addr, to_addr, body = found
        check("Email from operations@store.example.com",
              "operations@store.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}")
        check("Subject contains 'Monthly Shipping Performance Report'",
              "shipping" in (subj or "").lower() and "performance" in (subj or "").lower(),
              f"Subject: {subj}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("WC SHIPPING PERFORMANCE GSHEET WORD EMAIL - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_emails()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
