"""
Evaluation for 12306-canvas-fieldtrip-gcal-word-email.

Checks:
1. Field_Trip_Notice.docx exists in agent_workspace
2. Word doc has at least 4 headings
3. Word doc contains G235 and G236 train numbers (or departure times 17:30 and 15:00)
4. Word doc contains Qufu and Beijing
5. Word doc contains 1106 (round trip cost per person)
6. GCal has at least 2 new field trip events
7. Canvas announcement created with field trip content
8. Email sent to students@university.edu
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
from zoneinfo import ZoneInfo

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# Source of truth: docs/task.md line 1: field trip by high-speed rail from
# "Beijing" to "Qufu" on March 12/15, 2026 (China rail, CN wall-clock). The
# Calendar MCP writes events as UTC instants into the DB (TIMESTAMPTZ), so all
# comparisons anchor to Asia/Shanghai. Never use bare `st.hour` / `st.minute`
# on rows read from `gcal.events`: psycopg2 returns them in the PG session
# TimeZone (case-study: compute node default was Asia/Shanghai, masking the CN
# wall-clock in other sessions). Use `utils.evaluation.gcal_helpers` instead
# (session-tz-independent, DST-aware).
# ──────────────────────────────────────────────────────────────────────────
EXPECTED_TIMEZONE = ZoneInfo("Asia/Shanghai")
from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Field_Trip_Notice.docx ===")

    # Require exact filename Field_Trip_Notice.docx (case-insensitive match
    # to allow minor casing variations) per task spec line 9.
    target_basename = "field_trip_notice.docx"
    doc_path = None
    try:
        for fname in os.listdir(agent_workspace):
            if fname.lower() == target_basename:
                doc_path = os.path.join(agent_workspace, fname)
                break
    except Exception:
        pass

    if not doc_path:
        record("Field_Trip_Notice.docx exists with exact filename", False,
               f"Required exact filename Field_Trip_Notice.docx not found in {agent_workspace}")
        return
    record("Field_Trip_Notice.docx exists with exact filename", True)

    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    heading_texts_lower = [p.text.strip().lower() for p in headings]
    required_headings = ["trip overview", "schedule", "student requirements", "cost information"]
    missing = [h for h in required_headings if not any(h in ht for ht in heading_texts_lower)]
    record(
        "Word doc has all 4 required headings (Trip Overview, Schedule, Student Requirements, Cost Information)",
        len(missing) == 0,
        f"Headings found: {heading_texts_lower}; missing: {missing}",
    )

    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    # Tighten train number check: require both G235 and G236 (the
    # outbound and return train numbers) - departure times alone are
    # too coincidental to qualify as evidence of train numbers.
    has_g235 = "g235" in full_text
    has_g236 = "g236" in full_text
    record("Contains both G235 and G236 train numbers",
           has_g235 and has_g236,
           f"G235:{has_g235}, G236:{has_g236}")

    has_qufu = "qufu" in full_text
    has_beijing = "beijing" in full_text
    record("Contains Qufu and Beijing", has_qufu and has_beijing,
           f"Qufu:{has_qufu}, Beijing:{has_beijing}")

    # Cost check: 1106 as round-trip total. Require it to appear in
    # vicinity of the words 'total', 'round', 'cost', or 'per student'
    # to avoid generic '1106' substring coincidences (e.g. dates, page nums).
    cost_ok = False
    for p in doc.paragraphs:
        t = p.text.lower()
        if ("1106" in t or "1,106" in t) and (
            "total" in t or "round" in t or "cost" in t or "per student" in t
        ):
            cost_ok = True
            break
    record(
        "Contains round-trip cost 1106 (with cost/total/round-trip context)",
        cost_ok,
        f"Text snippet: {full_text[:300]}",
    )


def check_gcal():
    print("\n=== Check 2: Calendar field trip events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime, end_datetime FROM gcal.events
        WHERE start_datetime >= '2026-03-12' AND start_datetime < '2026-03-13'
        AND summary NOT ILIKE '%regular class%'
        ORDER BY start_datetime
    """)
    depart_events = cur.fetchall()
    cur.execute("""
        SELECT summary, start_datetime, end_datetime FROM gcal.events
        WHERE start_datetime >= '2026-03-15' AND start_datetime < '2026-03-16'
        AND summary NOT ILIKE '%regular class%'
        ORDER BY start_datetime
    """)
    return_events = cur.fetchall()
    cur.close()
    conn.close()

    # Outbound: G235 17:30 Beijing -> 19:16 Qufu (per GT doc).
    # Spec: 1 hr before departure (16:30) to 1 hr after arrival (20:16).
    # Tolerance: 15 min slack on each side.
    has_depart = False
    for s, st, et in depart_events:
        if not s or not st or not et:
            continue
        sl = s.lower()
        title_ok = ("field trip" in sl) and ("depart" in sl or "outbound" in sl or "qufu" in sl)
        if not title_ok:
            continue
        # gcal.events.start_datetime is TIMESTAMPTZ (UTC instant). Use the
        # session-tz-independent helper to extract CN wall-clock components.
        # Bare st.hour / st.minute silently compares against the PG session tz.
        _, st_h, st_m = get_zone_components(st, EXPECTED_TIMEZONE)
        _, et_h, et_m = get_zone_components(et, EXPECTED_TIMEZONE)
        if st_h is None or et_h is None:
            continue
        start_min = st_h * 60 + st_m
        end_min = et_h * 60 + et_m
        # Dep 17:30 (1050), 1 hr before = 16:30 (990).
        # Arr 19:16 (1156), 1 hr after = 20:16 (1216).
        if start_min <= 990 + 15 and end_min >= 1216 - 15:
            has_depart = True
            break
    record(
        "Calendar 'Field Trip: Depart for Qufu' on Mar 12 covers ~16:30 to ~20:16 window",
        has_depart,
        f"Mar 12 events: {[(e[0], str(e[1]), str(e[2])) for e in depart_events]}",
    )

    # Return: G236 15:00 Qufu -> 16:46 Beijing (per GT doc).
    # Window: 1 hr before departure (14:00) to 1 hr after arrival (17:46).
    has_return = False
    for s, st, et in return_events:
        if not s or not st or not et:
            continue
        sl = s.lower()
        title_ok = ("field trip" in sl) and ("return" in sl or "back" in sl or "beijing" in sl)
        if not title_ok:
            continue
        # Same TIMESTAMPTZ fix as the depart loop above.
        _, st_h, st_m = get_zone_components(st, EXPECTED_TIMEZONE)
        _, et_h, et_m = get_zone_components(et, EXPECTED_TIMEZONE)
        if st_h is None or et_h is None:
            continue
        start_min = st_h * 60 + st_m
        end_min = et_h * 60 + et_m
        # Dep 15:00 (900), 1 hr before = 14:00 (840).
        # Arr 16:46 (1006), 1 hr after = 17:46 (1066).
        if start_min <= 840 + 15 and end_min >= 1066 - 15:
            has_return = True
            break
    record(
        "Calendar 'Field Trip: Return to Beijing' on Mar 15 covers ~14:00 to ~17:46 window",
        has_return,
        f"Mar 15 events: {[(e[0], str(e[1]), str(e[2])) for e in return_events]}",
    )


def check_canvas_announcement():
    print("\n=== Check 3: Canvas announcement ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        # Tighten: title MUST contain 'field trip' (not just 'qufu') so that
        # the check is not tautologically satisfied by 'qufu' appearing in
        # both title and content. Content must include the train number AND
        # the travel date for substantive evidence.
        cur.execute("""
            SELECT COUNT(*) FROM canvas.announcements
            WHERE title ILIKE '%field%trip%'
            AND (message ILIKE '%g235%' OR message ILIKE '%g236%')
            AND (message ILIKE '%march 12%' OR message ILIKE '%2026-03-12%' OR message ILIKE '%03/12%')
        """)
        cnt = cur.fetchone()[0]
        record(
            "Canvas announcement for field trip created (title 'field trip' AND content has train number AND travel date)",
            cnt >= 1,
            f"Found {cnt} matching announcements",
        )
    except Exception as e:
        # Canvas may not have this table in test env
        record("Canvas announcement check (canvas table query)", False, str(e))
    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 4: Email to students@university.edu ===")
    # Tighten: require the email body mentions departure train, return train,
    # AND the round-trip cost (1106). Just count >=1 with any subject is
    # too lax. Use tables email.messages (outbound from professor) and
    # email.sent_log if present.
    matching_emails = []
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            "SELECT subject, body_text FROM email.messages"
            " WHERE to_addr::text ILIKE '%students@university.edu%'"
            "   AND from_addr NOT ILIKE '%students@university.edu%'"
        )
        matching_emails.extend(cur.fetchall())
        cur.close()
        conn.close()
    except Exception:
        pass
    try:
        conn2 = psycopg2.connect(**DB_CONFIG)
        cur2 = conn2.cursor()
        cur2.execute(
            "SELECT subject, body FROM email.sent_log"
            " WHERE to_addr::text ILIKE '%students@university.edu%'"
        )
        matching_emails.extend(cur2.fetchall())
        cur2.close()
        conn2.close()
    except Exception:
        pass

    # Email exists?
    record("Email sent to students@university.edu",
           len(matching_emails) >= 1,
           f"matching count: {len(matching_emails)}")

    # Body mentions both train numbers AND the round-trip cost
    body_ok = False
    for subj, body in matching_emails:
        full = ((subj or "") + " " + (body or "")).lower()
        has_g235 = "g235" in full
        has_g236 = "g236" in full
        has_cost = "1106" in full or "1,106" in full
        if has_g235 and has_g236 and has_cost:
            body_ok = True
            break
    record("Email body contains both train numbers (G235, G236) and round-trip cost (1106)",
           body_ok,
           f"emails inspected: {len(matching_emails)}")

    # Attachment check - per task line 15 'attaching the field trip notice'
    att_count = 0
    try:
        conn3 = psycopg2.connect(**DB_CONFIG)
        cur3 = conn3.cursor()
        # Try email.attachments joined with email.messages
        cur3.execute("""
            SELECT COUNT(*) FROM email.attachments a
            JOIN email.messages m ON a.message_id = m.id
            WHERE m.to_addr::text ILIKE '%students@university.edu%'
              AND m.from_addr NOT ILIKE '%students@university.edu%'
        """)
        att_count = cur3.fetchone()[0]
        cur3.close()
        conn3.close()
    except Exception:
        pass
    record("Email to students has at least one attachment",
           att_count >= 1,
           f"attachments: {att_count}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal()
    check_canvas_announcement()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if FAIL_COUNT == 0 and PASS_COUNT > 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
