"""Evaluation for yf-stock-comparison-word-gcal."""
import os
import re
import sys
import argparse
import psycopg2
from zoneinfo import ZoneInfo

# ──────────────────────────────────────────────────────────────────────────
# EVALUATION GROUND TRUTH SPEC (gcal tz root-fix v3, case-study 2026-08-13)
# gcal.events.start_datetime / end_datetime are TIMESTAMPTZ storing the UTC
# instant. psycopg2 returns a tz-aware datetime whose display tz follows the
# PG session TimeZone (case-study: compute node default Asia/Shanghai), so
# bare sd.hour / sd.date() / strftime('%H:%M') silently compares against the
# wrong wall-clock. Use utils.evaluation.gcal_helpers (session-tz-independent).
# ──────────────────────────────────────────────────────────────────────────
# task.md line 7: "Portfolio Review Meeting" on April 10, 2026 from 14:00 to
# 15:30 (NYSE / financial context → America/New_York / ET).
from utils.evaluation.gcal_helpers import get_zone_components  # noqa: E402

EXPECTED_TIMEZONE = ZoneInfo("America/New_York")
EXPECTED_REVIEW_DATE = "2026-04-10"
EXPECTED_REVIEW_START_HOUR = 14
EXPECTED_REVIEW_START_MINUTE = 0
EXPECTED_REVIEW_END_HOUR = 15
EXPECTED_REVIEW_END_MINUTE = 30

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, rel_tol=0.05, abs_tol=1.0):
    try:
        return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)
    except Exception:
        return False


SYMBOLS = ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]


def fetch_expected_stocks():
    """Fetch current price + sector from DB. Fallback to known snapshot."""
    fallback = {
        "GOOGL": {"price": 300.88, "sector": "Communication Services", "name_kw": ["alphabet", "google"]},
        "AMZN":  {"price": 218.94, "sector": "Consumer Cyclical",     "name_kw": ["amazon"]},
        "JPM":   {"price": 293.55, "sector": "Financial Services",    "name_kw": ["jpmorgan", "jp morgan", "chase"]},
        "JNJ":   {"price": 239.63, "sector": "Healthcare",            "name_kw": ["johnson"]},
        "XOM":   {"price": 150.76, "sector": "Energy",                "name_kw": ["exxon"]},
    }
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT symbol, data->>'currentPrice' AS price,
                   data->>'sector' AS sector, data->>'longName' AS lname
            FROM yf.stock_info
            WHERE symbol = ANY(%s)
        """, (SYMBOLS,))
        out = {}
        for sym, price, sector, lname in cur.fetchall():
            try:
                p = float(price)
            except Exception:
                p = None
            kw = []
            if lname:
                lname_l = lname.lower()
                for kk in ["alphabet", "google", "amazon", "jpmorgan", "jp morgan", "chase",
                           "johnson", "exxon"]:
                    if kk in lname_l:
                        kw.append(kk)
                if not kw:
                    kw = [lname_l.split()[0]]
            out[sym] = {"price": p if p is not None else fallback[sym]["price"],
                        "sector": (sector or fallback[sym]["sector"]),
                        "name_kw": kw or fallback[sym]["name_kw"]}
        cur.close()
        conn.close()
        # ensure all 5
        for s in SYMBOLS:
            if s not in out:
                out[s] = fallback[s]
        return out
    except Exception as e:
        print(f"[expected] DB fetch failed: {e}; using fallback")
        return fallback


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word document ===")
    doc_path = os.path.join(agent_workspace, "Stock_Comparison_Report.docx")
    if not os.path.exists(doc_path):
        record("Stock_Comparison_Report.docx exists", False, doc_path)
        for k in ["Word readable", "Title 'Stock Comparison Report'",
                  "Has 'Stock Overview' section",
                  "Has 'Analysis' section",
                  "Has 'Portfolio Strategy' section",
                  "Stock table has 5 rows",
                  "Stock table has 5 columns",
                  "All 5 symbols present in table",
                  "All 5 prices match DB"]:
            record(k, False, "skipped (no docx)")
        return
    record("Stock_Comparison_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        record("Word readable", False, str(e))
        return
    record("Word readable", True)

    para_texts = [p.text for p in doc.paragraphs]
    full_text = "\n".join(para_texts)
    full_text_lower = full_text.lower()

    # Title must be a paragraph whose stripped text equals 'Stock Comparison Report'
    title_ok = any(p.text.strip().lower() == "stock comparison report" for p in doc.paragraphs)
    record("Title 'Stock Comparison Report'", title_ok, full_text[:200])

    record("Has 'Stock Overview' section", "stock overview" in full_text_lower)
    record("Has 'Analysis' section", re.search(r"\banalysis\b", full_text_lower) is not None)
    record("Has 'Portfolio Strategy' section", "portfolio strategy" in full_text_lower)

    expected = fetch_expected_stocks()

    # Find the stock table
    target_table = None
    for t in doc.tables:
        if not t.rows:
            continue
        header = [c.text.strip().lower() for c in t.rows[0].cells]
        # detect symbol/company headers
        if any("symbol" in h for h in header) and any("price" in h for h in header):
            target_table = t
            break
    if target_table is None:
        record("Stock table has 5 rows", False, "no stock table found")
        record("Stock table has 5 columns", False, "skipped")
        record("All 5 symbols present in table", False, "skipped")
        record("All 5 prices match DB", False, "skipped")
        return

    rows_data = [[c.text.strip() for c in row.cells] for row in target_table.rows]
    # Header + 5 rows
    data_rows = rows_data[1:]
    record("Stock table has 5 rows", len(data_rows) == 5, f"found {len(data_rows)}")
    record("Stock table has 5 columns", len(rows_data[0]) >= 5, f"cols={len(rows_data[0])}")

    # Locate column indexes
    header = [h.lower() for h in rows_data[0]]
    def col_for(*kws):
        for i, h in enumerate(header):
            if any(k in h for k in kws):
                return i
        return -1
    ci_sym = col_for("symbol")
    ci_price = col_for("price")

    # Build map symbol -> row
    row_by_sym = {}
    if ci_sym >= 0:
        for r in data_rows:
            if ci_sym < len(r):
                row_by_sym[r[ci_sym].upper()] = r

    all_syms_present = all(s in row_by_sym for s in SYMBOLS)
    record("All 5 symbols present in table", all_syms_present,
           f"found: {sorted(row_by_sym.keys())}")

    # Verify each price matches DB (within tolerance)
    prices_ok = True
    price_detail = []
    for sym in SYMBOLS:
        row = row_by_sym.get(sym)
        if not row or ci_price < 0 or ci_price >= len(row):
            prices_ok = False
            price_detail.append(f"{sym}: no row/cell")
            continue
        cell = row[ci_price]
        # extract first number
        m = re.search(r"-?\d+(?:\.\d+)?", cell.replace(",", ""))
        if not m:
            prices_ok = False
            price_detail.append(f"{sym}: no number in '{cell}'")
            continue
        actual = float(m.group(0))
        target = expected[sym]["price"]
        if not num_close(actual, target):
            prices_ok = False
            price_detail.append(f"{sym}: {actual} vs {target}")
    record("All 5 prices match DB", prices_ok, "; ".join(price_detail)[:280])


def check_gcal_event():
    print("\n=== Check 2: GCal Portfolio Review Meeting ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, description, start_datetime, end_datetime
            FROM gcal.events
            WHERE start_datetime >= TIMESTAMPTZ '2026-04-09T00:00:00Z'
              AND start_datetime <  TIMESTAMPTZ '2026-04-12T00:00:00Z'
            ORDER BY start_datetime
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("GCal DB query", False, str(e))
        for k in ["GCal event 'Portfolio Review Meeting' on 2026-04-10",
                  "GCal event time 14:00-15:30",
                  "GCal event description mentions 5 stocks"]:
            record(k, False, "skipped")
        return

    matched = []
    for summary, desc, sdt, edt in rows:
        s_lower = (summary or "").lower()
        if "portfolio review" in s_lower or ("portfolio" in s_lower and "review" in s_lower):
            matched.append((summary, desc, sdt, edt))
    record("GCal event 'Portfolio Review Meeting' on 2026-04-10",
           len(matched) >= 1, f"events on date: {[r[0] for r in rows]}")

    if not matched:
        record("GCal event time 14:00-15:30", False, "skipped")
        record("GCal event description mentions 5 stocks", False, "skipped")
        return

    summary, desc, sdt, edt = matched[0]
    # gcal.events.start_datetime is TIMESTAMPTZ (UTC instant). strftime('%H:%M')
    # / str(sdt)[:10] would render the PG session-tz wall-clock (case-study:
    # Asia/Shanghai), so compare via the session-tz-independent helper anchored
    # to task.md's ET (NYSE hours).
    ok_time = False
    if sdt is not None:
        sd_date, sd_hour, sd_minute = get_zone_components(sdt, EXPECTED_TIMEZONE)
        ed_hour = ed_minute = None
        if edt is not None:
            _, ed_hour, ed_minute = get_zone_components(edt, EXPECTED_TIMEZONE)
        ok_time = (
            sd_date is not None
            and sd_date.isoformat() == EXPECTED_REVIEW_DATE
            and sd_hour == EXPECTED_REVIEW_START_HOUR
            and sd_minute == EXPECTED_REVIEW_START_MINUTE
            and ed_hour == EXPECTED_REVIEW_END_HOUR
            and ed_minute == EXPECTED_REVIEW_END_MINUTE
        )
    record("GCal event time 14:00-15:30", ok_time, f"start={sdt} end={edt}")

    desc_lower = (desc or "").lower()
    summary_lower = (summary or "").lower()
    combined = desc_lower + " " + summary_lower
    sym_hits = sum(1 for s in SYMBOLS if s.lower() in combined)
    record("GCal event description mentions 5 stocks", sym_hits >= 5,
           f"hits={sym_hits}, desc={desc_lower[:200]}")


def check_email():
    print("\n=== Check 3: Email to analyst@investment.com ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, to_addr, body_text, body_html FROM email.messages
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
    except Exception as e:
        record("Email DB query", False, str(e))
        for k in ["Email to analyst@investment.com",
                  "Email subject 'Stock Comparison Report - Q2 2026'",
                  "Email body summarizes findings"]:
            record(k, False, "skipped")
        return

    matched = []
    for subj, to_addr, btxt, bhtml in rows:
        to_str = str(to_addr or "").lower()
        if "analyst@investment.com" in to_str:
            matched.append((subj or "", (btxt or "") + "\n" + (bhtml or "")))
    record("Email to analyst@investment.com", len(matched) >= 1, f"count={len(matched)}")
    if not matched:
        record("Email subject 'Stock Comparison Report - Q2 2026'", False, "skipped")
        record("Email body summarizes findings", False, "skipped")
        return

    subj, body = matched[0]
    subj_l = subj.strip().lower()
    expected_subj = "stock comparison report - q2 2026"
    ok_subj = (
        expected_subj in subj_l
        or all(tok in subj_l for tok in ["stock comparison report", "q2 2026"])
        or all(tok in subj_l for tok in ["stock comparison", "q2"]) and "2026" in subj_l
    )
    record("Email subject 'Stock Comparison Report - Q2 2026'", ok_subj, f"subj={subj!r}")

    body_l = body.lower()
    # Body should mention at least 3 of the 5 symbols and at least one of recommendation/buy/strong_buy
    sym_hits = sum(1 for s in SYMBOLS if s.lower() in body_l)
    has_kw = any(k in body_l for k in ["recommend", "buy", "strong_buy", "strong buy"])
    ok_body = sym_hits >= 3 and has_kw
    record("Email body summarizes findings", ok_body, f"sym_hits={sym_hits}, has_kw={has_kw}, body={body_l[:200]}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal_event()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed [FAIL_COUNT={FAIL_COUNT}]")
    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
