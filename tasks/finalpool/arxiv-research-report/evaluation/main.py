"""
Evaluation script for arxiv-research-report task.
Checks that LLM_Reasoning_Survey.docx and .pdf exist with expected content.

Usage:
  python -m evaluation.main --agent_workspace <path> --groundtruth_workspace <path> --launch_time <time>
"""
import argparse
import os
import re
import sys


PASS_COUNT = 0
FAIL_COUNT = 0


def check(name: str, condition: bool, detail: str = ""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_truncated = (detail[:200] + "...") if len(detail) > 200 else detail
        print(f"  [FAIL] {name}: {detail_truncated}")


def normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace."""
    return re.sub(r'\s+', ' ', text.lower().strip())


def collect_body_text(doc):
    """Collect all visible text from a Word document, including table cell text
    (and nested tables), in document order.  A faithful agent may legitimately
    place the per-paper title/author/citation details in a table (word MCP's
    add_table), so the evaluator must see that content too."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parts = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith('}p'):
            parts.append(Paragraph(child, doc).text)
        elif child.tag.endswith('}tbl'):
            parts.append(_table_text(Table(child, doc)))
    return "\n".join(parts)


def _table_text(table):
    parts = []
    for row in table.rows:
        for cell in row.cells:
            parts.append("\n".join(p.text for p in cell.paragraphs))
            for nested in cell.tables:
                parts.append(_table_text(nested))
    return "\n".join(parts)


# ── Section heading detection ─────────────────────────────────────────────────
# A faithful agent may write section headings either with Word heading styles
# (word MCP add_paragraph with style) or as plain short lines (word MCP's default
# add_paragraph).  Body-text mentions of e.g. "methodology"/"conclusion" in the
# introduction must NEVER influence the order check, so positions are always
# anchored to detected heading *paragraphs*, never to raw-text find().
_SECTION_PATTERNS = {
    'introduction': [re.compile(r'\bintroduction\b'), re.compile(r'\bintro\b')],
    'literature review': [re.compile(r'\bliterature\b'), re.compile(r'\brelated\s+work\b')],
    'methodology': [re.compile(r'\bmethodolog\w*\b'),
                    re.compile(r'\bmethods?\s+comparison\b'),
                    re.compile(r'\bcomparison\s+of\s+methods?\b')],
    'conclusion': [re.compile(r'\bconclusion\w*\b'), re.compile(r'\bsummary\b'),
                   re.compile(r'\bfuture\s+work\b')],
}

_SECTION_START_WORDS = [
    'introduction', 'intro',
    'literature', 'related work',
    'methodology', 'methodologies', 'methodological', 'methods', 'method',
    'conclusion', 'concluding', 'summary', 'future work',
]

_SECTIONS_ORDER = ['introduction', 'literature review', 'methodology', 'conclusion']


def _strip_numbering(text):
    """Remove leading section/list numbering such as '1. ', '1.1 ', '(2) ', 'A) ', 'III.'."""
    m = re.match(r'^\s*(?:\(?\d+(?:\.\d+)*\)?[.):\-]?\s*|[A-Za-z]\)\s*|[IVXLCDM]+[.):]?\s*)', text)
    return text[m.end():] if m else text


def _leading_section_word(core):
    """If `core` starts with a section keyword or phrase (word boundary), return (keyword, rest)."""
    for kw in _SECTION_START_WORDS:
        if core.startswith(kw) and (len(core) == len(kw) or not core[len(kw)].isalnum()):
            return kw, core[len(kw):]
    return None, ''


def _is_heading_line(text_norm):
    """Heuristic: does an unstyled paragraph read like a section heading?

    A heading is a short line (<= 80 chars) that starts with a section keyword
    (after stripping list numbering) and does not continue into a full prose
    sentence.  Very short keyword lines with a trailing period ("Introduction.")
    are accepted; longer sentences ("Introduction of each paper is summarized.")
    are rejected so body text cannot be mistaken for headings.
    """
    if not text_norm or len(text_norm) > 80:
        return False
    core = _strip_numbering(text_norm).strip()
    if not core:
        return False
    kw, rest = _leading_section_word(core)
    if kw is None:
        return False
    # prose that continues well past the keyword phrase reads like a sentence
    if len(rest) > 20:
        return False
    if core.endswith(('.', '。', '!', '?', ';')):
        body = core[:-1].strip()
        if len(body) > 30:
            return False
    return True


def detect_section_headings(doc):
    """Return a list of (paragraph_index, normalized_text) for heading paragraphs."""
    headings = []
    for i, para in enumerate(doc.paragraphs):
        styled = bool(para.style and para.style.name and 'heading' in para.style.name.lower())
        t = normalize(para.text)
        if styled or _is_heading_line(t):
            headings.append((i, t))
    return headings


def section_positions(headings, order):
    """First heading-paragraph index for each section in `order` (-1 if absent)."""
    pos = []
    for sec in order:
        found = -1
        pats = _SECTION_PATTERNS[sec]
        for idx, ht in headings:
            if any(p.search(ht) for p in pats):
                found = idx
                break
        pos.append(found)
    return pos


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", type=str, required=True)
    parser.add_argument("--groundtruth_workspace", type=str, required=True)
    parser.add_argument("--launch_time", type=str, required=False)
    parser.add_argument("--res_log_file", type=str, required=False)
    args = parser.parse_args()

    docx_path = os.path.join(args.agent_workspace, "LLM_Reasoning_Survey.docx")
    pdf_path = os.path.join(args.agent_workspace, "LLM_Reasoning_Survey.pdf")

    # ── Check 1: Word document exists ────────────────────────────────────────
    check("LLM_Reasoning_Survey.docx exists", os.path.exists(docx_path),
          f"File not found at {docx_path}")

    if not os.path.exists(docx_path):
        print(f"\nResults: {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} passed, {FAIL_COUNT} failed")
        sys.exit(1)

    # Read the Word document
    try:
        from docx import Document
        doc = Document(docx_path)
        full_text = collect_body_text(doc)
    except Exception as e:
        check("Word document readable", False, str(e))
        print(f"\nResults: {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} passed, {FAIL_COUNT} failed")
        sys.exit(1)

    normalized = normalize(full_text)

    # ── Check 2: Minimum content length ──────────────────────────────────────
    check("Document has at least 500 characters",
          len(full_text.strip()) >= 500,
          f"Document has {len(full_text.strip())} characters")

    # ── Check 3: Title contains "Survey" or "Reasoning" ─────────────────────
    check("Document title mentions survey or reasoning",
          "survey" in normalized or "reasoning" in normalized,
          "Neither 'survey' nor 'reasoning' found")

    # ── Check 4: Has required section headings ───────────────────────────────
    # Existence checks are intentionally lenient and consistent with the keywords
    # used by the order check below (e.g. a heading "Methods Comparison" or a
    # "Summary" conclusion must satisfy the same predicates both places).
    has_intro = "introduction" in normalized
    has_lit_review = "literature" in normalized or "related work" in normalized
    has_methodology = ("methodolog" in normalized
                       or "methods comparison" in normalized
                       or "method comparison" in normalized)
    has_conclusion = ("conclusion" in normalized or "summary" in normalized
                      or "future work" in normalized)

    check("Has Introduction section", has_intro, "No 'Introduction' heading or text found")
    check("Has Literature Review section", has_lit_review, "No 'Literature Review' heading or text found")
    check("Has Methodology Comparison section", has_methodology, "No 'Methodology' heading or text found")
    check("Has Conclusion section", has_conclusion, "No 'Conclusion' heading or text found")

    # Section order validation: Introduction -> Literature Review -> Methodology Comparison -> Conclusion
    # Positions are anchored to detected heading paragraphs only (styled headings
    # or short heading-like lines), so body-text mentions of "methodology" or
    # "conclusion" inside the introduction can never break the order.
    headings = detect_section_headings(doc)
    positions = section_positions(headings, _SECTIONS_ORDER)
    sections_in_order = all(p >= 0 for p in positions) and positions == sorted(positions)
    check("Sections appear in correct order", sections_in_order,
          f"positions={positions}")

    # Sample citation counts from Literature Review (all 5 papers: 850/420/650/280/190)
    CITATION_SAMPLES = ['850', '420', '650', '280', '190']
    for cnt in CITATION_SAMPLES:
        check(f"Citation count '{cnt}' referenced",
              cnt in normalized,
              f"citation_count {cnt} not found in document")

    # ── Check 5: Has at least 4 headings ─────────────────────────────────────
    # Count detected heading paragraphs (styled headings OR short heading-like
    # lines), which covers both stylized and plain-text faithful documents.
    heading_count = len(headings)
    check("Has at least 4 headings/sections",
          heading_count >= 4,
          f"Found only {heading_count} headings")

    # ── Check 6: All 5 target paper titles present ───────────────────────────
    paper_titles = [
        "Chain-of-Thought Prompting Elicits Reasoning in Large Language Models",
        "Tree of Thoughts: Deliberate Problem Solving with Large Language Models",
        "Self-Consistency Improves Chain of Thought Reasoning in Language Models",
        "Process Supervision for Mathematical Reasoning",
        "Scaling LLM Reasoning with Reinforcement Learning",
    ]
    for title in paper_titles:
        check(f"Paper title present: {title[:60]}",
              title.lower() in normalized,
              "Title not found in document text")

    # ── Check 7: Key author names present ────────────────────────────────────
    key_authors = ["Jason Wei", "Shunyu Yao", "Xuezhi Wang"]
    for author in key_authors:
        check(f"Author present: {author}",
              author.lower() in normalized,
              f"Author '{author}' not found")

    # ── Check 8: Key domain terms present ────────────────────────────────────
    key_terms = ["chain-of-thought", "tree of thoughts", "self-consistency", "process supervision"]
    for term in key_terms:
        check(f"Key term present: {term}",
              term.lower() in normalized,
              f"Term '{term}' not found")

    # ── Check 9: Noise papers NOT reviewed ────────────────────────────────────
    # Check the exact noise paper titles (injected by preprocess) rather than generic
    # topic phrases, so a passing mention of e.g. "federated learning" in a comparison
    # sentence does not cause a false FAIL.
    noise_paper_titles = [
        "Efficient Image Classification with Vision Transformers",
        "Federated Learning for Privacy-Preserving NLP",
        "Protein Structure Prediction Using Deep Learning",
    ]
    for title in noise_paper_titles:
        check(f"Noise paper NOT reviewed: {title[:50]}",
              title.lower() not in normalized,
              f"Noise paper title '{title}' found in document -- should not be included")

    # ── Check 10: PDF file exists and has reasonable size ────────────────────
    check("LLM_Reasoning_Survey.pdf exists", os.path.exists(pdf_path),
          f"PDF not found at {pdf_path}")
    if os.path.exists(pdf_path):
        pdf_size = os.path.getsize(pdf_path)
        check("PDF file size > 5KB",
              pdf_size > 5000,
              f"PDF is only {pdf_size} bytes")

    # ── Summary ──────────────────────────────────────────────────────────────
    total = PASS_COUNT + FAIL_COUNT
    print(f"\nResults: {PASS_COUNT}/{total} passed, {FAIL_COUNT} failed")

    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
