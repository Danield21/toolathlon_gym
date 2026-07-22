"""
Evaluation for yt-veritasium-howtocook-wellness-excel-gcal task.

Checks:
1. Wellness_Plan.xlsx exists with "Videos" sheet having >= 4 rows
2. "Recipes" sheet has >= 8 rows
3. "Weekly_Meal_Plan" sheet has exactly 7 rows
4. Wellness_Guide.docx exists
5. Word doc has >= 3 heading-level sections
6. Word doc contains health/nutrition/recipe/wellness keywords
7. GCal has >= 3 new Wellness Check-in events in April 2026
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Check 1-3: Wellness_Plan.xlsx ===")
    # Require exact filename
    xlsx_path = os.path.join(agent_workspace, "Wellness_Plan.xlsx")
    if not os.path.isfile(xlsx_path):
        xlsx_path = None
    record("Wellness_Plan.xlsx exists (exact filename)", xlsx_path is not None,
           f"Expected file at {agent_workspace}/Wellness_Plan.xlsx")

    if not xlsx_path:
        for chk in ["Videos sheet >= 4 rows", "Recipes sheet >= 8 rows", "Weekly_Meal_Plan sheet = 7 rows"]:
            record(chk, False, "xlsx not found")
        return

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path)

        def _find_sheet(needle_exact):
            # Exact (case-insensitive)
            for name in wb.sheetnames:
                if name.strip().lower() == needle_exact.lower():
                    return wb[name]
            return None

        # Videos sheet (exact)
        videos_sheet = _find_sheet("Videos")
        record("Sheet 'Videos' exists (exact name)",
               videos_sheet is not None, f"Sheets: {wb.sheetnames}")
        if videos_sheet:
            data_rows = [r for r in videos_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Videos sheet has >= 4 data rows", len(data_rows) >= 4,
                   f"Found {len(data_rows)} rows")

        # Recipes sheet (exact)
        recipe_sheet = _find_sheet("Recipes")
        record("Sheet 'Recipes' exists (exact name)",
               recipe_sheet is not None, f"Sheets: {wb.sheetnames}")
        if recipe_sheet:
            data_rows = [r for r in recipe_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Recipes sheet has >= 8 data rows", len(data_rows) >= 8,
                   f"Found {len(data_rows)} rows")

        # Weekly_Meal_Plan sheet (exact)
        meal_sheet = _find_sheet("Weekly_Meal_Plan")
        if meal_sheet is None:
            # Allow space variant
            meal_sheet = _find_sheet("Weekly Meal Plan")
        record("Sheet 'Weekly_Meal_Plan' exists (exact or space variant)",
               meal_sheet is not None, f"Sheets: {wb.sheetnames}")
        if meal_sheet:
            data_rows = [r for r in meal_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Weekly_Meal_Plan sheet has exactly 7 rows", len(data_rows) == 7,
                   f"Found {len(data_rows)} rows (expected 7)")
            # Verify Day column has Monday..Sunday
            if len(data_rows) >= 7:
                expected_days = ["monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday"]
                seen_days = [str(r[0] or "").strip().lower() for r in data_rows]
                missing = [d for d in expected_days if d not in seen_days]
                record("Weekly_Meal_Plan has Mon-Sun (all 7 days)",
                       len(missing) == 0,
                       f"Got days: {seen_days}; missing: {missing}")
                # All Total_Recipes == 3
                tr_ok = all(num_close(r[4], 3, 0.01) for r in data_rows[:7] if r and r[4] is not None)
                record("Weekly_Meal_Plan Total_Recipes column == 3 for all 7 days",
                       tr_ok,
                       f"Total_Recipes column: {[r[4] for r in data_rows[:7] if r]}")

        # --- Groundtruth XLSX value comparison ---
        gt_path = os.path.join(groundtruth_workspace, "Wellness_Plan.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                record(f"GT '{gt_sname}' row count", len(a_rows) == len(gt_rows),
                       f"Expected {len(gt_rows)}, got {len(a_rows)}")
                for ri in range(min(3, len(gt_rows))):
                    if ri >= len(a_rows):
                        break
                    ok = True
                    for ci in range(min(len(gt_rows[ri]), len(a_rows[ri]))):
                        gv, av = gt_rows[ri][ci], a_rows[ri][ci]
                        if gv is None:
                            continue
                        if isinstance(gv, (int, float)):
                            if not num_close(av, gv, max(abs(gv) * 0.1, 1.0)):
                                ok = False
                                break
                        else:
                            if not str_match(av, gv):
                                ok = False
                                break
                    record(f"GT '{gt_sname}' row {ri+1} values", ok,
                           f"gt={gt_rows[ri][:4]}, agent={a_rows[ri][:4] if ri < len(a_rows) else 'missing'}")
            gt_wb.close()

    except Exception as e:
        for chk in ["Videos sheet >= 4 rows", "Recipes sheet >= 8 rows", "Weekly_Meal_Plan = 7 rows"]:
            record(chk, False, str(e))


def check_word(agent_workspace):
    print("\n=== Check 4-6: Wellness_Guide.docx ===")
    # Tighten: require exact filename 'Wellness_Guide.docx'
    expected = os.path.join(agent_workspace, "Wellness_Guide.docx")
    docx_path = expected if os.path.exists(expected) else None

    record("Wellness_Guide.docx exists (exact filename)",
           docx_path is not None,
           f"Expected: {expected}")

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found")
        record("Word doc contains wellness keywords", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings")

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        keywords = ["health", "nutrition", "recipe", "wellness", "diet", "meal", "biology", "science"]
        found = [k for k in keywords if k in full_text]
        # Tighten: comprehensive guide should hit >=5 of 8 keywords
        record("Word doc contains >=5 of 8 health/wellness/recipe keywords",
               len(found) >= 5,
               f"Found keywords: {found}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e))
        record("Word doc contains wellness keywords", False, str(e))


def check_gcal():
    print("\n=== Check 7: GCal Wellness Check-in Events in April 2026 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, description, start_datetime, end_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        AND summary ILIKE '%wellness check%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Tighten: exactly 4 wellness check-in events
    record("Exactly 4 'Wellness Check-in' events in April 2026",
           len(events) == 4,
           f"Found {len(events)} events: {[(e[0], str(e[2])) for e in events]}")

    # Verify dates: 2026-04-01, 04-08, 04-15, 04-22
    expected_dates = ["2026-04-01", "2026-04-08", "2026-04-15", "2026-04-22"]
    seen_dates = set()
    for s, d, sd, ed in events:
        if sd:
            sd_str = str(sd)
            for ed_iso in expected_dates:
                if ed_iso in sd_str:
                    seen_dates.add(ed_iso)
    record("All 4 Wellness Check-in events on Apr 1/8/15/22",
           len(seen_dates) == 4,
           f"Got dates: {sorted(seen_dates)}; expected: {expected_dates}")

    # Verify all start at 07:00 (or 11:00/12:00 UTC for NY DST)
    valid_starts = 0
    for _, _, sd, _ in events:
        s = str(sd or "")
        if len(s) >= 13:
            try:
                hr = int(s[11:13])
                if hr in (7, 11, 12):
                    valid_starts += 1
            except ValueError:
                pass
    record("All 4 events start at 07:00 (local) or UTC equivalent",
           valid_starts == len(events) and len(events) == 4,
           f"Valid starts: {valid_starts}")

    # Verify duration ~30 min
    valid_dur = 0
    for _, _, sd, ed in events:
        if sd and ed:
            try:
                if abs((ed - sd).total_seconds() - 1800) <= 60:
                    valid_dur += 1
            except Exception:
                pass
    record("All 4 events are 30 minutes long",
           valid_dur == len(events) and len(events) == 4,
           f"Valid durations: {valid_dur}")

    # Verify description mentions 'meal plan' and 'health topic'
    desc_ok = 0
    for _, desc, _, _ in events:
        dl = (desc or "").lower()
        if ("meal plan" in dl or "meal-plan" in dl) and ("health topic" in dl or "health" in dl):
            desc_ok += 1
    record("All 4 event descriptions mention meal plan + health topic",
           desc_ok == len(events) and len(events) == 4,
           f"Got {desc_ok} valid descriptions")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace, args.groundtruth_workspace)
    check_word(args.agent_workspace)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    # Use FAIL_COUNT == 0 (replaces 70% accuracy threshold)
    if FAIL_COUNT == 0:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
