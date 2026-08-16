"""Evaluation for canvas-quiz-analysis-gsheet-email."""
import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def _get_expected_quiz_count():
    """Compute expected quiz count from canvas DB."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(DISTINCT q.id)
            FROM canvas.quizzes q
            JOIN canvas.quiz_submissions qs ON qs.quiz_id = q.id
            WHERE qs.score IS NOT NULL
        """)
        n = cur.fetchone()[0] or 0
        cur.close()
        conn.close()
        return int(n)
    except (psycopg2.OperationalError, Exception):
        return 76


def _get_expected_quiz_stats():
    """Compute per-quiz expected stats keyed by (course_name, quiz_title)."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT c.name AS course_name,
                   q.title AS quiz_title,
                   COUNT(qs.id) AS total_attempts,
                   ROUND(AVG(qs.score)::numeric, 1) AS avg_score,
                   MAX(q.points_possible) AS max_score,
                   ROUND(100.0 * SUM(CASE WHEN qs.score >= 0.6 * q.points_possible
                                          THEN 1 ELSE 0 END)
                         / NULLIF(COUNT(qs.id), 0), 1) AS pass_rate
            FROM canvas.courses c
            JOIN canvas.quizzes q ON q.course_id = c.id
            JOIN canvas.quiz_submissions qs ON qs.quiz_id = q.id
            WHERE qs.score IS NOT NULL
            GROUP BY c.name, q.title
        """)
        out = {}
        for cname, qtitle, tot, avg, mx, pr in cur.fetchall():
            out[(cname.lower().strip(), qtitle.lower().strip())] = {
                "total_attempts": int(tot or 0),
                "avg_score": float(avg) if avg is not None else None,
                "max_score": float(mx) if mx is not None else None,
                "pass_rate": float(pr) if pr is not None else None,
            }
        cur.close()
        conn.close()
        return out
    except (psycopg2.OperationalError, Exception):
        return {}


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        cur.execute("""
            SELECT id, title FROM gsheet.spreadsheets
            WHERE LOWER(TRIM(title)) = LOWER(%s)
        """, ("Quiz Performance Tracker",))
        sheets = cur.fetchall()
        check("Spreadsheet titled exactly 'Quiz Performance Tracker' exists",
              len(sheets) >= 1,
              f"Found {len(sheets)} matching spreadsheets (exact title)")

        if sheets:
            ss_id = sheets[0][0]
            # Check that sheet named 'Quiz Scores' exists
            cur.execute("""
                SELECT id, title FROM gsheet.sheets
                WHERE spreadsheet_id = %s AND LOWER(title) = 'quiz scores'
            """, (ss_id,))
            quiz_sheets = cur.fetchall()
            check("Sheet titled exactly 'Quiz Scores' exists in spreadsheet",
                  len(quiz_sheets) >= 1,
                  f"Found {len(quiz_sheets)} matching sheets")

            # Build grid for value-level checks
            cur.execute("""
                SELECT row_index, col_index, value FROM gsheet.cells c
                WHERE c.spreadsheet_id = %s
                ORDER BY row_index, col_index
            """, (ss_id,))
            grid = {}
            for r, c, v in cur.fetchall():
                grid.setdefault(r, {})[c] = v

            # Number of data rows: tighten to expected DB-derived count
            expected_n = _get_expected_quiz_count()
            data_rows = sum(1 for r in grid if r > 0)
            check(f"GSheet has expected number of quiz data rows ({expected_n})",
                  data_rows >= expected_n,
                  f"Found {data_rows} data rows; expected at least {expected_n}")

            # Per-quiz value validation: spot-check up to 5 known quizzes
            stats = _get_expected_quiz_stats()
            if grid and stats:
                # find header row
                header_row = grid.get(0) or grid.get(min(grid.keys()))
                header_lower = {i: str(v or "").lower() for i, v in (header_row or {}).items()}
                col_idx = {}
                for i, h in header_lower.items():
                    if "course" in h:
                        col_idx["course"] = i
                    elif "quiz_title" in h or h.strip() == "quiz title":
                        col_idx["quiz_title"] = i
                    elif "total_attempts" in h or "attempts" in h:
                        col_idx["total"] = i
                    elif "avg_score" in h or h.strip() == "avg score":
                        col_idx["avg"] = i
                    elif "max_score" in h or h.strip() == "max score":
                        col_idx["max"] = i
                    elif "pass_rate" in h:
                        col_idx["pr"] = i
                # Build (course, title) -> row from sheet (preserve original
                # ordering for sort-validation).
                row_lookup = {}
                ordered_pairs = []
                for r in sorted(grid.keys()):
                    if r == 0:
                        continue
                    row = grid[r]
                    if "course" in col_idx and "quiz_title" in col_idx:
                        c_v = str(row.get(col_idx["course"], "") or "").lower().strip()
                        q_v = str(row.get(col_idx["quiz_title"], "") or "").lower().strip()
                        if c_v and q_v:
                            row_lookup[(c_v, q_v)] = row
                            ordered_pairs.append((c_v, q_v))

                # Validate sort order: by Course_Name asc, then Quiz_Title asc.
                if len(ordered_pairs) >= 2:
                    expected_sorted = sorted(ordered_pairs, key=lambda p: (p[0], p[1]))
                    sort_ok = ordered_pairs == expected_sorted
                    check(
                        "Quiz Scores sorted by Course_Name then Quiz_Title (ascending)",
                        sort_ok,
                        f"first 5 actual: {ordered_pairs[:5]}; first 5 expected: {expected_sorted[:5]}",
                    )

                checked = 0
                mismatches = 0
                for key, exp in list(stats.items())[:8]:
                    row = row_lookup.get(key)
                    if not row:
                        continue
                    checked += 1
                    if "avg" in col_idx:
                        try:
                            v = float(row[col_idx["avg"]])
                            if exp["avg_score"] is not None and abs(v - exp["avg_score"]) > 0.5:
                                mismatches += 1
                        except (TypeError, ValueError):
                            mismatches += 1
                    if "pr" in col_idx:
                        try:
                            v = float(row[col_idx["pr"]])
                            if exp["pass_rate"] is not None and abs(v - exp["pass_rate"]) > 1.0:
                                mismatches += 1
                        except (TypeError, ValueError):
                            mismatches += 1
                    if "total" in col_idx:
                        try:
                            v = int(float(row[col_idx["total"]]))
                            if v != exp["total_attempts"]:
                                mismatches += 1
                        except (TypeError, ValueError):
                            mismatches += 1
                check(f"Per-quiz value spot checks (avg/pass/total) match GT",
                      mismatches == 0,
                      f"{mismatches} mismatches in {checked} spot-checked quizzes")

        cur.close()
        conn.close()
    except Exception as e:
        check("GSheet check", False, str(e))


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Quiz_Performance_Summary.docx")
    if not os.path.isfile(docx_path):
        check("Quiz_Performance_Summary.docx exists", False, f"Not found: {docx_path}")
        return
    check("Quiz_Performance_Summary.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        para_text = " ".join(p.text for p in doc.paragraphs).lower()
        cell_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text += " " + cell.text.lower()
        all_text = para_text + " " + cell_text

        check("Word doc has meaningful content (>= 200 chars)",
              len(all_text.strip()) >= 200,
              f"Content length: {len(all_text)}")
        check("Word doc contains quiz-related content",
              "quiz" in all_text and ("score" in all_text or "performance" in all_text),
              f"Sample: {all_text[:200]}")
        # Stricter: require BOTH 'recommend' word AND ('support' OR
        # 'intervention') -- task asks for a recommendations section
        # suggesting which courses need additional student support.
        has_recommend_section = "recommend" in all_text
        has_support_term = "support" in all_text or "intervention" in all_text
        check("Word doc contains recommendations section with support guidance",
              has_recommend_section and has_support_term,
              f"recommend={has_recommend_section}, support/intervention={has_support_term}")
        # Require a summary section
        check("Word doc contains summary section",
              "summary" in all_text or "overview" in all_text or "trend" in all_text,
              f"Sample: {all_text[:200]}")
        # Require a table for top-performing quizzes -- table must have >=2 rows
        # (header + at least one data row) and reference quiz/course-related text.
        has_substantive_table = False
        for table in doc.tables:
            if len(table.rows) >= 2:
                tbl_text = " ".join(
                    cell.text.lower() for row in table.rows for cell in row.cells
                )
                if "quiz" in tbl_text or "course" in tbl_text or "score" in tbl_text:
                    has_substantive_table = True
                    break
        check("Word doc has at least one substantive table (top-performing quizzes)",
              has_substantive_table,
              f"Found {len(doc.tables)} tables, none with >=2 rows + quiz/course/score keywords")
    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def check_email():
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # AND on recipient and exact subject
        cur.execute("""
            SELECT id, subject, to_addr, body_text
            FROM email.messages
            WHERE to_addr::text ILIKE '%%academic_coordinator@university.edu%%'
              AND LOWER(subject) = LOWER(%s)
        """, ("Quiz Performance Analysis",))
        emails = cur.fetchall()
        check("Email to academic_coordinator@university.edu with subject 'Quiz Performance Analysis'",
              len(emails) >= 1,
              "No exact-subject email to coordinator found")
        if emails:
            body = str(emails[0][3] or "")
            check("Email body has substantial content (>= 100 chars)",
                  len(body) >= 100,
                  f"Body length: {len(body)}")
            # Body must mention key findings: quiz/score/pass/performance terms.
            body_lower = body.lower()
            has_quiz_finding = (
                ("quiz" in body_lower or "performance" in body_lower)
                and ("score" in body_lower or "pass" in body_lower or "average" in body_lower)
            )
            check("Email body mentions key findings (quiz/performance + score/pass/average)",
                  has_quiz_finding,
                  f"Sample body: {body[:200]}")
        cur.close()
        conn.close()
    except Exception as e:
        check("Email check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_gsheet()
    check_word(args.agent_workspace)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ===")
    if FAIL_COUNT > 0:
        print(f"{FAIL_COUNT} checks failed")
        sys.exit(1)
    else:
        print("All checks passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
