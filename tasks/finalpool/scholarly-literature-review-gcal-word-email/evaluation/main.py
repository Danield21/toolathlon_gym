"""Evaluation for scholarly-literature-review-gcal-word-email.

Checks:
1. Literature_Review.docx exists with title, at least 6 paper sections, and Summary section
2. GCal event "Literature Review Presentation" on 2026-03-21
3. Email to research-team@university.example.com with "Literature Review" in subject
"""
import argparse
import json
import os
import sys

import psycopg2

DB = {"host": os.environ.get("PGHOST", "localhost"), "port": int(os.environ.get("PGPORT", "5432")), "dbname": "toolathlon_gym", "user": "eigent", "password": "camel"}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300] + "...") if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def check_word(agent_ws):
    print("\n=== Check 1: Literature_Review.docx ===")
    path = os.path.join(agent_ws, "Literature_Review.docx")
    check("File Literature_Review.docx exists", os.path.isfile(path))
    if not os.path.isfile(path):
        return

    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        check("Word doc is readable", False, str(e))
        return

    full_text = " ".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()

    check("Doc contains 'Literature' and 'Review' in title/heading",
          "literature" in full_text_lower and "review" in full_text_lower)

    # Count paper entries - look for citation-related keywords (stricter list, not just any year)
    citation_indicators = ["citation", "citations", "year:", "authors:", "pub_year", "cited by"]
    has_citations = sum(1 for kw in citation_indicators if kw in full_text_lower)
    check("Doc contains citation data (years/authors)", has_citations >= 2,
          f"Citation indicators found: {has_citations}")

    # Check key paper titles appear (at least 4 of 5 - close to required 6 papers)
    key_papers = ["attention", "bert", "gpt", "llama", "transformer"]
    found_papers = sum(1 for p in key_papers if p in full_text_lower)
    check("Doc mentions at least 4 key LLM papers", found_papers >= 4,
          f"Found {found_papers}/5 key paper keywords")

    # Check for Summary/Conclusion section
    check("Doc has Summary or Conclusion section",
          "summary" in full_text_lower or "conclusion" in full_text_lower)

    # Count headings to estimate number of paper sections
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    check("Doc has at least 7 headings (title + 6+ paper sections + summary)",
          heading_count >= 7,
          f"Found {heading_count} headings")

    # Reject noise papers (non-LLM): "ImageNet" / "Krizhevsky" or "Survey of Deep Learning" / "LeCun"+"Hinton"+"Bengio" + 2012/2015
    # Make this strict: any noise paper title / unique author present should fail
    noise_indicators = ["imagenet", "krizhevsky", "convolutional neural network",
                        "lecun", "a survey of deep learning"]
    found_noise = [n for n in noise_indicators if n in full_text_lower]
    check("Doc excludes non-LLM noise papers (ImageNet/Survey of DL)",
          len(found_noise) == 0,
          f"Found noise indicators: {found_noise}")

    # Verify ALL 5 of the key LLM papers are mentioned (require 5/5 not 4/5)
    key_papers_strict = ["attention", "bert", "gpt", "llama", "transformer"]
    found_strict = [p for p in key_papers_strict if p in full_text_lower]
    check("Doc mentions all 5 key LLM papers (attention/bert/gpt/llama/transformer)",
          len(found_strict) == 5,
          f"Missing: {[p for p in key_papers_strict if p not in found_strict]}")


def check_gcal():
    print("\n=== Check 2: Google Calendar Event ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, summary, start_datetime, description FROM gcal.events
        WHERE summary ILIKE '%literature%review%'
           OR summary ILIKE '%literature review%'
           OR (summary ILIKE '%literature%' AND summary ILIKE '%review%')
        LIMIT 10
    """)
    events = cur.fetchall()
    check("GCal event with 'Literature Review' in title found",
          len(events) > 0, "No matching events in gcal.events")

    if events:
        # Find best matching event - prefer one on 2026-03-21
        target_event = None
        for ev in events:
            if ev[2] is not None and "2026-03-21" in str(ev[2]):
                target_event = ev
                break
        if target_event is None:
            target_event = events[0]

        evt_start = target_event[2]
        check("Event scheduled on 2026-03-21",
              evt_start is not None and "2026-03-21" in str(evt_start),
              f"Event start: {evt_start}")

        # Validate start time 14:00 and end time 15:30
        # also fetch end_datetime
        cur2 = conn.cursor()
        cur2.execute("SELECT start_datetime, end_datetime FROM gcal.events WHERE id=%s", (target_event[0],))
        row = cur2.fetchone()
        if row:
            start_dt = str(row[0]) if row[0] else ""
            end_dt = str(row[1]) if row[1] else ""
            # Tighten: anchor to common HH:MM:SS or HH:MM pattern boundaries
            import re as _re
            check("Event start time is exactly 14:00 UTC",
                  bool(_re.search(r"\b14:00(?::00)?\b", start_dt)),
                  f"start_dt: {start_dt}")
            check("Event end time is exactly 15:30 UTC",
                  bool(_re.search(r"\b15:30(?::00)?\b", end_dt)),
                  f"end_dt: {end_dt}")
            # Validate exact 90-minute duration
            try:
                from datetime import datetime as _dt
                # row[0]/row[1] may already be datetime objects
                sdt = row[0] if not isinstance(row[0], str) else _dt.fromisoformat(str(row[0]).replace("Z", "+00:00"))
                edt = row[1] if not isinstance(row[1], str) else _dt.fromisoformat(str(row[1]).replace("Z", "+00:00"))
                if sdt is not None and edt is not None:
                    dur_min = (edt - sdt).total_seconds() / 60.0
                    check("Event duration is exactly 90 minutes",
                          abs(dur_min - 90.0) < 0.5,
                          f"duration_min={dur_min}")
            except Exception as _exc:
                check("Event duration parseable", False, str(_exc))
        cur2.close()

        desc = str(target_event[3] or "").lower()
        # Per task.md: description must mention "LLM research" OR "large language models"
        check("Event description mentions LLM research or large language models",
              "llm" in desc or "large language" in desc,
              f"Description: {str(target_event[3])[:100]}")

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 3: Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    # Require BOTH subject contains 'Literature Review' AND to_addr contains 'research-team'
    cur.execute("""
        SELECT subject, to_addr, body_text FROM email.messages
        WHERE subject ILIKE '%%literature%%review%%'
          AND to_addr::text ILIKE '%%research-team%%'
        LIMIT 10
    """)
    rows = cur.fetchall()
    check("Email with 'Literature Review' in subject AND to research-team found",
          len(rows) > 0, "No matching email with subject 'Literature Review' AND to research-team")

    if rows:
        to_addrs = [str(r[1]) for r in rows]
        check("Email sent to research-team@university.example.com",
              any("research-team" in addr for addr in to_addrs),
              f"To addresses: {to_addrs}")
        bodies = [str(r[2] or "").lower() for r in rows]
        # Require: body mentions specific paper indicators (require at least 2 of 5 distinct paper keywords - exclude common substrings like 'brown')
        # 'brown' was too permissive; using paper-specific keywords only
        paper_keywords = ["attention", "bert", "gpt", "llama", "transformer", "vaswani", "devlin", "touvron"]
        match_count = sum(1 for kw in paper_keywords if any(kw in b for b in bodies))
        check("Email body mentions paper titles or authors (>=2 distinct)",
              match_count >= 2,
              f"Found {match_count} paper keywords; Body excerpt: {bodies[0][:200] if bodies else ''}")
        # Body should mention March 21 / 2026-03-21
        has_date = any("march 21" in b or "2026-03-21" in b or "march 21, 2026" in b for b in bodies)
        check("Email body mentions presentation date (March 21, 2026)", has_date,
              f"Body excerpt: {bodies[0][:200] if bodies else ''}")

    cur.close()
    conn.close()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=== Evaluation: scholarly-literature-review-gcal-word-email ===")

    check_word(args.agent_workspace)
    check_gcal()
    check_email()

    print(f"\n=== SUMMARY: {PASS_COUNT} passed, {FAIL_COUNT} failed ===")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT}, f)

    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
