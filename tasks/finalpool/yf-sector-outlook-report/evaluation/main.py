"""Evaluation for yf-sector-outlook-report."""
import argparse
import json
import os
import sys

import psycopg2

# Sentinel used when a cell is an Excel formula whose cached value is unavailable
# (e.g. written by openpyxl without a computed cache). Numeric/string checks on
# such cells are leniently skipped rather than treated as wrong answers.
_UNCOMPUTABLE = object()


def _to_float(v):
    """Robustly coerce a cell value to float.

    Supports int/float/str; strips thousands separators, currency symbols,
    percent signs and whitespace. Returns None when the value cannot be parsed
    (including None and formula strings).
    """
    if v is None:
        return None
    if isinstance(v, bool):
        return 1.0 if v else 0.0
    if isinstance(v, (int, float)):
        return float(v)
    if isinstance(v, str):
        s = v.strip()
        if not s or s.startswith("="):
            return None
        s = s.replace(",", "").replace("$", "").replace("¥", "").replace("€", "")
        s = s.replace("%", "").replace(" ", "").replace(" ", "")
        try:
            return float(s)
        except ValueError:
            return None
    return None


def num_close(a, b, abs_tol=1.0, rel_tol=0.05):
    """Numeric closeness with a lenient fallback for non-numeric cells.

    Formula cells with no cached value (_UNCOMPUTABLE) are always considered
    close (the check is skipped). If both sides parse as numbers, compare with
    tolerance. Otherwise fall back to a case-insensitive string comparison.
    """
    if a is _UNCOMPUTABLE or b is _UNCOMPUTABLE:
        return True
    a_f, b_f = _to_float(a), _to_float(b)
    if a_f is not None and b_f is not None:
        return abs(a_f - b_f) <= max(abs_tol, abs(b_f) * rel_tol)
    return str(a).strip().lower() == str(b).strip().lower()


def str_eq(a, b):
    """Case-insensitive string equality with formula-no-cache skip."""
    if a is _UNCOMPUTABLE or b is _UNCOMPUTABLE:
        return True
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def load_sheet_rows_resolved(wb_raw, wb_val, sheet_name):
    """Load a sheet's rows from a data_only=False workbook, resolving formulas.

    Cells that contain a formula (value starts with '=') are replaced by their
    cached value taken from the data_only=True workbook when available; when the
    cached value is missing they become the _UNCOMPUTABLE sentinel so the check
    can be leniently skipped.
    """
    raw_rows = None
    for name in wb_raw.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            raw_rows = [[cell.value for cell in row] for row in wb_raw[name].iter_rows()]
            break
    if raw_rows is None:
        return None

    cache = {}
    for name in wb_val.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            for row in wb_val[name].iter_rows():
                for c in row:
                    cache[(c.row, c.column)] = c.value
            break

    out = []
    for r_i, row in enumerate(raw_rows, start=1):
        nr = []
        for c_i, v in enumerate(row, start=1):
            if isinstance(v, str) and v.strip().startswith("="):
                cv = cache.get((r_i, c_i))
                nr.append(cv if cv is not None else _UNCOMPUTABLE)
            else:
                nr.append(v)
        out.append(nr)
    return out


def _agent_data_rows(rows):
    """Return the data rows of a sheet, skipping a detectable header row.

    A header row is the first non-empty row whose leading cell is 'stock',
    'metric' or 'sector', or that contains a cell equal to 'stock' together
    with a 'sector'/'metric' cell. If no header row is detected, every
    non-empty row is treated as a data row so that models which omit a header
    row (writing the data directly) are not penalised. This also tolerates a
    title row placed above the header.
    """
    if not rows:
        return []
    header_idx = None
    for i, row in enumerate(rows):
        cells = [str(c).strip().lower() for c in row if c is not None]
        if not cells:
            continue
        if len(cells) >= 2 and cells[0] in ("stock", "metric", "sector"):
            header_idx = i
            break
        if len(cells) >= 2 and "stock" in cells and ("sector" in cells or "metric" in cells):
            header_idx = i
            break
    start = 0 if header_idx is None else header_idx + 1
    return [r for r in rows[start:] if r and any(c is not None for c in r)]


def check_excel(agent_workspace, groundtruth_workspace=None):
    errors = []
    import openpyxl
    path = os.path.join(agent_workspace, "Sector_Outlook.xlsx")
    if not os.path.exists(path):
        return ["Sector_Outlook.xlsx not found"]

    # Load groundtruth values (preferred) for comparison
    gt_path = None
    if groundtruth_workspace:
        cand = os.path.join(groundtruth_workspace, "Sector_Outlook.xlsx")
        if os.path.exists(cand):
            gt_path = cand
    gt_perf_lookup = {}  # stock -> (sector, current_price, price_1y_ago, return_1y_pct, outlook, growth_forecast, risk_level)
    gt_summary = {}
    if gt_path:
        try:
            # GT is authored with literal values; read raw so numbers are literal
            gt_wb = openpyxl.load_workbook(gt_path, data_only=False)
            gt_perf_rows = load_sheet_rows(gt_wb, "Sector Performance") or []
            for r in gt_perf_rows[1:]:
                if r and len(r) > 7 and r[1]:
                    gt_perf_lookup[str(r[1]).strip().upper()] = tuple(r)
            gt_summary_rows = load_sheet_rows(gt_wb, "Cross-Sector Summary") or []
            for r in gt_summary_rows[1:]:
                if r and r[0] is not None:
                    gt_summary[str(r[0]).strip().lower()] = r[1]
        except Exception as e:
            errors.append(f"Error loading GT excel: {e}")
            return errors

    try:
        wb_raw = openpyxl.load_workbook(path, data_only=False)
        wb_val = openpyxl.load_workbook(path, data_only=True)

        rows = load_sheet_rows_resolved(wb_raw, wb_val, "Sector Performance")
        if rows is None:
            errors.append("Sheet 'Sector Performance' not found")
        else:
            data_rows = _agent_data_rows(rows)
            if len(data_rows) < 5:
                errors.append(f"Sector Performance has {len(data_rows)} rows, expected 5")
            # Check stocks present and validate each per groundtruth
            agent_lookup = {}
            for r in data_rows:
                if len(r) > 1 and r[1]:
                    agent_lookup[str(r[1]).strip().upper()] = r

            for sym in ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]:
                if sym not in agent_lookup:
                    errors.append(f"Stock {sym} missing from Sector Performance")
                    continue
                a_row = agent_lookup[sym]
                gt_row = gt_perf_lookup.get(sym)
                if gt_row is None:
                    continue
                # Sector (col 0)
                if not str_eq(a_row[0], gt_row[0]):
                    errors.append(f"{sym} Sector: '{a_row[0]}' vs expected '{gt_row[0]}'")
                # Current_Price (col 2)
                if not num_close(a_row[2], gt_row[2], abs_tol=1.0, rel_tol=0.02):
                    errors.append(f"{sym} Current_Price={a_row[2]} vs expected {gt_row[2]}")
                # Return_1Y_Pct (col 4)
                # Tolerance intentionally generous: absorbs the "closest trading
                # day ~1 year back / ~252 trading days" interpretations while
                # still rejecting a wall-clock-anchored 1y window (which would
                # drift by many percentage points given the frozen dataset).
                if not num_close(a_row[4], gt_row[4], abs_tol=3.0, rel_tol=0.08):
                    errors.append(f"{sym} Return_1Y_Pct={a_row[4]} vs expected {gt_row[4]}")
                # Outlook (col 5)
                if not str_eq(a_row[5], gt_row[5]):
                    errors.append(f"{sym} Outlook: '{a_row[5]}' vs expected '{gt_row[5]}'")
                # Growth_Forecast (col 6)
                if not num_close(a_row[6], gt_row[6], abs_tol=0.5):
                    errors.append(f"{sym} Growth_Forecast={a_row[6]} vs expected {gt_row[6]}")
                # Risk_Level (col 7)
                if not str_eq(a_row[7], gt_row[7]):
                    errors.append(f"{sym} Risk_Level: '{a_row[7]}' vs expected '{gt_row[7]}'")

        rows2 = load_sheet_rows_resolved(wb_raw, wb_val, "Cross-Sector Summary")
        if rows2 is None:
            errors.append("Sheet 'Cross-Sector Summary' not found")
        else:
            data_rows2 = [r for r in _agent_data_rows(rows2) if r and r[0] is not None]
            lookup = {str(r[0]).strip().lower(): r[1] for r in data_rows2 if r[0]}
            # Use groundtruth values as canonical
            for key in ("best_1y_sector", "worst_1y_sector"):
                if key in gt_summary:
                    if key not in lookup:
                        errors.append(f"Cross-Sector Summary missing {key}")
                    elif not str_eq(lookup[key], gt_summary[key]):
                        errors.append(
                            f"{key}: '{lookup[key]}' vs expected '{gt_summary[key]}'"
                        )
            for key in ("positive_outlook_count", "high_risk_count"):
                if key in gt_summary:
                    if key not in lookup:
                        errors.append(f"Cross-Sector Summary missing {key}")
                    elif not num_close(lookup[key], gt_summary[key], abs_tol=0):
                        errors.append(
                            f"{key}={lookup[key]} vs expected {gt_summary[key]}"
                        )
            if "avg_1y_return" in gt_summary:
                if "avg_1y_return" not in lookup:
                    errors.append("Cross-Sector Summary missing avg_1y_return")
                elif not num_close(lookup["avg_1y_return"], gt_summary["avg_1y_return"], abs_tol=3.0, rel_tol=0.05):
                    errors.append(
                        f"Avg_1Y_Return={lookup['avg_1y_return']} vs expected {gt_summary['avg_1y_return']}"
                    )

    except Exception as e:
        errors.append(f"Error reading Excel: {e}")
    return errors


def _doc_text(doc):
    """Concatenate all text in a Word document, including table cells."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# Company-name aliases accepted in the narrative report in place of the literal
# ticker symbol. A report may identify the representative stock by either its
# ticker (GOOGL/AMZN/JPM/JNJ/XOM) or its company name; both are legitimate ways
# to complete the task, so the check accepts either.
_STOCK_COMPANY_NAMES = {
    "GOOGL": ["alphabet", "google"],
    "AMZN": ["amazon"],
    "JPM": ["jpmorgan", "jp morgan", "j.p. morgan", "morgan chase"],
    "JNJ": ["johnson"],
    "XOM": ["exxon"],
}


def _mentioned_stock(sym, text_lower, text_upper):
    """True if a stock is discussed by ticker or by company name."""
    if sym in text_upper:
        return True
    return any(n in text_lower for n in _STOCK_COMPANY_NAMES.get(sym, []))


def check_word(agent_workspace):
    errors = []
    path = os.path.join(agent_workspace, "Sector_Report.docx")
    if not os.path.exists(path):
        return ["Sector_Report.docx not found"]
    try:
        from docx import Document
        doc = Document(path)
        text = _doc_text(doc).lower()
        text_upper = _doc_text(doc).upper()
        if len(text) < 400:
            errors.append(f"Sector_Report.docx too short ({len(text)} chars, need >=400)")
        # All 4 sectors must be discussed
        for kw in ["technology", "healthcare", "energy", "financial"]:
            if kw not in text:
                errors.append(f"Sector_Report.docx missing sector '{kw}'")
        # All 5 stocks must be discussed, by ticker OR company name
        for sym in ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]:
            if not _mentioned_stock(sym, text, text_upper):
                errors.append(f"Sector_Report.docx missing stock '{sym}' (ticker or company name)")
        # Document should include at least one numeric return value
        # (a decimal number, or any number followed by a % sign)
        import re
        if not re.search(r"[-]?\d+(?:\.\d+)?\s*%|[-]?\d+\.\d+", text):
            errors.append("Sector_Report.docx missing numeric return values")
        # Conclusion / synthesis section must appear
        if not any(kw in text for kw in ("recommend", "allocation", "conclu", "synth")):
            errors.append("Sector_Report.docx missing recommendation/conclusion section")
    except Exception as e:
        errors.append(f"Error reading Word doc: {e}")
    return errors


def check_notion():
    errors = []
    try:
        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"),
            port=int(os.environ.get("PGPORT", "5432")),
            dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
            user=os.environ.get("PGUSER", "eigent"),
            password=os.environ.get("PGPASSWORD", "camel"),
        )
        cur = conn.cursor()
        # Find pages whose title contains 'Sector Outlook' AND 'Q1 2026';
        # pick the most recently created one to be robust to duplicate pages
        # created by parallel sub-agents.
        cur.execute(
            """
            SELECT id, properties FROM notion.pages
            WHERE LOWER(properties::text) LIKE '%sector outlook%'
              AND LOWER(properties::text) LIKE '%q1 2026%'
            ORDER BY created_time DESC
            LIMIT 1
            """
        )
        rows = cur.fetchall()
        if len(rows) == 0:
            errors.append("Notion: no page titled 'Sector Outlook Report - Q1 2026' found")
        else:
            page_id = rows[0][0]
            cur.execute("""
                SELECT block_data FROM notion.blocks
                WHERE parent_id = %s
            """, (page_id,))
            blocks = cur.fetchall()
            body_text = " ".join(str(b[0]) for b in blocks).lower()
            # Check at least 2 sectors are mentioned in page body
            sectors_found = sum(1 for s in ("technology", "healthcare", "energy", "financial") if s in body_text)
            if sectors_found < 2:
                errors.append(f"Notion 'Sector Outlook' page body discusses only {sectors_found} sectors (need >=2)")
        cur.close()
        conn.close()
    except Exception as e:
        errors.append(f"Error checking Notion: {e}")
    return errors


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()
    agent_ws = args.agent_workspace or os.path.join(os.path.dirname(__file__), "..", "groundtruth_workspace")
    gt_ws = args.groundtruth_workspace or os.path.join(os.path.dirname(__file__), "..", "groundtruth_workspace")

    all_errors = []

    print("  Checking Excel file...")
    errs = check_excel(agent_ws, gt_ws)
    if errs:
        all_errors.extend(errs)
        for e in errs[:5]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    print("  Checking Word document...")
    errs = check_word(agent_ws)
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    print("  Checking Notion page...")
    errs = check_notion()
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
