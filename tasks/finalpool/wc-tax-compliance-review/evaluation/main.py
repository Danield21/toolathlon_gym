"""Evaluation for wc-tax-compliance-review."""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0

# Official rates from mock website (state base rates only).
# Kept as a fallback if the mock HTML cannot be parsed.
FALLBACK_OFFICIAL_RATES = {
    "CA": 7.25,
    "FL": 6.50,
    "NY": 8.875,
    "TX": 6.25,
    "WA": 6.50,
}

COMPLIANCE_THRESHOLD = 0.25


def load_official_rates_from_mock():
    """Parse the mock tax-authority website (localhost:30205 or the extracted
    tmp/mock_pages/index.html file) and return a dict of state_code -> base_rate.
    Falls back to FALLBACK_OFFICIAL_RATES if parsing fails."""
    import re

    html = None
    # Prefer live HTTP (most accurate to what the agent sees)
    try:
        import urllib.request
        with urllib.request.urlopen("http://localhost:30205/", timeout=2) as resp:
            html = resp.read().decode("utf-8", errors="ignore")
    except Exception:
        pass
    if not html:
        # Fall back to extracted file on disk
        task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        candidates = [
            os.path.join(task_root, "tmp", "mock_pages", "index.html"),
        ]
        for p in candidates:
            if os.path.isfile(p):
                try:
                    with open(p, "r", encoding="utf-8", errors="ignore") as f:
                        html = f.read()
                    break
                except Exception:
                    continue
    if not html:
        return dict(FALLBACK_OFFICIAL_RATES)

    # Extract <td> cells sequentially; then parse consecutive pairs (STATE, RATE)
    cells = re.findall(r"<td[^>]*>\s*([^<]+?)\s*</td>", html)
    rates = {}
    i = 0
    while i < len(cells) - 1:
        cell = cells[i].strip()
        if re.fullmatch(r"[A-Z]{2}", cell):
            # Next numeric cell within window of 2
            for j in (1, 2):
                if i + j < len(cells):
                    nxt = cells[i + j].strip().replace("%", "")
                    try:
                        rates[cell] = float(nxt)
                        break
                    except ValueError:
                        continue
        i += 1
    return rates or dict(FALLBACK_OFFICIAL_RATES)


OFFICIAL_RATES = load_official_rates_from_mock()


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=0.01):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def safe_str(v):
    return str(v).strip() if v is not None else ""


def get_wc_state_rates():
    """Get state-level tax rates from WC database."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT state, name, rate::numeric
        FROM wc.tax_rates
        WHERE country = 'US' AND state != '' AND class = 'standard'
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    # Keep only one rate per state (the state-level one, not district)
    state_rates = {}
    for state, name, rate in rows:
        # Use the one with "State Tax" in name, or the lower rate for a state
        if state not in state_rates:
            state_rates[state] = (name, float(rate))
        else:
            # Prefer the one with "State" in the name
            if "state" in name.lower():
                state_rates[state] = (name, float(rate))
    return state_rates


def compute_expected():
    """Compute expected comparison data."""
    wc_rates = get_wc_state_rates()
    results = []
    for state in sorted(wc_rates.keys()):
        if state in OFFICIAL_RATES:
            name, wc_rate = wc_rates[state]
            official = OFFICIAL_RATES[state]
            diff = round(abs(wc_rate - official), 4)
            compliant = "Yes" if diff <= COMPLIANCE_THRESHOLD else "No"
            results.append((state, name, wc_rate, official, diff, compliant))
    return results


def check_excel(agent_workspace):
    """Check Tax_Compliance.xlsx."""
    print("\n=== Checking Tax_Compliance.xlsx ===")

    excel_path = os.path.join(agent_workspace, "Tax_Compliance.xlsx")
    if not os.path.isfile(excel_path):
        record("Excel file exists", False, f"Not found: {excel_path}")
        return False
    record("Excel file exists", True)

    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True)
    except Exception as e:
        record("Excel readable", False, str(e))
        return False

    expected = compute_expected()
    all_ok = True

    # Check Rate Comparison sheet
    rc_sheet = None
    for name in wb.sheetnames:
        if "rate" in name.lower() and "compar" in name.lower():
            rc_sheet = wb[name]
            break
    if rc_sheet is None:
        record("Sheet 'Rate Comparison' exists", False, f"Sheets: {wb.sheetnames}")
        all_ok = False
    else:
        record("Sheet 'Rate Comparison' exists", True)

        # Check headers
        headers = [safe_str(rc_sheet.cell(1, c).value).lower() for c in range(1, 8)]
        record("Has State column", any("state" in h for h in headers))
        record("Has WC_Rate column", any("wc" in h and "rate" in h for h in headers))
        record("Has Official_Rate column", any("official" in h for h in headers))
        record("Has Difference column", any("diff" in h for h in headers))
        record("Has Compliant column", any("compli" in h for h in headers))

        rows = list(rc_sheet.iter_rows(min_row=2, values_only=True))
        record("Rate Comparison has correct row count",
               len(rows) >= len(expected),
               f"Expected {len(expected)}, got {len(rows)}")
        # Optional: verify states are sorted alphabetically (common convention)
        state_col = [safe_str(r[0]).upper() for r in rows if r and r[0]]
        sorted_ok = state_col == sorted(state_col)
        record("Rate Comparison states sorted alphabetically", sorted_ok,
               f"Got order: {state_col}")

        for exp_state, exp_name, exp_wc, exp_off, exp_diff, exp_compl in expected:
            found = False
            for r in rows:
                if r and r[0] and safe_str(r[0]).upper() == exp_state:
                    found = True
                    ok_wc = num_close(r[2], exp_wc, 0.01)
                    record(f"{exp_state} WC_Rate={exp_wc}", ok_wc,
                           f"Got {r[2]}")
                    if not ok_wc:
                        all_ok = False

                    ok_off = num_close(r[3], exp_off, 0.01)
                    record(f"{exp_state} Official_Rate={exp_off}", ok_off,
                           f"Got {r[3]}")
                    if not ok_off:
                        all_ok = False

                    ok_diff = num_close(r[4], exp_diff, 0.01)
                    record(f"{exp_state} Difference={exp_diff}", ok_diff,
                           f"Got {r[4]}")
                    if not ok_diff:
                        all_ok = False

                    ok_compl = safe_str(r[5]).lower() == exp_compl.lower()
                    record(f"{exp_state} Compliant={exp_compl}", ok_compl,
                           f"Got {r[5]}")
                    if not ok_compl:
                        all_ok = False
                    break
            if not found:
                record(f"State {exp_state} found in Rate Comparison", False)
                all_ok = False

    # Check Discrepancies sheet
    disc_sheet = None
    for name in wb.sheetnames:
        if "discrep" in name.lower():
            disc_sheet = wb[name]
            break
    if disc_sheet is None:
        record("Sheet 'Discrepancies' exists", False, f"Sheets: {wb.sheetnames}")
        all_ok = False
    else:
        record("Sheet 'Discrepancies' exists", True)
        disc_rows = list(disc_sheet.iter_rows(min_row=2, values_only=True))
        non_compliant = [e for e in expected if e[5] == "No"]
        record("Discrepancies has correct row count",
               len(disc_rows) >= len(non_compliant),
               f"Expected {len(non_compliant)}, got {len(disc_rows)}")

        for exp_state, _, _, _, _, _ in non_compliant:
            found = any(r and r[0] and safe_str(r[0]).upper() == exp_state for r in disc_rows)
            record(f"Discrepancy for {exp_state} listed", found)
            if not found:
                all_ok = False

    # Check Summary sheet
    sum_sheet = None
    for name in wb.sheetnames:
        if "summ" in name.lower():
            sum_sheet = wb[name]
            break
    if sum_sheet is None:
        record("Sheet 'Summary' exists", False, f"Sheets: {wb.sheetnames}")
        all_ok = False
    else:
        record("Sheet 'Summary' exists", True)
        summary = {}
        for row in sum_sheet.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                summary[safe_str(row[0]).lower().replace(" ", "_")] = row[1]

        total_reviewed = len(expected)
        compliant_count = sum(1 for e in expected if e[5] == "Yes")
        non_compliant_count = sum(1 for e in expected if e[5] == "No")
        max_diff = max(e[4] for e in expected)

        for key, val in summary.items():
            if "total" in key and "rate" in key or "reviewed" in key:
                ok = num_close(val, total_reviewed, 1)
                record(f"Summary Total_Rates_Reviewed={total_reviewed}", ok, f"Got {val}")
                if not ok:
                    all_ok = False
            elif "compliant_count" in key or ("compliant" in key and "non" not in key and "count" in key):
                ok = num_close(val, compliant_count, 1)
                record(f"Summary Compliant_Count={compliant_count}", ok, f"Got {val}")
                if not ok:
                    all_ok = False
            elif "non_compliant" in key or "noncompliant" in key:
                ok = num_close(val, non_compliant_count, 1)
                record(f"Summary Non_Compliant_Count={non_compliant_count}", ok, f"Got {val}")
                if not ok:
                    all_ok = False
            elif "max" in key and "diff" in key:
                ok = num_close(val, max_diff, 0.01)
                record(f"Summary Max_Difference={max_diff}", ok, f"Got {val}")
                if not ok:
                    all_ok = False

    return all_ok


def check_word(agent_workspace):
    """Check Tax_Compliance_Report.docx."""
    print("\n=== Checking Tax_Compliance_Report.docx ===")
    from docx import Document

    docx_path = os.path.join(agent_workspace, "Tax_Compliance_Report.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, f"Not found: {docx_path}")
        return False
    record("Word file exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word readable", False, str(e))
        return False

    import re
    full_text = " ".join(p.text for p in doc.paragraphs)
    full_lower = full_text.lower()

    record("Mentions 'compliance' or 'compliant'",
           "complian" in full_lower)
    record("Mentions 'florida' or 'FL'",
           "florida" in full_lower or re.search(r"\bfl\b", full_text) is not None)
    record("Mentions 'new york' or 'NY'",
           "new york" in full_lower or re.search(r"\bny\b", full_text) is not None)
    record("Mentions 'discrepan' or 'non-compliant'",
           "discrepan" in full_lower or "non-compliant" in full_lower or "non compliant" in full_lower)
    record("Mentions 'recommend'",
           "recommend" in full_lower)
    record("Has at least 3 paragraphs", len(doc.paragraphs) >= 3,
           f"Got {len(doc.paragraphs)} paragraphs")

    # Verify that at least one specific rate value appears (e.g., '6.0' for FL or '4.0' for NY)
    try:
        expected_disc = [e for e in compute_expected() if e[5] == "No"]
        for st, _, wc_rate, off_rate, diff, _ in expected_disc:
            wc_str = f"{wc_rate:.2f}"
            off_str = f"{off_rate:.2f}"
            has_rate = wc_str in full_text or off_str in full_text or f"{wc_rate:.3f}" in full_text
            record(f"Word mentions rate value for {st}", has_rate,
                   f"Neither {wc_str} nor {off_str} found")
    except Exception as e:
        print(f"  WARN: could not load expected discrepancies: {e}")

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    excel_ok = check_excel(args.agent_workspace)
    word_ok = check_word(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = excel_ok and word_ok
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
