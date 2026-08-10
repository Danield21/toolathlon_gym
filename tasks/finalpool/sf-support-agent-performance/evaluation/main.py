"""Evaluation for sf-support-agent-performance."""
import argparse
import os
import sys
import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def fetch_expected_agents():
    """Fetch active agents and their ticket stats from DB."""
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute('''
            SELECT a."AGENT_NAME", a."TEAM", a."SKILL_LEVEL"
            FROM sf_data."SUPPORT_CENTER__PUBLIC__AGENTS" a
            WHERE a."ACTIVE" = TRUE
            ORDER BY a."AGENT_NAME"
        ''')
        agents = []
        for name, team, skill in cur.fetchall():
            cur.execute('''
                SELECT COUNT(*)::int, AVG("CUSTOMER_SATISFACTION")::float
                FROM sf_data."SUPPORT_CENTER__PUBLIC__TICKETS"
                WHERE "RESOLVER" = %s
            ''', (name,))
            n, avg = cur.fetchone()
            agents.append({
                "name": name, "team": team, "skill": skill,
                "tickets": int(n or 0),
                "avg_sat": round(float(avg), 2) if avg is not None else None,
            })
        cur.close()
        conn.close()
        return agents
    except Exception as e:
        print(f"[fallback] DB query failed: {e}")
        return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Support_Agent_Performance.xlsx")
    gt_file = os.path.join(gt_dir, "Support_Agent_Performance.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    expected_agents = fetch_expected_agents()
    all_errors = []

    # === Agent Performance sheet ===
    print(f"  Checking Agent Performance...")
    a_rows = load_sheet_rows(agent_wb, "Agent Performance")
    g_rows = load_sheet_rows(gt_wb, "Agent Performance")
    if a_rows is None:
        all_errors.append("Sheet 'Agent Performance' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Agent Performance' not found in groundtruth")
    else:
        # Validate header
        if a_rows and a_rows[0]:
            a_header = [str(h or "").strip().lower().replace(" ", "_") for h in a_rows[0]]
            expected_headers = ["agent_name", "team", "skill_level", "tickets_handled", "avg_satisfaction"]
            for i, eh in enumerate(expected_headers):
                ah = a_header[i] if i < len(a_header) else "MISSING"
                if ah != eh:
                    all_errors.append(f"AgentPerformance header[{i}]: expected '{eh}' got '{ah}'")

        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_data = [r for r in a_data if r and r[0] is not None]
        g_data = [r for r in g_data if r and r[0] is not None]

        # Row count must equal active agent count
        if expected_agents:
            if len(a_data) != len(expected_agents):
                all_errors.append(f"Agent Performance row count: {len(a_data)} vs expected {len(expected_agents)}")
        else:
            if len(a_data) != len(g_data):
                all_errors.append(f"Agent Performance row count: {len(a_data)} vs GT {len(g_data)}")

        # Build expected lookup
        expected_lookup = {}
        if expected_agents:
            for ag in expected_agents:
                expected_lookup[ag["name"].lower()] = ag

        a_lookup = {}
        for row in a_data:
            a_lookup[str(row[0]).strip().lower()] = row

        # Check each GT agent has team/skill matching reality
        for g_row in g_data:
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                all_errors.append(f"Missing agent: {g_row[0]}")
                continue

            # Team (col 1)
            if expected_agents:
                exp = expected_lookup.get(key)
                if exp:
                    if not str_match(a_row[1], exp["team"]):
                        all_errors.append(f"{key}.Team: '{a_row[1]}' vs expected '{exp['team']}'")
                    if not str_match(a_row[2], exp["skill"]):
                        all_errors.append(f"{key}.Skill_Level: '{a_row[2]}' vs expected '{exp['skill']}'")
                    if not num_close(a_row[3], exp["tickets"], 0):
                        all_errors.append(f"{key}.Tickets_Handled: {a_row[3]} vs expected {exp['tickets']}")
                    # Avg_Satisfaction tol=0.05 (rounded to 2 decimals)
                    if not num_close(a_row[4], exp["avg_sat"], 0.05):
                        all_errors.append(f"{key}.Avg_Satisfaction: {a_row[4]} vs expected {exp['avg_sat']}")
            else:
                # Fall back to GT
                if not str_match(a_row[1], g_row[1]):
                    all_errors.append(f"{key}.Team: '{a_row[1]}' vs GT '{g_row[1]}'")
                if not str_match(a_row[2], g_row[2]):
                    all_errors.append(f"{key}.Skill_Level: '{a_row[2]}' vs GT '{g_row[2]}'")
                if not num_close(a_row[3], g_row[3], 0):
                    all_errors.append(f"{key}.Tickets_Handled: {a_row[3]} vs {g_row[3]}")
                if not num_close(a_row[4], g_row[4], 0.05):
                    all_errors.append(f"{key}.Avg_Satisfaction: {a_row[4]} vs {g_row[4]}")

        # Sort: Tickets_Handled descending
        try:
            tickets = [int(r[3]) if r[3] is not None else 0 for r in a_data]
            if tickets != sorted(tickets, reverse=True):
                all_errors.append(f"Agent Performance not sorted descending by Tickets_Handled: {tickets}")
        except (TypeError, ValueError):
            all_errors.append("Agent Performance Tickets_Handled non-numeric")

    # === Summary sheet ===
    print(f"  Checking Summary...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    g_rows = load_sheet_rows(gt_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Summary' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []

        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row[1] if len(row) > 1 else None
        g_lookup = {}
        for row in g_data:
            if row and row[0] is not None:
                g_lookup[str(row[0]).strip().lower()] = row[1] if len(row) > 1 else None

        # Compute expected summary from DB
        expected_total_active = 0
        expected_total_tickets = 0
        expected_top_agents = []
        expected_avg_sat = None
        if expected_agents:
            expected_total_active = len(expected_agents)
            expected_total_tickets = sum(a["tickets"] for a in expected_agents)
            # Top by tickets
            top_n = max(a["tickets"] for a in expected_agents)
            expected_top_agents = [a["name"] for a in expected_agents if a["tickets"] == top_n]
            # Weighted avg satisfaction
            num = den = 0
            for a in expected_agents:
                if a["avg_sat"] is not None and a["tickets"] > 0:
                    num += a["avg_sat"] * a["tickets"]
                    den += a["tickets"]
            if den > 0:
                expected_avg_sat = round(num / den, 2)
            else:
                expected_avg_sat = 0

        # Check Total_Active_Agents (tol 0)
        v = a_lookup.get("total_active_agents")
        if expected_agents:
            if v is None or not num_close(v, expected_total_active, 0):
                all_errors.append(f"Total_Active_Agents: {v} vs expected {expected_total_active}")
        else:
            gv = g_lookup.get("total_active_agents")
            if v is None or not num_close(v, gv, 0):
                all_errors.append(f"Total_Active_Agents: {v} vs GT {gv}")

        # Total_Tickets (tol 0)
        v = a_lookup.get("total_tickets")
        if expected_agents:
            if v is None or not num_close(v, expected_total_tickets, 0):
                all_errors.append(f"Total_Tickets: {v} vs expected {expected_total_tickets}")
        else:
            gv = g_lookup.get("total_tickets")
            if v is None or not num_close(v, gv, 0):
                all_errors.append(f"Total_Tickets: {v} vs GT {gv}")

        # Top_Agent: tied set (when all 0)
        v = a_lookup.get("top_agent")
        if expected_top_agents:
            ok = v is not None and any(str(v).strip().lower() == n.lower() for n in expected_top_agents)
            if not ok:
                all_errors.append(f"Top_Agent: '{v}' not in tied top set {expected_top_agents}")
        else:
            gv = g_lookup.get("top_agent")
            if not str_match(v, gv):
                all_errors.append(f"Top_Agent: '{v}' vs GT '{gv}'")

        # Avg_Satisfaction_Overall (tol 0.05)
        v = a_lookup.get("avg_satisfaction_overall")
        if expected_avg_sat is not None:
            if v is None or not num_close(v, expected_avg_sat, 0.05):
                all_errors.append(f"Avg_Satisfaction_Overall: {v} vs expected {expected_avg_sat}")
        else:
            gv = g_lookup.get("avg_satisfaction_overall")
            if v is None or not num_close(v, gv, 0.05):
                all_errors.append(f"Avg_Satisfaction_Overall: {v} vs GT {gv}")

    # === Word doc check ===
    docx_path = os.path.join(args.agent_workspace, "Agent_Review.docx")
    if not os.path.exists(docx_path):
        all_errors.append("Agent_Review.docx not found")
    else:
        try:
            from docx import Document
            doc = Document(docx_path)
            parts = [p.text for p in doc.paragraphs]
            # Include table cell text so agent names written into tables count.
            for _tbl in doc.tables:
                for _row in _tbl.rows:
                    for _cell in _row.cells:
                        parts.append(_cell.text)
            text = " ".join(parts)
            text_lower = text.lower()
            if len(text.strip()) < 200:
                all_errors.append(f"Agent_Review.docx too short ({len(text)} chars, need >=200)")
            # Required keywords
            for kw in ["agent", "performance"]:
                if kw not in text_lower:
                    all_errors.append(f"Agent_Review.docx missing keyword '{kw}'")
            # At least 3 agent names mentioned
            if expected_agents:
                names_mentioned = sum(1 for a in expected_agents if a["name"] in text)
                if names_mentioned < 3:
                    all_errors.append(f"Agent_Review.docx mentions only {names_mentioned} agents (need >=3)")

                # The "highlights" must name a real top performer (most tickets),
                # and the "needing attention" section must name a real low
                # performer — generic keywords alone are not sufficient. The
                # attention set accepts lowest-satisfaction OR fewest-tickets
                # agents, since "needing attention" can be read either way.
                max_tk = max(a["tickets"] for a in expected_agents)
                top_names = {a["name"] for a in expected_agents if a["tickets"] == max_tk}
                attn_names = set()
                sat_agents = [a for a in expected_agents
                              if a["avg_sat"] is not None and a["tickets"] > 0]
                if sat_agents:
                    min_sat = min(a["avg_sat"] for a in sat_agents)
                    attn_names |= {a["name"] for a in sat_agents
                                   if a["avg_sat"] == min_sat}
                min_tk = min(a["tickets"] for a in expected_agents)
                attn_names |= {a["name"] for a in expected_agents
                               if a["tickets"] == min_tk}
                if not any(n.lower() in text_lower for n in top_names):
                    all_errors.append(
                        f"Agent_Review.docx does not name a top performer "
                        f"(one of {sorted(top_names)})")
                if not any(n.lower() in text_lower for n in attn_names):
                    all_errors.append(
                        f"Agent_Review.docx does not name a low performer "
                        f"needing attention (one of {sorted(attn_names)})")
            # Highlights / attention section — drop 'performance' (already required globally)
            has_highlights = any(k in text_lower for k in ["highlight", "top", "best"])
            has_attention = any(k in text_lower for k in ["attention", "needs", "concern", "improve"])
            if not has_highlights:
                all_errors.append("Agent_Review.docx missing 'highlight'/'top'/'best' content")
            if not has_attention:
                all_errors.append("Agent_Review.docx missing 'attention'/'needs'/'concern'/'improve' content")
        except Exception as e:
            all_errors.append(f"Error reading Agent_Review.docx: {e}")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
