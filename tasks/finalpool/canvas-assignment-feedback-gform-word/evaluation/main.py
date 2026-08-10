"""Evaluation for canvas-assignment-feedback-gform-word."""
import argparse
import email.utils
import json
import os
import re
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"),
          port=int(os.environ.get("PGPORT", "5432")),
          dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
          user=os.environ.get("PGUSER", "eigent"),
          password=os.environ.get("PGPASSWORD", "camel"))

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def _to_float(s):
    """Robust numeric parse: handles int/float/None, strips %, currency symbols,
    thousands separators and whitespace. Returns None when unparseable."""
    if s is None:
        return None
    if isinstance(s, bool):
        return None
    if isinstance(s, (int, float)):
        return float(s)
    t = str(s).strip()
    if not t:
        return None
    t = re.sub(r'[^\d.+\-]', '', t.replace(',', ''))
    if not t:
        return None
    try:
        return float(t)
    except ValueError:
        return None


def _to_int(s):
    f = _to_float(s)
    return int(f) if f is not None else None


def num_close(a, b, tol):
    fa = _to_float(a)
    fb = _to_float(b)
    if fa is None or fb is None:
        return False
    return abs(fa - fb) <= tol


def _load_email_config():
    """Derive the expected From address from this task's email_config.json (R7)."""
    here = os.path.dirname(os.path.abspath(__file__))
    task_dir = os.path.dirname(here)  # evaluation/ -> task root
    cfg_path = os.path.join(task_dir, "email_config.json")
    email = "analytics@university.example.com"
    if os.path.isfile(cfg_path):
        try:
            with open(cfg_path) as f:
                cfg = json.load(f)
            email = cfg.get("email") or email
        except Exception:
            pass
    return email


def _from_matches(from_addr, expected_email):
    """True when the stored From header carries the expected email address,
    tolerating a display name prefix (e.g. 'Analytics Team <...>')."""
    if not from_addr:
        return False
    s = str(from_addr).strip()
    if not s:
        return False
    _, addr = email.utils.parseaddr(s)
    if addr and addr.lower() == expected_email.lower():
        return True
    if expected_email.lower() in s.lower():
        return True
    return False


def _norm_alpha(s):
    """Lowercase and strip all non-alphanumeric chars, for tolerant name/title
    matching (handles 'CMA 34879' vs 'CMA34879', en/em dashes, etc.)."""
    return re.sub(r'[^a-z0-9]', '', (s or "").lower())


def _extract_address(addr):
    """Return the bare email address from a possibly display-name-qualified
    string (e.g. 'Instructor <instructor@...>' -> 'instructor@...')."""
    if not addr:
        return ""
    s = str(addr).strip()
    if not s:
        return ""
    _, parsed = email.utils.parseaddr(s)
    return (parsed or s).strip()


def _recipient_in(to_addr, target):
    """True when any recipient element in to_addr matches target address,
    tolerating display-name forms like 'Instructor <instructor@...>'."""
    t = target.lower()
    if to_addr is None:
        return False
    elems = []
    if isinstance(to_addr, list):
        elems = [str(r) for r in to_addr]
    elif isinstance(to_addr, str):
        try:
            parsed = json.loads(to_addr)
            if isinstance(parsed, list):
                elems = [str(r) for r in parsed]
            else:
                elems = [to_addr]
        except (json.JSONDecodeError, TypeError):
            elems = [to_addr]
    for e in elems:
        a = _extract_address(e).lower()
        if a == t or t in a:
            return True
    return False


_NUM_WORDS = {
    0: "zero", 1: "one", 2: "two", 3: "three", 4: "four", 5: "five", 6: "six",
    7: "seven", 8: "eight", 9: "nine", 10: "ten", 11: "eleven", 12: "twelve",
    13: "thirteen", 14: "fourteen", 15: "fifteen", 16: "sixteen", 17: "seventeen",
    18: "eighteen", 19: "nineteen", 20: "twenty", 30: "thirty", 40: "forty",
    50: "fifty", 60: "sixty", 70: "seventy", 80: "eighty", 90: "ninety",
}


def _num_word(n):
    """English word form for small integers (12 -> 'twelve'); None if unsupported."""
    if n in _NUM_WORDS:
        return _NUM_WORDS[n]
    if 21 <= n < 100:
        t = (n // 10) * 10
        o = n % 10
        return _NUM_WORDS[t] + ("-" + _NUM_WORDS[o] if o else "")
    return None


def _mentions_number(body, expected_val, tol=0.01):
    """True when body contains a number numerically close to expected_val."""
    if expected_val is None or not body:
        return False
    for m in re.finditer(r'\d+(?:[.,]\d+)?', str(body)):
        v = _to_float(m.group(0))
        if v is not None and abs(v - expected_val) <= tol:
            return True
    return False


def _fetch_expected_from_db():
    """Fetch expected assignment statistics from canvas DB."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        # Course 16, exclude assignments with 0 submissions (e.g. Final Exam never submitted)
        cur.execute("""
            SELECT a.name,
                   COUNT(s.id) AS total,
                   ROUND(AVG(s.score)::numeric, 2)::float AS avg_score,
                   SUM(CASE WHEN s.late=true THEN 1 ELSE 0 END)::int AS late,
                   ROUND(SUM(CASE WHEN s.late=true THEN 1 ELSE 0 END) * 100.0 / NULLIF(COUNT(s.id),0), 2)::float AS late_rate
            FROM canvas.assignments a
            LEFT JOIN canvas.submissions s ON s.assignment_id=a.id
            WHERE a.course_id=16
            GROUP BY a.name
            HAVING COUNT(s.id) > 0
            ORDER BY a.name
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        return [(r[0], int(r[1]), float(r[2]) if r[2] is not None else None,
                 int(r[3]), float(r[4]) if r[4] is not None else None) for r in rows]
    except Exception as e:
        print(f"[warn] DB fetch fallback: {e}")
        return None


# Fallback in case DB unavailable
EXPECTED_ASSIGNMENTS_FALLBACK = [
    ("CMA 34878", 1470, 83.32, 36, 2.45),
    ("CMA 34879", 1352, 87.97, 43, 3.18),
    ("CMA 34880", 1252, 76.50, 61, 4.87),
    ("CMA 34881", 1224, 78.30, 59, 4.82),
    ("CMA 34882", 1193, 78.65, 75, 6.29),
    ("CMA 34883", 1196, 76.80, 119, 9.95),
    ("CMA 34884", 1160, 76.16, 88, 7.59),
    ("TMA 34873", 1859, 78.16, 157, 8.45),
    ("TMA 34874", 1661, 72.45, 193, 11.62),
    ("TMA 34875", 1402, 70.49, 376, 26.82),
    ("TMA 34876", 1313, 71.05, 268, 20.41),
    ("TMA 34877", 1158, 75.89, 148, 12.78),
]


def _norm_header(h):
    return re.sub(r'[^a-z0-9]', '', (h or "").lower())


def check_word(agent_workspace, expected):
    """Check Word document output."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx available", False, "python-docx not installed")
        return

    agent_file = os.path.join(agent_workspace, "Assignment_Analysis.docx")
    if not os.path.isfile(agent_file):
        check("Word file exists", False, f"Expected {agent_file}")
        # Subordinate failures
        for label in ["Document title present", "Document has at least one table",
                      "Column 'Assignment_Name' present", "Column 'Total_Submissions' present",
                      "Column 'Avg_Score' present", "Column 'Late_Submissions' present",
                      "Column 'Late_Rate(%)' present",
                      "Table has expected number of data rows",
                      "All expected assignments present with correct numeric values",
                      "Rows are sorted alphabetically by Assignment_Name"]:
            check(label, False, "Word file missing")
        return None
    check("Word file exists", True)

    try:
        doc = Document(agent_file)
    except Exception as e:
        check("Word file readable", False, str(e))
        return None

    # Check title
    full_text = "\n".join(p.text for p in doc.paragraphs)
    title_ok = (
        "foundations of finance" in full_text.lower() and
        "fall 2013" in full_text.lower() and
        "assignment analysis" in full_text.lower()
    )
    check("Document title present", title_ok,
          f"Title not found in: {full_text[:200]}")

    # Find table
    if not doc.tables:
        check("Document has at least one table", False, "no table found")
        return None
    check("Document has at least one table", True)

    tbl = doc.tables[0]

    # Check headers (exact normalized match)
    if not tbl.rows:
        check("Table has rows", False, "Empty table")
        return None

    headers = [c.text.strip() for c in tbl.rows[0].cells]
    headers_norm = [_norm_header(h) for h in headers]

    expected_headers_norm = ["assignmentname", "totalsubmissions", "avgscore",
                             "latesubmissions", "laterate"]
    for col_norm, label in zip(expected_headers_norm,
                               ["Assignment_Name", "Total_Submissions", "Avg_Score",
                                "Late_Submissions", "Late_Rate(%)"]):
        present = any(h.startswith(col_norm) for h in headers_norm) or col_norm in headers_norm
        check(f"Column '{label}' present", present, f"Headers: {headers}")

    # Find column indices for header-based extraction
    def _col_index(target_norm_prefix):
        for i, h in enumerate(headers_norm):
            if h == target_norm_prefix or h.startswith(target_norm_prefix):
                return i
        return None

    col_name_idx = _col_index("assignmentname") or 0
    col_total_idx = _col_index("totalsubmissions")
    col_avg_idx = _col_index("avgscore")
    col_late_idx = _col_index("latesubmissions")
    col_rate_idx = _col_index("laterate")

    # Check rows
    data_rows = list(tbl.rows)[1:]
    expected_count = len(expected)
    # DB query filters HAVING COUNT > 0 (the Final Exam has 0 submissions). An agent
    # following "include all assignments" may still add that 0-submission row, so
    # require at least expected_count rows instead of an exact match.
    check("Table has expected number of data rows",
          len(data_rows) >= expected_count,
          f"Found {len(data_rows)} rows, expected at least {expected_count}")

    # Build row map by name
    expected_by_name = {a[0].strip().lower(): a for a in expected}
    matched_names = set()
    bad_values = []
    for row in data_rows:
        if not row.cells:
            continue
        name = row.cells[col_name_idx].text.strip()
        ekey = name.lower()
        if ekey not in expected_by_name:
            continue  # could be the optional zero-submission row
        e_name, e_total, e_avg, e_late, e_rate = expected_by_name[ekey]
        # Read agent values (robust parse of %, thousands separators, currency)
        a_total = _to_int(row.cells[col_total_idx].text) if col_total_idx is not None else None
        a_avg = _to_float(row.cells[col_avg_idx].text) if col_avg_idx is not None else None
        a_late = _to_int(row.cells[col_late_idx].text) if col_late_idx is not None else None
        a_rate = _to_float(row.cells[col_rate_idx].text) if col_rate_idx is not None else None

        ok = True
        if a_total is None or a_total != e_total:
            ok = False
        if e_avg is None:
            pass
        elif a_avg is None or not num_close(a_avg, e_avg, 0.05):
            ok = False
        if a_late is None or a_late != e_late:
            ok = False
        if e_rate is None:
            pass
        elif a_rate is None or not num_close(a_rate, e_rate, 0.05):
            ok = False
        if ok:
            matched_names.add(ekey)
        else:
            bad_values.append((name, (a_total, a_avg, a_late, a_rate),
                               (e_total, e_avg, e_late, e_rate)))

    # Set-based: every expected assignment must appear at least once with correct
    # values. Duplicate rows (e.g. from multi-agent writes) don't inflate the count.
    check("All expected assignments present with correct numeric values",
          len(matched_names) == expected_count,
          f"matched={len(matched_names)}/{expected_count}, bad={bad_values[:3]}")

    # Sort check (alphabetical by name) over the rows that carry an expected
    # assignment name. Empty trailing rows and non-data rows (common python-docx
    # artifacts, or the optional zero-submission row) are ignored so they cannot
    # break the ordering check.
    data_names = []
    for r in data_rows:
        if not r.cells:
            continue
        nm = r.cells[col_name_idx].text.strip()
        if nm and nm.lower() in expected_by_name:
            data_names.append(nm)
    sorted_data_names = sorted(data_names, key=lambda s: s.lower())
    check("Rows are sorted alphabetically by Assignment_Name",
          data_names == sorted_data_names,
          f"Order: {data_names[:5]}")

    return len(data_rows)


def check_gform():
    """Check Google Form creation."""
    print("\n=== Checking Google Form ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT id, title FROM gform.forms ORDER BY created_at, id")
        forms = cur.fetchall()
    except Exception as e:
        check("GForm DB query OK", False, str(e))
        return

    check("At least one form created", len(forms) >= 1, f"Found {len(forms)} forms")

    expected_title = "foundations of finance - assignment feedback survey"
    expected_norm = _norm_alpha(expected_title)
    # Tolerant title match (dashes/spacing/case-insensitive). An agent may have
    # created a same-titled abandoned form before the real one, so consider ALL
    # matching forms below.
    matching_forms = [(fid, title) for fid, title in forms
                      if _norm_alpha(title) == expected_norm]

    check("Form title matches 'Foundations of Finance - Assignment Feedback Survey' exactly",
          len(matching_forms) >= 1,
          f"Forms: {[t for _, t in forms]}")

    if matching_forms:
        # The delivered survey is complete if ANY matching form carries exactly
        # 5 questions; an abandoned/re-created incomplete form must not penalise
        # a correct agent that also delivered a complete form.
        q_ok = False
        for fid, _title in matching_forms:
            try:
                cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (fid,))
                q_count = cur.fetchone()[0]
            except Exception:
                q_count = 0
            if q_count == 5:
                q_ok = True
                break
        check("Form has exactly 5 questions", q_ok,
              f"Matched forms: {[t for _, t in matching_forms]}")

    try:
        cur.close()
        conn.close()
    except Exception:
        pass


def check_emails(expected, doc_rows=None):
    """Check that summary email was sent."""
    print("\n=== Checking Emails ===")
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT id, subject, from_addr, to_addr, body_text "
                    "FROM email.messages ORDER BY id")
        all_emails = cur.fetchall()
        conn.close()
    except Exception as e:
        check("Emails DB query OK", False, str(e))
        return

    target = "instructor@financeou.example.com"
    # Recipient match tolerates display-name forms like 'Instructor <instructor@...>'.
    emails_to = [e for e in all_emails if _recipient_in(e[3], target)]
    check(f"Summary email sent to {target}", len(emails_to) >= 1,
          f"Total emails found: {len(all_emails)}")

    # Compute expected highlights
    if expected:
        # highest avg score
        with_avg = [a for a in expected if a[2] is not None]
        top_avg = max(with_avg, key=lambda a: a[2]) if with_avg else None
        # most late
        top_late = max(expected, key=lambda a: a[3]) if expected else None
    else:
        top_avg = None
        top_late = None

    if not emails_to:
        for label in ["Email subject exact match", "Email from analytics",
                      "Email mentions count of assignments", "Email mentions highest-avg assignment",
                      "Email mentions most-late assignment", "Email contains form link"]:
            check(label, False, "no email")
        return

    expected_subject = "Assignment Analysis Report - Foundations of Finance (Fall 2013)"

    def _subject_matches(subj):
        return (subj or "").strip().lower() == expected_subject.lower()

    # An agent may have sent a test/retry email first; require that SOME email to
    # the recipient carries the exact expected subject (the summary report).
    subject_ok = any(_subject_matches(e[1]) for e in emails_to)
    check("Email subject exact match", subject_ok,
          f"Subjects: {[e[1] for e in emails_to]}")

    # Evaluate body/from checks on the summary email: prefer the latest email to
    # the recipient with the expected subject, else the latest email to the
    # recipient (rows are ordered by id).
    primary = None
    with_subject = [e for e in emails_to if _subject_matches(e[1])]
    if with_subject:
        primary = with_subject[-1]
    else:
        primary = emails_to[-1]

    _pid, subj, from_addr, to_addr, body = primary
    expected_from = _load_email_config()
    check(f"Email from {expected_from}",
          _from_matches(from_addr, expected_from),
          f"From: {from_addr}")

    body_lower = (body or "").lower()
    n = len(expected)
    body_clean = re.sub(r'\s+', ' ', body_lower)
    # Accept the expected count, the count actually written into the Word doc, and
    # the all-assignments count (expected + the optional 0-submission row). Also
    # accept English word forms of the number (e.g. 'twelve assignments').
    accepted_counts = {n, n + 1}
    if doc_rows is not None:
        accepted_counts.add(doc_rows)
    count_alts = []
    for c in sorted(accepted_counts):
        count_alts.append(str(c))
        w = _num_word(c)
        if w:
            count_alts.append(w.replace("-", r"[-\s]"))
    count_pat = r'\b(?:' + '|'.join(count_alts) + r')\b'
    has_count = re.search(count_pat, body_clean) is not None
    check(f"Email mentions count of {n} assignments", has_count, f"Body[:200]: {body_lower[:200]}")

    if top_avg:
        # mention the assignment: tolerant name match (whitespace/punctuation-free)
        # or the unique avg-score value for that assignment.
        name_ok = _norm_alpha(top_avg[0]) in _norm_alpha(body)
        val_ok = _mentions_number(body, top_avg[2])
        has_top_avg = name_ok or val_ok
        check(f"Email mentions highest-avg assignment ({top_avg[0]})",
              has_top_avg, f"Body[:200]: {body_lower[:200]}")
    if top_late:
        name_ok = _norm_alpha(top_late[0]) in _norm_alpha(body)
        val_ok = _mentions_number(body, top_late[3])
        has_top_late = name_ok or val_ok
        check(f"Email mentions most-late assignment ({top_late[0]})",
              has_top_late, f"Body[:200]: {body_lower[:200]}")

    # Form link check: should have the form URL or a UUID-style id, mentioned as
    # a form/survey link.
    has_form_link = (
        ("form" in body_lower or "survey" in body_lower) and
        (re.search(r'https?://[^\s]+', body or "") is not None or
         re.search(r'\b[a-f0-9-]{16,}\b', body or "") is not None)
    )
    check("Email contains form link", has_form_link,
          f"Body[:300]: {body[:300] if body else ''}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("CANVAS ASSIGNMENT FEEDBACK GFORM WORD - EVALUATION")
    print("=" * 70)

    expected = _fetch_expected_from_db() or EXPECTED_ASSIGNMENTS_FALLBACK
    print(f"\nExpected {len(expected)} assignments from DB.")

    doc_rows = check_word(args.agent_workspace, expected)
    check_gform()
    check_emails(expected, doc_rows)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": overall}, f)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
