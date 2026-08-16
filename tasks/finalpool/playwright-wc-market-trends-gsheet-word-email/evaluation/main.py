"""Evaluation for playwright-wc-market-trends-gsheet-word-email."""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")
    agent_ws = args.agent_workspace or task_root

    all_errors = []

    # --- Check 1: Google Sheet (check via DB) ---
    print("Checking Google Sheet...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Check spreadsheet exists
        cur.execute("""
            SELECT id, title FROM gsheet.spreadsheets
            WHERE title ILIKE '%market%trend%'
        """)
        rows = cur.fetchall()
        if not rows:
            all_errors.append("Google Sheet 'Market Trend Analysis 2026' not found")
        else:
            ss_id = rows[0][0]
            # Check sheets exist
            cur.execute("""
                SELECT title FROM gsheet.sheets
                WHERE spreadsheet_id = %s
                ORDER BY "index"
            """, (ss_id,))
            sheet_names = [r[0].strip().lower() for r in cur.fetchall()]

            # Sheet names per task.md: Industry Trends, Internal Performance, Gap Analysis
            if not any("industry" in s and "trend" in s for s in sheet_names):
                all_errors.append(f"Sheet 'Industry Trends' not found; got {sheet_names}")
            if not any("internal" in s and "performance" in s for s in sheet_names):
                all_errors.append(f"Sheet 'Internal Performance' not found; got {sheet_names}")
            if not any("gap" in s and "analysis" in s for s in sheet_names):
                all_errors.append(f"Sheet 'Gap Analysis' not found; got {sheet_names}")
            print(f"    Found sheets: {sheet_names}")

            # Check cells for key data
            cur.execute("""
                SELECT COUNT(DISTINCT value) FROM gsheet.cells c
                JOIN gsheet.sheets s ON c.sheet_id = s.id
                WHERE s.spreadsheet_id = %s
                AND LOWER(value) IN ('audio', 'cameras', 'electronics', 'home appliances', 'tv & home theater', 'watches')
            """, (ss_id,))
            cat_count = cur.fetchone()[0]
            if cat_count < 6:
                all_errors.append(f"GSheet: found {cat_count}/6 categories in cells")

            # --- Tightened: per-category Internal_Performance value validation ---
            # Reference values from GT Market_Trend_Analysis.xlsx Internal Performance sheet
            ref_map = {}
            try:
                import openpyxl
                gt_xlsx = os.path.join(gt_dir, "Market_Trend_Analysis.xlsx")
                if os.path.isfile(gt_xlsx):
                    gt_wb = openpyxl.load_workbook(gt_xlsx, data_only=True)
                    for sn in gt_wb.sheetnames:
                        if "internal" in sn.lower() and "performance" in sn.lower():
                            ws = gt_wb[sn]
                            # Header row: Category, Product_Count, Avg_Price, Avg_Rating, Review_Count, Total_Revenue
                            for row in list(ws.iter_rows(min_row=2, values_only=True)):
                                if row and row[0]:
                                    key = str(row[0]).strip().lower()
                                    ref_map[key] = ("ref", row[1], row[2], row[3], row[4], row[5])
                            break
                    gt_wb.close()
            except Exception:
                pass

            # Find Internal Performance sheet id
            cur.execute("""
SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s
""", (ss_id,))
            sheet_meta = cur.fetchall()
            ip_sheet_id = None
            ga_sheet_id = None
            for sid, t in sheet_meta:
                tl = (t or "").lower()
                if "internal" in tl and "performance" in tl:
                    ip_sheet_id = sid
                if "gap" in tl and "analysis" in tl:
                    ga_sheet_id = sid

            if ip_sheet_id is not None:
                # Fetch all cells of Internal Performance into row map
                cur.execute("""
SELECT row_index, col_index, value FROM gsheet.cells
WHERE sheet_id = %s ORDER BY row_index, col_index
""", (ip_sheet_id,))
                ip_cells = cur.fetchall()
                ip_rows = {}
                for r, c, v in ip_cells:
                    ip_rows.setdefault(r, {})[c] = v
                # Header row
                header = ip_rows.get(0, {})
                hdr_lower = {c: str(v or "").lower() for c, v in header.items()}
                def find_col(*keys):
                    for c, v in hdr_lower.items():
                        if all(k in v for k in keys):
                            return c
                    return None
                cat_col = find_col("category")
                pc_col = find_col("product", "count")
                ap_col = find_col("avg", "price")
                ar_col = find_col("avg", "rating")
                rc_col = find_col("review", "count")
                tr_col = find_col("revenue")
                # Compare values per row
                matched_rows = 0
                for r, cols in ip_rows.items():
                    if r == 0:
                        continue
                    cat_val = cols.get(cat_col) if cat_col is not None else None
                    if not cat_val:
                        continue
                    key = str(cat_val).strip().lower()
                    ref = ref_map.get(key)
                    if ref is None:
                        continue
                    # ref = (cat, product_count, avg_price, avg_rating, review_count, total_revenue)
                    try:
                        pc = int(float(cols.get(pc_col) or 0)) if pc_col is not None else None
                        ap = float(cols.get(ap_col) or 0) if ap_col is not None else None
                        ar = float(cols.get(ar_col) or 0) if ar_col is not None else None
                        rc = int(float(cols.get(rc_col) or 0)) if rc_col is not None else None
                        tr = float(cols.get(tr_col) or 0) if tr_col is not None else None
                    except (TypeError, ValueError):
                        continue
                    ref_pc = int(ref[1])
                    ref_ar = float(ref[3])
                    ref_rc = int(ref[4])
                    # Product_Count and Review_Count are deterministic - must match GT
                    # Avg_Rating tolerance 0.2 (formulas vary)
                    ok_pc = pc is None or pc == ref_pc
                    ok_ar = ar is None or abs(ar - ref_ar) <= 0.2
                    ok_rc = rc is None or abs(rc - ref_rc) <= max(2, ref_rc * 0.05)
                    if ok_pc and ok_ar and ok_rc:
                        matched_rows += 1
                if matched_rows < 4:
                    all_errors.append(
                        f"Internal Performance value-match insufficient: {matched_rows}/6 categories matched (Product_Count + Avg_Rating + Review_Count) within tolerance"
                    )

            # --- Tightened: Gap Analysis Market_Position value validation ---
            if ga_sheet_id is not None:
                cur.execute("""
SELECT row_index, col_index, value FROM gsheet.cells
WHERE sheet_id = %s ORDER BY row_index, col_index
""", (ga_sheet_id,))
                ga_cells = cur.fetchall()
                ga_rows = {}
                for r, c, v in ga_cells:
                    ga_rows.setdefault(r, {})[c] = v
                header = ga_rows.get(0, {})
                hdr_lower = {c: str(v or "").lower() for c, v in header.items()}
                def find_col2(*keys):
                    for c, v in hdr_lower.items():
                        if all(k in v for k in keys):
                            return c
                    return None
                pos_col = find_col2("market", "position")
                ir_col = find_col2("internal", "rating")
                is_col = find_col2("industry", "sentiment")
                gap_col = find_col2("gap")
                # Verify position labels appear and are valid
                position_correct = 0
                position_total = 0
                if pos_col is not None and gap_col is not None:
                    for r, cols in ga_rows.items():
                        if r == 0:
                            continue
                        pos_val = str(cols.get(pos_col) or "").strip().lower()
                        try:
                            gap = float(cols.get(gap_col))
                        except (TypeError, ValueError):
                            continue
                        position_total += 1
                        if gap >= 0 and pos_val == "strong":
                            position_correct += 1
                        elif -1 <= gap < 0 and pos_val == "needs attention":
                            position_correct += 1
                        elif gap < -1 and pos_val == "critical":
                            position_correct += 1
                if position_total < 1 or position_correct < min(4, position_total):
                    all_errors.append(
                        f"Gap Analysis Market_Position labels incorrect: {position_correct}/{position_total} match gap rules"
                    )

        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking GSheet: {e}")

    # --- Check 2: Word document ---
    print("Checking Word document...")
    doc_path = os.path.join(agent_ws, "Market_Insights_Report.docx")
    if not os.path.exists(doc_path):
        all_errors.append("Market_Insights_Report.docx not found")
    else:
        try:
            from docx import Document
            doc = Document(doc_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            full_lower = full_text.lower()

            # Check for key content
            categories = ["audio", "cameras", "electronics", "home appliances", "watches"]
            found_cats = sum(1 for c in categories if c in full_lower)
            if found_cats < 4:
                all_errors.append(f"Word doc mentions only {found_cats}/5 categories")

            if "trend" not in full_lower and "market" not in full_lower:
                all_errors.append("Word doc missing market/trend references")

            if "recommendation" not in full_lower and "strategic" not in full_lower:
                all_errors.append("Word doc missing recommendations section")

            if "emerging" not in full_lower:
                all_errors.append("Word doc missing emerging trends section")

        except Exception as e:
            all_errors.append(f"Error reading Word doc: {e}")

    # --- Check 3: Email sent ---
    print("Checking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        # Per task.md: subject = "Market Trends Analysis - March 2026"
        # to product-team@company.com
        cur.execute("""
            SELECT subject, body_text FROM email.messages
            WHERE to_addr::text ILIKE '%product-team@company.com%'
            AND subject ILIKE '%market%trend%analysis%'
        """)
        rows = cur.fetchall()
        if not rows:
            all_errors.append(
                "No email sent to product-team@company.com with subject "
                "containing 'Market Trends Analysis'"
            )
        else:
            print(f"    Email found ({len(rows)})")
            subj, body = rows[0]
            body_lower = (body or "").lower()
            # Body should mention top opportunity, strongest internal category,
            # and emerging trends
            if not any(k in body_lower for k in ["opportunity", "top category", "gap"]):
                all_errors.append("Email body lacks reference to top opportunity / gap")
            if not any(k in body_lower for k in ["emerging", "trend"]):
                all_errors.append("Email body lacks reference to emerging trends")
        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking email: {e}")

    # NOTE: task.md does NOT request an Excel file - only a Google Sheet,
    # a Word doc, and an email. XLSX check intentionally removed.

    # --- Final result ---
    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
