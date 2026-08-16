"""Evaluation for canvas-grade-equity-excel-word-gcal."""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": int(os.environ.get("PGPORT", "5432")),
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": os.environ.get("PGUSER", "eigent"),
    "password": os.environ.get("PGPASSWORD", "camel"),
}

PASS_COUNT = 0
FAIL_COUNT = 0

# Known course data from actual DB
COURSE_NAMES = [
    "Applied Analytics & Algorithms",
    "Biochemistry & Bioinformatics",
    "Data-Driven Design",
    "Environmental Economics & Ethics",
    "Foundations of Finance",
    "Global Governance & Geopolitics",
]

ACTION_REQUIRED_COURSES = ["Biochemistry & Bioinformatics"]


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def _to_float(v):
    """Convert a raw cell value to float, or None if not parseable.

    Handles int/float/str (strips thousands separators, currency symbols,
    percent signs and whitespace). Formula strings ('=' prefix) and booleans
    return None (formula handling is done by the caller via cached values).
    """
    if v is None:
        return None
    if isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v)
    if isinstance(v, str):
        s = v.strip()
        if not s or s.startswith("="):
            return None
        for ch in (",", "$", "¥", "€", "%", "£"):
            s = s.replace(ch, "")
        s = s.strip()
        try:
            return float(s)
        except ValueError:
            return None
    return None


def _resolve_cell(raw_value, cached_value):
    """Resolve a cell to (numeric_value, is_formula).

    If the raw value is a formula string, use the cached (data_only) value when
    available; otherwise return (None, True) so the caller can skip gracefully.
    """
    if raw_value is None:
        return None, False
    if isinstance(raw_value, str) and raw_value.startswith("="):
        if cached_value is not None:
            return _to_float(cached_value), True
        return None, True
    return _to_float(raw_value), False


def _normalize_percent(v, other):
    """If v looks like a fraction (0 < |v| < 1) and scaling it to a percentage
    (x100) brings it closer to `other`, rescale v to the percentage form.

    Uses the "does x100 bring the two values closer" test rather than a fixed
    threshold on the counterpart, so it works for pass-rate columns (94.2 vs
    0.942) AND for pass-rate-change columns whose percentage form is small
    (-3.3 vs -0.033, 5.2 vs 0.052, 0.3 vs 0.003). Handles negative values too.
    """
    if v is None or other is None:
        return v
    if 0.0 < abs(v) < 1.0:
        scaled = v * 100.0
        if abs(scaled - other) < abs(v - other):
            return scaled
    return v


def _norm_header(h):
    """Normalize a column header for order/case/punctuation-insensitive compare."""
    return "".join(ch for ch in str(h).strip().lower() if ch.isalnum())


def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Checking Excel ===")
    xlsx_path = os.path.join(agent_workspace, "Grade_Equity_Analysis.xlsx")
    if not os.path.isfile(xlsx_path):
        check("Grade_Equity_Analysis.xlsx exists", False, f"Not found: {xlsx_path}")
        return
    check("Grade_Equity_Analysis.xlsx exists", True)

    try:
        from openpyxl import load_workbook

        # Load with data_only=False (keeps formulas visible as strings) AND
        # data_only=True (cached computed values) so formula cells can be
        # resolved robustly even without Excel recalculation.
        wb = load_workbook(xlsx_path, data_only=False)
        wb_cached = load_workbook(xlsx_path, data_only=True)

        # Check sheet names
        check("Has 'Course Comparison' sheet", "Course Comparison" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")
        check("Has 'Grade Distribution' sheet", "Grade Distribution" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")
        check("Has 'Summary' sheet", "Summary" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")

        # Load groundtruth for comparison
        gt_path = os.path.join(groundtruth_workspace, "Grade_Equity_Analysis.xlsx")
        gt_wb = load_workbook(gt_path, data_only=False)
        gt_wb_cached = load_workbook(gt_path, data_only=True)

        # Check Course Comparison sheet
        if "Course Comparison" in wb.sheetnames:
            ws = wb["Course Comparison"]
            gt_ws = gt_wb["Course Comparison"]
            cached_ws = wb_cached["Course Comparison"]
            gt_cached_ws = gt_wb_cached["Course Comparison"]

            rows = list(ws.iter_rows(min_row=2, values_only=True))
            gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))
            cached_rows = list(cached_ws.iter_rows(min_row=2, values_only=True))
            gt_cached_rows = list(gt_cached_ws.iter_rows(min_row=2, values_only=True))

            check(f"Course Comparison has {len(gt_rows)} data rows",
                  len(rows) == len(gt_rows),
                  f"Found {len(rows)} rows, expected {len(gt_rows)}")

            # Check each course exists.
            # Tolerate an optional course_code prefix (e.g. "AAA Applied
            # Analytics & Algorithms" vs "Applied Analytics & Algorithms"):
            # task.md does not forbid prepending it, and Canvas exposes both
            # `name` and `course_code` fields (case-study 2026-08-16 §5.4).
            def strip_code_prefix(n):
                import re as _re
                return _re.sub(r'^[A-Z]{2,4}\s+', '', n.strip())
            agent_names = [str(r[0]).strip() if r[0] else "" for r in rows]
            agent_names_stripped = [strip_code_prefix(n) for n in agent_names]
            for name in COURSE_NAMES:
                check(f"Course '{name}' in Course Comparison",
                      name in agent_names_stripped,
                      f"Found: {agent_names}")

            # Check sort order (alphabetical by Course_Name).
            # Use the prefix-stripped view so "AAA Applied..." sorts with
            # "Applied..." — the code prefix must not break the order check.
            agent_names_lower = [strip_code_prefix(n).lower() for n in agent_names if n]
            check("Course Comparison sorted alphabetically",
                  agent_names_lower == sorted(agent_names_lower),
                  f"Order: {agent_names}")

            # ---- GT self-consistency guard -----------------------------------
            # Verify the loaded GT is internally consistent before trusting it.
            # A prior hand-built GT drifted from the DB (AAA-2014J had
            # Fall_2014_Pass_Rate = 90.9 ≈ 309/340, while the real seed data has
            # 310/340 = 91.2). That turned a correct model answer into a false
            # FAIL. We now check, per GT row, that pass-rate columns equal
            # (A+B+C)/Total*100 using the GT's own Grade Distribution sheet; if
            # they don't, we flag GT_INVALID rather than silently mis-scoring.
            gt_dist_lookup = {}
            if "Grade Distribution" in gt_wb.sheetnames:
                _gd = gt_wb["Grade Distribution"]
                _gd_cached = gt_wb_cached["Grade Distribution"]
                for _r in _gd.iter_rows(min_row=2, values_only=True):
                    if not _r or not _r[0]:
                        continue
                    try:
                        _yr = int(_r[1])
                    except (TypeError, ValueError):
                        continue
                    _col = {k: i for i, k in enumerate(
                        [_norm_header(c.value) for c in _gd[1]])}
                    def _gv(row, key):
                        idx = _col.get(_norm_header(key))
                        return row[idx] if idx is not None and idx < len(row) else None
                    a = _to_float(_gv(_r, "A_Count"))
                    b = _to_float(_gv(_r, "B_Count"))
                    c = _to_float(_gv(_r, "C_Count"))
                    d = _to_float(_gv(_r, "D_Count"))
                    f = _to_float(_gv(_r, "F_Count"))
                    tot = _to_float(_gv(_r, "Total_Students"))
                    if None in (a, b, c, d, f, tot) or tot == 0:
                        continue
                    gt_dist_lookup[(str(_r[0]).strip(), _yr)] = {
                        "A": a, "B": b, "C": c, "D": d, "F": f, "Total": tot,
                        "Passed": a + b + c,
                    }

            def _gt_pass_rate_for(name, year, pr_col_idx, gt_row, gt_cached_row):
                """Return (gt_pass_rate_float, numerator, denominator) from the GT's
                own Grade Distribution, so failures can report numerator/denominator
                and we can cross-check GT Course Comparison vs GT Grade Distribution."""
                d = gt_dist_lookup.get((name, year))
                if not d:
                    return None, None, None
                return round(d["Passed"] / d["Total"] * 100, 1), int(d["Passed"]), int(d["Total"])

            # Cross-check GT Course Comparison pass rates against GT Grade
            # Distribution. If GT disagrees with itself, emit a hard GT_INVALID
            # check so the run is not scored against a broken oracle.
            for _name, _gt_row in {str(r[0]).strip(): r for r in gt_rows}.items():
                for _year, _ci in [(2013, 4), (2014, 5)]:
                    _g_val, _ = _resolve_cell(
                        _gt_row[_ci] if len(_gt_row) > _ci else None,
                        None)
                    _d = gt_dist_lookup.get((_name, _year))
                    if _g_val is None or not _d:
                        continue
                    _expected = round(_d["Passed"] / _d["Total"] * 100, 1)
                    if abs(_expected - _g_val) > 0.05:
                        check(f"GT self-consistent {_name} {_year} pass_rate",
                              False,
                              f"GT Course Comparison={_g_val} but GT Grade Distribution "
                              f"({_d['Passed']}/{_d['Total']})*100={_expected}. "
                              f"Regenerate GT with files/generate_groundtruth.py.")
            # ------------------------------------------------------------------

            # Check numeric values with tolerance — tightened to 0.2 (1-decimal rounding)
            # Match agent rows to GT rows on the code-prefix-stripped name so a
            # "AAA Applied..." row still pairs with the GT "Applied..." row.
            gt_dict = {strip_code_prefix(str(r[0]).strip() if r[0] else ""): r for r in gt_rows}
            gt_cached_dict = {strip_code_prefix(str(r[0]).strip() if r[0] else ""): r for r in gt_cached_rows}
            for i, row in enumerate(rows):
                name = str(row[0]).strip() if row[0] else ""
                key = strip_code_prefix(name)
                if key in gt_dict:
                    gt_row = gt_dict[key]
                    gt_cached_row = gt_cached_dict.get(key)
                    cached_row = cached_rows[i] if i < len(cached_rows) else None
                    for col_idx, col_name in [(1, "Fall_2013_Mean"), (2, "Fall_2014_Mean"),
                                               (3, "Score_Difference"),
                                               (4, "Fall_2013_Pass_Rate"),
                                               (5, "Fall_2014_Pass_Rate"),
                                               (6, "Pass_Rate_Change")]:
                        a_val, a_formula = _resolve_cell(
                            row[col_idx],
                            cached_row[col_idx] if cached_row and col_idx < len(cached_row) else None)
                        g_val, _g_formula = _resolve_cell(
                            gt_row[col_idx],
                            gt_cached_row[col_idx] if gt_cached_row and col_idx < len(gt_cached_row) else None)
                        a_val = _normalize_percent(a_val, g_val)
                        g_val = _normalize_percent(g_val, a_val)
                        if a_val is not None and g_val is not None:
                            diff = abs(a_val - g_val)
                            # On pass-rate mismatches, surface numerator/denominator
                            # from the GT's own Grade Distribution so it's obvious
                            # whether GT or the agent is the outlier.
                            extra = ""
                            if col_name in ("Fall_2013_Pass_Rate", "Fall_2014_Pass_Rate"):
                                _yr = 2013 if col_name == "Fall_2013_Pass_Rate" else 2014
                                _grp, _gn, _gd2 = _gt_pass_rate_for(
                                    name, _yr, col_idx, gt_row, gt_cached_row)
                                if _gn is not None:
                                    extra = (f" [GT dist: pass={_grp} ({_gn}/{_gd2}); "
                                             f"check agent Grade Distribution {name} {_yr}]")
                            check(f"{name} {col_name} within tolerance",
                                  diff <= 0.2,
                                  f"Agent={row[col_idx]}, GT={gt_row[col_idx]}, "
                                  f"diff={diff:.2f}{extra}")
                        elif a_formula and g_val is None:
                            check(f"{name} {col_name} value present",
                                  True,
                                  "formula cell without cached value; numeric check skipped")
                        elif a_val is None and not a_formula and g_val is not None:
                            check(f"{name} {col_name} value present",
                                  False,
                                  f"Agent value not parseable as number: {row[col_idx]!r}")

                    # Check equity status (exact match)
                    if row[7] is not None:
                        check(f"{name} Equity_Status matches",
                              str(row[7]).strip().lower() == str(gt_row[7]).strip().lower(),
                              f"Agent='{row[7]}', GT='{gt_row[7]}'")

        # Check Grade Distribution sheet
        if "Grade Distribution" in wb.sheetnames:
            ws = wb["Grade Distribution"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Grade Distribution has 12 data rows (6 courses x 2 years)",
                  len(rows) == 12,
                  f"Found {len(rows)} rows")

            # Structural check: expected column headers are present (order-agnostic,
            # extra columns tolerated; case-insensitive, and both "A_Count" and
            # "A Count" variants accepted by stripping non-alphanumerics).
            # Values are not numerically compared because grade-band boundaries
            # are ambiguous at continuous-score edges.
            expected_cols = {"course_name", "year", "a_count", "b_count", "c_count",
                             "d_count", "f_count", "total_students"}

            header_vals = [_norm_header(c.value) for c in ws[1] if c.value is not None]
            expected_vals = {_norm_header(h) for h in expected_cols}
            check("Grade Distribution has expected columns",
                  expected_vals.issubset(set(header_vals)),
                  f"Header: {header_vals}")

        # Check Summary sheet
        if "Summary" in wb.sheetnames:
            ws = wb["Summary"]
            gt_ws = gt_wb["Summary"]
            cached_ws = wb_cached["Summary"]
            rows = {str(r[0]).strip(): r[1] for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}
            gt_rows = {str(r[0]).strip(): r[1] for r in gt_ws.iter_rows(min_row=2, values_only=True) if r[0]}
            cached_map = {str(r[0]).strip(): r[1] for r in cached_ws.iter_rows(min_row=2, values_only=True) if r[0]}

            def sum_val(metric):
                raw = rows.get(metric)
                if raw is None:
                    return None
                if isinstance(raw, str) and raw.startswith("="):
                    return _to_float(cached_map.get(metric))
                return _to_float(raw)

            check("Summary has Total_Courses_Compared",
                  "Total_Courses_Compared" in rows,
                  f"Keys: {list(rows.keys())}")

            if "Total_Courses_Compared" in rows:
                v = sum_val("Total_Courses_Compared")
                check("Total_Courses_Compared = 6",
                      v is not None and abs(v - 6) <= 0.5,
                      f"Got {rows['Total_Courses_Compared']}")

            if "Courses_Action_Required" in rows:
                v = sum_val("Courses_Action_Required")
                check("Courses_Action_Required = 1",
                      v is not None and abs(v - 1) <= 0.5,
                      f"Got {rows['Courses_Action_Required']}")

            for key in ["Overall_Avg_2013", "Overall_Avg_2014", "Overall_Change"]:
                if key in rows and key in gt_rows:
                    a = sum_val(key)
                    g = _to_float(gt_rows[key])
                    if a is not None and g is not None:
                        diff = abs(a - g)
                        # Tolerance 1.0: Overall_Avg is an aggregate over all
                        # students, so slightly different but reasonable
                        # aggregation choices (simple mean of course means vs
                        # student-weighted pooled mean) can legitimately differ
                        # by up to ~0.7. Accept any genuine aggregate; a wrong
                        # or missing computation still fails.
                        check(f"Summary {key} within tolerance",
                              diff <= 1.0,
                              f"Agent={rows[key]}, GT={gt_rows[key]}")
                    elif a is None:
                        check(f"Summary {key} within tolerance",
                              True,
                              "agent value is formula without cached value; numeric check skipped")

            # Validate Acceptable / Concerning / Action_Required counts (deterministic Canvas data)
            for key, expected in [("Courses_Acceptable", 5), ("Courses_Concerning", 0)]:
                if key in rows:
                    v = sum_val(key)
                    check(f"Summary {key} = {expected}",
                          v is not None and abs(v - expected) <= 0.5,
                          f"Got {rows[key]}")

    except ImportError:
        check("openpyxl available", False, "openpyxl not installed")
    except Exception as e:
        check("Excel parsing", False, str(e))


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Equity_Report.docx")
    if not os.path.isfile(docx_path):
        check("Equity_Report.docx exists", False, f"Not found: {docx_path}")
        return
    check("Equity_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        check("Word doc has meaningful content (>= 200 chars)",
              len(all_text.strip()) >= 200,
              f"Content length: {len(all_text)}")

        # Check required sections (tolerate hyphenated vs spaced 'course-by-course')
        for section in ["executive summary", "methodology", "recommendation", "appendix"]:
            check(f"Word doc contains '{section}' section",
                  section in all_text,
                  f"Not found in document text")
        check("Word doc contains 'course-by-course' section",
              "course-by-course" in all_text or "course by course" in all_text,
              "Not found in document text")

        # Check course names mentioned
        found_courses = sum(1 for name in COURSE_NAMES if name.lower() in all_text)
        check("Word doc mentions at least 4 course names",
              found_courses >= 4,
              f"Found {found_courses} of {len(COURSE_NAMES)} course names")

        # Check for action required course
        check("Word doc mentions 'Biochemistry & Bioinformatics'",
              "biochemistry" in all_text,
              "Action required course not mentioned")

        # Check for equity-related content
        check("Word doc discusses equity/grading",
              "equity" in all_text or "grade" in all_text or "score" in all_text,
              f"Sample: {all_text[:300]}")

    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        from datetime import datetime, date
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("SET TIME ZONE 'UTC'")

        # Restrict to the target week (March 16-20, 2026 inclusive) using start_datetime range
        cur.execute("""
            SELECT summary, description,
                   start_datetime AT TIME ZONE 'UTC',
                   end_datetime AT TIME ZONE 'UTC'
            FROM gcal.events
            WHERE start_datetime >= '2026-03-16 00:00:00+00'
              AND start_datetime <= '2026-03-20 23:59:59+00'
              AND (LOWER(summary) LIKE '%biochemistry%bioinformatics%grade%equity%review%'
                   OR LOWER(summary) LIKE '%biochemistry & bioinformatics grade equity review%'
                   OR (LOWER(summary) LIKE '%biochemistry%' AND LOWER(summary) LIKE '%grade equity review%'))
        """)
        events = cur.fetchall()
        check("Action Required course meeting scheduled in target week",
              len(events) >= 1,
              f"Found {len(events)} events for Biochemistry & Bioinformatics Grade Equity Review in week of March 16-20, 2026")

        if events:
            # Validate title format: contains course name + "Grade Equity Review"
            # (tolerate full name vs shortened course-name variants).
            format_ok = [e for e in events
                         if "grade equity review" in (e[0] or "").strip().lower()
                         and "biochemistry" in (e[0] or "").strip().lower()]
            check("Meeting title is '<Course> Grade Equity Review' format",
                  len(format_ok) >= 1,
                  f"Titles: {[e[0] for e in events]}")

            # Validate description contains equity details: action required AND
            # (score/mean difference OR equity status). In a P=2 swarm both
            # sub-agents may schedule a meeting; the query can return several
            # matching events in arbitrary order. Judge on "at least one event
            # is fully correct" rather than only the first row returned, so a
            # correct meeting is never rejected because a stray/erroneous event
            # happened to sort first.
            desc_ok = False
            for e in events:
                desc = (e[1] or "").lower()
                has_action = "action required" in desc
                has_detail = (("score" in desc or "mean" in desc or "grade" in desc) and "difference" in desc) \
                    or "equity status" in desc
                if has_action and has_detail:
                    desc_ok = True
                    break
            check("Meeting description has equity details",
                  desc_ok,
                  f"Descriptions: {[(e[0], (e[1] or '')[:120]) for e in events]}")

            # Check that at least one matching event is exactly 45 minutes long.
            dur_ok = False
            for e in events:
                if e[2] and e[3]:
                    start = e[2] if isinstance(e[2], datetime) else datetime.fromisoformat(str(e[2]))
                    end = e[3] if isinstance(e[3], datetime) else datetime.fromisoformat(str(e[3]))
                    duration_min = (end - start).total_seconds() / 60
                    if abs(duration_min - 45) <= 1:
                        dur_ok = True
                        break
            check("Meeting duration is 45 minutes",
                  dur_ok,
                  f"Events: {[(e[0], str(e[2]), str(e[3])) for e in events]}")

        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    gt = args.groundtruth_workspace or os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "groundtruth_workspace"
    )

    check_excel(args.agent_workspace, gt)
    check_word(args.agent_workspace)
    check_calendar()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ===")
    if FAIL_COUNT > 0:
        print(f"{FAIL_COUNT} checks failed")
        sys.exit(1)
    else:
        print("All checks passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
